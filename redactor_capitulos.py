"""
redactor_capitulos.py — Redactor Profesional de Capítulos ETP (MIGRADO · Paso 18)
Genera capítulos técnicos profundos con estructura editorial, analogías
didácticas y exportación Word profesional (portada, índice TOC, paginación,
tablas, bloques de código y formato enriquecido).
• IA vía core/ia (modo texto/Markdown, reintento por truncamiento, auditoría).
• Super interfaz: hero, métricas, toasts, editor/vista previa, width="stretch".
"""
import re
import logging
from io import BytesIO
from datetime import datetime

import streamlit as st
from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn

from core import ia

logger = logging.getLogger("redactor_capitulos")
logging.basicConfig(level=logging.INFO)

# ═══════════════════════════════════════════════════════════════════════════
# ESTILOS (SUPER INTERFAZ)
# ═══════════════════════════════════════════════════════════════════════════
st.markdown("""
<style>
@import url('https://fonts.googleapis.com/css2?family=Inter:wght@300;400;500;600;700;800&display=swap');
html, body, [class*="css"] { font-family: 'Inter', sans-serif; background-color: #F8FAFC; color: #1E293B; }
.red-hero { background: linear-gradient(135deg, #0F172A 0%, #4C1D95 55%, #7C3AED 100%); color: #fff;
    padding: 1.8rem; border-radius: 18px; margin-bottom: 1.25rem; box-shadow: 0 18px 35px rgba(15,23,42,.18); }
.red-title { font-size: 2rem; font-weight: 800; letter-spacing: -0.02em; }
.red-sub { opacity: .88; font-size: 1rem; margin-top: .3rem; }
.section-title { color: #1D4ED8; font-weight: 700; font-size: 1.12rem; border-bottom: 2px solid #DBEAFE;
    padding-bottom: 8px; margin: 1.2rem 0 .9rem 0; }
.editor-container { background-color: #FFFFFF; padding: 20px; border-radius: 8px; border: 1px solid #E2E8F0; }
.metric-card { background:#fff; border:1px solid #E2E8F0; border-top:4px solid #7C3AED; border-radius:12px;
    padding:14px 16px; box-shadow:0 4px 12px rgba(15,23,42,.06); text-align:center; }
.metric-value { font-size:1.9rem; font-weight:800; color:#0F172A; }
.metric-label { font-size:.78rem; font-weight:800; color:#64748B; text-transform:uppercase; letter-spacing:.05em; }
</style>
""", unsafe_allow_html=True)

ia.panel_sidebar_ia("Redactor Profundo")

st.markdown("""
<div class="red-hero">
    <div class="red-title">✒️ Redactor Profesional de Capítulos ETP</div>
    <div class="red-sub">Capítulos técnicos con estructura editorial, analogías y exportación Word profesional</div>
</div>
""", unsafe_allow_html=True)

# ═══════════════════════════════════════════════════════════════════════════
# PROMPT MAESTRO (conservado del original)
# ═══════════════════════════════════════════════════════════════════════════
def construir_prompt_capitulo(titulo_capitulo: str, subtemas: str, nivel_tecnico: str) -> str:
    return f"""Actúa como Autor y Catedrático Senior de Libros de Texto para Ingeniería y Educación Técnico Profesional.
Tu tarea es redactar un capítulo completo, profundo, didáctico y con calidad editorial profesional.
IMPORTANTE:
Debes responder en español formal.
Devuelve exclusivamente Markdown válido.
No envuelvas todo el capítulo en una cerca de código.
Si necesitas incluir ejemplos de código, usa bloques Markdown con ``` con moderación.
No escribas comentarios, explicaciones previas ni texto fuera del capítulo.
No uses HTML.
No dejes secciones vacías ni placeholders.
DATOS DEL CAPÍTULO:
TEMA DEL CAPÍTULO: {titulo_capitulo}
SUBTEMAS A DESARROLLAR:
{subtemas}
NIVEL TÉCNICO: {nivel_tecnico}
ESTRUCTURA OBLIGATORIA DEL CAPÍTULO:
Debes respetar exactamente esta estructura de libro profesional:
# {titulo_capitulo}
## Resumen del capítulo
## Objetivos de aprendizaje
## Prerrequisitos
## Introducción
## Desarrollo técnico
## Caso práctico aplicado
## Errores frecuentes y cómo evitarlos
## Síntesis del capítulo
## Glosario técnico
## Autoevaluación
## Actividades prácticas propuestas
## Referencias y fuentes recomendadas
REGLAS DE REDACCIÓN DEL DESARROLLO TÉCNICO:
En la sección "Desarrollo técnico", convierte cada subtema en una sección numerada con profundidad real (1.1, 1.2, 1.3...).
Cada sección debe comenzar con una definición técnica formal y luego explicar el concepto con claridad pedagógica.
Por cada concepto técnico complejo, DEBES incluir un subapartado con el formato "> La analogía de [metáfora apropiada]".
Usa analogías potentes y memorables, sin sacrificar rigor técnico.
Incluye ejemplos técnicos concretos vinculados a ingeniería, informática, automatización, redes, sistemas operativos, bases de datos o ETP según el tema.
Incluye al menos una tabla comparativa en Markdown si el tema lo permite (ventajas vs desventajas, concepto A vs B, buenas vs malas prácticas, componentes/funciones/ejemplos).
Usa negritas para términos clave. No escribas relleno; cada párrafo debe aportar contenido.
Mantén un flujo narrativo profesional, coherente y bien conectado entre secciones.
Adapta la profundidad al nivel técnico indicado:
- Técnico Básico: lenguaje accesible, ejemplos simples y explicaciones paso a paso.
- Bachillerato Técnico: equilibrio entre fundamentos y aplicación.
- Técnico Superior / Universitario: mayor formalismo, arquitectura, modelos, buenas prácticas y casos técnicos.
En "Autoevaluación", incluye entre 5 y 10 preguntas variadas (comprensión conceptual, aplicación técnica y análisis).
En "Actividades prácticas propuestas", incluye actividades realizables en taller, laboratorio, aula virtual o proyecto práctico.
En "Referencias y fuentes recomendadas", sugiere libros, normas, documentación técnica o fuentes confiables relacionadas con el tema.
REDACCIÓN FINAL:
Redacta ahora el capítulo completo, listo para ser convertido en un documento Word profesional.
"""

def limpiar_markdown_respuesta(texto: str) -> str:
    if not texto:
        return ""
    texto = texto.strip()
    texto = re.sub(r"^```(?:markdown|md)?\s*", "", texto, flags=re.IGNORECASE)
    texto = re.sub(r"\s*```$", "", texto)
    texto = texto.replace("\r\n", "\n")
    texto = re.sub(r"\n{3,}", "\n\n", texto)
    return texto.strip()

# ═══════════════════════════════════════════════════════════════════════════
# UTILIDADES WORD PROFESIONAL (conservadas del original)
# ═══════════════════════════════════════════════════════════════════════════
INLINE_MD = re.compile(r"(\*\*\*.+?\*\*\*|\*\*.+?\*\*|\*.+?\*|__.+?__|_.+?_|`[^`]+`)")

def set_paragraph_shading(paragraph, fill_hex: str):
    pPr = paragraph._p.get_or_add_pPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear"); shd.set(qn("w:color"), "auto"); shd.set(qn("w:fill"), fill_hex)
    pPr.append(shd)

def set_cell_shading(cell, fill_hex: str):
    tcPr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear"); shd.set(qn("w:color"), "auto"); shd.set(qn("w:fill"), fill_hex)
    tcPr.append(shd)

def agregar_numero_pagina(footer):
    p = footer.paragraphs[0] if footer.paragraphs else footer.add_paragraph()
    p.text = ""
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run()
    fldChar1 = OxmlElement("w:fldChar"); fldChar1.set(qn("w:fldCharType"), "begin")
    instrText = OxmlElement("w:instrText"); instrText.text = "PAGE"
    fldChar2 = OxmlElement("w:fldChar"); fldChar2.set(qn("w:fldCharType"), "end")
    run._r.append(fldChar1); run._r.append(instrText); run._r.append(fldChar2)

def configurar_documento_libro(doc, titulo_capitulo: str):
    doc.core_properties.title = titulo_capitulo
    doc.core_properties.author = "Redactor Profesional ETP"
    for section in doc.sections:
        section.top_margin = Cm(2.5); section.bottom_margin = Cm(2.5)
        section.left_margin = Cm(3.0); section.right_margin = Cm(2.5)
        section.header_distance = Cm(1.5); section.footer_distance = Cm(1.25)
        section.different_first_page_header_footer = True
        header = section.header
        hp = header.paragraphs[0] if header.paragraphs else header.add_paragraph()
        hp.text = ""; hp.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        run = hp.add_run(titulo_capitulo)
        run.font.size = Pt(9); run.font.color.rgb = RGBColor(100, 116, 139)
        first_header = section.first_page_header
        fhp = first_header.paragraphs[0] if first_header.paragraphs else first_header.add_paragraph()
        fhp.text = ""
        agregar_numero_pagina(section.footer)
        first_footer = section.first_page_footer
        ffp = first_footer.paragraphs[0] if first_footer.paragraphs else first_footer.add_paragraph()
        ffp.text = ""
    styles = doc.styles
    normal = styles["Normal"]
    normal.font.name = "Calibri"; normal.font.size = Pt(11)
    normal.font.color.rgb = RGBColor(30, 41, 59)
    normal.paragraph_format.line_spacing = 1.15
    normal.paragraph_format.space_after = Pt(6); normal.paragraph_format.space_before = Pt(0)
    headings = {1: (22, RGBColor(15, 23, 42)), 2: (16, RGBColor(29, 78, 216)),
                3: (13, RGBColor(30, 58, 138)), 4: (12, RGBColor(51, 65, 85))}
    for level, (size, color) in headings.items():
        try:
            style = styles[f"Heading {level}"]
            style.font.name = "Calibri"; style.font.size = Pt(size)
            style.font.bold = True; style.font.color.rgb = color
            style.paragraph_format.space_before = Pt(18 if level == 1 else 12)
            style.paragraph_format.space_after = Pt(10 if level == 1 else 6)
            style.paragraph_format.keep_with_next = True
        except KeyError:
            pass

def agregar_portada(doc, titulo_capitulo: str):
    for _ in range(7):
        doc.add_paragraph()
    p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("CAPÍTULO"); run.font.size = Pt(14); run.bold = True
    run.font.color.rgb = RGBColor(29, 78, 216)
    p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(titulo_capitulo); run.font.size = Pt(28); run.bold = True
    run.font.color.rgb = RGBColor(15, 23, 42)
    p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("Documento técnico-profesional"); run.font.size = Pt(12)
    run.font.color.rgb = RGBColor(71, 85, 105)
    p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(datetime.now().strftime("%d/%m/%Y")); run.font.size = Pt(10)
    run.font.color.rgb = RGBColor(100, 116, 139)

def agregar_indice(doc):
    doc.add_heading("Índice", level=1)
    p = doc.add_paragraph()
    run = p.add_run()
    fldChar1 = OxmlElement("w:fldChar"); fldChar1.set(qn("w:fldCharType"), "begin")
    instrText = OxmlElement("w:instrText"); instrText.text = 'TOC \\o "1-3" \\h \\z \\u'
    fldChar2 = OxmlElement("w:fldChar"); fldChar2.set(qn("w:fldCharType"), "separate")
    t = OxmlElement("w:t"); t.text = "Actualiza este campo para mostrar el índice."
    fldChar3 = OxmlElement("w:fldChar"); fldChar3.set(qn("w:fldCharType"), "end")
    run._r.append(fldChar1); run._r.append(instrText); run._r.append(fldChar2)
    run._r.append(t); run._r.append(fldChar3)
    nota = doc.add_paragraph("Nota: en Word, haz clic derecho sobre el índice y selecciona 'Actualizar campos'.")
    if nota.runs:
        nota.runs[0].font.size = Pt(9); nota.runs[0].font.color.rgb = RGBColor(100, 116, 139)

def add_rich_text(paragraph, text: str):
    text = re.sub(r"\[([^\]]+)\]\([^)]+\)", r"\1", text)
    parts = INLINE_MD.split(text)
    for part in parts:
        if not part:
            continue
        if part.startswith("***") and part.endswith("***") and len(part) > 6:
            run = paragraph.add_run(part[3:-3]); run.bold = True; run.italic = True
        elif part.startswith("**") and part.endswith("**") and len(part) > 4:
            run = paragraph.add_run(part[2:-2]); run.bold = True
        elif part.startswith("__") and part.endswith("__") and len(part) > 4:
            run = paragraph.add_run(part[2:-2]); run.bold = True
        elif part.startswith("*") and part.endswith("*") and len(part) > 2 and not part.startswith("**"):
            run = paragraph.add_run(part[1:-1]); run.italic = True
        elif part.startswith("_") and part.endswith("_") and len(part) > 2 and not part.startswith("__"):
            run = paragraph.add_run(part[1:-1]); run.italic = True
        elif part.startswith("`") and part.endswith("`") and len(part) > 2:
            run = paragraph.add_run(part[1:-1]); run.font.name = "Consolas"; run.font.size = Pt(10)
        else:
            paragraph.add_run(part)

def add_code_block(doc, code: str):
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Cm(0.5); p.paragraph_format.right_indent = Cm(0.5)
    p.paragraph_format.space_after = Pt(10)
    set_paragraph_shading(p, "F1F5F9")
    run = p.add_run(code)
    run.font.name = "Consolas"; run.font.size = Pt(9.5)
    run.font.color.rgb = RGBColor(15, 23, 42)

def split_table_row(row: str):
    row = row.strip()
    if row.startswith("|"): row = row[1:]
    if row.endswith("|"): row = row[:-1]
    return [cell.strip() for cell in row.split("|")]

def add_markdown_table(doc, lines, i: int) -> int:
    header_line = lines[i].strip()
    if "|" not in header_line:
        return i + 1
    header = split_table_row(header_line)
    if len(header) == 0:
        return i + 1
    i += 2
    rows = []
    while i < len(lines) and "|" in lines[i] and lines[i].strip():
        rows.append(split_table_row(lines[i])); i += 1
    table = doc.add_table(rows=1 + len(rows), cols=len(header))
    try:
        table.style = "Table Grid"
    except KeyError:
        pass
    table.autofit = True
    for j, h in enumerate(header):
        cell = table.rows[0].cells[j]
        cell.text = ""
        p = cell.paragraphs[0]
        add_rich_text(p, h)
        for run in p.runs: run.bold = True
        set_cell_shading(cell, "DBEAFE")
    for r, row in enumerate(rows):
        for j in range(len(header)):
            value = row[j] if j < len(row) else ""
            cell = table.rows[r + 1].cells[j]
            cell.text = ""
            add_rich_text(cell.paragraphs[0], value)
    spacer = doc.add_paragraph(); spacer.paragraph_format.space_after = Pt(2)
    return i

def renderizar_markdown(doc, texto_markdown: str):
    lineas = texto_markdown.split("\n")
    i = 0
    h1_seen = False
    while i < len(lineas):
        linea = lineas[i].rstrip()
        if not linea.strip():
            i += 1; continue
        if linea.strip().startswith("```"):
            i += 1; codigo = []
            while i < len(lineas) and not lineas[i].strip().startswith("```"):
                codigo.append(lineas[i]); i += 1
            i += 1
            add_code_block(doc, "\n".join(codigo)); continue
        if ("|" in linea and i + 1 < len(lineas)
                and re.match(r"^\s*\|?[\s:|-]+\|?\s*$", lineas[i + 1]) and "-" in lineas[i + 1]):
            i = add_markdown_table(doc, lineas, i); continue
        if "|" in linea and re.match(r"^\s*\|?[\s:|-]+\|?\s*$", linea) and "-" in linea:
            i += 1; continue
        heading_match = re.match(r"^\s*(#{1,6})\s+(.*)", linea)
        if heading_match:
            level = min(len(heading_match.group(1)), 4)
            texto_heading = heading_match.group(2).strip()
            texto_heading = re.sub(r"[*_`]", "", texto_heading)
            if level == 1:
                if h1_seen: doc.add_page_break()
                else: h1_seen = True
            try:
                p = doc.add_heading("", level=level); p.add_run(texto_heading)
            except Exception:
                p = doc.add_paragraph(); run = p.add_run(texto_heading); run.bold = True
            i += 1; continue
        if linea.strip() in ("---", "***", "___"):
            p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            p.add_run("• • •"); i += 1; continue
        if linea.strip().startswith(">"):
            texto_cita = linea.strip().lstrip(">").strip()
            p = doc.add_paragraph()
            p.paragraph_format.left_indent = Cm(0.75); p.paragraph_format.space_after = Pt(8)
            set_paragraph_shading(p, "F8FAFC")
            add_rich_text(p, texto_cita)
            for run in p.runs: run.italic = True
            i += 1; continue
        bullet_match = re.match(r"^(\s*)([-*+])\s+(.*)", linea)
        if bullet_match:
            nivel = len(bullet_match.group(1)) // 2
            contenido = bullet_match.group(3).strip()
            p = doc.add_paragraph()
            p.paragraph_format.left_indent = Cm(0.75 + nivel * 0.6)
            p.paragraph_format.first_line_indent = Cm(-0.4)
            p.paragraph_format.space_after = Pt(3)
            add_rich_text(p, "•  " + contenido)
            i += 1; continue
        num_match = re.match(r"^(\s*)(\d+)[.)]\s+(.*)", linea)
        if num_match:
            nivel = len(num_match.group(1)) // 3
            numero = num_match.group(2); contenido = num_match.group(3).strip()
            p = doc.add_paragraph()
            p.paragraph_format.left_indent = Cm(0.75 + nivel * 0.6)
            p.paragraph_format.first_line_indent = Cm(-0.4)
            p.paragraph_format.space_after = Pt(3)
            add_rich_text(p, f"{numero}.  {contenido}")
            i += 1; continue
        p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        add_rich_text(p, linea.strip()); i += 1

def compilar_a_word(texto_markdown: str, titulo_capitulo: str) -> BytesIO:
    texto_markdown = limpiar_markdown_respuesta(texto_markdown)
    doc = Document()
    configurar_documento_libro(doc, titulo_capitulo)
    agregar_portada(doc, titulo_capitulo)
    doc.add_page_break()
    agregar_indice(doc)
    doc.add_page_break()
    renderizar_markdown(doc, texto_markdown)
    buffer = BytesIO(); doc.save(buffer); buffer.seek(0)
    return buffer

# ═══════════════════════════════════════════════════════════════════════════
# FORMULARIO
# ═══════════════════════════════════════════════════════════════════════════
if "texto_generado_capitulo" not in st.session_state:
    st.session_state.texto_generado_capitulo = ""
if "titulo_actual" not in st.session_state:
    st.session_state.titulo_actual = ""
if "flags_capitulo" not in st.session_state:
    st.session_state.flags_capitulo = {}

with st.form("form_redactor", clear_on_submit=False):
    st.markdown('<div class="section-title">✍️ 1. Definición del Capítulo</div>', unsafe_allow_html=True)
    col1, col2 = st.columns([2, 1])
    with col1:
        titulo = st.text_input("Título del Capítulo / RA", placeholder="Ej: Capítulo 3: Procesos e Hilos")
    with col2:
        nivel = st.selectbox("Nivel Técnico",
                             ["Técnico Básico", "Bachillerato Técnico", "Técnico Superior / Universitario"], index=1)
    subtemas = st.text_area(
        "Subtemas a desarrollar (Pega aquí los contenidos conceptuales/procedimentales):", height=140,
        placeholder="Ej:\n- El Concepto de Proceso frente a Programa\n- El Bloque de Control de Proceso (PCB)\n- El Cambio de Contexto\n- Estados de un proceso")
    max_tokens, temperature = ia.control_avanzado(default_tokens=16384, tope=32000, default_temp=0.3)
    st.markdown("<br>", unsafe_allow_html=True)
    submit_btn = st.form_submit_button("✒️ Redactar Capítulo Completo", type="primary", width="stretch")

if submit_btn:
    if not titulo or not subtemas:
        st.warning("⚠️ Debes proporcionar un título y los subtemas a desarrollar.")
    else:
        with st.spinner("🧠 Escribiendo contenido técnico con pedagogía, analogías y estructura editorial..."):
            texto_crudo = None
            try:
                prompt = construir_prompt_capitulo(titulo, subtemas, nivel)
                texto, flags = ia.solicitar_texto(prompt, max_tokens=max_tokens,
                                                  temperature=temperature, modulo="redactor_capitulos")
                resultado = limpiar_markdown_respuesta(texto)
                if not resultado:
                    raise ValueError("La IA devolvió una respuesta vacía.")
                st.session_state.texto_generado_capitulo = resultado
                st.session_state.titulo_actual = titulo
                st.session_state.flags_capitulo = flags
                st.toast("✅ ¡Capítulo redactado exitosamente!", icon="✒️")
                logger.info("Capítulo generado para: %s", titulo)
            except ValueError as ve:
                ia.render_error_ia(ve, texto_crudo)
            except Exception as e:
                ia.render_error_ia(e, texto_crudo)

# ═══════════════════════════════════════════════════════════════════════════
# ÁREA DE EDICIÓN, VISTA PREVIA Y DESCARGA
# ═══════════════════════════════════════════════════════════════════════════
if st.session_state.texto_generado_capitulo:
    st.markdown('<div class="section-title">📄 2. Revisión, Vista Previa y Descarga</div>', unsafe_allow_html=True)
    flags = st.session_state.flags_capitulo or {}
    if flags.get("reintento"):
        st.info("ℹ️ La primera llamada se cortó por tokens; el reintento automático tuvo éxito.")
    word_count = len(st.session_state.texto_generado_capitulo.split())
    secciones = len(re.findall(r"^#{1,4}\s", st.session_state.texto_generado_capitulo, flags=re.MULTILINE))
    m1, m2, m3 = st.columns(3)
    with m1:
        st.markdown(f'<div class="metric-card"><div class="metric-label">Palabras</div><div class="metric-value">{word_count}</div></div>', unsafe_allow_html=True)
    with m2:
        st.markdown(f'<div class="metric-card"><div class="metric-label">Lectura aprox.</div><div class="metric-value">{max(1, word_count // 200)} min</div></div>', unsafe_allow_html=True)
    with m3:
        st.markdown(f'<div class="metric-card"><div class="metric-label">Secciones</div><div class="metric-value">{secciones}</div></div>', unsafe_allow_html=True)

    tab_editor, tab_preview = st.tabs(["📝 Editor", "👁️ Vista previa"])
    with tab_editor:
        texto_editado = st.text_area("Puedes afinar detalles del texto antes de descargarlo:",
                                     value=st.session_state.texto_generado_capitulo, height=450)
        if st.button("🗑️ Limpiar texto generado", width="stretch"):
            st.session_state.texto_generado_capitulo = ""
            st.session_state.titulo_actual = ""
            st.session_state.flags_capitulo = {}
            st.rerun()
    with tab_preview:
        st.markdown(texto_editado)

    buffer_word = compilar_a_word(texto_editado, st.session_state.get("titulo_actual", "Capitulo"))
    nombre_archivo = ia.sanear_nombre_archivo(st.session_state.get("titulo_actual", "Capitulo"), "capitulo")
    st.download_button(
        label="📥 Descargar Capítulo (.docx)", data=buffer_word,
        file_name=f"{nombre_archivo}.docx",
        mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
        type="primary", width="stretch")
    st.caption("Al abrir el documento en Word, actualiza el índice haciendo clic derecho y seleccionando 'Actualizar campos'.")