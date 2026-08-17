"""
feedback.py — Feedback del Portal ETP (Nivel Dios)
Sistema de retroalimentación multidimensional con análisis IA.
• Evaluación en 6 dimensiones con estrellas interactivas.
• Almacenamiento SQLite + historial completo.
• Dashboard de analítica para coordinación.
• Análisis IA de comentarios y generación de insights.
• Exportación Word profesional.
"""
import sqlite3
import re
from datetime import datetime, date, timedelta
from io import BytesIO
from typing import Any, Dict, List, Optional

import pandas as pd
import streamlit as st
from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import parse_xml
from docx.oxml.ns import nsdecls
from docx.shared import Inches, Pt, RGBColor

from core import ia

# ═══════════════════════════════════════════════════════════════════════════
# BASE DE DATOS
# ═══════════════════════════════════════════════════════════════════════════
DB_NAME = "gestion_etp.db"

def init_feedback_db():
    conn = sqlite3.connect(DB_NAME)
    cursor = conn.cursor()
    cursor.execute('''
    CREATE TABLE IF NOT EXISTS feedback (
        id INTEGER PRIMARY KEY AUTOINCREMENT,
        usuario TEXT,
        rol TEXT,
        fecha TEXT,
        facilidad_uso INTEGER,
        utilidad INTEGER,
        diseno INTEGER,
        velocidad INTEGER,
        contenido INTEGER,
        satisfaccion_general INTEGER,
        comentario TEXT,
        sugerencia TEXT,
        modulo_favorito TEXT,
        modulo_mejorar TEXT,
        puntuacion_total REAL
    )
    ''')
    conn.commit()
    return conn

def guardar_feedback(datos: Dict) -> None:
    conn = sqlite3.connect(DB_NAME)
    cursor = conn.cursor()
    cursor.execute('''
    INSERT INTO feedback (usuario, rol, fecha, facilidad_uso, utilidad, diseno, velocidad,
                          contenido, satisfaccion_general, comentario, sugerencia,
                          modulo_favorito, modulo_mejorar, puntuacion_total)
    VALUES (?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?)
    ''', (
        datos["usuario"], datos["rol"], datos["fecha"],
        datos["facilidad_uso"], datos["utilidad"], datos["diseno"],
        datos["velocidad"], datos["contenido"], datos["satisfaccion_general"],
        datos["comentario"], datos["sugerencia"],
        datos["modulo_favorito"], datos["modulo_mejorar"],
        datos["puntuacion_total"]
    ))
    conn.commit()
    conn.close()

def obtener_feedbacks() -> List[Dict]:
    conn = sqlite3.connect(DB_NAME)
    cursor = conn.cursor()
    cursor.execute('''
    SELECT id, usuario, rol, fecha, facilidad_uso, utilidad, diseno, velocidad,
           contenido, satisfaccion_general, comentario, sugerencia,
           modulo_favorito, modulo_mejorar, puntuacion_total
    FROM feedback ORDER BY fecha DESC
    ''')
    rows = cursor.fetchall()
    conn.close()
    return [dict(zip([
        "id", "usuario", "rol", "fecha", "facilidad_uso", "utilidad", "diseno",
        "velocidad", "contenido", "satisfaccion_general", "comentario",
        "sugerencia", "modulo_favorito", "modulo_mejorar", "puntuacion_total"
    ], row)) for row in rows]

def obtener_estadisticas_feedback() -> Optional[Dict]:
    feedbacks = obtener_feedbacks()
    if not feedbacks:
        return None
    df = pd.DataFrame(feedbacks)
    return {
        "total": len(df),
        "promedio_general": round(df["puntuacion_total"].mean(), 2),
        "promedios_dimension": {
            "Facilidad de Uso": round(df["facilidad_uso"].mean(), 2),
            "Utilidad": round(df["utilidad"].mean(), 2),
            "Diseño": round(df["diseno"].mean(), 2),
            "Velocidad": round(df["velocidad"].mean(), 2),
            "Contenido": round(df["contenido"].mean(), 2),
            "Satisfacción General": round(df["satisfaccion_general"].mean(), 2),
        },
        "modulos_favoritos": df["modulo_favorito"].value_counts().head(5).to_dict(),
        "modulos_mejorar": df["modulo_mejorar"].value_counts().head(5).to_dict(),
        "comentarios": df["comentario"].dropna().tolist(),
        "sugerencias": df["sugerencia"].dropna().tolist(),
        "por_rol": df["rol"].value_counts().to_dict(),
        "ultimos_7_dias": len(df[df["fecha"] >= (date.today() - timedelta(days=7)).strftime("%Y-%m-%d")]),
    }

# ═══════════════════════════════════════════════════════════════════════════
# SUPER INTERFAZ — ESTILOS
# ═══════════════════════════════════════════════════════════════════════════
st.markdown("""
<style>
@import url('https://fonts.googleapis.com/css2?family=Inter:wght@300;400;500;600;700;800;900&display=swap');

html, body, [class*="css"] {
    font-family: 'Inter', sans-serif;
    background-color: #F0F4F8;
    color: #1E293B;
}

.feedback-hero {
    background: linear-gradient(135deg, #0F172A 0%, #7C3AED 40%, #A78BFA 70%, #C4B5FD 100%);
    color: #fff;
    padding: 2.2rem;
    border-radius: 20px;
    margin-bottom: 1.5rem;
    box-shadow: 0 25px 50px rgba(124, 58, 237, 0.3);
    position: relative;
    overflow: hidden;
}

.feedback-hero::before {
    content: '⭐';
    position: absolute;
    right: 2rem;
    top: 50%;
    transform: translateY(-50%);
    font-size: 6rem;
    opacity: 0.15;
}

.feedback-hero-title { font-size: 2.4rem; font-weight: 900; letter-spacing: -0.03em; margin-bottom: 0.4rem; }
.feedback-hero-sub { font-size: 1.05rem; opacity: 0.9; line-height: 1.5; }

.feedback-hero-badge {
    display: inline-block;
    background: rgba(255,255,255,0.15);
    border: 1px solid rgba(255,255,255,0.25);
    border-radius: 8px;
    padding: 4px 12px;
    font-size: 0.8rem;
    font-weight: 600;
    margin-top: 0.8rem;
    margin-right: 8px;
}

.feedback-section-title {
    color: #7C3AED;
    font-weight: 700;
    font-size: 1.12rem;
    border-bottom: 2px solid #EDE9FE;
    padding-bottom: 8px;
    margin: 1.2rem 0 0.9rem 0;
}

.feedback-dimension-card {
    background: #fff;
    border: 2px solid #E2E8F0;
    border-radius: 14px;
    padding: 1.2rem;
    transition: all 0.25s ease;
    text-align: center;
    box-shadow: 0 2px 8px rgba(0,0,0,0.04);
}

.feedback-dimension-card:hover {
    transform: translateY(-3px);
    box-shadow: 0 8px 20px rgba(124, 58, 237, 0.15);
    border-color: #A78BFA;
}

.feedback-dimension-icono { font-size: 2.5rem; margin-bottom: 0.6rem; }
.feedback-dimension-nombre { font-weight: 700; font-size: 0.92rem; color: #0F172A; margin-bottom: 0.4rem; }
.feedback-dimension-desc { font-size: 0.78rem; color: #64748B; line-height: 1.3; }

.feedback-stat {
    background: #fff;
    border-radius: 12px;
    padding: 1rem;
    text-align: center;
    border: 1px solid #E2E8F0;
    box-shadow: 0 2px 6px rgba(0,0,0,0.04);
}

.feedback-stat-value { font-size: 1.8rem; font-weight: 800; color: #7C3AED; }
.feedback-stat-label { font-size: 0.72rem; font-weight: 700; color: #94A3B8; text-transform: uppercase; letter-spacing: 0.05em; }

.feedback-progress-bar {
    height: 10px;
    border-radius: 5px;
    background: #E2E8F0;
    overflow: hidden;
    margin-top: 8px;
}

.feedback-progress-fill {
    height: 100%;
    border-radius: 5px;
    background: linear-gradient(90deg, #7C3AED, #A78BFA);
    transition: width 0.8s ease;
}

.feedback-comment-card {
    background: #fff;
    border: 1px solid #E2E8F0;
    border-left: 4px solid #A78BFA;
    border-radius: 10px;
    padding: 1rem;
    margin-bottom: 0.8rem;
}

.feedback-comment-usuario { font-weight: 700; font-size: 0.88rem; color: #0F172A; }
.feedback-comment-fecha { font-size: 0.75rem; color: #94A3B8; }
.feedback-comment-texto { font-size: 0.88rem; color: #334155; margin-top: 0.5rem; line-height: 1.4; }
</style>
""", unsafe_allow_html=True)

# ═══════════════════════════════════════════════════════════════════════════
# ESTADO DE SESIÓN
# ═══════════════════════════════════════════════════════════════════════════
def init_estado():
    if "feedback_enviado" not in st.session_state:
        st.session_state.feedback_enviado = False
    if "feedback_analisis_ia" not in st.session_state:
        st.session_state.feedback_analisis_ia = None

init_estado()

# Inicializar BD
init_feedback_db()

# ═══════════════════════════════════════════════════════════════════════════
# DIMENSIONES DE EVALUACIÓN
# ═══════════════════════════════════════════════════════════════════════════
DIMENSIONES = [
    {"clave": "facilidad_uso", "nombre": "Facilidad de Uso", "icono": "🖱️",
     "descripcion": "¿Qué tan fácil es navegar y usar el portal?"},
    {"clave": "utilidad", "nombre": "Utilidad", "icono": "🎯",
     "descripcion": "¿El portal te ayuda en tu trabajo diario?"},
    {"clave": "diseno", "nombre": "Diseño Visual", "icono": "🎨",
     "descripcion": "¿El diseño es atractivo y profesional?"},
    {"clave": "velocidad", "nombre": "Velocidad", "icono": "⚡",
     "descripcion": "¿El portal responde rápido?"},
    {"clave": "contenido", "nombre": "Contenido", "icono": "📚",
     "descripcion": "¿El contenido generado es de calidad?"},
    {"clave": "satisfaccion_general", "nombre": "Satisfacción General", "icono": "⭐",
     "descripcion": "¿Qué tan satisfecho estás con el portal?"},
]

MODULOS_DISPONIBLES = [
    "Ponderación RA", "Planificación Modular", "Plan Diario ETP",
    "Generador de Contenidos", "Generador de Libros/Guías",
    "Redactor Profundo", "Fábrica de Simuladores", "Banco de Ítems",
    "Recuperación R.A", "Auditor de Calificaciones", "Auditor de Planificaciones",
    "Acompañamiento Docente", "Gestor de Acuerdos", "Directorio de Docentes",
    "Registro de Incidencias", "Otro"
]

# ═══════════════════════════════════════════════════════════════════════════
# COMPONENTE DE ESTRELLAS
# ═══════════════════════════════════════════════════════════════════════════
def render_star_rating(clave: str, label: str, icono: str, key_prefix: str) -> int:
    """Renderiza un selector de estrellas interactivo."""
    col_label, col_stars = st.columns([2, 3])

    with col_label:
        st.markdown(f"**{icono} {label}**")

    with col_stars:
        # Usar selectbox con emojis de estrellas para simular rating
        opciones = ["⭐", "⭐⭐", "⭐⭐⭐", "⭐⭐⭐⭐", "⭐⭐⭐⭐⭐"]
        seleccion = st.selectbox(
            "Puntuación",
            opciones,
            index=4,  # Default: 5 estrellas
            key=f"{key_prefix}_{clave}",
            label_visibility="collapsed"
        )
        return opciones.index(seleccion) + 1

# ═══════════════════════════════════════════════════════════════════════════
# PROMPT ANÁLISIS IA
# ═══════════════════════════════════════════════════════════════════════════
def prompt_analisis_feedback(estadisticas: Dict) -> str:
    comentarios_txt = "\n".join([f"- {c}" for c in estadisticas.get("comentarios", [])[:20]])
    sugerencias_txt = "\n".join([f"- {s}" for s in estadisticas.get("sugerencias", [])[:20]])

    promedios_txt = "\n".join([
        f"- {dim}: {valor}/5"
        for dim, valor in estadisticas.get("promedios_dimension", {}).items()
    ])

    return f"""Actúa como un Analista de Experiencia de Usuario (UX) especializado en plataformas educativas.

Tu tarea es analizar los feedbacks recopilados del Portal de Gestión Docente ETP y generar insights accionables.

ESTADÍSTICAS GENERALES:
Total de feedbacks: {estadisticas.get('total', 0)}
Puntuación promedio general: {estadisticas.get('promedio_general', 0)}/5
Feedbacks últimos 7 días: {estadisticas.get('ultimos_7_dias', 0)}

PROMEDIOS POR DIMENSIÓN:
{promedios_txt}

MÓDULOS MÁS VALORADOS:
{estadisticas.get('modulos_favoritos', {})}

MÓDULOS A MEJORAR:
{estadisticas.get('modulos_mejorar', {})}

COMENTARIOS DE USUARIOS:
{comentarios_txt if comentarios_txt else 'Sin comentarios'}

SUGERENCIAS DE USUARIOS:
{sugerencias_txt if sugerencias_txt else 'Sin sugerencias'}

REGLAS:
- Sé objetivo y constructivo.
- Identifica patrones en los comentarios.
- Las recomendaciones deben ser específicas y priorizadas.
- Destaca las fortalezas genuinas.
- Identifica las áreas críticas que requieren atención inmediata.

Devuelve ÚNICAMENTE JSON válido:
{{
  "RESUMEN_EJECUTIVO": "Resumen de 2-3 párrafos del estado general del feedback...",
  "FORTALEZAS": ["Fortaleza 1", "Fortaleza 2", "Fortaleza 3"],
  "AREAS_CRITICAS": ["Área crítica 1", "Área crítica 2"],
  "PATRONES_IDENTIFICADOS": ["Patrón 1", "Patrón 2"],
  "RECOMENDACIONES_PRIORITARIAS": [
    {{"PRIORIDAD": "Alta", "ACCION": "Acción específica", "MODULO_AFECTADO": "Módulo"}},
    {{"PRIORIDAD": "Media", "ACCION": "Acción específica", "MODULO_AFECTADO": "Módulo"}}
  ],
  "SENTIMIENTO_GENERAL": "Positivo/Negativo/Mixto",
  "NPS_ESTIMADO": "Número estimado de Net Promoter Score"
}}
"""

# ═══════════════════════════════════════════════════════════════════════════
# GENERACIÓN WORD — REPORTE DE FEEDBACK
# ═══════════════════════════════════════════════════════════════════════════
def shade_cell(cell, color):
    shd = parse_xml(r'<w:shd {} w:fill="{}"/>'.format(nsdecls("w"), color))
    cell._tc.get_or_add_tcPr().append(shd)

def set_cell_text(cell, text, bold=False, center=False, color=None):
    cell.text = ""
    p = cell.paragraphs[0]
    run = p.add_run(str(text))
    run.bold = bold
    if center:
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    if color:
        shade_cell(cell, color)

def build_feedback_docx(estadisticas: Dict, analisis_ia: Optional[Dict]) -> BytesIO:
    doc = Document()
    doc.styles["Normal"].font.name = "Calibri"
    doc.styles["Normal"].font.size = Pt(11)

    for section in doc.sections:
        section.left_margin = Inches(0.7)
        section.right_margin = Inches(0.7)

    # ─── PORTADA ───
    for _ in range(3):
        doc.add_paragraph()

    p_titulo = doc.add_paragraph()
    p_titulo.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run_t = p_titulo.add_run("REPORTE DE FEEDBACK\nPORTAL DE GESTIÓN DOCENTE ETP")
    run_t.bold = True
    run_t.font.size = Pt(22)
    run_t.font.color.rgb = RGBColor(0x7C, 0x3A, 0xED)

    doc.add_paragraph()

    p_fecha = doc.add_paragraph()
    p_fecha.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p_fecha.add_run(f"Fecha: {datetime.now().strftime('%d/%m/%Y')}").font.size = Pt(12)

    p_total = doc.add_paragraph()
    p_total.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p_total.add_run(f"Total de feedbacks: {estadisticas.get('total', 0)}").font.size = Pt(12)

    doc.add_page_break()

    # ─── RESUMEN EJECUTIVO ───
    doc.add_heading("1. Resumen Ejecutivo", level=1)

    p_score = doc.add_paragraph()
    p_score.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run_score = p_score.add_run(f"Puntuación Promedio: {estadisticas.get('promedio_general', 0)}/5")
    run_score.bold = True
    run_score.font.size = Pt(18)

    doc.add_paragraph()

    # ─── PROMEDIOS POR DIMENSIÓN ───
    doc.add_heading("2. Promedios por Dimensión", level=1)

    tabla_dim = doc.add_table(rows=1, cols=3)
    tabla_dim.style = "Table Grid"

    hdr = tabla_dim.rows[0].cells
    set_cell_text(hdr[0], "Dimensión", bold=True, color="EDE9FE")
    set_cell_text(hdr[1], "Promedio", bold=True, center=True, color="EDE9FE")
    set_cell_text(hdr[2], "Nivel", bold=True, center=True, color="EDE9FE")

    for dim, valor in estadisticas.get("promedios_dimension", {}).items():
        row = tabla_dim.add_row().cells
        set_cell_text(row[0], dim)
        set_cell_text(row[1], f"{valor}/5", center=True)

        if valor >= 4.5:
            nivel = "Excelente"
            color_nivel = "D1FAE5"
        elif valor >= 3.5:
            nivel = "Bueno"
            color_nivel = "DBEAFE"
        elif valor >= 2.5:
            nivel = "Aceptable"
            color_nivel = "FEF3C7"
        else:
            nivel = "Requiere Mejora"
            color_nivel = "FEE2E2"

        set_cell_text(row[2], nivel, center=True, color=color_nivel)

    doc.add_paragraph()

    # ─── MÓDULOS ───
    doc.add_heading("3. Módulos Más Valorados", level=1)
    for modulo, count in estadisticas.get("modulos_favoritos", {}).items():
        doc.add_paragraph(f"• {modulo}: {count} menciones", style="List Bullet")

    doc.add_heading("4. Módulos a Mejorar", level=1)
    for modulo, count in estadisticas.get("modulos_mejorar", {}).items():
        doc.add_paragraph(f"• {modulo}: {count} menciones", style="List Bullet")

    doc.add_paragraph()

    # ─── ANÁLISIS IA ───
    if analisis_ia:
        doc.add_heading("5. Análisis IA de Feedbacks", level=1)

        doc.add_heading("Resumen Ejecutivo", level=2)
        doc.add_paragraph(analisis_ia.get("RESUMEN_EJECUTIVO", ""))

        doc.add_heading("Fortalezas", level=2)
        for fortaleza in analisis_ia.get("FORTALEZAS", []):
            doc.add_paragraph(f"✅ {fortaleza}", style="List Bullet")

        doc.add_heading("Áreas Críticas", level=2)
        for area in analisis_ia.get("AREAS_CRITICAS", []):
            doc.add_paragraph(f"⚠️ {area}", style="List Bullet")

        doc.add_heading("Recomendaciones Prioritarias", level=2)
        for rec in analisis_ia.get("RECOMENDACIONES_PRIORITARIAS", []):
            doc.add_paragraph(
                f"[{rec.get('PRIORIDAD', '')}] {rec.get('ACCION', '')} - Módulo: {rec.get('MODULO_AFECTADO', '')}",
                style="List Bullet"
            )

        doc.add_paragraph()
        doc.add_paragraph(f"Sentimiento General: {analisis_ia.get('SENTIMIENTO_GENERAL', 'N/A')}")
        doc.add_paragraph(f"NPS Estimado: {analisis_ia.get('NPS_ESTIMADO', 'N/A')}")

    # ─── FIRMAS ───
    doc.add_paragraph()
    doc.add_paragraph()
    tabla_firmas = doc.add_table(rows=2, cols=2)
    tabla_firmas.cell(0, 0).text = "_________________________"
    tabla_firmas.cell(0, 1).text = "_________________________"
    tabla_firmas.cell(1, 0).text = "Coordinación ETP"
    tabla_firmas.cell(1, 1).text = "Dirección del Centro"

    for row in tabla_firmas.rows:
        for cell in row.cells:
            cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER

    buffer = BytesIO()
    doc.save(buffer)
    buffer.seek(0)
    return buffer

# ═══════════════════════════════════════════════════════════════════════════
# HERO
# ═══════════════════════════════════════════════════════════════════════════
ia.panel_sidebar_ia("Feedback del Portal")

st.markdown("""
<div class="feedback-hero">
    <div class="feedback-hero-title">⭐ Feedback del Portal ETP</div>
    <div class="feedback-hero-sub">
        Tu opinión nos ayuda a mejorar · Evaluación multidimensional · Análisis IA de comentarios
    </div>
    <div>
        <span class="feedback-hero-badge">🖱️ 6 Dimensiones</span>
        <span class="feedback-hero-badge">⭐ Ratings con Estrellas</span>
        <span class="feedback-hero-badge">🤖 Análisis IA</span>
        <span class="feedback-hero-badge">📊 Dashboard Analítica</span>
    </div>
</div>
""", unsafe_allow_html=True)

# ═══════════════════════════════════════════════════════════════════════════
# VERIFICAR AUTENTICACIÓN
# ═══════════════════════════════════════════════════════════════════════════
usuario_actual = ""
rol_actual = ""

if st.session_state.get("coordinador_autenticado", False):
    usuario_actual = st.session_state.get("coordinador_nombre", "Coordinador")
    rol_actual = "Coordinador"
elif st.session_state.get("docente_autenticado", False):
    usuario_actual = st.session_state.get("nombre_docente", "Docente")
    rol_actual = "Docente"
else:
    st.warning("⚠️ Debes iniciar sesión para enviar feedback.")
    st.stop()

# ═══════════════════════════════════════════════════════════════════════════
# TAB PRINCIPAL
# ═══════════════════════════════════════════════════════════════════════════
es_coordinador = rol_actual == "Coordinador"

if es_coordinador:
    tab_feedback, tab_analitica = st.tabs(["⭐ Enviar Feedback", "📊 Dashboard de Analítica"])
else:
    tab_feedback = st.container()
    tab_analitica = None

# ═══════════════════════════════════════════════════════════════════════════
# TAB: ENVIAR FEEDBACK
# ═══════════════════════════════════════════════════════════════════════════
with tab_feedback:
    st.markdown('<div class="feedback-section-title">⭐ Evalúa tu Experiencia</div>', unsafe_allow_html=True)

    # Mostrar dimensiones como cards
    cols_dims = st.columns(3)
    for i, dim in enumerate(DIMENSIONES):
        with cols_dims[i % 3]:
            st.markdown(f"""
            <div class="feedback-dimension-card">
                <div class="feedback-dimension-icono">{dim['icono']}</div>
                <div class="feedback-dimension-nombre">{dim['nombre']}</div>
                <div class="feedback-dimension-desc">{dim['descripcion']}</div>
            </div>
            """, unsafe_allow_html=True)

    st.markdown("---")

    # Formulario de feedback
    with st.form("form_feedback"):
        st.markdown('<div class="feedback-section-title">📋 Puntuación por Dimensión</div>', unsafe_allow_html=True)

        ratings = {}
        for dim in DIMENSIONES:
            ratings[dim["clave"]] = render_star_rating(
                dim["clave"], dim["nombre"], dim["icono"], "rating"
            )

        st.markdown("---")
        st.markdown('<div class="feedback-section-title">📚 Módulos</div>', unsafe_allow_html=True)

        col_mod1, col_mod2 = st.columns(2)
        with col_mod1:
            modulo_favorito = st.selectbox(
                "🏆 ¿Cuál es tu módulo favorito?",
                ["— Seleccionar —"] + MODULOS_DISPONIBLES
            )
        with col_mod2:
            modulo_mejorar = st.selectbox(
                "🔧 ¿Qué módulo crees que necesita mejorar?",
                ["— Seleccionar —"] + MODULOS_DISPONIBLES
            )

        st.markdown("---")
        st.markdown('<div class="feedback-section-title">💬 Comentarios y Sugerencias</div>', unsafe_allow_html=True)

        comentario = st.text_area(
            "💭 ¿Qué te gusta del portal? ¿Qué experiencia has tenido?",
            height=100,
            placeholder="Comparte tu experiencia..."
        )

        sugerencia = st.text_area(
            "💡 ¿Qué sugerencias tienes para mejorar?",
            height=100,
            placeholder="Tus ideas nos ayudan a crecer..."
        )

        st.markdown("---")

        col_btn1, col_btn2 = st.columns(2)
        with col_btn1:
            submit_feedback = st.form_submit_button(
                "📤 Enviar Feedback",
                type="primary",
                use_container_width=True
            )
        with col_btn2:
            clear_feedback = st.form_submit_button(
                "🗑️ Limpiar Formulario",
                use_container_width=True
            )

    if submit_feedback:
        # Calcular puntuación total (promedio de las 6 dimensiones)
        puntuacion_total = round(sum(ratings.values()) / len(ratings), 2)

        feedback_data = {
            "usuario": usuario_actual,
            "rol": rol_actual,
            "fecha": date.today().strftime("%Y-%m-%d"),
            "facilidad_uso": ratings["facilidad_uso"],
            "utilidad": ratings["utilidad"],
            "diseno": ratings["diseno"],
            "velocidad": ratings["velocidad"],
            "contenido": ratings["contenido"],
            "satisfaccion_general": ratings["satisfaccion_general"],
            "comentario": comentario,
            "sugerencia": sugerencia,
            "modulo_favorito": modulo_favorito if modulo_favorito != "— Seleccionar —" else None,
            "modulo_mejorar": modulo_mejorar if modulo_mejorar != "— Seleccionar —" else None,
            "puntuacion_total": puntuacion_total,
        }

        guardar_feedback(feedback_data)
        st.session_state.feedback_enviado = True
        st.toast(f"✅ ¡Feedback enviado! Puntuación total: {puntuacion_total}/5", icon="⭐")
        st.rerun()

    if clear_feedback:
        st.rerun()

    # Mostrar feedbacks recientes del usuario
    feedbacks_usuario = [f for f in obtener_feedbacks() if f["usuario"] == usuario_actual]
    if feedbacks_usuario:
        st.markdown("---")
        st.markdown('<div class="feedback-section-title">📝 Tus Feedbacks Anteriores</div>', unsafe_allow_html=True)

        for fb in feedbacks_usuario[:5]:
            st.markdown(f"""
            <div class="feedback-comment-card">
                <div class="feedback-comment-usuario">⭐ {fb['puntuacion_total']}/5 · {fb['fecha']}</div>
                <div class="feedback-comment-texto">{fb['comentario'] or 'Sin comentario'}</div>
            </div>
            """, unsafe_allow_html=True)

# ═══════════════════════════════════════════════════════════════════════════
# TAB: DASHBOARD DE ANALÍTICA (Solo Coordinador)
# ═══════════════════════════════════════════════════════════════════════════
if es_coordinador and tab_analitica:
    with tab_analitica:
        estadisticas = obtener_estadisticas_feedback()

        if not estadisticas:
            st.info("📭 Aún no hay feedbacks registrados. Los usuarios pueden enviar feedback desde la pestaña '⭐ Enviar Feedback'.")
        else:
            st.markdown('<div class="feedback-section-title">📊 Métricas Generales</div>', unsafe_allow_html=True)

            # Métricas principales
            col_m1, col_m2, col_m3, col_m4 = st.columns(4)
            with col_m1:
                st.markdown(f"""
                <div class="feedback-stat">
                    <div class="feedback-stat-value">{estadisticas['total']}</div>
                    <div class="feedback-stat-label">Feedbacks Totales</div>
                </div>
                """, unsafe_allow_html=True)
            with col_m2:
                st.markdown(f"""
                <div class="feedback-stat">
                    <div class="feedback-stat-value">{estadisticas['promedio_general']}/5</div>
                    <div class="feedback-stat-label">Puntuación Promedio</div>
                </div>
                """, unsafe_allow_html=True)
            with col_m3:
                st.markdown(f"""
                <div class="feedback-stat">
                    <div class="feedback-stat-value">{estadisticas['ultimos_7_dias']}</div>
                    <div class="feedback-stat-label">Últimos 7 Días</div>
                </div>
                """, unsafe_allow_html=True)
            with col_m4:
                roles = estadisticas.get('por_rol', {})
                st.markdown(f"""
                <div class="feedback-stat">
                    <div class="feedback-stat-value">{roles.get('Docente', 0)}/{roles.get('Coordinador', 0)}</div>
                    <div class="feedback-stat-label">Docentes/Coord.</div>
                </div>
                """, unsafe_allow_html=True)

            st.markdown("---")

            # Promedios por dimensión con barras de progreso
            st.markdown('<div class="feedback-section-title">📈 Promedios por Dimensión</div>', unsafe_allow_html=True)

            for dim, valor in estadisticas.get("promedios_dimension", {}).items():
                porcentaje = (valor / 5) * 100
                st.markdown(f"""
                <div style="margin-bottom: 12px;">
                    <div style="display: flex; justify-content: space-between; margin-bottom: 4px;">
                        <span style="font-weight: 600;">{dim}</span>
                        <span style="font-weight: 700; color: #7C3AED;">{valor}/5</span>
                    </div>
                    <div class="feedback-progress-bar">
                        <div class="feedback-progress-fill" style="width: {porcentaje}%;"></div>
                    </div>
                </div>
                """, unsafe_allow_html=True)

            st.markdown("---")

            # Módulos favoritos y a mejorar
            col_mod1, col_mod2 = st.columns(2)
            with col_mod1:
                st.markdown('<div class="feedback-section-title">🏆 Módulos Más Valorados</div>', unsafe_allow_html=True)
                for modulo, count in estadisticas.get("modulos_favoritos", {}).items():
                    st.markdown(f"• **{modulo}**: {count} menciones")

            with col_mod2:
                st.markdown('<div class="feedback-section-title">🔧 Módulos a Mejorar</div>', unsafe_allow_html=True)
                for modulo, count in estadisticas.get("modulos_mejorar", {}).items():
                    st.markdown(f"• **{modulo}**: {count} menciones")

            st.markdown("---")

            # Comentarios recientes
            st.markdown('<div class="feedback-section-title">💬 Comentarios Recientes</div>', unsafe_allow_html=True)

            feedbacks = obtener_feedbacks()
            for fb in feedbacks[:10]:
                if fb["comentario"]:
                    st.markdown(f"""
                    <div class="feedback-comment-card">
                        <div class="feedback-comment-usuario">
                            {fb['usuario']} · {fb['rol']} · ⭐ {fb['puntuacion_total']}/5
                        </div>
                        <div class="feedback-comment-fecha">{fb['fecha']}</div>
                        <div class="feedback-comment-texto">{fb['comentario']}</div>
                    </div>
                    """, unsafe_allow_html=True)

            st.markdown("---")

            # Análisis IA
            st.markdown('<div class="feedback-section-title">🤖 Análisis IA de Feedbacks</div>', unsafe_allow_html=True)

            if st.button("🧠 Generar Análisis IA", type="primary", use_container_width=True):
                with st.spinner("🧠 Analizando feedbacks con IA..."):
                    try:
                        prompt = prompt_analisis_feedback(estadisticas)
                        resultado_ia, flags = ia.solicitar_json(
                            prompt, max_tokens=16384, temperature=0.3, modulo="feedback_analisis"
                        )
                        st.session_state.feedback_analisis_ia = resultado_ia
                        st.toast("✅ Análisis IA generado.", icon="🤖")
                        st.rerun()
                    except Exception as e:
                        ia.render_error_ia(e)

            if st.session_state.feedback_analisis_ia:
                analisis = st.session_state.feedback_analisis_ia

                with st.expander("📄 Resumen Ejecutivo", expanded=True):
                    st.write(analisis.get("RESUMEN_EJECUTIVO", ""))

                with st.expander("💪 Fortalezas"):
                    for fortaleza in analisis.get("FORTALEZAS", []):
                        st.markdown(f"✅ {fortaleza}")

                with st.expander("⚠️ Áreas Críticas"):
                    for area in analisis.get("AREAS_CRITICAS", []):
                        st.markdown(f"🚨 {area}")

                with st.expander("🔍 Patrones Identificados"):
                    for patron in analisis.get("PATRONES_IDENTIFICADOS", []):
                        st.markdown(f"📊 {patron}")

                with st.expander("💡 Recomendaciones Prioritarias"):
                    for rec in analisis.get("RECOMENDACIONES_PRIORITARIAS", []):
                        st.markdown(
                            f"**[{rec.get('PRIORIDAD', '')}]** {rec.get('ACCION', '')} "
                            f"- Módulo: {rec.get('MODULO_AFECTADO', '')}"
                        )

                st.markdown(f"**Sentimiento General:** {analisis.get('SENTIMIENTO_GENERAL', 'N/A')}")
                st.markdown(f"**NPS Estimado:** {analisis.get('NPS_ESTIMADO', 'N/A')}")

            st.markdown("---")

            # Exportar Word
            st.markdown('<div class="feedback-section-title">📥 Exportar Reporte</div>', unsafe_allow_html=True)

            if st.button("📄 Generar Reporte Word (.docx)", type="primary", use_container_width=True):
                with st.spinner("📄 Construyendo reporte..."):
                    buffer = build_feedback_docx(estadisticas, st.session_state.feedback_analisis_ia)
                    st.session_state.feedback_buffer = buffer

            if st.session_state.get("feedback_buffer"):
                st.download_button(
                    label="⬇️ Descargar Reporte de Feedback (.docx)",
                    data=st.session_state.feedback_buffer,
                    file_name=f"Reporte_Feedback_{date.today().strftime('%Y%m%d')}.docx",
                    mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                    type="primary",
                    use_container_width=True,
                )