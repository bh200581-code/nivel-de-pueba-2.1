"""
incidencias.py — Registro y Análisis de Incidencias (MIGRADO a core.ia · Paso 9)
• IA vía motor unificado: marcadores seguros, parseo robusto, reintento por
  truncamiento, auditoría de llamadas y errores estándar en español.
• Persistencia SQLite (tabla incidencias) que alimenta la Sala de Situación.
• Interfaz mejorada: hero, tarjetas KPI, cascada Gravedad→Falta en vivo,
  búsqueda/filtros, borrado con confirmación, export CSV y toasts.
• Informe Word apaisado profesional con análisis IA editable.
"""
import datetime
import json
import sqlite3
from io import BytesIO

import pandas as pd
import streamlit as st
from docx import Document
from docx.enum.section import WD_ORIENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import nsdecls
from docx.oxml import parse_xml
from docx.shared import Inches, Pt, RGBColor

from core import ia

# ═══════════════════════════════════════════════════════════════════════════
# BASE DE DATOS
# ═══════════════════════════════════════════════════════════════════════════
DB_NAME = "gestion_etp.db"

COL_FECHA = "Fecha"
COL_ESTUDIANTE = "Estudiante"
COL_GENERO = "Género"
COL_CICLO = "Ciclo"
COL_GRADO = "Grado"
COL_SECCION = "Sección"
COL_GRAVEDAD = "Gravedad"
COL_FALTA = "Falta"

GRAVEDAD_CRITICA = ("Grave", "Muy Grave")

MAPA_GRAVEDAD_UI = {
    "Leve": "🟢 Leve",
    "Moderada": "🟡 Moderada",
    "Grave": "🟠 Grave",
    "Muy Grave": "🔴 Muy Grave",
}


def init_db():
    conn = sqlite3.connect(DB_NAME, check_same_thread=False)
    cursor = conn.cursor()
    cursor.execute('''
    CREATE TABLE IF NOT EXISTS incidencias (
        id INTEGER PRIMARY KEY AUTOINCREMENT,
        fecha TEXT, estudiante TEXT, genero TEXT, ciclo TEXT,
        grado TEXT, seccion TEXT, gravedad TEXT, falta TEXT
    )
    ''')
    conn.commit()
    return conn


conn = init_db()


def get_todas_incidencias():
    cursor = conn.cursor()
    cursor.execute(
        "SELECT id, fecha, estudiante, genero, ciclo, grado, seccion, gravedad, falta "
        "FROM incidencias ORDER BY id DESC"
    )
    rows = cursor.fetchall()
    columnas = ["id", COL_FECHA, COL_ESTUDIANTE, COL_GENERO, COL_CICLO,
                COL_GRADO, COL_SECCION, COL_GRAVEDAD, COL_FALTA]
    return [dict(zip(columnas, row)) for row in rows]


def insertar_incidencia(fecha, estudiante, genero, ciclo, grado, seccion, gravedad, falta):
    cursor = conn.cursor()
    cursor.execute('''
    INSERT INTO incidencias (fecha, estudiante, genero, ciclo, grado, seccion, gravedad, falta)
    VALUES (?, ?, ?, ?, ?, ?, ?, ?)
    ''', (fecha, estudiante, genero, ciclo, grado, seccion, gravedad, falta))
    conn.commit()


def eliminar_incidencia(id_incidencia):
    cursor = conn.cursor()
    cursor.execute("DELETE FROM incidencias WHERE id = ?", (id_incidencia,))
    conn.commit()


# ═══════════════════════════════════════════════════════════════════════════
# CATÁLOGO MINERD
# ═══════════════════════════════════════════════════════════════════════════
FALTAS_MINERD = {
    "Leve": [
        "Llegadas tarde injustificadas",
        "No portar el uniforme correctamente",
        "Interrupción del orden en clases",
        "No traer materiales o tareas asignadas",
        "Uso de dispositivos electrónicos sin autorización",
        "Permanecer en aulas o pasillos durante recreos sin autorización",
    ],
    "Moderada": [
        "Falta de respeto a compañeros o personal (verbal o gestual)",
        "Evasión de clases (escaqueo) dentro del centro",
        "Daño leve a la propiedad escolar o pertenencias ajenas",
        "Uso de vocabulario soez o inapropiado",
        "Falsificación de justificantes o comunicaciones",
        "Comportamiento inadecuado en actos cívicos o actividades",
    ],
    "Grave": [
        "Agresión física o verbal a compañeros o personal",
        "Hurto o robo de pertenencias",
        "Daño intencional y grave a la propiedad escolar",
        "Acoso escolar (Bullying) o ciberacoso",
        "Incumplimiento de medidas disciplinarias previas",
        "Fomento de desorden o riñas dentro del centro",
    ],
    "Muy Grave": [
        "Portación de armas blancas o de fuego",
        "Consumo o distribución de sustancias controladas (drogas/alcohol)",
        "Agresión física grave a personal del centro",
        "Actos que atenten contra la vida o integridad moral",
        "Abuso sexual o acoso grave",
        "Inducción a conductas que dañen la moral o la salud pública",
    ],
}

OPCIONES_GENERO = ["Masculino", "Femenino"]
OPCIONES_CICLO = ["Primer Ciclo", "Segundo Ciclo"]
OPCIONES_GRADO = ["1ro Grado", "2do Grado", "3ro Grado", "4to Grado", "5to Grado", "6to Grado"]

# ═══════════════════════════════════════════════════════════════════════════
# ESTILOS (interfaz mejorada)
# ═══════════════════════════════════════════════════════════════════════════
st.markdown("""
<style>
@import url('https://fonts.googleapis.com/css2?family=Inter:wght@300;400;500;600;700;800&display=swap');
html, body, [class*="css"] { font-family: 'Inter', sans-serif; background-color: #F8FAFC; color: #1E293B; }
.inc-hero { background: linear-gradient(135deg, #0F172A 0%, #7F1D1D 55%, #DC2626 100%); color: #fff;
    padding: 1.7rem 1.8rem; border-radius: 18px; margin-bottom: 1.25rem; box-shadow: 0 18px 35px rgba(15,23,42,.18); }
.inc-title { font-size: 2rem; font-weight: 800; letter-spacing: -0.02em; }
.inc-sub { opacity: .88; font-size: .98rem; margin-top: .3rem; }
.section-title { color: #1D4ED8; font-weight: 700; font-size: 1.12rem; border-bottom: 2px solid #DBEAFE;
    padding-bottom: 8px; margin: 1.2rem 0 .9rem 0; }
.metric-card { background:#fff; border:1px solid #E2E8F0; border-top:4px solid #2563EB; border-radius:12px;
    padding:14px 16px; box-shadow:0 4px 12px rgba(15,23,42,.06); text-align:center; }
.metric-value { font-size:2rem; font-weight:800; color:#0F172A; }
.metric-label { font-size:.78rem; font-weight:800; color:#64748B; text-transform:uppercase; letter-spacing:.05em; }
</style>
""", unsafe_allow_html=True)

# ═══════════════════════════════════════════════════════════════════════════
# GUARDIA + SIDEBAR + HERO
# ═══════════════════════════════════════════════════════════════════════════
if not st.session_state.get("coordinador_autenticado", False):
    st.error("🔒 Esta página es exclusiva de Coordinación.")
    st.stop()

ia.panel_sidebar_ia("Registro de Incidencias")

st.markdown("""
<div class="inc-hero">
    <div class="inc-title">🚨 Registro y Análisis de Incidencias</div>
    <div class="inc-sub">Manual de Convivencia MINERD · Base de datos permanente · Informes con IA</div>
</div>
""", unsafe_allow_html=True)

if "analisis_ia_temp" not in st.session_state:
    st.session_state.analisis_ia_temp = ""
if "plan_ia_temp" not in st.session_state:
    st.session_state.plan_ia_temp = ""

db_incidencias = get_todas_incidencias()


def calcular_metricas(df: pd.DataFrame) -> dict:
    return {
        "total": len(df),
        "criticos": int(df[COL_GRAVEDAD].isin(GRAVEDAD_CRITICA).sum()) if not df.empty else 0,
        "masculino": int((df[COL_GENERO] == "Masculino").sum()) if not df.empty else 0,
        "femenino": int((df[COL_GENERO] == "Femenino").sum()) if not df.empty else 0,
    }


def _metric_card(label, value, color=None) -> str:
    style = f' style="color: {color};"' if color else ""
    return (
        f'<div class="metric-card"><div class="metric-label">{label}</div>'
        f'<div class="metric-value"{style}>{value}</div></div>'
    )


# ═══════════════════════════════════════════════════════════════════════════
# WORD: INFORME APAISADO
# ═══════════════════════════════════════════════════════════════════════════
def _shade_cell(cell, color_hex: str) -> None:
    shd = parse_xml(r'<w:shd {} w:fill="{}"/>'.format(nsdecls('w'), color_hex))
    cell._tc.get_or_add_tcPr().append(shd)


def _fijar_anchos(tabla, anchos):
    tabla.autofit = False
    for row in tabla.rows:
        for i, w in enumerate(anchos):
            if i < len(row.cells):
                row.cells[i].width = Inches(w)


def generar_informe_docx(datos: dict, df_filtrado: pd.DataFrame) -> BytesIO:
    doc = Document()
    section = doc.sections[0]
    section.orientation = WD_ORIENT.LANDSCAPE
    section.page_width, section.page_height = section.page_height, section.page_width
    section.left_margin = Inches(0.5)
    section.right_margin = Inches(0.5)
    section.top_margin = Inches(0.5)
    section.bottom_margin = Inches(0.5)
    doc.styles["Normal"].font.name = "Calibri"
    doc.styles["Normal"].font.size = Pt(10)

    p_enc = doc.add_paragraph()
    p_enc.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r1 = p_enc.add_run(f"{datos['centro'].upper()}\n")
    r1.bold = True
    r1.font.size = Pt(12)
    r2 = p_enc.add_run("COORDINACIÓN TÉCNICO-PEDAGÓGICA | DEPARTAMENTO DE DISCIPLINA Y CONVIVENCIA\n")
    r2.bold = True
    r2.font.size = Pt(10)
    p_tit = doc.add_paragraph()
    p_tit.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r_tit = p_tit.add_run(f"INFORME ESTADÍSTICO Y CUALITATIVO DE INCIDENCIAS\n{datos['periodo'].upper()}\n")
    r_tit.bold = True
    r_tit.font.size = Pt(13)
    r_tit.font.color.rgb = RGBColor(37, 99, 235)
    alcance_txt = "Estudiantes Seleccionados" if datos['alcance_es_seleccion'] else "Reporte General del Centro"
    p_alc = doc.add_paragraph()
    p_alc.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p_alc.add_run(
        f"Alcance del Reporte: {alcance_txt} | Fecha Emisión: {datetime.date.today().strftime('%d/%m/%Y')}"
    ).italic = True
    doc.add_paragraph("_" * 110)

    doc.add_heading("I. Resumen Estadístico", level=2)
    metricas = calcular_metricas(df_filtrado)
    t_stat = doc.add_table(rows=2, cols=4)
    t_stat.style = "Table Grid"
    hdr_s = t_stat.rows[0].cells
    for i, h in enumerate(["Total Casos", "Graves/Muy Graves", "Masculino", "Femenino"]):
        hdr_s[i].text = h
        hdr_s[i].paragraphs[0].runs[0].bold = True
        hdr_s[i].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
        _shade_cell(hdr_s[i], "EFF6FF")
    dat_s = t_stat.rows[1].cells
    for cell, valor in zip(dat_s, [metricas["total"], metricas["criticos"],
                                   metricas["masculino"], metricas["femenino"]]):
        cell.text = str(valor)
        cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
        cell.paragraphs[0].runs[0].font.size = Pt(12)
        cell.paragraphs[0].runs[0].bold = True
    _fijar_anchos(t_stat, [2.5, 2.5, 2.5, 2.5])

    doc.add_heading("II. Análisis Cualitativo del Clima Escolar", level=2)
    p_an = doc.add_paragraph(datos['analisis'])
    p_an.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

    doc.add_heading("III. Recomendaciones y Plan de Acción", level=2)
    p_plan = doc.add_paragraph(datos['plan_accion'])
    p_plan.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

    doc.add_heading("IV. Registro Detallado de Incidencias", level=2)
    t_det = doc.add_table(rows=1, cols=8)
    t_det.style = "Table Grid"
    hdr_d = t_det.rows[0].cells
    for i, txt in enumerate(["Fecha", "Estudiante", "G.", "Ciclo", "Grado", "Sec.", "Gravedad", "Falta Cometida"]):
        hdr_d[i].text = txt
        hdr_d[i].paragraphs[0].runs[0].bold = True
        hdr_d[i].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
        _shade_cell(hdr_d[i], "DBEAFE")
    for _, row in df_filtrado.iterrows():
        r = t_det.add_row().cells
        r[0].text = str(row[COL_FECHA])
        r[1].text = str(row[COL_ESTUDIANTE])
        r[2].text = str(row[COL_GENERO])[0]
        r[3].text = str(row[COL_CICLO]).replace(" Ciclo", "")
        r[4].text = str(row[COL_GRADO]).split(" ")[0]
        r[5].text = str(row[COL_SECCION])
        r[6].text = str(row[COL_GRAVEDAD])
        r[7].text = str(row[COL_FALTA])
        if str(row[COL_GRAVEDAD]) in GRAVEDAD_CRITICA:
            _shade_cell(r[6], "FEE2E2")
    _fijar_anchos(t_det, [0.8, 1.8, 0.3, 0.8, 0.6, 0.4, 0.9, 3.9])

    doc.add_paragraph("\n\n")
    t_firmas = doc.add_table(rows=2, cols=2)
    t_firmas.cell(0, 0).text = "____________________________________"
    t_firmas.cell(0, 1).text = "____________________________________"
    t_firmas.cell(1, 0).text = f"Elaborado por:\n{datos['coordinador']}"
    t_firmas.cell(1, 1).text = f"Recibido por:\n{datos['dirigido_a']}"
    for row in t_firmas.rows:
        for cell in row.cells:
            cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER

    buffer = BytesIO()
    doc.save(buffer)
    buffer.seek(0)
    return buffer


# ═══════════════════════════════════════════════════════════════════════════
# TAB 1: REGISTRO (sin st.form → la cascada Gravedad→Falta se actualiza en vivo)
# ═══════════════════════════════════════════════════════════════════════════
def render_tab_registro() -> None:
    st.markdown('<div class="section-title">📝 Registrar nueva incidencia</div>', unsafe_allow_html=True)
    c1, c2 = st.columns(2)
    with c1:
        fecha = st.date_input("Fecha", value=datetime.date.today())
        estudiante = st.text_input("Nombre del estudiante *")
        genero = st.selectbox("Género", OPCIONES_GENERO)
        ciclo = st.selectbox("Ciclo", OPCIONES_CICLO)
    with c2:
        grado = st.selectbox("Grado", OPCIONES_GRADO)
        seccion = st.text_input("Sección *", placeholder="Ej. A, B, C")
        gravedad = st.selectbox("Nivel de gravedad", list(FALTAS_MINERD.keys()))
        opciones_faltas = FALTAS_MINERD[gravedad] + ["Otra (Especificar)"]
        # La key dinámica fuerza el refresco de opciones al cambiar la gravedad
        falta_sel = st.selectbox("Tipo de falta (Manual MINERD)", opciones_faltas, key=f"falta_{gravedad}")
        if falta_sel == "Otra (Especificar)":
            falta_desc = st.text_area("Describe la falta:", height=70)
        else:
            falta_desc = falta_sel

    if st.button("➕ Registrar incidencia en base de datos", type="primary", width="stretch"):
        if not estudiante.strip() or not seccion.strip() or not falta_desc.strip():
            st.error("⚠️ Estudiante, Sección y Falta son obligatorios.")
        else:
            insertar_incidencia(
                fecha.strftime("%Y-%m-%d"), estudiante.strip().upper(), genero, ciclo,
                grado, seccion.strip().upper(), gravedad, falta_desc,
            )
            st.toast(f"Incidencia de {estudiante.strip()} registrada.", icon="✅")
            st.rerun()

    st.markdown("---")
    st.markdown('<div class="section-title">📋 Incidencias registradas</div>', unsafe_allow_html=True)
    if not db_incidencias:
        st.info("Aún no hay incidencias registradas en la base de datos.")
        return
    df = pd.DataFrame(db_incidencias)
    col_b, col_g = st.columns([2, 1])
    with col_b:
        busq = st.text_input("🔎 Buscar estudiante", placeholder="Nombre o parte del nombre")
    with col_g:
        filtro_grav = st.selectbox("Filtrar por gravedad", ["Todas"] + list(FALTAS_MINERD.keys()))
    df_v = df.copy()
    if busq.strip():
        df_v = df_v[df_v[COL_ESTUDIANTE].astype(str).str.contains(busq.strip(), case=False, na=False)]
    if filtro_grav != "Todas":
        df_v = df_v[df_v[COL_GRAVEDAD] == filtro_grav]
    df_show = df_v.drop(columns=["id"]).copy()
    df_show[COL_GRAVEDAD] = df_show[COL_GRAVEDAD].map(lambda g: MAPA_GRAVEDAD_UI.get(g, g))
    st.dataframe(df_show, width="stretch", hide_index=True, height=380)

    with st.expander("🗑️ Eliminar un registro"):
        st.warning("⚠️ Esta acción borra la constancia de la incidencia.")
        opciones = {
            f"ID {row['id']}: {row[COL_FECHA]} — {row[COL_ESTUDIANTE]} ({str(row[COL_FALTA])[:30]})": row["id"]
            for row in db_incidencias
        }
        seleccion = st.selectbox("Selecciona el registro a eliminar", list(opciones.keys()),
                                 index=None, placeholder="Elige...")
        confirmar = st.checkbox("Confirmo la eliminación", key="chk_del_inc")
        if st.button("Eliminar seleccionado", type="primary", disabled=not (seleccion and confirmar)):
            if seleccion:
                eliminar_incidencia(opciones[seleccion])
                st.toast("Registro eliminado.", icon="🗑️")
                st.rerun()


# ═══════════════════════════════════════════════════════════════════════════
# TAB 2: ANALÍTICA
# ═══════════════════════════════════════════════════════════════════════════
def render_tab_analitica() -> None:
    st.markdown('<div class="section-title">📊 Analítica institucional</div>', unsafe_allow_html=True)
    if not db_incidencias:
        st.info("⚠️ Registra datos para visualizar la analítica.")
        return
    df = pd.DataFrame(db_incidencias)
    m = calcular_metricas(df)
    c1, c2, c3, c4 = st.columns(4)
    with c1:
        st.markdown(_metric_card("Total", m["total"]), unsafe_allow_html=True)
    with c2:
        st.markdown(_metric_card("Casos críticos", m["criticos"], "#DC2626"), unsafe_allow_html=True)
    with c3:
        st.markdown(_metric_card("Masculino", m["masculino"], "#2563EB"), unsafe_allow_html=True)
    with c4:
        st.markdown(_metric_card("Femenino", m["femenino"], "#DB2777"), unsafe_allow_html=True)
    col_g1, col_g2 = st.columns(2)
    with col_g1:
        st.markdown("##### 📈 Incidencias por grado")
        st.bar_chart(df[COL_GRADO].value_counts())
    with col_g2:
        st.markdown("##### 📊 Faltas por sección")
        st.bar_chart(df[COL_SECCION].value_counts())
    st.markdown("##### ⚠️ Distribución de gravedad")
    st.bar_chart(df[COL_GRAVEDAD].value_counts())
    csv = df.drop(columns=["id"]).to_csv(index=False).encode("utf-8-sig")
    st.download_button("📄 Exportar historial (.csv)", data=csv,
                       file_name="Incidencias_historial.csv", mime="text/csv", width="stretch")


# ═══════════════════════════════════════════════════════════════════════════
# TAB 3: REPORTES E IA
# ═══════════════════════════════════════════════════════════════════════════
def render_tab_reportes() -> None:
    st.markdown('<div class="section-title">🤖 Inteligencia institucional y reportes</div>', unsafe_allow_html=True)
    if not db_incidencias:
        st.warning("Necesitas registrar incidencias para usar este módulo.")
        return
    df = pd.DataFrame(db_incidencias)
    estudiantes_sel = st.multiselect("Filtrar por estudiante (vacío = todo el centro)",
                                     df[COL_ESTUDIANTE].unique().tolist())
    df_filtrado = df[df[COL_ESTUDIANTE].isin(estudiantes_sel)] if estudiantes_sel else df

    max_tokens, temperature = ia.control_avanzado(default_tokens=8192, tope=16384, default_temp=0.2)

    if st.button("✨ Generar análisis con IA", type="primary", width="stretch"):
        cfg = ia.config_ia()
        if not cfg["api_key"]:
            st.error("🔒 Configura tu API Key en la página de Inicio.")
        else:
            with st.spinner("Analizando datos con IA..."):
                texto_crudo = None
                try:
                    resumen_datos = {
                        "Total_Incidencias": len(df_filtrado),
                        "Por_Gravedad": df_filtrado[COL_GRAVEDAD].value_counts().to_dict(),
                        "Por_Genero": df_filtrado[COL_GENERO].value_counts().to_dict(),
                        "Faltas_Mas_Comunes": df_filtrado[COL_FALTA].value_counts().head(3).to_dict(),
                    }
                    alcance = (f"los estudiantes: {', '.join(estudiantes_sel)}"
                               if estudiantes_sel else "todo el centro educativo")
                    prompt = f"""Actúa como Asesor Psicopedagógico y Coordinador del MINERD.
Analiza las siguientes estadísticas disciplinarias correspondientes a {alcance}:
DATOS ESTADÍSTICOS:
{json.dumps(resumen_datos, ensure_ascii=False, indent=2)}
TAREA:
Redacta un "Análisis Cualitativo del Clima Escolar" formal, explicando tendencias y significado de los datos.
Redacta "Recomendaciones y Plan de Acción" específicas basadas en enfoques de disciplina positiva del MINERD.
CODIFICACIÓN OBLIGATORIA: salto de línea → {ia.MARKER_NL} · comilla doble → {ia.MARKER_DQ} ·
tabulación → {ia.MARKER_TAB}. Nunca uses saltos de línea literales dentro de un valor JSON.
Devuelve ÚNICAMENTE un JSON válido:
{{
  "analisis_cualitativo": "Texto redactado...",
  "plan_de_accion": "Texto redactado..."
}}
"""
                    texto_crudo, flags = ia.solicitar_ia(
                        prompt, modo="json", max_tokens=max_tokens,
                        temperature=temperature, modulo="incidencias",
                    )
                    datos_ia = ia.decodificar_marcadores(ia.parsear_json_robusto(texto_crudo))
                    analisis = datos_ia.get("analisis_cualitativo", "")
                    plan = datos_ia.get("plan_de_accion", "")
                    if not analisis or not plan:
                        raise ValueError("La respuesta de la IA estaba incompleta.")
                    st.session_state.analisis_ia_temp = analisis
                    st.session_state.plan_ia_temp = plan
                    st.toast("Análisis completado. Revísalo abajo.", icon="✅")
                    if flags.get("reintento"):
                        st.info("ℹ️ La primera llamada se cortó por tokens; el reintento automático tuvo éxito.")
                except ValueError as ve:
                    ia.render_error_ia(ve, texto_crudo)
                except Exception as e:
                    ia.render_error_ia(e, texto_crudo)

    with st.form("form_imprimir"):
        c_r1, c_r2 = st.columns(2)
        with c_r1:
            centro_rep = st.text_input("Centro educativo", value="Politécnico Salesiano Arquides Calderón")
            mes_rep = st.text_input("Período / Motivo", value="Reporte Mensual")
        with c_r2:
            coordinador_rep = st.text_input("Elaborado por", value="Coordinación Técnico-Pedagógica")
            dir_rep = st.text_input("Dirigido a", value="Dirección Académica")
        txt_analisis = st.text_area("I. Análisis cualitativo",
                                    value=st.session_state.get("analisis_ia_temp", ""), height=120)
        txt_plan = st.text_area("II. Plan de acción",
                                value=st.session_state.get("plan_ia_temp", ""), height=120)
        btn_word = st.form_submit_button("📄 Generar y descargar informe Word", type="primary", width="stretch")

    if btn_word:
        datos_informe = {
            "centro": centro_rep, "periodo": mes_rep, "coordinador": coordinador_rep,
            "dirigido_a": dir_rep, "analisis": txt_analisis, "plan_accion": txt_plan,
            "alcance_es_seleccion": bool(estudiantes_sel),
        }
        buffer = generar_informe_docx(datos_informe, df_filtrado)
        st.download_button(
            label="📥 Descargar informe completo (.docx)",
            data=buffer,
            file_name=f"Informe_Incidencias_{datetime.date.today().strftime('%Y%m%d')}.docx",
            mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
            type="primary",
            width="stretch",
        )


# ═══════════════════════════════════════════════════════════════════════════
# MAIN
# ═══════════════════════════════════════════════════════════════════════════
tab1, tab2, tab3 = st.tabs(["📝 Registro", "📊 Analítica", "📄 Reportes e IA"])
with tab1:
    render_tab_registro()
with tab2:
    render_tab_analitica()
with tab3:
    render_tab_reportes()