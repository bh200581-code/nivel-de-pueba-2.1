"""
acuerdos.py — Gestor de Acuerdos y Compromisos Institucionales (MÁXIMO NIVEL EJECUTIVO)
• Eje normativo: Manual de Convivencia MINERD, Ley 66-97 y Ley 136-03.
• Tipo de Acta y Clasificación de falta se AJUSTAN DINÁMICAMENTE a los Actores a intervenir.
• Catálogos de faltas por perfil: Estudiante, Maestro y Familia/Tutor.
• Acta ejecutiva profesional: expediente, fundamentos legales, categorización,
  compromisos por rol, plan de seguimiento y cláusula de confidencialidad.
• IA vía core/ia (marcadores, parseo robusto, reintento, auditoría).
• Persistencia SQLite + historial con filtros, búsqueda y gráficos.
"""
import datetime
import json
import sqlite3
from io import BytesIO

import pandas as pd
import streamlit as st
from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import nsdecls
from docx.oxml import parse_xml, OxmlElement
from docx.shared import Inches, Pt, RGBColor

from core import ia

# ═══════════════════════════════════════════════════════════════════════════
# BASE DE DATOS
# ═══════════════════════════════════════════════════════════════════════════
DB_NAME = "gestion_etp.db"

def init_db():
    conn = sqlite3.connect(DB_NAME, check_same_thread=False)
    cursor = conn.cursor()
    cursor.execute('''
    CREATE TABLE IF NOT EXISTS acuerdos (
        id INTEGER PRIMARY KEY AUTOINCREMENT,
        fecha TEXT, actores TEXT, tipo_acta TEXT, estudiante TEXT,
        docente TEXT, reincidente TEXT, falta TEXT, contexto TEXT,
        nivel_falta TEXT, seguimiento TEXT
    )
    ''')
    conn.commit()
    for col, tipo in [("nivel_falta", "TEXT"), ("seguimiento", "TEXT")]:
        try:
            cursor.execute(f"ALTER TABLE acuerdos ADD COLUMN {col} {tipo}")
            conn.commit()
        except Exception:
            pass
    return conn

conn = init_db()

def get_todos_acuerdos():
    cursor = conn.cursor()
    cursor.execute("""SELECT id, fecha, actores, tipo_acta, estudiante, docente,
                             reincidente, falta, contexto, nivel_falta, seguimiento
                      FROM acuerdos""")
    rows = cursor.fetchall()
    columnas = ["id", "Fecha", "Actores", "Tipo de Acta", "Estudiante", "Docente",
                "Reincidente", "Falta", "Contexto", "Nivel Falta", "Seguimiento"]
    return [dict(zip(columnas, row)) for row in rows]

def insertar_acuerdo(fecha, actores, tipo_acta, estudiante, docente, reincidente,
                     falta, contexto, nivel_falta="", seguimiento=""):
    cursor = conn.cursor()
    cursor.execute('''
    INSERT INTO acuerdos (fecha, actores, tipo_acta, estudiante, docente, reincidente,
                          falta, contexto, nivel_falta, seguimiento)
    VALUES (?, ?, ?, ?, ?, ?, ?, ?, ?, ?)
    ''', (fecha, actores, tipo_acta, estudiante, docente, reincidente,
          falta, contexto, nivel_falta, seguimiento))
    conn.commit()
    return cursor.lastrowid

def eliminar_acuerdo(id_acuerdo):
    cursor = conn.cursor()
    cursor.execute("DELETE FROM acuerdos WHERE id = ?", (id_acuerdo,))
    conn.commit()

# ═══════════════════════════════════════════════════════════════════════════
# CATÁLOGOS DE FALTAS POR PERFIL (Manual de Convivencia MINERD)
# ═══════════════════════════════════════════════════════════════════════════
NINGUNA_FALTA = "Ninguna / No aplica (Acuerdo preventivo)"

FALTAS_ESTUDIANTE = [
    NINGUNA_FALTA,
    "Falta Leve: Impuntualidad reiterada al centro o a clases",
    "Falta Leve: Uso injustificado de dispositivos electrónicos en horario de clases",
    "Falta Leve: Incumplimiento reiterado de tareas o normas del uniforme",
    "Falta Grave: Irrespeto verbal a docentes, directivos o personal de apoyo",
    "Falta Grave: Evasión de clases o abandono del centro educativo sin autorización",
    "Falta Grave: Deterioro o daño intencional a los equipos, mobiliario o instalaciones",
    "Falta Grave: Reincidencia en faltas leves a pesar de intervenciones previas",
    "Falta Muy Grave: Agresión física, acoso (bullying) o peleas dentro del plantel",
    "Falta Muy Grave: Sustracción, robo o daño severo a bienes de la institución o de terceros",
    "Falta Muy Grave: Porte o consumo de sustancias prohibidas / armas en la institución",
    "Falta Muy Grave: Falsificación de firmas, documentos oficiales, registros o evaluaciones",
]

FALTAS_MAESTRO = [
    NINGUNA_FALTA,
    "Falta Leve Docente: Incumplimiento reiterado en la entrega de planificaciones o registros",
    "Falta Leve Docente: Impuntualidad reiterada al inicio de clases o actividades docentes",
    "Falta Leve Docente: Uso inadecuado del tiempo de clase o de los recursos institucionales",
    "Falta Grave Docente: Evasión de responsabilidades docentes sin justificación",
    "Falta Grave Docente: Irrespeto verbal a directivos, colegas, estudiantes o familias",
    "Falta Grave Docente: Negativa reiterada a aplicar adecuaciones curriculares (NEAE)",
    "Falta Muy Grave Docente: Agresión física o maltrato a estudiantes",
    "Falta Muy Grave Docente: Acoso o discriminación a estudiantes o colegas",
    "Falta Muy Grave Docente: Falsificación de documentos oficiales o registros de evaluación",
]

FALTAS_FAMILIA = [
    NINGUNA_FALTA,
    "Falta Leve Familiar: Inasistencia reiterada a reuniones de seguimiento",
    "Falta Leve Familiar: Incumplimiento en el seguimiento de asignaciones en el hogar",
    "Falta Grave Familiar: Negligencia en la atención de requerimientos del centro",
    "Falta Grave Familiar: Irrespeto a docentes o personal del centro",
    "Falta Muy Grave Familiar: Agresión a docentes o personal del centro",
]

# Catálogo enriquecido con nivel, color y consecuencia base normativa
CATALOGO_FALTAS = [
    {"texto": NINGUNA_FALTA, "nivel": "Preventivo", "color": "64748B",
     "consecuencia_base": "Seguimiento, acompañamiento pedagógico y monitoreo preventivo."},
    # ── Estudiante ──
    {"texto": "Falta Leve: Impuntualidad reiterada al centro o a clases", "nivel": "Leve", "color": "F59E0B",
     "consecuencia_base": "Amonestación verbal, registro de seguimiento y compromiso escrito."},
    {"texto": "Falta Leve: Uso injustificado de dispositivos electrónicos en horario de clases", "nivel": "Leve", "color": "F59E0B",
     "consecuencia_base": "Retención temporal del dispositivo, amonestación y compromiso escrito."},
    {"texto": "Falta Leve: Incumplimiento reiterado de tareas o normas del uniforme", "nivel": "Leve", "color": "F59E0B",
     "consecuencia_base": "Amonestación, plan de cumplimiento y notificación a la familia."},
    {"texto": "Falta Grave: Irrespeto verbal a docentes, directivos o personal de apoyo", "nivel": "Grave", "color": "EA580C",
     "consecuencia_base": "Amonestación escrita, citación a la familia y plan de mejora conductual."},
    {"texto": "Falta Grave: Evasión de clases o abandono del centro educativo sin autorización", "nivel": "Grave", "color": "EA580C",
     "consecuencia_base": "Citación urgente a la familia, registro de incidencia y seguimiento de asistencia."},
    {"texto": "Falta Grave: Deterioro o daño intencional a los equipos, mobiliario o instalaciones", "nivel": "Grave", "color": "EA580C",
     "consecuencia_base": "Reposición o reparación del daño, amonestación escrita y compromiso familiar."},
    {"texto": "Falta Grave: Reincidencia en faltas leves a pesar de intervenciones previas", "nivel": "Grave", "color": "EA580C",
     "consecuencia_base": "Intervención de Orientación/Psicología, suspensión de actividades recreativas y plan integral."},
    {"texto": "Falta Muy Grave: Agresión física, acoso (bullying) o peleas dentro del plantel", "nivel": "Muy Grave", "color": "DC2626",
     "consecuencia_base": "Suspensión temporal, intervención multidisciplinaria y remisión según protocolo MINERD."},
    {"texto": "Falta Muy Grave: Sustracción, robo o daño severo a bienes de la institución o de terceros", "nivel": "Muy Grave", "color": "DC2626",
     "consecuencia_base": "Suspensión, reposición del bien y notificación a las autoridades competentes."},
    {"texto": "Falta Muy Grave: Porte o consumo de sustancias prohibidas / armas en la institución", "nivel": "Muy Grave", "color": "DC2626",
     "consecuencia_base": "Suspensión inmediata y remisión a la Fiscalía de Niños, Niñas y Adolescentes (Ley 136-03)."},
    {"texto": "Falta Muy Grave: Falsificación de firmas, documentos oficiales, registros o evaluaciones", "nivel": "Muy Grave", "color": "DC2626",
     "consecuencia_base": "Anulación del documento, suspensión y proceso disciplinario conforme al Reglamento del Centro."},
    # ── Maestro ──
    {"texto": "Falta Leve Docente: Incumplimiento reiterado en la entrega de planificaciones o registros", "nivel": "Leve", "color": "F59E0B",
     "consecuencia_base": "Amonestación verbal, plan de entrega con plazos y seguimiento de Coordinación."},
    {"texto": "Falta Leve Docente: Impuntualidad reiterada al inicio de clases o actividades docentes", "nivel": "Leve", "color": "F59E0B",
     "consecuencia_base": "Amonestación verbal, registro de asistencia y compromiso escrito."},
    {"texto": "Falta Leve Docente: Uso inadecuado del tiempo de clase o de los recursos institucionales", "nivel": "Leve", "color": "F59E0B",
     "consecuencia_base": "Amonestación, plan de mejora y acompañamiento pedagógico."},
    {"texto": "Falta Grave Docente: Evasión de responsabilidades docentes sin justificación", "nivel": "Grave", "color": "EA580C",
     "consecuencia_base": "Amonestación escrita, citación a Dirección y plan de mejora supervisado."},
    {"texto": "Falta Grave Docente: Irrespeto verbal a directivos, colegas, estudiantes o familias", "nivel": "Grave", "color": "EA580C",
     "consecuencia_base": "Amonestación escrita, mediación y compromiso de trato respetuoso."},
    {"texto": "Falta Grave Docente: Negativa reiterada a aplicar adecuaciones curriculares (NEAE)", "nivel": "Grave", "color": "EA580C",
     "consecuencia_base": "Intervención de Orientación, capacitación obligatoria y seguimiento de adecuaciones."},
    {"texto": "Falta Muy Grave Docente: Agresión física o maltrato a estudiantes", "nivel": "Muy Grave", "color": "DC2626",
     "consecuencia_base": "Suspensión inmediata, proceso disciplinario y remisión a las autoridades competentes."},
    {"texto": "Falta Muy Grave Docente: Acoso o discriminación a estudiantes o colegas", "nivel": "Muy Grave", "color": "DC2626",
     "consecuencia_base": "Suspensión, investigación disciplinaria y remisión según protocolo MINERD."},
    {"texto": "Falta Muy Grave Docente: Falsificación de documentos oficiales o registros de evaluación", "nivel": "Muy Grave", "color": "DC2626",
     "consecuencia_base": "Anulación del documento, proceso disciplinario y remisión a las autoridades competentes."},
    # ── Familia ──
    {"texto": "Falta Leve Familiar: Inasistencia reiterada a reuniones de seguimiento", "nivel": "Leve", "color": "F59E0B",
     "consecuencia_base": "Notificación escrita, nueva citación y registro de inasistencia."},
    {"texto": "Falta Leve Familiar: Incumplimiento en el seguimiento de asignaciones en el hogar", "nivel": "Leve", "color": "F59E0B",
     "consecuencia_base": "Compromiso escrito de corresponsabilidad y seguimiento periódico."},
    {"texto": "Falta Grave Familiar: Negligencia en la atención de requerimientos del centro", "nivel": "Grave", "color": "EA580C",
     "consecuencia_base": "Citación formal a Dirección, acta de compromiso y remisión a Orientación."},
    {"texto": "Falta Grave Familiar: Irrespeto a docentes o personal del centro", "nivel": "Grave", "color": "EA580C",
     "consecuencia_base": "Amonestación escrita, mediación y compromiso de trato respetuoso."},
    {"texto": "Falta Muy Grave Familiar: Agresión a docentes o personal del centro", "nivel": "Muy Grave", "color": "DC2626",
     "consecuencia_base": "Denuncia ante las autoridades competentes y restricción de acceso al centro."},
]

FALTAS_TEXTOS = [f["texto"] for f in CATALOGO_FALTAS]

def get_info_falta(texto_falta):
    for f in CATALOGO_FALTAS:
        if f["texto"] == texto_falta:
            return f
    return {"texto": texto_falta, "nivel": "No clasificada", "color": "94A3B8",
            "consecuencia_base": "Según determinación de la autoridad competente."}

# ═══════════════════════════════════════════════════════════════════════════
# MATRIZ DE ACTORES: tipos de acta y faltas ajustadas por actor
# ═══════════════════════════════════════════════════════════════════════════
OPCIONES_ACUERDOS = {
    "👨🏫 Maestro(a) + Coordinación": {
        "perfil_faltas": "Docente",
        "tipos_acta": [
            "Compromiso de entrega de planificaciones atrasadas",
            "Acuerdo de mejora en la gestión de la disciplina en el aula",
            "Acta de compromiso de puntualidad y asistencia",
            "Acuerdo de implementación de estrategias ETP (simuladores/talleres)",
            "Compromiso de participación en formaciones pedagógicas",
            "Acta de compromiso de actualización y uso del registro de evaluación",
        ],
        "faltas": FALTAS_MAESTRO,
    },
    "👨‍🎓 Estudiante + Coordinación": {
        "perfil_faltas": "Estudiante",
        "tipos_acta": [
            "Compromiso de recuperación académica (R.A. no logrados)",
            "Acta de compromiso conductual (disciplina, respeto, uniforme)",
            "Acuerdo de cuidado de equipos e infraestructura del Politécnico",
            "Compromiso de mejora en la asistencia y puntualidad",
            "Acta de compromiso por uso inadecuado de dispositivos electrónicos",
        ],
        "faltas": FALTAS_ESTUDIANTE,
    },
    "👨‍👩‍👧 Tutor(a) + Coordinación": {
        "perfil_faltas": "Familia / Tutor",
        "tipos_acta": [
            "Compromiso de seguimiento a las asignaciones en el hogar",
            "Acta de asistencia obligatoria a reuniones de seguimiento",
            "Acuerdo de corresponsabilidad por faltas disciplinarias del estudiante",
            "Compromiso de reposición/pago por daños a equipos causados por el estudiante",
        ],
        "faltas": FALTAS_FAMILIA,
    },
    "👨‍👩‍👧 Tutor + 👨‍🎓 Estudiante + Coordinación": {
        "perfil_faltas": "Estudiante",
        "tipos_acta": [
            "Acuerdo integral de recuperación académica y apoyo familiar",
            "Acta de compromiso disciplinario definitivo (Última advertencia)",
            "Acuerdo de reposición de ausencias prácticas en módulos formativos",
        ],
        "faltas": FALTAS_ESTUDIANTE,
    },
    "👨‍🏫 Maestro(a) + 👨‍🎓 Estudiante + Coordinación": {
        "perfil_faltas": "Estudiante",
        "tipos_acta": [
            "Acuerdo de mediación y resolución de conflicto en el aula/taller",
            "Compromiso mutuo de nivelación y entrega de prácticas atrasadas",
            "Acta de compromiso de respeto mutuo y mejora del clima escolar",
        ],
        "faltas": FALTAS_ESTUDIANTE,
    },
    "🧠 Psicología + 👨‍🎓 Estudiante + Coordinación": {
        "perfil_faltas": "Estudiante",
        "tipos_acta": [
            "Acuerdo de seguimiento conductual y apoyo socioemocional",
            "Acta de intervención por bajo rendimiento académico o desmotivación",
            "Compromiso de asistencia a sesiones de apoyo psicopedagógico",
        ],
        "faltas": FALTAS_ESTUDIANTE,
    },
    "🧠 Psicología + 👨‍👩‍👧 Tutor + 👨‍🎓 Estudiante + Coordinación": {
        "perfil_faltas": "Estudiante",
        "tipos_acta": [
            "Acuerdo integral de intervención psicoeducativa y apoyo familiar",
            "Acta de compromiso por faltas graves a la convivencia escolar",
            "Acuerdo de seguimiento a necesidades específicas de apoyo educativo (NEAE)",
            "Compromiso familiar para evaluación externa (psicológica/neurológica)",
        ],
        "faltas": FALTAS_ESTUDIANTE,
    },
    "🧠 Psicología + 👨‍👩‍👧 Tutor + 👨‍🏫 Maestro(a) + Coordinación": {
        "perfil_faltas": "Estudiante",
        "tipos_acta": [
            "Acta de mediación familiar-docente para manejo conjunto del estudiante",
            "Acuerdo de adaptación curricular y estrategias de seguimiento",
            "Compromiso de comunicación asertiva entre familia y escuela",
        ],
        "faltas": FALTAS_ESTUDIANTE,
    },
    "🧠 Psicología + 👨‍🏫 Maestro(a) + Coordinación": {
        "perfil_faltas": "Docente",
        "tipos_acta": [
            "Acuerdo de implementación de estrategias de manejo conductual en el aula",
            "Acta de orientación pedagógica para el abordaje de estudiantes con NEAE",
            "Compromiso de derivación oportuna de casos al departamento de orientación",
        ],
        "faltas": FALTAS_MAESTRO,
    },
    "🧠 Psicología + 👨‍🏫 Maestro + 👨‍👩‍👧 Tutor + 👨‍🎓 Estudiante + Coordinación": {
        "perfil_faltas": "Estudiante",
        "tipos_acta": [
            "Acta plenaria de intervención, mediación y compromiso integral (Caso Crítico)",
            "Acuerdo general de convivencia y plan de acción psicoeducativo unificado",
        ],
        "faltas": FALTAS_ESTUDIANTE,
    },
}

FUNDAMENTOS_LEGALES = [
    "Ley General de Educación No. 66-97 y sus modificaciones, que establece el marco normativo del sistema educativo dominicano.",
    "Ley No. 136-03, Código para el Sistema de Protección y los Derechos Fundamentales de Niños, Niñas y Adolescentes.",
    "Manual de Convivencia Escolar del Ministerio de Educación de la República Dominicana (MINERD).",
    "Reglamento Interno del Centro Educativo y normas institucionales vigentes.",
]

# ═══════════════════════════════════════════════════════════════════════════
# ESTILOS (Máximo Nivel)
# ═══════════════════════════════════════════════════════════════════════════
st.markdown("""
<style>
@import url('https://fonts.googleapis.com/css2?family=Inter:wght@300;400;500;600;700;800;900&display=swap');

html, body, [class*="css"] {
    font-family: 'Inter', sans-serif;
    background-color: #F0F4F8;
    color: #1E293B;
}

.acuerdos-hero {
    background: linear-gradient(135deg, #0F172A 0%, #1E3A8A 40%, #4338CA 70%, #6366F1 100%);
    color: #fff;
    padding: 2.2rem;
    border-radius: 20px;
    margin-bottom: 1.5rem;
    box-shadow: 0 25px 50px rgba(30, 58, 138, 0.3);
    position: relative;
    overflow: hidden;
}

.acuerdos-hero::before {
    content: '';
    position: absolute;
    top: -50%; left: -50%; width: 200%; height: 200%;
    background: radial-gradient(circle, rgba(255,255,255,0.1) 0%, transparent 60%);
    animation: acuerdosPulse 6s ease-in-out infinite;
}

@keyframes acuerdosPulse {
    0%, 100% { transform: scale(1); opacity: 0.5; }
    50% { transform: scale(1.1); opacity: 0.8; }
}

.acuerdos-hero-title { font-size: 2.4rem; font-weight: 900; letter-spacing: -0.03em; margin-bottom: 0.4rem; position: relative; }
.acuerdos-hero-sub { font-size: 1.05rem; opacity: 0.9; line-height: 1.5; position: relative; }

.acuerdos-hero-badge {
    display: inline-block;
    background: rgba(255,255,255,0.15);
    border: 1px solid rgba(255,255,255,0.25);
    border-radius: 8px;
    padding: 4px 12px;
    font-size: 0.8rem;
    font-weight: 600;
    margin-top: 0.8rem;
    margin-right: 8px;
    position: relative;
}

.acuerdos-section-title {
    color: #1E3A8A;
    font-weight: 700;
    font-size: 1.12rem;
    border-bottom: 2px solid #DBEAFE;
    padding-bottom: 8px;
    margin: 1.2rem 0 0.9rem 0;
}

.acuerdos-kpi-card {
    background: #fff;
    border-radius: 14px;
    padding: 1.2rem;
    text-align: center;
    border: 1px solid #E2E8F0;
    box-shadow: 0 4px 12px rgba(0,0,0,0.06);
    transition: all 0.3s ease;
}

.acuerdos-kpi-card:hover { transform: translateY(-4px); box-shadow: 0 12px 24px rgba(0,0,0,0.12); }
.acuerdos-kpi-icon { font-size: 2rem; margin-bottom: 0.5rem; }
.acuerdos-kpi-value { font-size: 2rem; font-weight: 800; color: #1E3A8A; }
.acuerdos-kpi-label { font-size: 0.75rem; font-weight: 700; color: #94A3B8; text-transform: uppercase; letter-spacing: 0.05em; margin-top: 0.3rem; }

.acuerdos-severidad-box {
    border-radius: 12px;
    padding: 1rem 1.2rem;
    margin-top: 0.8rem;
    border-left: 5px solid;
    transition: all 0.3s ease;
}

.acuerdos-severidad-nivel { font-weight: 800; font-size: 1rem; }
.acuerdos-severidad-desc { font-size: 0.85rem; margin-top: 0.4rem; line-height: 1.4; }

.acuerdos-perfil-badge {
    display: inline-block;
    background: #E0E7FF;
    color: #3730A3;
    border-radius: 6px;
    padding: 3px 10px;
    font-size: 0.75rem;
    font-weight: 700;
    margin-top: 0.5rem;
}

.acuerdos-legal-box {
    background: #F8FAFC;
    border: 1px solid #E2E8F0;
    border-radius: 10px;
    padding: 1rem;
    font-size: 0.85rem;
    color: #475569;
}

.acuerdos-flash {
    padding: 12px 16px;
    border-radius: 10px;
    margin-bottom: 12px;
    font-weight: 500;
    animation: slideIn 0.3s ease;
}

@keyframes slideIn {
    from { opacity: 0; transform: translateY(-10px); }
    to { opacity: 1; transform: translateY(0); }
}

.acuerdos-flash-success { background: #D1FAE5; color: #065F46; border-left: 4px solid #10B981; }
.acuerdos-flash-error { background: #FEE2E2; color: #991B1B; border-left: 4px solid #EF4444; }
</style>
""", unsafe_allow_html=True)

# ═══════════════════════════════════════════════════════════════════════════
# GUARDIA + SIDEBAR + HERO
# ═══════════════════════════════════════════════════════════════════════════
if not st.session_state.get("coordinador_autenticado", False):
    st.error("🔒 Esta página es exclusiva de Coordinación.")
    st.stop()

ia.panel_sidebar_ia("Gestor de Acuerdos")

st.markdown("""
<div class="acuerdos-hero">
    <div class="acuerdos-hero-title">⚖️ Gestor de Acuerdos y Compromisos Institucionales</div>
    <div class="acuerdos-hero-sub">
        Actas ejecutivas alineadas al Manual de Convivencia MINERD · Tipo de Acta y Clasificación ajustados a los Actores
    </div>
    <div>
        <span class="acuerdos-hero-badge">📖 Manual de Convivencia MINERD</span>
        <span class="acuerdos-hero-badge">⚖️ Ley 66-97 · Ley 136-03</span>
        <span class="acuerdos-hero-badge">🎯 Faltas por Perfil de Actor</span>
        <span class="acuerdos-hero-badge">🗂️ Archivo Permanente</span>
    </div>
</div>
""", unsafe_allow_html=True)

# Flash messages
if "flash" in st.session_state:
    st.markdown(f'<div class="acuerdos-flash acuerdos-flash-success">✅ {st.session_state.pop("flash")}</div>', unsafe_allow_html=True)
if "flash_error" in st.session_state:
    st.markdown(f'<div class="acuerdos-flash acuerdos-flash-error">❌ {st.session_state.pop("flash_error")}</div>', unsafe_allow_html=True)

# ═══════════════════════════════════════════════════════════════════════════
# KPIs
# ═══════════════════════════════════════════════════════════════════════════
db_acuerdos = get_todos_acuerdos()

if db_acuerdos:
    df_kpi = pd.DataFrame(db_acuerdos)
    total_actas = len(df_kpi)
    reincidentes = int((df_kpi["Reincidente"] == "Sí").sum())
    hoy = datetime.date.today()
    semana_atras = hoy - datetime.timedelta(days=7)
    recientes = int((pd.to_datetime(df_kpi["Fecha"], errors="coerce").dt.date >= semana_atras).sum())
    niveles = df_kpi["Nivel Falta"].dropna()
    nivel_top = niveles.mode().iloc[0] if not niveles.empty else "—"
else:
    total_actas = reincidentes = recientes = 0
    nivel_top = "—"

col_k1, col_k2, col_k3, col_k4 = st.columns(4)
with col_k1:
    st.markdown(f"""
    <div class="acuerdos-kpi-card">
        <div class="acuerdos-kpi-icon">📄</div>
        <div class="acuerdos-kpi-value">{total_actas}</div>
        <div class="acuerdos-kpi-label">Actas Totales</div>
    </div>
    """, unsafe_allow_html=True)
with col_k2:
    st.markdown(f"""
    <div class="acuerdos-kpi-card">
        <div class="acuerdos-kpi-icon">🔄</div>
        <div class="acuerdos-kpi-value">{reincidentes}</div>
        <div class="acuerdos-kpi-label">Reincidentes</div>
    </div>
    """, unsafe_allow_html=True)
with col_k3:
    st.markdown(f"""
    <div class="acuerdos-kpi-card">
        <div class="acuerdos-kpi-icon">🕐</div>
        <div class="acuerdos-kpi-value">{recientes}</div>
        <div class="acuerdos-kpi-label">Últimos 7 Días</div>
    </div>
    """, unsafe_allow_html=True)
with col_k4:
    st.markdown(f"""
    <div class="acuerdos-kpi-card">
        <div class="acuerdos-kpi-icon">🎯</div>
        <div class="acuerdos-kpi-value" style="font-size:1.3rem;">{nivel_top}</div>
        <div class="acuerdos-kpi-label">Nivel Frecuente</div>
    </div>
    """, unsafe_allow_html=True)

# ═══════════════════════════════════════════════════════════════════════════
# TABS
# ═══════════════════════════════════════════════════════════════════════════
tab1, tab2 = st.tabs(["📝 Redactar Acta Ejecutiva", "📋 Historial de Acuerdos"])

# ═══════════════════════════════════════════════════════════════════════════
# GENERACIÓN DEL ACTA EJECUTIVA (Word profesional)
# ═══════════════════════════════════════════════════════════════════════════
def shade_cell(cell, color):
    shd = parse_xml(r'<w:shd {} w:fill="{}"/>'.format(nsdecls('w'), color))
    cell._tc.get_or_add_tcPr().append(shd)

def set_cell_text(cell, text, bold=False, center=False, color=None, size=None, font_color=None):
    cell.text = ""
    p = cell.paragraphs[0]
    run = p.add_run(str(text))
    run.bold = bold
    if size:
        run.font.size = Pt(size)
    if font_color:
        run.font.color.rgb = font_color
    if center:
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    if color:
        shade_cell(cell, color)

def add_page_footer(doc, centro):
    for section in doc.sections:
        footer = section.footer
        footer.is_linked_to_previous = False
        p = footer.paragraphs[0] if footer.paragraphs else footer.add_paragraph()
        p.text = ""
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run(f"{centro} · Documento confidencial conforme a la Ley 136-03 · Página ")
        run.font.size = Pt(7)
        run.font.color.rgb = RGBColor(0x64, 0x74, 0x8B)
        fld1 = OxmlElement("w:fldChar"); fld1.set(qn("w:fldCharType"), "begin")
        instr = OxmlElement("w:instrText"); instr.set(qn("xml:space"), "preserve"); instr.text = "PAGE"
        fld2 = OxmlElement("w:fldChar"); fld2.set(qn("w:fldCharType"), "end")
        run2 = p.add_run()
        run2._r.append(fld1); run2._r.append(instr); run2._r.append(fld2)
        run2.font.size = Pt(7)
        run2.font.color.rgb = RGBColor(0x64, 0x74, 0x8B)

def build_acta_ejecutiva(datos_json, meta, expediente):
    doc = Document()
    doc.styles["Normal"].font.name = "Calibri"
    doc.styles["Normal"].font.size = Pt(10.5)
    for section in doc.sections:
        section.left_margin = Inches(0.9)
        section.right_margin = Inches(0.9)
        section.top_margin = Inches(0.8)
        section.bottom_margin = Inches(0.8)

    NAVY = RGBColor(0x1F, 0x38, 0x64)
    GRIS = RGBColor(0x59, 0x59, 0x59)

    def encabezado_seccion(texto):
        p = doc.add_paragraph()
        p.paragraph_format.space_before = Pt(12)
        p.paragraph_format.space_after = Pt(5)
        run = p.add_run(texto)
        run.bold = True
        run.font.size = Pt(12)
        run.font.color.rgb = NAVY
        pPr = p._p.get_or_add_pPr()
        pBdr = parse_xml(
            r'<w:pBdr {}><w:bottom w:val="single" w:sz="8" w:space="1" w:color="1F3864"/></w:pBdr>'.format(nsdecls('w'))
        )
        pPr.append(pBdr)
        return p

    p_minerd = doc.add_paragraph()
    p_minerd.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run_m = p_minerd.add_run("MINISTERIO DE EDUCACIÓN DE LA REPÚBLICA DOMINICANA")
    run_m.bold = True
    run_m.font.size = Pt(11)
    run_m.font.color.rgb = NAVY

    p_centro = doc.add_paragraph()
    p_centro.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run_c = p_centro.add_run(f"{meta['centro']}\nCoordinación Técnico-Pedagógica y Departamento de Orientación")
    run_c.bold = True
    run_c.font.size = Pt(10)

    doc.add_paragraph()

    tabla_exp = doc.add_table(rows=1, cols=2)
    tabla_exp.style = "Table Grid"
    set_cell_text(tabla_exp.rows[0].cells[0], f"EXPEDIENTE: {expediente}", bold=True, color="E2E8F0", size=9)
    set_cell_text(tabla_exp.rows[0].cells[1], f"FECHA: {meta['fecha']}", bold=True, color="E2E8F0", size=9)

    doc.add_paragraph()

    p_titulo = doc.add_paragraph()
    p_titulo.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run_t = p_titulo.add_run(datos_json.get("titulo_acta", "ACTA DE COMPROMISO INSTITUCIONAL"))
    run_t.bold = True
    run_t.font.size = Pt(14)
    run_t.font.color.rgb = NAVY

    p_sub = doc.add_paragraph()
    p_sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run_sub = p_sub.add_run(meta["tipo_acta"])
    run_sub.italic = True
    run_sub.font.size = Pt(10)
    run_sub.font.color.rgb = GRIS

    doc.add_paragraph()

    encabezado_seccion("I. IDENTIFICACIÓN DE LAS PARTES")
    tabla_partes = doc.add_table(rows=1, cols=2)
    tabla_partes.style = "Table Grid"
    hdr = tabla_partes.rows[0].cells
    set_cell_text(hdr[0], "Actor", bold=True, center=True, color="1F3864", font_color=RGBColor(0xFF, 0xFF, 0xFF))
    set_cell_text(hdr[1], "Nombre", bold=True, center=True, color="1F3864", font_color=RGBColor(0xFF, 0xFF, 0xFF))

    partes = [
        ("Actores intervinientes", meta["actores"]),
        ("Perfil de faltas aplicado", meta.get("perfil_faltas", "N/A")),
        ("Estudiante", meta.get("estudiante", "N/A")),
        ("Docente / Maestro(a)", meta.get("docente", "N/A")),
        ("Coordinador(a) Técnico-Pedagógico", meta["coordinador"]),
        ("Condición de reincidencia", meta["reincidencia"]),
    ]
    for actor, nombre in partes:
        row = tabla_partes.add_row().cells
        set_cell_text(row[0], actor, bold=True, color="F1F5F9")
        set_cell_text(row[1], nombre)

    doc.add_paragraph()

    encabezado_seccion("II. ANTECEDENTES Y CONTEXTUALIZACIÓN")
    p_ant = doc.add_paragraph(datos_json.get("antecedentes", ""))
    p_ant.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

    encabezado_seccion("III. FUNDAMENTOS LEGALES Y NORMATIVOS")
    p_intro_legal = doc.add_paragraph(
        "El presente acuerdo se fundamenta en el marco normativo vigente del sistema educativo dominicano:"
    )
    p_intro_legal.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    for fundamento in FUNDAMENTOS_LEGALES:
        p_f = doc.add_paragraph(style="List Bullet")
        p_f.add_run(fundamento)
        p_f.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    fundamentos_ia = datos_json.get("fundamentos_legales", [])
    for f_ia in fundamentos_ia:
        p_fia = doc.add_paragraph(style="List Bullet")
        p_fia.add_run(f_ia)
        p_fia.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

    doc.add_paragraph()

    encabezado_seccion("IV. CATEGORIZACIÓN DE LA FALTA")
    info_falta = get_info_falta(meta["falta"])
    tabla_falta = doc.add_table(rows=3, cols=2)
    tabla_falta.style = "Table Grid"
    set_cell_text(tabla_falta.rows[0].cells[0], "Falta imputada", bold=True, color="F1F5F9")
    set_cell_text(tabla_falta.rows[0].cells[1], info_falta["texto"])
    set_cell_text(tabla_falta.rows[1].cells[0], "Nivel de severidad", bold=True, color="F1F5F9")
    set_cell_text(tabla_falta.rows[1].cells[1], info_falta["nivel"], bold=True,
                  color=info_falta["color"], font_color=RGBColor(0xFF, 0xFF, 0xFF))
    set_cell_text(tabla_falta.rows[2].cells[0], "Consecuencia base normativa", bold=True, color="F1F5F9")
    set_cell_text(tabla_falta.rows[2].cells[1], info_falta["consecuencia_base"])

    analisis_falta = datos_json.get("analisis_falta", "")
    if analisis_falta:
        p_an = doc.add_paragraph(analisis_falta)
        p_an.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

    doc.add_paragraph()

    encabezado_seccion("V. ACUERDOS Y COMPROMISOS ADOPTADOS")
    p_intro_ac = doc.add_paragraph(
        "Las partes intervinientes acuerdan de manera libre, voluntaria e informada los siguientes compromisos:"
    )
    p_intro_ac.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    for comp in datos_json.get("compromisos", []):
        p_comp = doc.add_paragraph(style="List Bullet")
        p_comp.add_run(comp)
        p_comp.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

    doc.add_paragraph()

    encabezado_seccion("VI. CONSECUENCIAS EN CASO DE INCUMPLIMIENTO")
    p_cons = doc.add_paragraph(datos_json.get("consecuencias", ""))
    p_cons.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

    encabezado_seccion("VII. PLAN DE SEGUIMIENTO Y VERIFICACIÓN")
    p_seg = doc.add_paragraph(datos_json.get("plan_seguimiento", ""))
    p_seg.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    if meta.get("seguimiento"):
        p_fecha_seg = doc.add_paragraph()
        p_fecha_seg.add_run("Próxima fecha de revisión: ").bold = True
        p_fecha_seg.add_run(meta["seguimiento"])

    doc.add_paragraph()

    encabezado_seccion("VIII. DISPOSICIONES FINALES")
    p_disp = doc.add_paragraph(datos_json.get("disposiciones_finales", ""))
    p_disp.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

    doc.add_paragraph()

    encabezado_seccion("FIRMAS DE CONFORMIDAD")
    p_firmas_intro = doc.add_paragraph(
        "En señal de conformidad con lo establecido en la presente acta, las partes intervinientes firman "
        "al pie del presente documento, en el entendido de que su incumplimiento dará lugar a las "
        "consecuencias previstas en la normativa institucional y en el Manual de Convivencia del MINERD."
    )
    p_firmas_intro.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

    doc.add_paragraph("\n\n")
    firmas = datos_json.get("firmantes", [])
    cols_count = 2
    t_firmas = doc.add_table(rows=0, cols=cols_count)
    for i in range(0, len(firmas), cols_count):
        row_line = t_firmas.add_row().cells
        row_text = t_firmas.add_row().cells
        row_line[0].text = "______________________________"
        row_text[0].text = firmas[i]
        row_line[0].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
        row_text[0].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
        if i + 1 < len(firmas):
            row_line[1].text = "______________________________"
            row_text[1].text = firmas[i + 1]
            row_line[1].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
            row_text[1].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
        t_firmas.add_row()

    add_page_footer(doc, meta["centro"])

    buffer = BytesIO()
    doc.save(buffer)
    buffer.seek(0)
    return buffer

# ═══════════════════════════════════════════════════════════════════════════
# TAB 1: REDACTAR ACTA EJECUTIVA
# ═══════════════════════════════════════════════════════════════════════════
# ═══════════════════════════════════════════════════════════════════════════
# TAB 1: REDACTAR ACTA EJECUTIVA (selectores reactivos, SIN formulario)
# ═══════════════════════════════════════════════════════════════════════════
with tab1:
    st.markdown('<div class="acuerdos-section-title">🏛️ 1. Datos Institucionales</div>', unsafe_allow_html=True)
    col1, col2 = st.columns(2)
    with col1:
        centro = st.text_input("Centro Educativo", value="Politécnico Salesiano Arquides Calderón")
    with col2:
        coordinador = st.text_input("Coordinador Técnico-Pedagógico", value="Ing. Bernardo Antonio Hernández Batista")

    st.markdown('<div class="acuerdos-section-title">👥 2. Involucrados y Naturaleza del Acta</div>', unsafe_allow_html=True)
    st.info("💡 El **Tipo de Acta** y la **Clasificación de la Falta** se ajustan automáticamente según los **Actores a intervenir** seleccionados.")

    # ── SELECTOR DE ACTORES (reactivo, fuera de formulario) ──
    actor_seleccionado = st.selectbox("Actores a intervenir:", list(OPCIONES_ACUERDOS.keys()))
    config_actor = OPCIONES_ACUERDOS[actor_seleccionado]

    # Badge del perfil de faltas aplicado
    st.markdown(
        f'<span class="acuerdos-perfil-badge">🎯 Perfil de faltas aplicado: {config_actor["perfil_faltas"]}</span>',
        unsafe_allow_html=True,
    )

    # ── TIPO DE ACTA (depende del actor, reactivo) ──
    tipo_acuerdo = st.selectbox(
        "Tipo de Acta / Compromiso:",
        config_actor["tipos_acta"],
        key=f"tipo_acta_{actor_seleccionado}",
    )

    col_f1, col_f2, col_f3 = st.columns([1, 1, 1])
    with col_f1:
        reincidente = st.checkbox("🔄 ¿Es un caso de REINCIDENCIA?", value=False,
                                  help="Marca esta casilla si ya había firmado un acuerdo previo.")
    with col_f2:
        # ── CLASIFICACIÓN (depende del actor, reactivo) ──
        falta_manual = st.selectbox(
            "Clasificación según Manual de Convivencia MINERD:",
            config_actor["faltas"],
            key=f"falta_{actor_seleccionado}",
        )
    with col_f3:
        fecha_seguimiento = st.date_input(
            "📅 Próxima fecha de revisión",
            value=datetime.date.today() + datetime.timedelta(days=15),
            help="Fecha en la que se verificará el cumplimiento del acuerdo."
        )

    # Indicador visual de severidad
    info_falta_sel = get_info_falta(falta_manual)
    st.markdown(f"""
    <div class="acuerdos-severidad-box" style="background: #{info_falta_sel['color']}15; border-left-color: #{info_falta_sel['color']};">
        <div class="acuerdos-severidad-nivel" style="color: #{info_falta_sel['color']};">
            ⚖️ Nivel de Severidad: {info_falta_sel['nivel']}
        </div>
        <div class="acuerdos-severidad-desc">
            <b>Consecuencia base normativa:</b> {info_falta_sel['consecuencia_base']}
        </div>
    </div>
    """, unsafe_allow_html=True)

    # ── Campos condicionales según actores ──
    col_d1, col_d2 = st.columns(2)
    with col_d1:
        nombre_psicologia = st.text_input("Nombre del Psicólogo(a) u Orientador(a)", placeholder="Ej: Licda. Carmen Torres") if "Psicología" in actor_seleccionado else ""
        nombre_docente = st.text_input("Nombre del Maestro(a)", placeholder="Ej: Licda. María Pérez") if "Maestro" in actor_seleccionado else ""
        nombre_tutor = st.text_input("Nombre del Padre/Madre o Tutor", placeholder="Ej: Sr. Juan Rodríguez") if "Tutor" in actor_seleccionado else ""
    with col_d2:
        if "Estudiante" in actor_seleccionado or "Tutor" in actor_seleccionado:
            nombre_estudiante = st.text_input("Nombre del Estudiante", placeholder="Ej: Luis Rodríguez")
            modulo = st.text_input("Módulo Formativo / Grado", placeholder="Ej: 4to Informática")
        elif "Maestro" in actor_seleccionado:
            nombre_estudiante = ""
            modulo = st.text_input("Módulo Formativo impartido", placeholder="Ej: Redes LAN")
        else:
            nombre_estudiante = ""
            modulo = ""

    st.markdown('<div class="acuerdos-section-title">📝 3. Contexto de la Situación</div>', unsafe_allow_html=True)
    st.info("Explica brevemente qué originó la necesidad de este acuerdo. La IA usará esto junto con la clasificación del Manual de Convivencia para redactar los antecedentes y fundamentos legales.")
    contexto = st.text_area("Detalles del incidente o situación actual:", height=100,
                            placeholder="Ej: El estudiante presenta episodios de agresividad en el aula, negándose a acatar instrucciones del docente...")

    with st.expander("⚖️ Referencias legales que se aplicarán (Manual de Convivencia MINERD)"):
        st.markdown('<div class="acuerdos-legal-box">', unsafe_allow_html=True)
        for f in FUNDAMENTOS_LEGALES:
            st.markdown(f"📌 {f}")
        st.markdown(f"**Perfil de faltas aplicado:** {config_actor['perfil_faltas']}")
        st.markdown(f"**Nivel de la falta seleccionada:** {info_falta_sel['nivel']}")
        st.markdown(f"**Consecuencia base:** {info_falta_sel['consecuencia_base']}")
        st.markdown('</div>', unsafe_allow_html=True)

    max_tokens, temperature = ia.control_avanzado(default_tokens=8192, tope=16384, default_temp=0.2)

    st.markdown("<br>", unsafe_allow_html=True)
    # ── BOTÓN DE GENERACIÓN (botón normal, no de formulario) ──
    submit_button = st.button("⚖️ Generar Acta Ejecutiva y Guardar", type="primary", width="stretch")

    if submit_button:
        cfg = ia.config_ia()
        if not cfg["api_key"]:
            st.error("🔒 Configura tu API Key en la barra lateral (página de Inicio).")
        elif not contexto.strip():
            st.warning("⚠️ Debes proporcionar los detalles del contexto para redactar el acta.")
        else:
            with st.spinner(f'🧠 Redactando Acta Ejecutiva con {cfg["modelo"]}...'):
                texto_crudo = None
                try:
                    firmantes_lista = [f"Coordinación Técnico-Pedagógica:\n{coordinador}"]
                    if nombre_psicologia:
                        firmantes_lista.append(f"Orientación y Psicología:\n{nombre_psicologia}")
                    if nombre_docente:
                        firmantes_lista.append(f"Docente:\n{nombre_docente}")
                    if nombre_tutor:
                        firmantes_lista.append(f"Padre/Madre o Tutor:\n{nombre_tutor}")
                    if nombre_estudiante:
                        firmantes_lista.append(f"Estudiante:\n{nombre_estudiante}")
                    firmantes_lista.append("Dirección Académica")

                    estado_reincidencia = ("SÍ (Caso reincidente con historial previo de acuerdos)"
                                           if reincidente else "NO (Primera intervención por esta causa)")
                    info_falta = get_info_falta(falta_manual)
                    fecha_seg_str = fecha_seguimiento.strftime("%d/%m/%Y")
                    perfil_faltas = config_actor["perfil_faltas"]

                    prompt_maestro = f"""Actúa como un Coordinador Técnico-Pedagógico Nivel Máster y Especialista en Legislación Escolar del MINERD (República Dominicana).

Tu tarea es redactar un "Acta de Compromiso Institucional" formal, legal y fundamentada en el Manual de Convivencia del MINERD, la Ley General de Educación 66-97 y la Ley 136-03.

DATOS DEL ACTA:
TIPO DE ACTA: {tipo_acuerdo}
ACTORES INVOLUCRADOS: {actor_seleccionado}
PERFIL DE FALTAS APLICADO: {perfil_faltas}
Coordinador: {coordinador}
Psicólogo/a: {nombre_psicologia if nombre_psicologia else "N/A"}
Docente: {nombre_docente if nombre_docente else "N/A"}
Estudiante: {nombre_estudiante if nombre_estudiante else "N/A"}
Tutor: {nombre_tutor if nombre_tutor else "N/A"}
Módulo/Grado: {modulo if modulo else "N/A"}
¿ES REINCIDENTE?: {estado_reincidencia}
CLASIFICACIÓN SEGÚN MANUAL DE CONVIVENCIA: {falta_manual}
NIVEL DE SEVERIDAD: {info_falta['nivel']}
CONSECUENCIA BASE NORMATIVA: {info_falta['consecuencia_base']}
CONTEXTO / MOTIVO DE LA REUNIÓN:
{contexto}

REGLAS DE REDACCIÓN (ESTILO EJECUTIVO-LEGAL):
Usa un tono institucional, formal, firme y fundamentado en las normativas del MINERD y el Manual de Convivencia.
- "antecedentes": Describe los hechos de manera objetiva, cronológica y contextualizada.
- "fundamentos_legales": Genera 2-3 fundamentos legales ESPECÍFICOS aplicables a esta falta (referencias a artículos de la Ley 66-97, Ley 136-03, Manual de Convivencia o Reglamento del Centro).
- "analisis_falta": Analiza la gravedad de la falta y su impacto en la comunidad educativa, considerando el perfil del actor involucrado ({perfil_faltas}).
- "compromisos": Genera acuerdos claros y accionables divididos por roles (estudiante, familia, docente, centro).
- "consecuencias": Establece consecuencias claras, progresivas y de carácter reglamentario.
- "plan_seguimiento": Describe cómo se dará seguimiento al acuerdo (frecuencia, responsables, indicadores).
- "disposiciones_finales": Cláusulas finales sobre validez, confidencialidad y protección del menor (Ley 136-03).

CODIFICACIÓN OBLIGATORIA (evita errores de formato JSON):
Si cualquier valor de texto necesitara un salto de línea, tabulación o comilla doble,
usa en su lugar estos marcadores: salto de línea → {ia.MARKER_NL} · comilla doble → {ia.MARKER_DQ} ·
tabulación → {ia.MARKER_TAB}. Nunca uses saltos de línea literales dentro de un valor JSON.

FORMATO DE SALIDA ESTRICTO (JSON NATIVO OBLIGATORIO):
{{
 "titulo_acta": "ACTA DE COMPROMISO Y CONVIVENCIA ESCOLAR",
 "antecedentes": "En las instalaciones del centro...",
 "fundamentos_legales": ["Fundamento legal específico 1", "Fundamento legal específico 2"],
 "analisis_falta": "Análisis de la gravedad de la falta...",
 "compromisos": ["Por parte del estudiante: ...", "Por parte de la familia: ...", "Por parte del centro: ..."],
 "consecuencias": "En caso de reincidencia o incumplimiento...",
 "plan_seguimiento": "El seguimiento se realizará...",
 "disposiciones_finales": "La presente acta tiene validez...",
 "firmantes": {json.dumps(firmantes_lista)}
}}
"""
                    texto_crudo, flags = ia.solicitar_ia(
                        prompt_maestro, modo="json", max_tokens=max_tokens,
                        temperature=temperature, modulo="acuerdos",
                    )
                    datos_json = ia.parsear_json_robusto(texto_crudo)
                    datos_json = ia.decodificar_marcadores(datos_json)

                    fecha_actual = datetime.date.today().strftime('%Y-%m-%d')
                    id_acuerdo = insertar_acuerdo(
                        fecha=fecha_actual,
                        actores=actor_seleccionado,
                        tipo_acta=tipo_acuerdo,
                        estudiante=nombre_estudiante if nombre_estudiante else "N/A",
                        docente=nombre_docente if nombre_docente else "N/A",
                        reincidente="Sí" if reincidente else "No",
                        falta=falta_manual,
                        contexto=contexto.strip(),
                        nivel_falta=info_falta["nivel"],
                        seguimiento=fecha_seg_str,
                    )

                    anio = datetime.date.today().year
                    expediente = f"ACT-{anio}-{str(id_acuerdo).zfill(4)}"

                    meta_datos = {
                        "centro": centro,
                        "coordinador": coordinador,
                        "actores": actor_seleccionado,
                        "perfil_faltas": perfil_faltas,
                        "tipo_acta": tipo_acuerdo,
                        "estudiante": nombre_estudiante if nombre_estudiante else "N/A",
                        "docente": nombre_docente if nombre_docente else "N/A",
                        "reincidencia": estado_reincidencia,
                        "falta": falta_manual,
                        "fecha": datetime.date.today().strftime('%d/%m/%Y'),
                        "seguimiento": fecha_seg_str,
                    }

                    buffer = build_acta_ejecutiva(datos_json, meta_datos, expediente)

                    st.toast("✅ Acta ejecutiva generada y guardada en el archivo institucional.", icon="⚖️")
                    st.success(f"📁 **Expediente asignado:** {expediente}")

                    archivo_nombre = f"Acta_{expediente}"
                    if nombre_estudiante:
                        archivo_nombre += f"_{nombre_estudiante.split()[0]}"
                    elif nombre_docente:
                        archivo_nombre += f"_{nombre_docente.split()[0]}"
                    else:
                        archivo_nombre += "_Convivencia"

                    st.download_button(
                        label="📥 Descargar Acta Ejecutiva (.docx)",
                        data=buffer,
                        file_name=f"{archivo_nombre}.docx",
                        mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                        type="primary",
                        width="stretch",
                    )
                    if flags.get("reintento"):
                        st.info("ℹ️ La primera llamada se cortó por tokens; el reintento automático tuvo éxito.")
                except ValueError as ve:
                    ia.render_error_ia(ve, texto_crudo)
                except Exception as e:
                    ia.render_error_ia(e, texto_crudo)

# ═══════════════════════════════════════════════════════════════════════════
# TAB 2: HISTORIAL (con filtros, búsqueda y gráficos)
# ═══════════════════════════════════════════════════════════════════════════
with tab2:
    st.markdown('<div class="acuerdos-section-title">🗄️ Archivo Institucional de Acuerdos</div>', unsafe_allow_html=True)

    db_acuerdos_tab = get_todos_acuerdos()
    if not db_acuerdos_tab:
        st.info("Aún no hay acuerdos registrados en la base de datos.")
    else:
        df_acuerdos = pd.DataFrame(db_acuerdos_tab)

        col_b1, col_b2, col_b3 = st.columns([2, 1, 1])
        with col_b1:
            busqueda = st.text_input("🔍 Buscar", placeholder="Estudiante, docente, falta o contexto...")
        with col_b2:
            niveles_disp = sorted(df_acuerdos["Nivel Falta"].dropna().unique().tolist())
            filtro_nivel = st.selectbox("Filtrar por nivel", ["Todos"] + niveles_disp)
        with col_b3:
            filtro_reinc = st.selectbox("Reincidencia", ["Todos", "Sí", "No"])

        df_filtrado = df_acuerdos.copy()
        if busqueda.strip():
            texto = busqueda.strip().lower()
            df_filtrado = df_filtrado[
                df_filtrado.apply(
                    lambda row: texto in " ".join(str(v).lower() for v in row if pd.notna(v)),
                    axis=1,
                )
            ]
        if filtro_nivel != "Todos":
            df_filtrado = df_filtrado[df_filtrado["Nivel Falta"] == filtro_nivel]
        if filtro_reinc != "Todos":
            df_filtrado = df_filtrado[df_filtrado["Reincidente"] == filtro_reinc]

        st.markdown(f"**{len(df_filtrado)}** acta(s) encontrada(s)")

        col_g1, col_g2 = st.columns(2)
        with col_g1:
            st.markdown("#### 📊 Distribución por Nivel de Falta")
            df_nivel = df_filtrado["Nivel Falta"].value_counts()
            if not df_nivel.empty:
                st.bar_chart(df_nivel)
            else:
                st.info("Sin datos para graficar.")
        with col_g2:
            st.markdown("#### 📈 Actas por Mes")
            df_filtrado_graf = df_filtrado.copy()
            df_filtrado_graf["Mes"] = pd.to_datetime(df_filtrado_graf["Fecha"], errors="coerce").dt.strftime("%Y-%m")
            df_mes = df_filtrado_graf["Mes"].value_counts().sort_index()
            if not df_mes.empty:
                st.bar_chart(df_mes)
            else:
                st.info("Sin datos para graficar.")

        st.markdown("---")

        st.dataframe(
            df_filtrado.drop(columns=["id", "Contexto"]),
            use_container_width=True,
            hide_index=True,
        )

        st.markdown("---")

        with st.expander("🗑️ Gestionar Historial (Eliminar Registro)"):
            st.warning("⚠️ Precaución: Eliminar un registro borrará la constancia del acuerdo en el sistema.")
            opciones_del = {
                f"ID {row['id']}: {row['Fecha']} - {row['Estudiante']} ({str(row['Tipo de Acta'])[:30]}...)": row['id']
                for row in db_acuerdos_tab
            }
            seleccion_del = st.selectbox("Selecciona el acuerdo a eliminar", list(opciones_del.keys()),
                                         index=None, placeholder="Elige...")
            if seleccion_del is not None and st.button("Eliminar permanentemente", type="primary"):
                eliminar_acuerdo(opciones_del[seleccion_del])
                st.session_state.flash = "Acuerdo eliminado de la base de datos institucional."
                st.rerun()