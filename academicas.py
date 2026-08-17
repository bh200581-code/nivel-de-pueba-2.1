"""
academicas.py — Planificación de Unidad Académica (Paso 25 · Nivel Dios)
Planificación de Unidad por Situación de Aprendizaje y/o Proyecto,
alineada al formato oficial MINERD con competencias fundamentales,
contenidos tridimensionales y evaluación integral.
• IA vía core/ia (solicitar_json, marcadores, reintento, auditoría).
• Modo dual: Situación de Aprendizaje / Proyecto.
• Super interfaz: hero, stepper, cards de modo, métricas, vista previa.
• Word profesional con formato MINERD.
"""
import re
from datetime import datetime, date
from io import BytesIO
from typing import Any, Dict, List, Optional

import pandas as pd
import streamlit as st
from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement, parse_xml
from docx.oxml.ns import qn, nsdecls
from docx.shared import Inches, Pt, RGBColor

from core import ia

# ═══════════════════════════════════════════════════════════════════════════
# CONSTANTES MINERD
# ═══════════════════════════════════════════════════════════════════════════
COMPETENCIAS_FUNDAMENTALES = [
    "Relación con el Entorno",
    "Pensamiento Lógico, Crítico y Creativo",
    "Resolución de Problemas",
    "Comunicación Lingüística",
    "Desarrollo Personal y Espiritual",
    "Científica y Tecnológica",
    "Ambiental y de la Salud",
    "Ética y Ciudadana",
]

ESTRATEGIAS_EA = [
    "Estudio de casos",
    "Aprendizaje basado en problemas",
    "Aprendizaje colaborativo",
    "Investigación dirigida",
    "Juego de roles",
    "Debate",
    "Taller práctico",
    "Aprendizaje por descubrimiento",
    "Método de proyectos",
    "Exposición dialogada",
]

MODOS_PLANIFICACION = {
    "situacion_aprendizaje": {
        "nombre": "Situación de Aprendizaje",
        "icono": "🎯",
        "descripcion": "Contextualiza el aprendizaje en una situación real o simulada que los estudiantes deben resolver, generando un producto final.",
        "color": "#F59E0B",
    },
    "proyecto": {
        "nombre": "Proyecto",
        "icono": "🚀",
        "descripcion": "Organiza la unidad en torno a un proyecto con fases de Planificación, Ejecución y Comunicación, culminando en un producto tangible.",
        "color": "#8B5CF6",
    },
}

# ═══════════════════════════════════════════════════════════════════════════
# SUPER INTERFAZ — ESTILOS
# ═══════════════════════════════════════════════════════════════════════════
st.markdown("""
<style>
@import url('https://fonts.googleapis.com/css2?family=Inter:wght@300;400;500;600;700;800;900&display=swap');

html, body, [class*="css"] {
    font-family: 'Inter', sans-serif;
    background-color: #F0F4F8;
    color: #1E293B;
}

/* Hero premium */
.acad-hero {
    background: linear-gradient(135deg, #451A03 0%, #92400E 40%, #F59E0B 70%, #FCD34D 100%);
    color: #fff;
    padding: 2.2rem;
    border-radius: 20px;
    margin-bottom: 1.5rem;
    box-shadow: 0 25px 50px rgba(146, 64, 14, 0.3);
    position: relative;
    overflow: hidden;
}

.acad-hero::before {
    content: '📖';
    position: absolute;
    right: 2rem;
    top: 50%;
    transform: translateY(-50%);
    font-size: 6rem;
    opacity: 0.15;
}

.acad-hero-title {
    font-size: 2.4rem;
    font-weight: 900;
    letter-spacing: -0.03em;
    margin-bottom: 0.4rem;
}

.acad-hero-sub {
    font-size: 1.05rem;
    opacity: 0.9;
    line-height: 1.5;
}

.acad-hero-badge {
    display: inline-block;
    background: rgba(255,255,255,0.15);
    border: 1px solid rgba(255,255,255,0.25);
    border-radius: 8px;
    padding: 4px 12px;
    font-size: 0.8rem;
    font-weight: 600;
    margin-top: 0.8rem;
    margin-right: 8px;
}

/* Stepper */
.acad-stepper {
    display: flex;
    align-items: center;
    justify-content: center;
    gap: 0;
    margin: 1.5rem 0;
}

.acad-step {
    display: flex;
    flex-direction: column;
    align-items: center;
    gap: 6px;
}

.acad-step-circle {
    width: 48px;
    height: 48px;
    border-radius: 50%;
    display: flex;
    align-items: center;
    justify-content: center;
    font-weight: 800;
    font-size: 1.2rem;
    transition: all 0.3s ease;
}

.acad-step-circle.inactive {
    background: #E2E8F0;
    color: #94A3B8;
    border: 3px solid #CBD5E1;
}

.acad-step-circle.active {
    background: linear-gradient(135deg, #92400E, #F59E0B);
    color: #fff;
    border: 3px solid #92400E;
    box-shadow: 0 4px 15px rgba(146, 64, 14, 0.4);
}

.acad-step-circle.done {
    background: linear-gradient(135deg, #059669, #34D399);
    color: #fff;
    border: 3px solid #059669;
    box-shadow: 0 4px 15px rgba(5, 150, 105, 0.3);
}

.acad-step-label {
    font-size: 0.72rem;
    font-weight: 700;
    color: #64748B;
    text-transform: uppercase;
    letter-spacing: 0.03em;
}

.acad-step-line {
    width: 70px;
    height: 3px;
    background: #CBD5E1;
    margin: 0 4px;
    margin-bottom: 22px;
}

.acad-step-line.done {
    background: linear-gradient(90deg, #059669, #34D399);
}

/* Cards de modo */
.acad-mode-card {
    background: #fff;
    border: 3px solid #E2E8F0;
    border-radius: 16px;
    padding: 1.5rem;
    transition: all 0.25s ease;
    cursor: pointer;
    text-align: center;
    box-shadow: 0 2px 8px rgba(0,0,0,0.04);
}

.acad-mode-card:hover {
    transform: translateY(-4px);
    box-shadow: 0 10px 25px rgba(0,0,0,0.1);
}

.acad-mode-card.selected {
    border-color: #F59E0B;
    background: #FFFBEB;
    box-shadow: 0 6px 20px rgba(245, 158, 11, 0.2);
}

.acad-mode-icono { font-size: 3rem; margin-bottom: 0.8rem; }
.acad-mode-nombre { font-weight: 800; font-size: 1.1rem; color: #0F172A; margin-bottom: 0.5rem; }
.acad-mode-desc { font-size: 0.82rem; color: #64748B; line-height: 1.4; }

/* Secciones */
.acad-section-title {
    color: #92400E;
    font-weight: 700;
    font-size: 1.12rem;
    border-bottom: 2px solid #FEF3C7;
    padding-bottom: 8px;
    margin: 1.2rem 0 0.9rem 0;
}

/* Métricas */
.acad-stat {
    background: #fff;
    border-radius: 12px;
    padding: 1rem;
    text-align: center;
    border: 1px solid #E2E8F0;
    box-shadow: 0 2px 6px rgba(0,0,0,0.04);
}

.acad-stat-value {
    font-size: 1.8rem;
    font-weight: 800;
    color: #92400E;
}

.acad-stat-label {
    font-size: 0.72rem;
    font-weight: 700;
    color: #94A3B8;
    text-transform: uppercase;
    letter-spacing: 0.05em;
}
</style>
""", unsafe_allow_html=True)

# ═══════════════════════════════════════════════════════════════════════════
# ESTADO DE SESIÓN
# ═══════════════════════════════════════════════════════════════════════════
def init_estado():
    if "acad_resultado" not in st.session_state:
        st.session_state.acad_resultado = None
    if "acad_fase" not in st.session_state:
        st.session_state.acad_fase = 0  # 0=config, 1=generando, 2=resultados
    if "acad_modo" not in st.session_state:
        st.session_state.acad_modo = None

init_estado()

# ═══════════════════════════════════════════════════════════════════════════
# UTILIDADES WORD
# ═══════════════════════════════════════════════════════════════════════════
def shade_cell(cell, color):
    shd = parse_xml(r'<w:shd {} w:fill="{}"/>'.format(nsdecls("w"), color))
    cell._tc.get_or_add_tcPr().append(shd)

def set_cell_text(cell, text, bold=False, center=False, color=None):
    cell.text = ""
    p = cell.paragraphs[0]
    run = p.add_run(str(text))
    run.bold = bold
    if center:
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    if color:
        shade_cell(cell, color)

def add_page_number(doc):
    for section in doc.sections:
        footer = section.footer
        footer.is_linked_to_previous = False
        p = footer.paragraphs[0] if footer.paragraphs else footer.add_paragraph()
        p.text = ""
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run()
        fld1 = OxmlElement("w:fldChar"); fld1.set(qn("w:fldCharType"), "begin")
        instr = OxmlElement("w:instrText"); instr.set(qn("xml:space"), "preserve"); instr.text = "PAGE"
        fld2 = OxmlElement("w:fldChar"); fld2.set(qn("w:fldCharType"), "end")
        run._r.append(fld1); run._r.append(instr); run._r.append(fld2)
        run.add_run(" / ")
        run2 = p.add_run()
        fld3 = OxmlElement("w:fldChar"); fld3.set(qn("w:fldCharType"), "begin")
        instr2 = OxmlElement("w:instrText"); instr2.set(qn("xml:space"), "preserve"); instr2.text = "NUMPAGES"
        fld4 = OxmlElement("w:fldChar"); fld4.set(qn("w:fldCharType"), "end")
        run2._r.append(fld3); run2._r.append(instr2); run2._r.append(fld4)
        run.font.size = Pt(8)
        run.font.color.rgb = RGBColor(0x64, 0x74, 0x8B)
        run2.font.size = Pt(8)
        run2.font.color.rgb = RGBColor(0x64, 0x74, 0x8B)

# ═══════════════════════════════════════════════════════════════════════════
# PROMPTS
# ═══════════════════════════════════════════════════════════════════════════
def prompt_situacion_aprendizaje(centro, docente, asignatura, grado, periodo, num_semanas, competencias_sel, estrategias_sel, tema_unidad):
    competencias_txt = "\n".join([f"- {c}" for c in competencias_sel])
    estrategias_txt = "\n".join([f"- {e}" for e in estrategias_sel])
    
    return f"""Actúa como un Especialista Curricular del MINERD, Experto en Planificación de Unidad por Situación de Aprendizaje para el Nivel Secundario.

Tu tarea es diseñar una Planificación de Unidad completa y profesional basada en una SITUACIÓN DE APRENDIZAJE.

DATOS GENERALES:
Centro Educativo: {centro}
Docente: {docente}
Asignatura/Área: {asignatura}
Grado y Sección: {grado}
Período: {periodo}
Duración de la unidad: {num_semanas} semanas
Tema central de la unidad: {tema_unidad}

COMPETENCIAS FUNDAMENTALES A DESARROLLAR:
{competencias_txt}

ESTRATEGIAS DE ENSEÑANZA-APRENDIZAJE:
{estrategias_txt}

REGLAS DE DISEÑO:
- La SITUACIÓN DE APRENDIZAJE debe ser contextualizada, realista y relevante para los estudiantes.
- Debe incluir un PROBLEMA o NECESIDAD que los estudiantes deban resolver.
- El PRODUCTO FINAL debe ser tangible y demostrable.
- Las PREGUNTAS GUÍA deben orientar el proceso de aprendizaje.
- La SECUENCIA DIDÁCTICA debe distribuirse en {num_semanas} semanas con actividades de Inicio, Desarrollo y Cierre.
- Los CONTENIDOS deben ser tridimensionales: Conceptuales, Procedimentales y Actitudinales.
- La EVALUACIÓN debe incluir momentos Diagnóstica, Formativa y Sumativa.
- Incluye ADECUACIONES CURRICULARES para estudiantes con NEAE.

Devuelve ÚNICAMENTE JSON válido con este formato exacto:
{{
  "TITULO_UNIDAD": "Título creativo de la unidad",
  "SITUACION_APRENDIZAJE": {{
    "CONTEXTO": "Descripción del contexto donde ocurre la situación...",
    "PROBLEMA": "Descripción del problema o necesidad a resolver...",
    "PRODUCTO_FINAL": "Descripción del producto final que generarán los estudiantes...",
    "PREGUNTAS_GUIA": ["Pregunta guía 1", "Pregunta guía 2", "Pregunta guía 3"]
  }},
  "COMPETENCIAS_FUNDAMENTALES": [
    {{"COMPETENCIA": "Nombre de la competencia", "INDICADORES": "Indicadores de logro asociados"}}
  ],
  "COMPETENCIAS_ESPECIFICAS": ["Competencia específica 1", "Competencia específica 2"],
  "CONTENIDOS": {{
    "CONCEPTUALES": ["Contenido conceptual 1", "Contenido conceptual 2"],
    "PROCEDIMENTALES": ["Contenido procedimental 1", "Contenido procedimental 2"],
    "ACTITUDINALES": ["Contenido actitudinal 1", "Contenido actitudinal 2"]
  }},
  "ESTRATEGIAS": ["Estrategia 1", "Estrategia 2"],
  "SECUENCIA_DIDACTICA": [
    {{
      "SEMANA": 1,
      "TEMATICA": "Tema de la semana",
      "ACTIVIDADES_INICIO": "Actividades de inicio...",
      "ACTIVIDADES_DESARROLLO": "Actividades de desarrollo...",
      "ACTIVIDADES_CIERRE": "Actividades de cierre...",
      "RECURSOS": "Recursos necesarios...",
      "EVALUACION": "Evaluación de la semana..."
    }}
  ],
  "RECURSOS": ["Recurso 1", "Recurso 2"],
  "EVALUACION": {{
    "DIAGNOSTICA": "Descripción de la evaluación diagnóstica...",
    "FORMATIVA": "Descripción de la evaluación formativa...",
    "SUMATIVA": "Descripción de la evaluación sumativa..."
  }},
  "ADECUACIONES_NEAE": "Descripción de las adecuaciones curriculares para estudiantes con NEAE...",
  "BIBLIOGRAFIA": ["Referencia 1", "Referencia 2"]
}}
"""

def prompt_proyecto(centro, docente, asignatura, grado, periodo, num_semanas, competencias_sel, estrategias_sel, tema_unidad):
    competencias_txt = "\n".join([f"- {c}" for c in competencias_sel])
    estrategias_txt = "\n".join([f"- {e}" for e in estrategias_sel])
    
    return f"""Actúa como un Especialista Curricular del MINERD, Experto en Planificación de Unidad por Proyecto para el Nivel Secundario.

Tu tarea es diseñar una Planificación de Unidad completa y profesional basada en un PROYECTO.

DATOS GENERALES:
Centro Educativo: {centro}
Docente: {docente}
Asignatura/Área: {asignatura}
Grado y Sección: {grado}
Período: {periodo}
Duración de la unidad: {num_semanas} semanas
Tema central de la unidad: {tema_unidad}

COMPETENCIAS FUNDAMENTALES A DESARROLLAR:
{competencias_txt}

ESTRATEGIAS DE ENSEÑANZA-APRENDIZAJE:
{estrategias_txt}

REGLAS DE DISEÑO:
- El PROYECTO debe tener un título creativo, descripción clara y un producto final tangible.
- Debe incluir DESTINATARIOS reales (comunidad escolar, barrio, etc.).
- Las FASES DEL PROYECTO deben ser: Planificación, Ejecución y Comunicación/Evaluación.
- Cada fase debe tener actividades concretas, productos parciales, recursos y evaluación.
- Los CONTENIDOS deben ser tridimensionales: Conceptuales, Procedimentales y Actitudinales.
- La EVALUACIÓN debe incluir momentos Diagnóstica, Formativa y Sumativa.
- Incluye ADECUACIONES CURRICULARES para estudiantes con NEAE.
- Distribuye las {num_semanas} semanas entre las fases del proyecto.

Devuelve ÚNICAMENTE JSON válido con este formato exacto:
{{
  "TITULO_UNIDAD": "Título creativo de la unidad",
  "PROYECTO": {{
    "TITULO_PROYECTO": "Título del proyecto...",
    "DESCRIPCION": "Descripción del proyecto...",
    "PRODUCTO_FINAL": "Descripción del producto final...",
    "DESTINATARIOS": "Descripción de los destinatarios..."
  }},
  "COMPETENCIAS_FUNDAMENTALES": [
    {{"COMPETENCIA": "Nombre de la competencia", "INDICADORES": "Indicadores de logro asociados"}}
  ],
  "COMPETENCIAS_ESPECIFICAS": ["Competencia específica 1", "Competencia específica 2"],
  "CONTENIDOS": {{
    "CONCEPTUALES": ["Contenido conceptual 1", "Contenido conceptual 2"],
    "PROCEDIMENTALES": ["Contenido procedimental 1", "Contenido procedimental 2"],
    "ACTITUDINALES": ["Contenido actitudinal 1", "Contenido actitudinal 2"]
  }},
  "ESTRATEGIAS": ["Estrategia 1", "Estrategia 2"],
  "FASES_PROYECTO": [
    {{
      "FASE": "Planificación",
      "SEMANAS": "1-2",
      "ACTIVIDADES": ["Actividad 1", "Actividad 2"],
      "PRODUCTOS_PARCIALES": ["Producto parcial 1"],
      "RECURSOS": "Recursos necesarios...",
      "EVALUACION": "Evaluación de la fase..."
    }},
    {{
      "FASE": "Ejecución",
      "SEMANAS": "3-4",
      "ACTIVIDADES": ["Actividad 1", "Actividad 2"],
      "PRODUCTOS_PARCIALES": ["Producto parcial 1"],
      "RECURSOS": "Recursos necesarios...",
      "EVALUACION": "Evaluación de la fase..."
    }},
    {{
      "FASE": "Comunicación y Evaluación",
      "SEMANAS": "5",
      "ACTIVIDADES": ["Actividad 1", "Actividad 2"],
      "PRODUCTOS_PARCIALES": ["Producto final"],
      "RECURSOS": "Recursos necesarios...",
      "EVALUACION": "Evaluación de la fase..."
    }}
  ],
  "RECURSOS": ["Recurso 1", "Recurso 2"],
  "EVALUACION": {{
    "DIAGNOSTICA": "Descripción de la evaluación diagnóstica...",
    "FORMATIVA": "Descripción de la evaluación formativa...",
    "SUMATIVA": "Descripción de la evaluación sumativa..."
  }},
  "ADECUACIONES_NEAE": "Descripción de las adecuaciones curriculares para estudiantes con NEAE...",
  "BIBLIOGRAFIA": ["Referencia 1", "Referencia 2"]
}}
"""

# ═══════════════════════════════════════════════════════════════════════════
# GENERACIÓN WORD PROFESIONAL
# ═══════════════════════════════════════════════════════════════════════════
def build_acad_docx(datos, meta, modo):
    doc = Document()
    doc.styles["Normal"].font.name = "Calibri"
    doc.styles["Normal"].font.size = Pt(11)
    
    for section in doc.sections:
        section.left_margin = Inches(0.8)
        section.right_margin = Inches(0.8)
    
    # ─── PORTADA ───
    for _ in range(3):
        doc.add_paragraph()
    
    p_minerd = doc.add_paragraph()
    p_minerd.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run_m = p_minerd.add_run("MINISTERIO DE EDUCACIÓN DE LA REPÚBLICA DOMINICANA")
    run_m.bold = True
    run_m.font.size = Pt(12)
    
    p_centro = doc.add_paragraph()
    p_centro.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run_c = p_centro.add_run(meta.get("centro", ""))
    run_c.bold = True
    run_c.font.size = Pt(14)
    run_c.font.color.rgb = RGBColor(0x92, 0x40, 0x0E)
    
    doc.add_paragraph()
    
    p_titulo = doc.add_paragraph()
    p_titulo.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run_t = p_titulo.add_run("PLANIFICACIÓN DE UNIDAD ACADÉMICA")
    run_t.bold = True
    run_t.font.size = Pt(22)
    run_t.font.color.rgb = RGBColor(0x92, 0x40, 0x0E)
    
    modo_info = MODOS_PLANIFICACION.get(modo, {"nombre": "Académico"})
    p_modo = doc.add_paragraph()
    p_modo.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run_modo = p_modo.add_run(f"Modalidad: {modo_info.get('nombre', '')}")
    run_modo.font.size = Pt(14)
    run_modo.font.color.rgb = RGBColor(0x64, 0x74, 0x8B)
    
    doc.add_paragraph()
    
    p_unidad = doc.add_paragraph()
    p_unidad.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run_u = p_unidad.add_run(datos.get("TITULO_UNIDAD", ""))
    run_u.bold = True
    run_u.font.size = Pt(16)
    
    doc.add_paragraph()
    doc.add_paragraph()
    
    p_meta = doc.add_paragraph()
    p_meta.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p_meta.add_run(f"Docente: {meta.get('docente', '')}").font.size = Pt(12)
    
    p_meta2 = doc.add_paragraph()
    p_meta2.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p_meta2.add_run(f"Asignatura: {meta.get('asignatura', '')} | Grado: {meta.get('grado', '')}").font.size = Pt(11)
    
    p_meta3 = doc.add_paragraph()
    p_meta3.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p_meta3.add_run(f"Período: {meta.get('periodo', '')} | Duración: {meta.get('num_semanas', '')} semanas").font.size = Pt(11)
    
    doc.add_page_break()
    
    # ─── DATOS GENERALES ───
    doc.add_heading("1. Datos Generales", level=1)
    
    tabla_datos = doc.add_table(rows=4, cols=4)
    tabla_datos.style = "Table Grid"
    
    datos_gen = [
        ("Centro Educativo", meta.get("centro", ""), "Docente", meta.get("docente", "")),
        ("Asignatura/Área", meta.get("asignatura", ""), "Grado y Sección", meta.get("grado", "")),
        ("Período", meta.get("periodo", ""), "Duración", f"{meta.get('num_semanas', '')} semanas"),
        ("Modalidad", modo_info.get("nombre", ""), "Fecha", datetime.now().strftime("%d/%m/%Y")),
    ]
    
    for i, (l1, v1, l2, v2) in enumerate(datos_gen):
        row = tabla_datos.rows[i].cells
        set_cell_text(row[0], l1, bold=True, color="FEF3C7")
        set_cell_text(row[1], v1)
        set_cell_text(row[2], l2, bold=True, color="FEF3C7")
        set_cell_text(row[3], v2)
    
    doc.add_paragraph()
    
    # ─── SITUACIÓN DE APRENDIZAJE / PROYECTO ───
    if modo == "situacion_aprendizaje":
        doc.add_heading("2. Situación de Aprendizaje", level=1)
        
        sa = datos.get("SITUACION_APRENDIZAJE", {})
        
        doc.add_heading("Contexto", level=2)
        doc.add_paragraph(sa.get("CONTEXTO", ""))
        
        doc.add_heading("Problema / Necesidad", level=2)
        doc.add_paragraph(sa.get("PROBLEMA", ""))
        
        doc.add_heading("Producto Final", level=2)
        doc.add_paragraph(sa.get("PRODUCTO_FINAL", ""))
        
        doc.add_heading("Preguntas Guía", level=2)
        for i, preg in enumerate(sa.get("PREGUNTAS_GUIA", []), 1):
            doc.add_paragraph(f"{i}. {preg}", style="List Number")
    
    elif modo == "proyecto":
        doc.add_heading("2. Proyecto", level=1)
        
        proy = datos.get("PROYECTO", {})
        
        doc.add_heading("Título del Proyecto", level=2)
        doc.add_paragraph(proy.get("TITULO_PROYECTO", ""))
        
        doc.add_heading("Descripción", level=2)
        doc.add_paragraph(proy.get("DESCRIPCION", ""))
        
        doc.add_heading("Producto Final", level=2)
        doc.add_paragraph(proy.get("PRODUCTO_FINAL", ""))
        
        doc.add_heading("Destinatarios", level=2)
        doc.add_paragraph(proy.get("DESTINATARIOS", ""))
    
    doc.add_page_break()
    
    # ─── COMPETENCIAS FUNDAMENTALES ───
    doc.add_heading("3. Competencias Fundamentales", level=1)
    
    tabla_comp = doc.add_table(rows=1, cols=2)
    tabla_comp.style = "Table Grid"
    hdr = tabla_comp.rows[0].cells
    set_cell_text(hdr[0], "Competencia Fundamental", bold=True, color="FEF3C7")
    set_cell_text(hdr[1], "Indicadores de Logro", bold=True, color="FEF3C7")
    
    for comp in datos.get("COMPETENCIAS_FUNDAMENTALES", []):
        if isinstance(comp, dict):
            row = tabla_comp.add_row().cells
            set_cell_text(row[0], comp.get("COMPETENCIA", ""), bold=True)
            set_cell_text(row[1], comp.get("INDICADORES", ""))
    
    doc.add_paragraph()
    
    # ─── COMPETENCIAS ESPECÍFICAS ───
    doc.add_heading("4. Competencias Específicas", level=1)
    for comp in datos.get("COMPETENCIAS_ESPECIFICAS", []):
        doc.add_paragraph(comp, style="List Bullet")
    
    doc.add_paragraph()
    
    # ─── CONTENIDOS ───
    doc.add_heading("5. Contenidos", level=1)
    
    contenidos = datos.get("CONTENIDOS", {})
    
    doc.add_heading("5.1 Conceptuales", level=2)
    for c in contenidos.get("CONCEPTUALES", []):
        doc.add_paragraph(c, style="List Bullet")
    
    doc.add_heading("5.2 Procedimentales", level=2)
    for c in contenidos.get("PROCEDIMENTALES", []):
        doc.add_paragraph(c, style="List Bullet")
    
    doc.add_heading("5.3 Actitudinales", level=2)
    for c in contenidos.get("ACTITUDINALES", []):
        doc.add_paragraph(c, style="List Bullet")
    
    doc.add_page_break()
    
    # ─── ESTRATEGIAS ───
    doc.add_heading("6. Estrategias de Enseñanza-Aprendizaje", level=1)
    for e in datos.get("ESTRATEGIAS", []):
        doc.add_paragraph(e, style="List Bullet")
    
    doc.add_paragraph()
    
    # ─── SECUENCIA DIDÁCTICA / FASES DEL PROYECTO ───
    if modo == "situacion_aprendizaje":
        doc.add_heading("7. Secuencia Didáctica", level=1)
        
        for semana in datos.get("SECUENCIA_DIDACTICA", []):
            doc.add_heading(f"Semana {semana.get('SEMANA', '')}: {semana.get('TEMATICA', '')}", level=2)
            
            p_inicio = doc.add_paragraph()
            p_inicio.add_run("Inicio: ").bold = True
            p_inicio.add_run(semana.get("ACTIVIDADES_INICIO", ""))
            
            p_des = doc.add_paragraph()
            p_des.add_run("Desarrollo: ").bold = True
            p_des.add_run(semana.get("ACTIVIDADES_DESARROLLO", ""))
            
            p_cierre = doc.add_paragraph()
            p_cierre.add_run("Cierre: ").bold = True
            p_cierre.add_run(semana.get("ACTIVIDADES_CIERRE", ""))
            
            p_rec = doc.add_paragraph()
            p_rec.add_run("Recursos: ").bold = True
            p_rec.add_run(semana.get("RECURSOS", ""))
            
            p_eval = doc.add_paragraph()
            p_eval.add_run("Evaluación: ").bold = True
            p_eval.add_run(semana.get("EVALUACION", ""))
            
            doc.add_paragraph()
    
    elif modo == "proyecto":
        doc.add_heading("7. Fases del Proyecto", level=1)
        
        for fase in datos.get("FASES_PROYECTO", []):
            doc.add_heading(f"Fase: {fase.get('FASE', '')} (Semanas {fase.get('SEMANAS', '')})", level=2)
            
            doc.add_heading("Actividades", level=3)
            for act in fase.get("ACTIVIDADES", []):
                doc.add_paragraph(act, style="List Bullet")
            
            doc.add_heading("Productos Parciales", level=3)
            for prod in fase.get("PRODUCTOS_PARCIALES", []):
                doc.add_paragraph(prod, style="List Bullet")
            
            p_rec = doc.add_paragraph()
            p_rec.add_run("Recursos: ").bold = True
            p_rec.add_run(fase.get("RECURSOS", ""))
            
            p_eval = doc.add_paragraph()
            p_eval.add_run("Evaluación: ").bold = True
            p_eval.add_run(fase.get("EVALUACION", ""))
            
            doc.add_paragraph()
    
    doc.add_page_break()
    
    # ─── RECURSOS ───
    doc.add_heading("8. Recursos", level=1)
    for r in datos.get("RECURSOS", []):
        doc.add_paragraph(r, style="List Bullet")
    
    doc.add_paragraph()
    
    # ─── EVALUACIÓN ───
    doc.add_heading("9. Evaluación", level=1)
    
    evaluacion = datos.get("EVALUACION", {})
    
    doc.add_heading("9.1 Evaluación Diagnóstica", level=2)
    doc.add_paragraph(evaluacion.get("DIAGNOSTICA", ""))
    
    doc.add_heading("9.2 Evaluación Formativa", level=2)
    doc.add_paragraph(evaluacion.get("FORMATIVA", ""))
    
    doc.add_heading("9.3 Evaluación Sumativa", level=2)
    doc.add_paragraph(evaluacion.get("SUMATIVA", ""))
    
    doc.add_paragraph()
    
    # ─── ADECUACIONES NEAE ───
    doc.add_heading("10. Adecuaciones Curriculares (NEAE)", level=1)
    doc.add_paragraph(datos.get("ADECUACIONES_NEAE", "Sin adecuaciones adicionales requeridas."))
    
    doc.add_paragraph()
    
    # ─── BIBLIOGRAFÍA ───
    doc.add_heading("11. Bibliografía", level=1)
    for ref in datos.get("BIBLIOGRAFIA", []):
        doc.add_paragraph(ref, style="List Bullet")
    
    doc.add_paragraph()
    
    # ─── FIRMAS ───
    doc.add_paragraph()
    doc.add_paragraph()
    
    tabla_firmas = doc.add_table(rows=2, cols=3)
    tabla_firmas.cell(0, 0).text = "_________________________"
    tabla_firmas.cell(0, 1).text = "_________________________"
    tabla_firmas.cell(0, 2).text = "_________________________"
    tabla_firmas.cell(1, 0).text = "Docente"
    tabla_firmas.cell(1, 1).text = "Coordinador/a Académico"
    tabla_firmas.cell(1, 2).text = "Director/a Académico"
    
    for row in tabla_firmas.rows:
        for cell in row.cells:
            cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    add_page_number(doc)
    
    buffer = BytesIO()
    doc.save(buffer)
    buffer.seek(0)
    return buffer

# ═══════════════════════════════════════════════════════════════════════════
# RENDERIZADO DE STEPPER
# ═══════════════════════════════════════════════════════════════════════════
def render_stepper(fase_actual):
    pasos = [
        ("1", "Configurar"),
        ("2", "Generar"),
        ("3", "Descargar"),
    ]
    
    html = '<div class="acad-stepper">'
    for i, (num, label) in enumerate(pasos):
        if i + 1 < fase_actual:
            estado = "done"
        elif i + 1 == fase_actual:
            estado = "active"
        else:
            estado = "inactive"
        
        html += f'''
        <div class="acad-step">
            <div class="acad-step-circle {estado}">{"✓" if estado == "done" else num}</div>
            <div class="acad-step-label">{label}</div>
        </div>
        '''
        if i < len(pasos) - 1:
            line_class = "done" if i + 1 < fase_actual else ""
            html += f'<div class="acad-step-line {line_class}"></div>'
    
    html += '</div>'
    st.markdown(html, unsafe_allow_html=True)

# ═══════════════════════════════════════════════════════════════════════════
# HERO Y STEPPER
# ═══════════════════════════════════════════════════════════════════════════
ia.panel_sidebar_ia("Planificación de Unidad Académica")

st.markdown("""
<div class="acad-hero">
    <div class="acad-hero-title">📖 Planificación de Unidad Académica</div>
    <div class="acad-hero-sub">
        Situación de Aprendizaje · Proyecto · Competencias Fundamentales MINERD · Word profesional
    </div>
    <div>
        <span class="acad-hero-badge">🎯 Situación de Aprendizaje</span>
        <span class="acad-hero-badge">🚀 Proyecto</span>
        <span class="acad-hero-badge">🤖 Asistido por IA</span>
        <span class="acad-hero-badge">📄 Word MINERD</span>
    </div>
</div>
""", unsafe_allow_html=True)

render_stepper(st.session_state.acad_fase + 1)

# ═══════════════════════════════════════════════════════════════════════════
# FASE 1: CONFIGURACIÓN
# ═══════════════════════════════════════════════════════════════════════════
if st.session_state.acad_fase == 0:
    st.markdown('<div class="acad-section-title">🎯 Selecciona el Modo de Planificación</div>', unsafe_allow_html=True)
    
    col_m1, col_m2 = st.columns(2)
    
    for clave, info in MODOS_PLANIFICACION.items():
        col = col_m1 if clave == "situacion_aprendizaje" else col_m2
        with col:
            selected = st.session_state.acad_modo == clave
            css_class = "acad-mode-card selected" if selected else "acad-mode-card"
            
            st.markdown(f"""
            <div class="{css_class}">
                <div class="acad-mode-icono">{info['icono']}</div>
                <div class="acad-mode-nombre">{info['nombre']}</div>
                <div class="acad-mode-desc">{info['descripcion']}</div>
            </div>
            """, unsafe_allow_html=True)
            
            if st.button(f"Seleccionar {info['nombre']}", key=f"modo_{clave}", use_container_width=True,
                         type="primary" if selected else "secondary"):
                st.session_state.acad_modo = clave
                st.rerun()
    
    if st.session_state.acad_modo:
        st.markdown('<hr style="border: none; height: 2px; background: linear-gradient(90deg, #F59E0B, #FCD34D); margin: 1.5rem 0;">', unsafe_allow_html=True)
        st.markdown('<div class="acad-section-title">📋 Datos de la Unidad</div>', unsafe_allow_html=True)
        
        with st.form("form_academicas"):
            col1, col2 = st.columns(2)
            with col1:
                centro = st.text_input("Centro Educativo", value="Politécnico Salesiano Arquides Calderón")
                docente = st.text_input("Docente", value="Ing. Bernardo Antonio Hernández Batista")
                asignatura = st.text_input("Asignatura / Área", placeholder="Ej: Ciencias Sociales, Lengua Española")
            with col2:
                grado = st.text_input("Grado y Sección", placeholder="Ej: 3ro A")
                periodo = st.text_input("Período", placeholder="Ej: 2026-2027, Segundo Trimestre")
                num_semanas = st.slider("Duración de la unidad (semanas)", 1, 12, 4)
            
            tema_unidad = st.text_area(
                "Tema central de la unidad",
                height=70,
                placeholder="Ej: La contaminación del agua en mi comunidad y su impacto en la salud"
            )
            
            st.markdown('<div class="acad-section-title">🎓 Competencias y Estrategias</div>', unsafe_allow_html=True)
            
            col_c1, col_c2 = st.columns(2)
            with col_c1:
                competencias_sel = st.multiselect(
                    "Competencias Fundamentales a desarrollar",
                    COMPETENCIAS_FUNDAMENTALES,
                    default=["Relación con el Entorno", "Resolución de Problemas", "Comunicación Lingüística"]
                )
            with col_c2:
                estrategias_sel = st.multiselect(
                    "Estrategias de Enseñanza-Aprendizaje",
                    ESTRATEGIAS_EA,
                    default=["Aprendizaje basado en problemas", "Aprendizaje colaborativo"]
                )
            
            max_tokens, temperature = ia.control_avanzado(default_tokens=16384, tope=32000, default_temp=0.2)
            
            st.markdown("---")
            btn_generar = st.form_submit_button(
                "⚡ Generar Planificación de Unidad con IA",
                type="primary",
                use_container_width=True
            )
        
        if btn_generar:
            if not asignatura or not tema_unidad:
                st.warning("⚠️ Completa al menos la asignatura y el tema central.")
            else:
                with st.spinner("🧠 Generando planificación de unidad con IA..."):
                    try:
                        if st.session_state.acad_modo == "situacion_aprendizaje":
                            prompt = prompt_situacion_aprendizaje(
                                centro, docente, asignatura, grado, periodo,
                                num_semanas, competencias_sel, estrategias_sel, tema_unidad
                            )
                        else:
                            prompt = prompt_proyecto(
                                centro, docente, asignatura, grado, periodo,
                                num_semanas, competencias_sel, estrategias_sel, tema_unidad
                            )
                        
                        resultado, flags = ia.solicitar_json(
                            prompt, max_tokens=max_tokens, temperature=temperature, modulo="academicas"
                        )
                        
                        st.session_state.acad_resultado = {
                            "datos": resultado,
                            "flags": flags,
                            "meta": {
                                "centro": centro,
                                "docente": docente,
                                "asignatura": asignatura,
                                "grado": grado,
                                "periodo": periodo,
                                "num_semanas": num_semanas,
                                "tema_unidad": tema_unidad,
                                "modo": st.session_state.acad_modo,
                                "fecha": datetime.now().strftime("%Y-%m-%d"),
                            }
                        }
                        st.session_state.acad_fase = 2
                        st.toast("✅ Planificación de unidad generada.", icon="📖")
                        st.rerun()
                    
                    except Exception as e:
                        ia.render_error_ia(e)

# ═══════════════════════════════════════════════════════════════════════════
# FASE 2: RESULTADOS
# ═══════════════════════════════════════════════════════════════════════════
elif st.session_state.acad_fase == 2:
    resultado = st.session_state.acad_resultado
    datos = resultado["datos"]
    meta = resultado["meta"]
    modo = meta.get("modo", "situacion_aprendizaje")
    
    st.markdown('<div class="acad-section-title">🎉 Planificación de Unidad Generada</div>', unsafe_allow_html=True)
    
    # ─── MÉTRICAS ───
    modo_info = MODOS_PLANIFICACION.get(modo, {"nombre": "Académico", "icono": "📖"})
    
    col_s1, col_s2, col_s3, col_s4 = st.columns(4)
    with col_s1:
        st.markdown(f'<div class="acad-stat"><div class="acad-stat-value">{len(datos.get("COMPETENCIAS_FUNDAMENTALES", []))}</div><div class="acad-stat-label">Competencias Fund.</div></div>', unsafe_allow_html=True)
    with col_s2:
        contenidos = datos.get("CONTENIDOS", {})
        total_cont = len(contenidos.get("CONCEPTUALES", [])) + len(contenidos.get("PROCEDIMENTALES", [])) + len(contenidos.get("ACTITUDINALES", []))
        st.markdown(f'<div class="acad-stat"><div class="acad-stat-value">{total_cont}</div><div class="acad-stat-label">Contenidos</div></div>', unsafe_allow_html=True)
    with col_s3:
        if modo == "situacion_aprendizaje":
            st.markdown(f'<div class="acad-stat"><div class="acad-stat-value">{len(datos.get("SECUENCIA_DIDACTICA", []))}</div><div class="acad-stat-label">Semanas</div></div>', unsafe_allow_html=True)
        else:
            st.markdown(f'<div class="acad-stat"><div class="acad-stat-value">{len(datos.get("FASES_PROYECTO", []))}</div><div class="acad-stat-label">Fases</div></div>', unsafe_allow_html=True)
    with col_s4:
        st.markdown(f'<div class="acad-stat"><div class="acad-stat-value">{len(datos.get("RECURSOS", []))}</div><div class="acad-stat-label">Recursos</div></div>', unsafe_allow_html=True)
    
    st.markdown("---")
    
    # ─── TÍTULO Y MODO ───
    st.markdown(f"## {modo_info.get('icono', '📖')} {datos.get('TITULO_UNIDAD', '')}")
    st.caption(f"Modo: {modo_info.get('nombre', '')} | Asignatura: {meta.get('asignatura', '')} | Grado: {meta.get('grado', '')}")
    
    st.markdown("---")
    
    # ─── TABS ───
    tab_preview, tab_json, tab_debug = st.tabs(["👁️ Vista Previa", "🧾 JSON Completo", "🐛 Depuración"])
    
    with tab_preview:
        if modo == "situacion_aprendizaje":
            sa = datos.get("SITUACION_APRENDIZAJE", {})
            with st.expander("🎯 Situación de Aprendizaje", expanded=True):
                st.markdown(f"**Contexto:** {sa.get('CONTEXTO', '')}")
                st.markdown(f"**Problema:** {sa.get('PROBLEMA', '')}")
                st.markdown(f"**Producto Final:** {sa.get('PRODUCTO_FINAL', '')}")
                st.markdown("**Preguntas Guía:**")
                for preg in sa.get("PREGUNTAS_GUIA", []):
                    st.markdown(f"- {preg}")
        else:
            proy = datos.get("PROYECTO", {})
            with st.expander("🚀 Proyecto", expanded=True):
                st.markdown(f"**Título:** {proy.get('TITULO_PROYECTO', '')}")
                st.markdown(f"**Descripción:** {proy.get('DESCRIPCION', '')}")
                st.markdown(f"**Producto Final:** {proy.get('PRODUCTO_FINAL', '')}")
                st.markdown(f"**Destinatarios:** {proy.get('DESTINATARIOS', '')}")
        
        with st.expander("🎓 Competencias Fundamentales"):
            for comp in datos.get("COMPETENCIAS_FUNDAMENTALES", []):
                if isinstance(comp, dict):
                    st.markdown(f"**{comp.get('COMPETENCIA', '')}**: {comp.get('INDICADORES', '')}")
        
        with st.expander("📚 Contenidos"):
            contenidos = datos.get("CONTENIDOS", {})
            st.markdown("**Conceptuales:**")
            for c in contenidos.get("CONCEPTUALES", []):
                st.markdown(f"- {c}")
            st.markdown("**Procedimentales:**")
            for c in contenidos.get("PROCEDIMENTALES", []):
                st.markdown(f"- {c}")
            st.markdown("**Actitudinales:**")
            for c in contenidos.get("ACTITUDINALES", []):
                st.markdown(f"- {c}")
        
        if modo == "situacion_aprendizaje":
            with st.expander("📅 Secuencia Didáctica"):
                for semana in datos.get("SECUENCIA_DIDACTICA", []):
                    st.markdown(f"**Semana {semana.get('SEMANA', '')}: {semana.get('TEMATICA', '')}**")
                    st.markdown(f"Inicio: {semana.get('ACTIVIDADES_INICIO', '')[:150]}...")
                    st.markdown(f"Desarrollo: {semana.get('ACTIVIDADES_DESARROLLO', '')[:150]}...")
                    st.markdown("---")
        else:
            with st.expander("📅 Fases del Proyecto"):
                for fase in datos.get("FASES_PROYECTO", []):
                    st.markdown(f"**{fase.get('FASE', '')}** (Semanas {fase.get('SEMANAS', '')})")
                    for act in fase.get("ACTIVIDADES", []):
                        st.markdown(f"- {act}")
                    st.markdown("---")
    
    with tab_json:
        st.json(datos)
    
    with tab_debug:
        st.write("Meta:", meta)
        st.write("Flags:", resultado.get("flags", {}))
    
    st.markdown("---")
    
    # ─── DESCARGA WORD ───
    col_dl1, col_dl2 = st.columns(2)
    
    with col_dl1:
        if st.button("📥 Generar Word (.docx)", type="primary", use_container_width=True):
            with st.spinner("📄 Construyendo documento Word..."):
                buffer = build_acad_docx(datos, meta, modo)
                st.session_state.acad_buffer = buffer
    
    if hasattr(st.session_state, "acad_buffer") and st.session_state.acad_buffer:
        nombre_archivo = ia.sanear_nombre_archivo(f"Planificacion_Unidad_{meta.get('asignatura', 'unidad')}")
        st.download_button(
            label="⬇️ Descargar Planificación de Unidad (.docx)",
            data=st.session_state.acad_buffer,
            file_name=f"{nombre_archivo}.docx",
            mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
            type="primary",
            use_container_width=True
        )
    
    with col_dl2:
        if st.button("🔄 Nueva Planificación", use_container_width=True):
            st.session_state.acad_fase = 0
            st.session_state.acad_resultado = None
            st.session_state.acad_modo = None
            if hasattr(st.session_state, "acad_buffer"):
                del st.session_state.acad_buffer
            st.rerun()