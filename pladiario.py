"""
pladiario.py — Planificación de Clase Diaria ETP (MIGRADO · Paso 12 · MEJORADO)
• Genera el Plan Diario alineado al esquema oficial MINERD (réplica del PDF institucional).
• IA vía core/ia (marcadores, parseo robusto, reintento, auditoría).
• MEJORAS: Selectores de Tipo de Actividad y Estrategias, campo de Actividad del RA,
  y adaptación inteligente de la actividad al Tiempo Estimado del docente.
"""
import datetime
from io import BytesIO

import streamlit as st
from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement, parse_xml
from docx.oxml.ns import nsdecls, qn
from docx.shared import Inches, Pt, RGBColor

from core import ia

# ═══════════════════════════════════════════════════════════════════════════
# CONSTANTES DE DOMINIO ETP
# ═══════════════════════════════════════════════════════════════════════════
TIPOS_ACTIVIDAD = [
    "En equipos de 2 (parejas)",
    "En equipos de 3",
    "En equipos de 4",
    "En equipos de 5",
    "Individual",
    "Exposición oral",
    "Taller práctico",
    "Simulación",
    "Estudio de caso",
    "Proyecto",
    "Debate",
    "Mesa redonda",
    "Juego de roles",
    "Laboratorio",
    "Práctica de taller",
    "Demostración",
    "Trabajo de campo",
    "Aprendizaje por descubrimiento",
]

ESTRATEGIAS_EA = [
    "Estudio de casos",
    "Trabajo colaborativo",
    "Aprendizaje basado en problemas (ABP)",
    "Simulación",
    "Indagación dialógica",
    "Aprendizaje por descubrimiento",
    "Taller práctico",
    "Gamificación",
    "Aprendizaje basado en proyectos",
    "Aula invertida",
    "Aprendizaje cooperativo",
    "Técnica del rompecabezas",
    "Método expositivo",
    "Resolución de problemas",
    "Aprendizaje significativo",
    "Contrato de aprendizaje",
    "Enseñanza recíproca",
    "Mapa mental y conceptual",
]

TIEMPOS_ESTIMADOS = [
    "45 minutos",
    "50 minutos (estándar MINERD)",
    "60 minutos",
    "1 hora 15 minutos",
    "1 hora 30 minutos",
    "2 horas",
]

# ═══════════════════════════════════════════════════════════════════════════
# UTILIDADES WORD
# ═══════════════════════════════════════════════════════════════════════════
def shade_cell(cell, color):
    shd = parse_xml(r'<w:shd {} w:fill="{}"/>'.format(nsdecls('w'), color))
    cell._tc.get_or_add_tcPr().append(shd)

def fijar_anchos_columna(tabla, anchos_pulgadas):
    tabla.autofit = False
    for row in tabla.rows:
        for idx, ancho in enumerate(anchos_pulgadas):
            if idx < len(row.cells):
                row.cells[idx].width = Inches(ancho)

def agregar_numeracion_pagina(doc):
    pie = doc.sections[0].footer.paragraphs[0]
    pie.alignment = WD_ALIGN_PARAGRAPH.CENTER
    pie.text = ""
    run = pie.add_run()
    def _campo(instruccion):
        inicio = OxmlElement("w:fldChar")
        inicio.set(qn("w:fldCharType"), "begin")
        instr = OxmlElement("w:instrText")
        instr.set(qn("xml:space"), "preserve")
        instr.text = instruccion
        fin = OxmlElement("w:fldChar")
        fin.set(qn("w:fldCharType"), "end")
        run._r.append(inicio)
        run._r.append(instr)
        run._r.append(fin)
    run.add_text("Página ")
    _campo("PAGE")
    run.add_text(" de ")
    _campo("NUMPAGES")
    run.font.size = Pt(8)
    run.font.color.rgb = RGBColor(0x64, 0x74, 0x8B)

def _get(d, clave, default=""):
    if isinstance(d, dict):
        return d.get(clave, default)
    return default

# ═══════════════════════════════════════════════════════════════════════════
# GENERADOR WORD — RÉPLICA DEL ESQUEMA OFICIAL
# ═══════════════════════════════════════════════════════════════════════════
def generar_documento_plandiario(datos, f):
    doc = Document()
    doc.styles['Normal'].font.name = 'Calibri'
    doc.styles['Normal'].font.size = Pt(10)
    for section in doc.sections:
        section.left_margin = Inches(0.5)
        section.right_margin = Inches(0.5)
        section.top_margin = Inches(0.5)
        section.bottom_margin = Inches(0.5)

    # ── Encabezado oficial ──
    p_titulo = doc.add_paragraph()
    p_titulo.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r1 = p_titulo.add_run("MINISTERIO DE EDUCACIÓN DE LA REPÚBLICA DOMINICANA\n")
    r1.bold = True
    r1.font.size = Pt(14)
    r2 = p_titulo.add_run("PLANIFICACIÓN DE CLASE DIARIA\n")
    r2.bold = True
    r2.font.size = Pt(12)
    r3 = p_titulo.add_run("Modalidad Técnico Profesional (ETP)")
    r3.italic = True
    r3.font.size = Pt(11)

    def add_table_header(text):
        p = doc.add_paragraph()
        run = p.add_run(text)
        run.bold = True
        run.font.color.rgb = RGBColor(37, 99, 235)

    # ── DATOS GENERALES ──
    add_table_header("DATOS GENERALES")
    t1 = doc.add_table(rows=7, cols=4)
    t1.style = 'Table Grid'
    filas = [
        ("Nombre completo", f['docente'], "Cédula", f['cedula']),
        ("Regional", f['regional'], "Distrito", f['distrito']),
        ("Centro Educativo", f['centro'], "Código del Centro", f['codigo_centro']),
        ("Nivel / Subsistema", "Secundaria", "Ciclo", "Segundo"),
        ("Grado y Sección", f['grado'], "Modalidad", "Técnico Profesional (ETP)"),
        ("Área / Asignatura / Módulo", f['modulo'], "Fecha", f['fecha'].strftime('%d/%m/%Y')),
        ("Duración", f['tiempo_estimado'], "Tanda", f['tanda']),
    ]
    for i, (l1, v1, l2, v2) in enumerate(filas):
        t1.cell(i, 0).text = l1
        t1.cell(i, 1).text = str(v1)
        t1.cell(i, 2).text = l2
        t1.cell(i, 3).text = str(v2)
    for row in t1.rows:
        for c_idx, cell in enumerate(row.cells):
            if c_idx % 2 == 0 and cell.paragraphs[0].runs:
                cell.paragraphs[0].runs[0].bold = True
                shade_cell(cell, "F1F5F9")
    fijar_anchos_columna(t1, [1.5, 2.5, 1.5, 2.0])
    doc.add_paragraph()

    # ── MATRIZ DE PLANIFICACIÓN DIARIA ──
    add_table_header("Matriz de Planificación Diaria o Por Actividad - Bachillerato Técnico")
    comp = datos.get("COMPONENTES", {})
    t2 = doc.add_table(rows=8, cols=2)
    t2.style = 'Table Grid'
    t2.cell(0, 0).text = "Características del grupo de estudiantes"
    t2.cell(0, 1).text = f['caracteristicas']
    t2.cell(1, 0).text = "Módulo Formativo (MF)"
    t2.cell(1, 1).text = f['modulo']
    t2.cell(2, 0).text = "Resultado de Aprendizaje (RA)"
    t2.cell(2, 1).text = f['ra']
    t2.cell(3, 0).text = "Criterio de Evaluación (CE)"
    t2.cell(3, 1).text = f['ce']
    t2.cell(4, 0).text = "Elemento de Capacidad (EC)"
    t2.cell(4, 1).text = f['ec']
    t2.cell(5, 0).text = "Tipo / Tiempo Estimado / Estrategias / Valor"
    t2.cell(5, 1).text = (
        f"Tipo: {f['tipo_actividad']}\nTiempo: {f['tiempo_estimado']}\n"
        f"Estrategias: {f['estrategias']}\nValor: {f['valor']}"
    )
    t2.cell(6, 0).text = "Componentes Curriculares (Contenidos)"
    t2.cell(6, 1).text = (
        f"Conceptuales:\n{_get(comp, 'CONCEPTUALES')}\n\n"
        f"Procedimentales:\n{_get(comp, 'PROCEDIMENTALES')}\n\n"
        f"Actitudinales:\n{_get(comp, 'ACTITUDINALES')}"
    )
    t2.cell(7, 0).text = "Roles en el aula"
    t2.cell(7, 1).text = (
        f"Docente (modelado y mediación): {_get(datos.get('ROLES', {}), 'DOCENTE')}\n"
        f"Estudiante (práctica y construcción): {_get(datos.get('ROLES', {}), 'ESTUDIANTE')}"
    )
    for row in t2.rows:
        if row.cells[0].paragraphs[0].runs:
            row.cells[0].paragraphs[0].runs[0].bold = True
            shade_cell(row.cells[0], "F1F5F9")
    fijar_anchos_columna(t2, [2.0, 5.5])
    doc.add_paragraph()

    # ── ENUNCIADO E INTENCIÓN ──
    add_table_header("Enunciado de la Actividad e Intención Educativa")
    act = datos.get("ACTIVIDAD", {})
    t3 = doc.add_table(rows=2, cols=1)
    t3.style = 'Table Grid'
    t3.cell(0, 0).text = f"Enunciado de la Actividad:\n{_get(act, 'ENUNCIADO')}"
    t3.cell(1, 0).text = f"Intención Educativa:\n{_get(act, 'INTENCION')}"
    for row in t3.rows:
        if row.cells[0].paragraphs[0].runs:
            row.cells[0].paragraphs[0].runs[0].bold = True
    fijar_anchos_columna(t3, [7.5])
    doc.add_paragraph()

    # ── MOMENTOS PEDAGÓGICOS ──
    add_table_header("Momentos Pedagógicos")
    ini, des, cie = datos.get("INICIO", {}), datos.get("DESARROLLO", {}), datos.get("CIERRE", {})
    t4 = doc.add_table(rows=3, cols=1)
    t4.style = 'Table Grid'
    p_inicio = t4.cell(0, 0).paragraphs[0]
    p_inicio.add_run(f"INICIO ({_get(ini, 'TIEMPO', '10 min')})\n").bold = True
    p_inicio.add_run(
        f"FASE 1 — Motivación y activación de conocimientos previos:\n{_get(ini, 'FASE1')}\n\n"
        f"FASE 2 — Recuperación de saberes previos:\n{_get(ini, 'FASE2')}\n\n"
        f"FASE 3 — Presentación de la intención educativa:\n{_get(ini, 'FASE3')}"
    )
    p_des = t4.cell(1, 0).paragraphs[0]
    p_des.add_run(f"DESARROLLO ({_get(des, 'TIEMPO', '30 min')}) — Construcción del aprendizaje\n").bold = True
    p_des.add_run(
        f"FASE 1 (modelado docente):\n{_get(des, 'FASE1')}\n\n"
        f"FASE 2 (práctica colaborativa del estudiante):\n{_get(des, 'FASE2')}\n\n"
        f"FASE 3 (socialización y retroalimentación):\n{_get(des, 'FASE3')}"
    )
    p_cie = t4.cell(2, 0).paragraphs[0]
    p_cie.add_run(f"CIERRE ({_get(cie, 'TIEMPO', '10 min')}) — Reflexión, consolidación y metacognición\n").bold = True
    p_cie.add_run(
        f"FASE 1 (actividad de cierre interactivo):\n{_get(cie, 'FASE1')}\n\n"
        f"FASE 2 (metacognición):\n{_get(cie, 'FASE2')}\n\n"
        f"FASE 3 (preguntas de reflexión):\n{_get(cie, 'FASE3')}"
    )
    fijar_anchos_columna(t4, [7.5])
    doc.add_paragraph()

    # ── RECURSOS, EVALUACIÓN, NEAE, OBSERVACIONES ──
    add_table_header("Recursos, Adaptaciones y Observaciones")
    t5 = doc.add_table(rows=4, cols=1)
    t5.style = 'Table Grid'
    p_rec = t5.cell(0, 0).paragraphs[0]
    p_rec.add_run("Recursos:\n").bold = True
    p_rec.add_run(str(datos.get("RECURSOS", "")))
    p_eval = t5.cell(1, 0).paragraphs[0]
    p_eval.add_run("Instrumento/s de evaluación e Indicadores:\n").bold = True
    p_eval.add_run(str(datos.get("INDICADORES_TEXT", "")))
    p_neae = t5.cell(2, 0).paragraphs[0]
    p_neae.add_run("Adaptaciones para NEAE:\n").bold = True
    p_neae.add_run(str(datos.get("NEAE", "")))
    p_obs = t5.cell(3, 0).paragraphs[0]
    p_obs.add_run("Observaciones:\n").bold = True
    p_obs.add_run(str(datos.get("OBSERVACIONES", "")))
    fijar_anchos_columna(t5, [7.5])
    doc.add_paragraph()

    # ── LISTA DE COTEJO ──
    add_table_header("Instrumento de Evaluación: Lista de Cotejo")
    t6 = doc.add_table(rows=1, cols=6)
    t6.style = 'Table Grid'
    hdr = t6.rows[0].cells
    for i, h in enumerate(["No.", "Criterios de Evaluación", "L", "EP", "NA", "Observaciones"]):
        hdr[i].text = h
        hdr[i].paragraphs[0].runs[0].bold = True
        hdr[i].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
        shade_cell(hdr[i], "DBEAFE")
    criterios = datos.get("COTEJO", [])
    while len(criterios) < 5:
        criterios.append("Criterio pendiente de definir")
    for i, crit in enumerate(criterios[:5], start=1):
        row = t6.add_row().cells
        row[0].text = str(i)
        row[1].text = str(crit)
        for j in range(2, 5):
            row[j].text = ""
            row[j].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
        row[0].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
    fijar_anchos_columna(t6, [0.4, 3.8, 0.4, 0.4, 0.4, 2.1])

    # ── ESCALA DE VALORACIÓN ──
    p_esc = doc.add_paragraph()
    p_esc.add_run("Escala de Valoración\n").bold = True
    t7 = doc.add_table(rows=4, cols=2)
    t7.style = 'Table Grid'
    t7.cell(0, 0).text = "Sigla"
    t7.cell(0, 1).text = "Descripción"
    for c in t7.rows[0].cells:
        c.paragraphs[0].runs[0].bold = True
        shade_cell(c, "F1F5F9")
    t7.cell(1, 0).text = "L"
    t7.cell(1, 1).text = "Logrado (4 o 5, Máximo 2 indicadores en EP y los demás en L)"
    t7.cell(2, 0).text = "EP"
    t7.cell(2, 1).text = "En proceso (3 o 4, Máximo 2 indicadores en L, pero 3 en EP)"
    t7.cell(3, 0).text = "NA"
    t7.cell(3, 1).text = "Necesita apoyo (3 o más indicadores en NA, con 1 o 2 EP y sin apenas ninguna L)"
    fijar_anchos_columna(t7, [0.8, 6.7])

    # ── FIRMAS ──
    doc.add_paragraph("\n\n")
    t_firmas = doc.add_table(rows=2, cols=3)
    t_firmas.cell(0, 0).text = "__________________________"
    t_firmas.cell(0, 1).text = "__________________________"
    t_firmas.cell(0, 2).text = "__________________________"
    t_firmas.cell(1, 0).text = "Director/a de Centro Educativo"
    t_firmas.cell(1, 1).text = "Coordinador/a ETP"
    t_firmas.cell(1, 2).text = "Docente ETP"
    for row in t_firmas.rows:
        for cell in row.cells:
            cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER

    agregar_numeracion_pagina(doc)
    buffer = BytesIO()
    doc.save(buffer)
    buffer.seek(0)
    return buffer

# ═══════════════════════════════════════════════════════════════════════════
# INTERFAZ (SUPER UI)
# ═══════════════════════════════════════════════════════════════════════════
st.markdown("""
<style>
@import url('https://fonts.googleapis.com/css2?family=Inter:wght@300;400;500;600;700;800&display=swap');
html, body, [class*="css"] { font-family: 'Inter', sans-serif; background-color: #F8FAFC; color: #1E293B; }
.pd-hero { background: linear-gradient(135deg, #0F172A 0%, #1E3A8A 55%, #0EA5E9 100%); color: #fff;
padding: 1.8rem; border-radius: 18px; margin-bottom: 1.25rem; box-shadow: 0 18px 35px rgba(15,23,42,.18); }
.pd-title { font-size: 2rem; font-weight: 800; letter-spacing: -0.02em; }
.pd-sub { opacity: .88; font-size: 1rem; margin-top: .3rem; }
.section-title { color: #1D4ED8; font-weight: 700; font-size: 1.12rem; border-bottom: 2px solid #DBEAFE;
padding-bottom: 8px; margin: 1.2rem 0 .9rem 0; }
.pd-info-box { background: #EFF6FF; border-left: 4px solid #3B82F6; padding: 12px 16px; border-radius: 6px; margin-bottom: 12px; }
</style>
""", unsafe_allow_html=True)

ia.panel_sidebar_ia("Plan Diario ETP")

st.markdown("""
<div class="pd-hero">
    <div class="pd-title">📅 Planificación de Clase Diaria ETP</div>
    <div class="pd-sub">Estructura Oficial MINERD · Adaptación inteligente de la actividad al tiempo estimado · Nace "apto para auditoría"</div>
</div>
""", unsafe_allow_html=True)

with st.form("form_plandiario", clear_on_submit=False):
    st.markdown('<div class="section-title">🏫 1. Datos Generales</div>', unsafe_allow_html=True)
    col1, col2, col3 = st.columns(3)
    with col1:
        docente = st.text_input("Nombre completo", placeholder="Nombre del Docente")
        regional = st.text_input("Regional", value="06")
        centro = st.text_input("Centro Educativo", placeholder="Nombre del Centro")
        grado = st.text_input("Grado y Sección", value="5to B")
        fecha = st.date_input("Fecha")
    with col2:
        cedula = st.text_input("Cédula", placeholder="000-0000000-0")
        distrito = st.text_input("Distrito", value="06")
        codigo_centro = st.text_input("Código del Centro", placeholder="00000")
        tanda = st.text_input("Tanda", placeholder="Ej: Matutina, Vespertina, JEE")
    with col3:
        modulo = st.text_area("Área / Asignatura / Módulo", height=100,
                              placeholder="Ej: MF 358-3 Impuestos al Consumo y Vehículos de Motor")
        tipo_actividad = st.selectbox("Tipo de Actividad", TIPOS_ACTIVIDAD, index=2)
        tiempo_estimado = st.selectbox("Tiempo Estimado", TIEMPOS_ESTIMADOS, index=1)
        estrategias = st.multiselect("Estrategias de Enseñanza-Aprendizaje", ESTRATEGIAS_EA,
                                     default=["Estudio de casos", "Trabajo colaborativo"])
        valor = st.text_input("Valor", value="5 pts.")

    st.markdown('<div class="section-title">🎯 2. Parámetros Curriculares</div>', unsafe_allow_html=True)
    col_c1, col_c2 = st.columns(2)
    with col_c1:
        ra = st.text_area("Resultado de Aprendizaje (RA)", height=100)
        ec = st.text_area("Elemento de Capacidad (EC)", height=100)
    with col_c2:
        ce = st.text_area("Criterio de Evaluación (CE)", height=100)
        caracteristicas = st.text_area("Características del grupo de estudiantes", height=100,
                                       placeholder="Describa el perfil sociocognitivo del grupo...")

    st.markdown('<div class="section-title">🎯 3. Actividad Correspondiente al RA</div>', unsafe_allow_html=True)
    st.markdown("""
    <div class="pd-info-box">
        💡 <b>Instrucción:</b> Describe la actividad que deseas desarrollar para este RA.
        La IA analizará la actividad y la <b>adaptará al tiempo estimado</b> que seleccionaste,
        distribuyendo proporcionalmente los momentos pedagógicos (Inicio, Desarrollo, Cierre).
    </div>
    """, unsafe_allow_html=True)
    actividad_ra = st.text_area(
        "Describe la actividad que deseas desarrollar para este RA:",
        height=120,
        placeholder="Ej: Los estudiantes analizarán un caso práctico de importación de vehículos de lujo, identificarán los documentos requeridos, los impuestos aplicables y las instituciones responsables. Trabajarán en equipos para completar una tabla de clasificación y luego socializarán sus conclusiones.",
    )

    max_tokens, temperature = ia.control_avanzado(default_tokens=16384, tope=32000, default_temp=0.2)
    modo_debug = st.checkbox("🐛 Modo depuración (ver JSON crudo)", value=False)
    st.markdown("<br>", unsafe_allow_html=True)
    submit_button = st.form_submit_button("⚙️ Generar Planificación Diaria", type="primary", width="stretch")

if submit_button:
    cfg = ia.config_ia()
    if not cfg["api_key"]:
        st.error("🔒 Debes ingresar tu API Key en la página de Inicio.")
    elif not docente or not modulo or not ra or not ce or not ec or not caracteristicas or not actividad_ra:
        st.warning("📝 Por favor, completa los datos básicos, curriculares y la actividad del RA.")
    else:
        # Convertir estrategias de lista a texto
        estrategias_texto = " · ".join(estrategias) if estrategias else "Estudio de casos · Trabajo colaborativo"

        with st.spinner(f'🧠 Diseñando matriz oficial con {cfg["modelo"]}...'):
            texto_crudo = None
            try:
                prompt_maestro = f"""Actúa como experto en diseño curricular de la ETP (MINERD).

Objetivo: Diseñar el contenido pedagógico de una "Planificación de Clase Diaria" para Bachillerato Técnico.

El resultado DEBE aprobar una auditoría de 14 criterios oficiales, por lo que cumple ESTRICTAMENTE:
- Momentos dosificados según el TIEMPO ESTIMADO: distribuye proporcionalmente Inicio, Desarrollo y Cierre.
- El CIERRE incluye preguntas de metacognición en su FASE 3.
- El DESARROLLO diferencia el rol del docente (modelado) y del estudiante (práctica).
- Tridimensionalidad completa: contenidos conceptuales, procedimentales y actitudinales.
- Actividad práctica anclada al RA/EC e intención educativa orientada al perfil profesional.
- Recursos técnicos verídicos (PDI, fichas técnicas, documentos simulados, plataformas).
- Lista de cotejo con EXACTAMENTE 5 criterios observables derivados del CE.
- Adaptaciones NEAE y observaciones con plan alternativo ante contingencias.
- PROHIBIDO dejar campos vacíos.

REGLA CRÍTICA DE ADAPTACIÓN AL TIEMPO:
El TIEMPO ESTIMADO de la clase es: {tiempo_estimado}
ANALIZA la actividad propuesta por el docente y ADÁPTALA a este tiempo:
- Si el tiempo es corto (45-50 min), simplifica las actividades y prioriza lo esencial.
- Si el tiempo es largo (90 min o más), puedes agregar más detalle, actividades adicionales o profundizar.
- Distribuye el tiempo proporcionalmente entre Inicio (20%), Desarrollo (60%) y Cierre (20%).
- Asegúrate de que todas las fases quepan en el tiempo estimado.
- En el campo TIEMPO de cada momento, indica los minutos asignados (ej: "10 min", "30 min").

INSUMOS:
Características del grupo: {caracteristicas}
Módulo: {modulo}
RA: {ra}
CE: {ce}
EC: {ec}
Tipo de Actividad: {tipo_actividad}
Tiempo Estimado: {tiempo_estimado}
Estrategias: {estrategias_texto}
Valor: {valor}

ACTIVIDAD PROPUESTA POR EL DOCENTE (analízala y adáptala al tiempo estimado):
{actividad_ra}

CODIFICACIÓN OBLIGATORIA: salto de línea → {ia.MARKER_NL} · comilla doble → {ia.MARKER_DQ} · tabulación → {ia.MARKER_TAB}.
Nunca uses saltos de línea literales dentro de un valor JSON.

FORMATO JSON:
{{
 "COMPONENTES": {{ "CONCEPTUALES": "...", "PROCEDIMENTALES": "...", "ACTITUDINALES": "..." }},
 "ACTIVIDAD": {{ "ENUNCIADO": "...", "INTENCION": "..." }},
 "ROLES": {{ "DOCENTE": "...", "ESTUDIANTE": "..." }},
 "INICIO": {{ "TIEMPO": "X min", "FASE1": "...", "FASE2": "...", "FASE3": "..." }},
 "DESARROLLO": {{ "TIEMPO": "X min", "FASE1": "...", "FASE2": "...", "FASE3": "..." }},
 "CIERRE": {{ "TIEMPO": "X min", "FASE1": "...", "FASE2": "...", "FASE3": "..." }},
 "RECURSOS": "...",
 "INDICADORES_TEXT": "...",
 "NEAE": "...",
 "OBSERVACIONES": "...",
 "COTEJO": ["criterio 1", "criterio 2", "criterio 3", "criterio 4", "criterio 5"]
}}
"""
                texto_crudo, flags = ia.solicitar_ia(
                    prompt_maestro, modo="json", max_tokens=max_tokens,
                    temperature=temperature, modulo="pladiario",
                )
                datos = ia.decodificar_marcadores(ia.parsear_json_robusto(texto_crudo))

                datos_formulario = {
                    "docente": docente, "cedula": cedula, "regional": regional, "distrito": distrito,
                    "centro": centro, "codigo_centro": codigo_centro, "grado": grado, "tanda": tanda,
                    "modulo": modulo, "tipo_actividad": tipo_actividad, "tiempo_estimado": tiempo_estimado,
                    "estrategias": estrategias_texto, "valor": valor, "fecha": fecha,
                    "ra": ra, "ce": ce, "ec": ec, "caracteristicas": caracteristicas,
                }
                buffer_docx = generar_documento_plandiario(datos, datos_formulario)

                n_fases = sum(
                    1 for m in ("INICIO", "DESARROLLO", "CIERRE")
                    for k in ("FASE1", "FASE2", "FASE3")
                    if _get(datos.get(m, {}), k)
                )
                m1, m2, m3 = st.columns(3)
                with m1:
                    st.metric("Momentos", 3)
                with m2:
                    st.metric("Fases redactadas", n_fases)
                with m3:
                    st.metric("Criterios de cotejo", len(datos.get("COTEJO", [])))

                st.toast("✅ Planificación diaria generada con el formato oficial.", icon="📅")
                if flags.get("reintento"):
                    st.info("ℹ️ La primera llamada se cortó por tokens; el reintento automático tuvo éxito.")

                with st.expander("👁️ Vista previa del contenido generado"):
                    st.json(datos)
                if modo_debug:
                    with st.expander("🐛 JSON crudo de la IA"):
                        st.text((texto_crudo or "")[:5000])

                st.download_button(
                    label="📥 Descargar Planificación (.docx)",
                    data=buffer_docx,
                    file_name=f"Plan_Diario_{ia.sanear_nombre_archivo(docente)}_{fecha.strftime('%Y%m%d')}.docx",
                    mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                    type="primary",
                    width="stretch",
                )
            except ValueError as ve:
                ia.render_error_ia(ve, texto_crudo)
            except Exception as e:
                ia.render_error_ia(e, texto_crudo)