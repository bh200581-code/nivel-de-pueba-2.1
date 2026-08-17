"""
planifiacionra.py — Matriz de Planificación por R.A. (Depurado)
• IA vía motor unificado alineado a la Guía MINERD.
• Interfaz, métricas, y documento Word profesional.
• Código optimizado sin errores de linting ni parámetros obsoletos.
"""
import re
import unicodedata
from io import BytesIO

import PyPDF2
import streamlit as st
from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import parse_xml
from docx.oxml.ns import nsdecls
from docx.shared import Inches, Pt, RGBColor

from core import ia


# ═══════════════════════════════════════════════════════════════════════════
# ESTILOS (INTERFAZ MEJORADA)
# ═══════════════════════════════════════════════════════════════════════════
st.markdown("""
<style>
@import url('https://fonts.googleapis.com/css2?family=Inter:wght@300;400;500;600;700;800;900&display=swap');
html, body, [class*="css"] { font-family: 'Inter', sans-serif; background-color: #F8FAFC; color: #1E293B; }

.planra-hero {
    background: linear-gradient(135deg, #0F172A 0%, #1E3A8A 50%, #3B82F6 100%);
    color: white;
    padding: 2rem;
    border-radius: 18px;
    margin-bottom: 1.5rem;
    box-shadow: 0 18px 35px rgba(15, 23, 42, 0.18);
    position: relative;
    overflow: hidden;
}
.planra-hero::before {
    content: '';
    position: absolute;
    top: -50%; left: -50%; width: 200%; height: 200%;
    background: radial-gradient(circle, rgba(255,255,255,0.12) 0%, transparent 60%);
    animation: planraPulse 7s ease-in-out infinite;
}
@keyframes planraPulse { 0%,100%{transform:scale(1)} 50%{transform:scale(1.05)} }
.planra-hero-title { font-size: 2.2rem; font-weight: 900; letter-spacing: -0.03em; position: relative; }
.planra-hero-sub { font-size: 1.05rem; opacity: 0.88; line-height: 1.5; position: relative; margin-top: 0.4rem; }
.planra-hero-badge {
    display: inline-block;
    background: rgba(255,255,255,0.15);
    border: 1px solid rgba(255,255,255,0.25);
    border-radius: 10px;
    padding: 5px 13px;
    font-size: 0.8rem;
    font-weight: 600;
    margin-right: 8px;
    margin-top: 8px;
    position: relative;
}

.planra-section-title {
    color: #1D4ED8;
    font-weight: 700;
    font-size: 1.15rem;
    border-bottom: 2px solid #DBEAFE;
    padding-bottom: 8px;
    margin-top: 1.4rem;
    margin-bottom: 1rem;
}

.planra-kpi {
    background: #fff;
    border-radius: 14px;
    padding: 1.1rem;
    text-align: center;
    border: 1px solid #E2E8F0;
    box-shadow: 0 4px 12px rgba(0,0,0,0.05);
    transition: all 0.25s ease;
}
.planra-kpi:hover { transform: translateY(-3px); box-shadow: 0 10px 22px rgba(59,130,246,0.12); }
.planra-kpi-icon { font-size: 1.8rem; margin-bottom: 0.3rem; }
.planra-kpi-value { font-size: 1.9rem; font-weight: 800; color: #1D4ED8; }
.planra-kpi-label { font-size: 0.72rem; font-weight: 700; color: #64748B; text-transform: uppercase; letter-spacing: 0.05em; margin-top: 0.3rem; }

.planra-preview-box {
    background: #F0F9FF;
    border-left: 4px solid #0EA5E9;
    padding: 12px 16px;
    border-radius: 6px;
    margin-bottom: 10px;
}
.planra-preview-ec { font-weight: 700; color: #0C4A6E; font-size: 0.95rem; }
.planra-preview-act { color: #334155; font-size: 0.88rem; margin-top: 4px; }
.planra-preview-nivel {
    display: inline-block;
    padding: 2px 10px;
    border-radius: 999px;
    font-size: 0.72rem;
    font-weight: 700;
    margin-left: 6px;
}
</style>
""", unsafe_allow_html=True)

# ═══════════════════════════════════════════════════════════════════════════
# SIDEBAR + HERO
# ═══════════════════════════════════════════════════════════════════════════
ia.panel_sidebar_ia("Planificación por R.A.")

st.markdown("""
<div class="planra-hero">
    <div class="planra-hero-title">Matriz de Planificación por R.A.</div>
    <div class="planra-hero-sub">
        Compilación curricular ETP alineada a auditoría MINERD. Extrae el Módulo Formativo
        del PDF oficial y diseña la matriz de actividades por Elementos de Capacidad bajo el enfoque por competencias.
    </div>
    <div>
        <span class="planra-hero-badge">📄 Extracción del PDF</span>
        <span class="planra-hero-badge">🎯 3 Niveles de Desempeño</span>
        <span class="planra-hero-badge">📊 Instrumentos Basados en Evidencias</span>
        <span class="planra-hero-badge">📥 Word Profesional</span>
    </div>
</div>
""", unsafe_allow_html=True)


# ═══════════════════════════════════════════════════════════════════════════
# FUNCIONES AUXILIARES DE EXTRACCIÓN Y FORMATO
# ═══════════════════════════════════════════════════════════════════════════
def extraer_texto_pdf(archivo):
    """Extrae texto de un archivo PDF ignorando páginas vacías."""
    pdf_reader = PyPDF2.PdfReader(archivo)
    total_pag = len(pdf_reader.pages)
    texto = "".join([pagina.extract_text() or "" for pagina in pdf_reader.pages])
    return texto[:60000], total_pag


def sanear_nombre(texto):
    """Elimina caracteres especiales para generar nombres de archivo seguros."""
    if not texto:
        return "archivo"
    txt_norm = unicodedata.normalize('NFKD', texto).encode('ascii', 'ignore').decode('ascii')
    txt_norm = re.sub(r'[^A-Za-z0-9_\- ]', '', txt_norm).strip()
    return re.sub(r'\s+', '_', txt_norm)[:60]


# ═══════════════════════════════════════════════════════════════════════════
# GENERACIÓN DEL DOCUMENTO WORD
# ═══════════════════════════════════════════════════════════════════════════
NIVEL_COLORES = {
    "1": ("DBEAFE", "Nivel 1 — Conocimiento"),
    "2": ("D1FAE5", "Nivel 2 — Aplicación"),
    "3": ("EDE9FE", "Nivel 3 — Dominio / Autonomía"),
}


def shade_cell(cell, color):
    """Aplica un color de fondo a una celda de Word."""
    shd = parse_xml(rf'<w:shd {nsdecls("w")} w:fill="{color}"/>')
    cell._tc.get_or_add_tcPr().append(shd)


def set_cell_text(cell, text, bold=False, center=False, size=9, color=None, font_color=None):
    """Configura el texto, alineación, color de fuente y fondo de una celda."""
    cell.text = ""
    parrafo = cell.paragraphs[0]
    run = parrafo.add_run(str(text))
    run.bold = bold
    run.font.size = Pt(size)
    if font_color:
        run.font.color.rgb = RGBColor(*font_color)
    if center:
        parrafo.alignment = WD_ALIGN_PARAGRAPH.CENTER
    if color:
        shade_cell(cell, color)


def fijar_anchos_columna(tabla, anchos_pulgadas):
    """Fija el ancho exacto de cada columna en una tabla de Word."""
    tabla.autofit = False
    for row in tabla.rows:
        for idx, ancho in enumerate(anchos_pulgadas):
            if idx < len(row.cells):
                row.cells[idx].width = Inches(ancho)


def generar_documento_word(datos_json, form_data):
    """Construye y formatea el archivo DOCX con los resultados de la IA."""
    doc = Document()
    for section in doc.sections:
        section.left_margin = Inches(0.5)
        section.right_margin = Inches(0.5)
        section.top_margin = Inches(0.5)
        section.bottom_margin = Inches(0.5)
    
    style = doc.styles['Normal']
    style.font.name = 'Calibri'
    style.font.size = Pt(10)

    # ── ENCABEZADO OFICIAL MINERD ──
    p_enc = doc.add_paragraph()
    p_enc.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r_minerd = p_enc.add_run("MINISTERIO DE EDUCACIÓN DE LA REPÚBLICA DOMINICANA\n")
    r_minerd.bold = True
    r_minerd.font.size = Pt(12)
    
    r_poli = p_enc.add_run(f"{form_data['politecnico']}\n")
    r_poli.bold = True
    r_poli.font.size = Pt(11)
    
    if form_data['eslogan']:
        r_eslogan = p_enc.add_run(f"{form_data['eslogan']}\n")
        r_eslogan.italic = True
        r_eslogan.font.size = Pt(9)
        
    r_titulo = p_enc.add_run("MATRIZ DE PLANIFICACIÓN POR RESULTADOS DE APRENDIZAJE")
    r_titulo.bold = True
    r_titulo.font.size = Pt(12)
    r_titulo.font.color.rgb = RGBColor(0x1D, 0x4E, 0xD8)

    doc.add_paragraph()

    # ── I. DATOS GENERALES DEL MÓDULO ──
    doc.add_heading("I. Datos Generales del Módulo Formativo", level=2)
    datos_gen = datos_json.get("DATOS_GENERALES", {})

    tabla_gen = doc.add_table(rows=3, cols=4)
    tabla_gen.style = 'Table Grid'

    filas_gen = [
        [("Familia Profesional", datos_gen.get("FAMILIA", "N/E")),
         ("Denominación", datos_gen.get("DENOMINACION", "N/E"))],
        [("Módulo Formativo", form_data['modulo']),
         ("Código", datos_gen.get("CODIGO_MODULO", "N/E"))],
        [("Nivel", datos_gen.get("NIVEL", "N/E")),
         ("Horas Totales", datos_gen.get("HORAS", "N/E"))],
    ]
    for i, fila in enumerate(filas_gen):
        col_idx = 0
        for label, valor in fila:
            set_cell_text(tabla_gen.cell(i, col_idx), label, bold=True, color="E2E8F0")
            set_cell_text(tabla_gen.cell(i, col_idx + 1), valor)
            col_idx += 2
    fijar_anchos_columna(tabla_gen, [1.4, 2.3, 1.4, 2.3])
    doc.add_paragraph()

    # ── II. DATOS DE LA SESIÓN ──
    doc.add_heading("II. Datos de la Sesión de Planificación", level=2)
    tabla_sesion = doc.add_table(rows=3, cols=4)
    tabla_sesion.style = 'Table Grid'

    fechas_split = form_data['fechas'].split("-")
    f_inicio = fechas_split[0].replace("Inicio:", "").strip() if len(fechas_split) > 0 else form_data['fechas']
    f_final = fechas_split[1].replace("Final:", "").strip() if len(fechas_split) > 1 else ""

    filas_sesion = [
        [("Docente", form_data['docente']), ("Sesión", "Única")],
        [("Fecha de Inicio", f_inicio), ("Fecha Final", f_final)],
        [("Unidad de Competencia (UC)", form_data['uc_input'][:80]), ("", "")],
    ]
    for i, fila in enumerate(filas_sesion):
        col_idx = 0
        for label, valor in fila:
            if label:
                set_cell_text(tabla_sesion.cell(i, col_idx), label, bold=True, color="E2E8F0")
                set_cell_text(tabla_sesion.cell(i, col_idx + 1), valor)
            col_idx += 2
    fijar_anchos_columna(tabla_sesion, [1.6, 2.1, 1.6, 2.1])

    # RA
    p_ra = doc.add_paragraph()
    p_ra.add_run("\nResultado de Aprendizaje (RA): ").bold = True
    p_ra.add_run(form_data['ra'])
    doc.add_paragraph()

    # ── III. MATRIZ DE PLANIFICACIÓN ──
    doc.add_heading("III. Matriz de Planificación por Elementos de Capacidad", level=2)

    tabla_matriz = doc.add_table(rows=1, cols=6)
    tabla_matriz.style = 'Table Grid'
    encabezados = ["Elementos de Capacidad", "Nivel", "Fechas",
                   "Actividades de Enseñanza-Aprendizaje",
                   "Instrumento de Evaluación", "Contenidos"]
    hdr_cells = tabla_matriz.rows[0].cells
    for i, nombre in enumerate(encabezados):
        set_cell_text(hdr_cells[i], nombre, bold=True, center=True, color="2563EB", font_color=(255, 255, 255))

    anchos = [1.5, 0.5, 1.0, 2.5, 1.5, 1.5]
    for fila in datos_json.get("TABLA_MATRIZ", []):
        row_cells = tabla_matriz.add_row().cells
        nivel = str(fila.get("NIVEL", "")).strip()
        color_nivel, _ = NIVEL_COLORES.get(nivel, ("F1F5F9", nivel))

        set_cell_text(row_cells[0], str(fila.get("EC", "")), bold=True)
        set_cell_text(row_cells[1], nivel, center=True, color=color_nivel, bold=True)
        set_cell_text(row_cells[2], str(fila.get("FECHAS", "")), center=True)
        set_cell_text(row_cells[3], str(fila.get("ACTIVIDAD", "")))
        set_cell_text(row_cells[4], str(fila.get("INSTRUMENTO", "")))
        set_cell_text(row_cells[5], str(fila.get("CONTENIDOS", "")))

    fijar_anchos_columna(tabla_matriz, anchos)
    doc.add_paragraph()

    # ── IV. RESUMEN DE VALIDACIÓN ──
    doc.add_heading("IV. Resumen de Validación", level=2)
    tabla_val = doc.add_table(rows=1, cols=3)
    tabla_val.style = 'Table Grid'
    hdr_val = tabla_val.rows[0].cells
    for i, h in enumerate(["Concepto", "Valor", "Estado"]):
        set_cell_text(hdr_val[i], h, bold=True, center=True, color="E2E8F0")

    matriz = datos_json.get("TABLA_MATRIZ", [])
    total_act = len(matriz)
    niveles_usados = sorted(set(str(f.get("NIVEL", "")).strip() for f in matriz))
    instrumentos_unicos = len(set(str(f.get("INSTRUMENTO", "")).strip().lower() for f in matriz))

    filas_val = [
        ("Total de Actividades Diseñadas", str(total_act),
         "✅ Correcto" if total_act == form_data['cantidad_actividades'] else "⚠️ Verificar"),
        ("Niveles de Desempeño Cubiertos", ", ".join(niveles_usados),
         "✅ Correcto" if len(niveles_usados) >= 2 else "⚠️ Ampliar"),
        ("Instrumentos Únicos", str(instrumentos_unicos),
         "✅ Variado" if instrumentos_unicos >= 2 else "⚠️ Diversificar"),
    ]
    for concepto, valor, estado in filas_val:
        row_val = tabla_val.add_row().cells
        set_cell_text(row_val[0], concepto, bold=True)
        set_cell_text(row_val[1], valor, center=True)
        set_cell_text(row_val[2], estado, center=True, color="D1FAE5" if "✅" in estado else "FEF3C7")
        
    fijar_anchos_columna(tabla_val, [2.5, 2.0, 2.0])
    doc.add_paragraph()

    # ── V. OBSERVACIONES DEL COORDINADOR ──
    doc.add_heading("V. Observaciones del Coordinador", level=2)
    for _ in range(3):
        doc.add_paragraph("_" * 90)

    # ── FIRMAS (3 columnas) ──
    doc.add_paragraph("\n")
    t_firmas = doc.add_table(rows=2, cols=3)
    t_firmas.cell(0, 0).text = "__________________________"
    t_firmas.cell(0, 1).text = "__________________________"
    t_firmas.cell(0, 2).text = "__________________________"
    t_firmas.cell(1, 0).text = "Director/a de Centro"
    t_firmas.cell(1, 1).text = "Coordinador/a Módulos Formativos ETP"
    t_firmas.cell(1, 2).text = "Docente"
    for row in t_firmas.rows:
        for cell in row.cells:
            cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER

    buffer = BytesIO()
    doc.save(buffer)
    buffer.seek(0)
    return buffer


# ═══════════════════════════════════════════════════════════════════════════
# FORMULARIO
# ═══════════════════════════════════════════════════════════════════════════
with st.form("form_planificacion", clear_on_submit=False):
    st.markdown('<div class="planra-section-title">📄 1. Fuente de Conocimiento Curricular</div>', unsafe_allow_html=True)
    archivo_pdf = st.file_uploader(
        "Cargue el documento PDF oficial del diseño curricular",
        type=["pdf"],
        help="RECOMENDACIÓN: Sube solo las páginas del módulo a trabajar.",
    )

    st.markdown('<div class="planra-section-title">🏛️ 2. Arquitectura Institucional</div>', unsafe_allow_html=True)
    col_inst1, col_inst2 = st.columns(2)
    with col_inst1:
        politecnico = st.text_input("Nombre del Politécnico", value="Politécnico Salesiano Arquides Calderón")
        docente = st.text_input("Nombre del Docente", value="Ing. Bernardo Antonio Hernández Batista")
    with col_inst2:
        eslogan = st.text_input("Eslogan del Politécnico", placeholder="Ingrese el eslogan institucional")
        coordinador = st.text_input("Coordinador Módulos Formativos", placeholder="Nombre del coordinador")

    st.markdown('<div class="planra-section-title">📝 3. Parámetros de Operación</div>', unsafe_allow_html=True)
    col1, col2 = st.columns(2)
    with col1:
        modulo = st.text_input("Módulo Formativo", placeholder="Ej: MF_358_3 Impuestos al Consumo...")
        fechas = st.text_input("Fechas estimadas", placeholder="Ej: Inicio: 12/11/2026 - Final: 18/12/2026")
    with col2:
        cantidad_ec = st.number_input("Cantidad EXACTA de Elementos de Capacidad (EC)", min_value=1, value=3)
        cantidad_actividades = st.number_input("Cantidad EXACTA de Actividades a diseñar", min_value=1, value=6)

    uc_input = st.text_area("🔗 Unidad de Competencia (UC)", height=80,
                            placeholder="Pega aquí la Unidad de Competencia asociada...")
    ra = st.text_area("🎯 Resultado de Aprendizaje (RA)", height=80,
                      placeholder="Pega aquí el RA completo a planificar...")

    max_tokens, temperature = ia.control_avanzado(default_tokens=8192, tope=16384, default_temp=0.1)

    st.markdown("<br>", unsafe_allow_html=True)
    
    # FIX: Eliminado width="stretch". Se utiliza el parámetro nativo correcto de Streamlit: use_container_width=True
    submit_button = st.form_submit_button("⚙️ Iniciar Compilación de Matriz Oficial", type="primary", use_container_width=True)


# ═══════════════════════════════════════════════════════════════════════════
# LÓGICA CORE
# ═══════════════════════════════════════════════════════════════════════════
if submit_button:
    cfg = ia.config_ia()
    if not cfg["api_key"]:
        st.error("🔒 Debes ingresar tu API Key en la página de Inicio (barra lateral).")
    elif not archivo_pdf or not modulo or not ra or not uc_input or not politecnico or not docente:
        st.warning("📝 Parámetros Incompletos: Carga el PDF y completa todos los datos institucionales y curriculares.")
    else:
        with st.spinner(f'🧠 Ejecutando análisis curricular experto con {cfg["modelo"]}...'):
            texto_crudo = None
            try:
                # 1. Extracción del texto del PDF
                texto_curriculo, total_paginas = extraer_texto_pdf(archivo_pdf)

                with st.expander(f"📖 Preview del PDF ({total_paginas} págs · {len(texto_curriculo):,} chars)"):
                    st.text(texto_curriculo[:3000])
                    if len(texto_curriculo) > 3000:
                        st.caption(f"... y {len(texto_curriculo) - 3000:,} caracteres más")

                # 2. PROMPT MAESTRO (ADAPTADO A GUÍA MINERD)
                MARKER_NL = getattr(ia, "MARKER_NL", "\\n")
                prompt_maestro = f"""Actúa como un Especialista Curricular de Alto Nivel de la ETP (MINERD), experto en la "Guía de Planificación & Evaluación por Competencias".
He extraído el texto del Diseño Curricular oficial. Debes buscar el Módulo Formativo (MF) y el Resultado de Aprendizaje (RA), extraer los Criterios de Evaluación (CE) y Contenidos, y diseñar la matriz de planificación modular garantizando la alineación constructiva.

INSUMOS:
Módulo Formativo: {modulo}
Fechas estimadas: {fechas}
Cantidad de EC a crear: {cantidad_ec}
Cantidad de Actividades a diseñar: {cantidad_actividades}

REGLAS ESTRICTAS DE DISEÑO BASADO EN COMPETENCIAS (MINERD):
1. EXTRACCIÓN AUTÓNOMA: Localiza en el PDF la Familia Profesional, Denominación, Nivel y Horas totales.
2. ALINEACIÓN CONSTRUCTIVA: Asegura la coherencia absoluta entre el Resultado de Aprendizaje (RA), los Elementos de Capacidad (EC), las Actividades y los Instrumentos de Evaluación.
3. NIVELES DE DESEMPEÑO: Las actividades DEBEN evidenciar una progresión clara: Nivel 1 (Conocimiento/Comprensión), Nivel 2 (Aplicación/Análisis) y Nivel 3 (Dominio/Autonomía/Resolución de problemas).
4. ACTIVIDADES SIGNIFICATIVAS Y PRÁCTICAS: Diseña secuencias didácticas basadas en el "saber hacer" en contextos reales. Queda terminantemente prohibido proponer presentaciones de diapositivas estáticas. Prioriza el uso de simuladores interactivos online, configuración práctica, programación y resolución de casos reales del sector socioproductivo.
5. EVALUACIÓN BASADA EN EVIDENCIAS: Evita los exámenes teóricos tradicionales como único medio. Varía los instrumentos objetivos de evaluación por competencias (Rúbricas de desempeño, Listas de cotejo, Portafolios de evidencias, Guías de observación directa).
6. TEXTO PLANO: No utilices formato Markdown. Si necesitas salto de línea en un texto, usa {MARKER_NL}.

FORMATO DE SALIDA ESTRICTO (JSON NATIVO):
{{
  "DATOS_GENERALES": {{
    "FAMILIA": "[Extracción del PDF]",
    "DENOMINACION": "[Extracción del PDF]",
    "NIVEL": "[Extracción del PDF]",
    "CODIGO_MODULO": "[Extracción del PDF]",
    "HORAS": "[Extracción del PDF]"
  }},
  "TABLA_MATRIZ": [
    {{
      "EC": "[Redacción clara del Elemento de Capacidad]",
      "NIVEL": "[1, 2 o 3]",
      "FECHAS": "{fechas}",
      "ACTIVIDAD": "[Descripción técnica y significativa de la actividad práctica]",
      "INSTRUMENTO": "[Instrumento basado en evidencias pertinente a la actividad]",
      "CONTENIDOS": "[Contenidos curriculares específicos abordados]"
    }}
  ]
}}

El arreglo "TABLA_MATRIZ" debe contener exactamente {cantidad_actividades} objetos.

DOCUMENTO CURRICULAR A ANALIZAR:
{texto_curriculo}
"""
                # 3. Petición a la IA vía core.ia
                texto_crudo, flags = ia.solicitar_ia(
                    prompt_maestro, modo="json", max_tokens=max_tokens,
                    temperature=temperature, modulo="planificacionra",
                )

                # 4. PARSEO JSON ROBUSTO
                datos_json = ia.parsear_json_robusto(texto_crudo)
                datos_json = ia.decodificar_marcadores(datos_json)

                if not datos_json.get("TABLA_MATRIZ", []):
                    st.error("❌ El JSON se procesó pero la tabla de actividades está vacía.")
                    st.stop()

                # 5. PREPARAR DATOS DEL FORMULARIO
                datos_formulario = {
                    "politecnico": politecnico,
                    "eslogan": eslogan,
                    "modulo": modulo,
                    "docente": docente,
                    "fechas": fechas,
                    "uc_input": uc_input,
                    "ra": ra,
                    "cantidad_actividades": int(cantidad_actividades),
                }

                # 6. MÉTRICAS POST-GENERACIÓN
                matriz = datos_json.get("TABLA_MATRIZ", [])
                niveles_usados = sorted(set(str(f.get("NIVEL", "")).strip() for f in matriz))
                instrumentos_unicos = len(set(str(f.get("INSTRUMENTO", "")).strip().lower() for f in matriz))

                k1, k2, k3, k4 = st.columns(4)
                with k1:
                    st.markdown(f'<div class="planra-kpi"><div class="planra-kpi-icon">🎯</div><div class="planra-kpi-value">{len(matriz)}</div><div class="planra-kpi-label">Actividades</div></div>', unsafe_allow_html=True)
                with k2:
                    st.markdown(f'<div class="planra-kpi"><div class="planra-kpi-icon">📊</div><div class="planra-kpi-value">{len(niveles_usados)}</div><div class="planra-kpi-label">Niveles Cubiertos</div></div>', unsafe_allow_html=True)
                with k3:
                    st.markdown(f'<div class="planra-kpi"><div class="planra-kpi-icon">📋</div><div class="planra-kpi-value">{instrumentos_unicos}</div><div class="planra-kpi-label">Instrumentos Únicos</div></div>', unsafe_allow_html=True)
                with k4:
                    st.markdown(f'<div class="planra-kpi"><div class="planra-kpi-icon">📄</div><div class="planra-kpi-value">{int(cantidad_ec)}</div><div class="planra-kpi-label">Elementos Capacidad</div></div>', unsafe_allow_html=True)

                # 7. PREVIEW DE LA MATRIZ GENERADA
                st.markdown('<div class="planra-section-title">👁️ Preview de la Matriz Generada</div>', unsafe_allow_html=True)
                for i, fila in enumerate(matriz):
                    nivel = str(fila.get("NIVEL", "")).strip()
                    color_nivel = {"1": "#DBEAFE", "2": "#D1FAE5", "3": "#EDE9FE"}.get(nivel, "#F1F5F9")
                    st.markdown(f"""
                    <div class="planra-preview-box">
                        <div class="planra-preview-ec">{fila.get("EC", "")}
                            <span class="planra-preview-nivel" style="background:{color_nivel};">Nivel {nivel}</span>
                        </div>
                        <div class="planra-preview-act"><b>Actividad:</b> {fila.get("ACTIVIDAD", "")}</div>
                        <div class="planra-preview-act"><b>Instrumento:</b> {fila.get("INSTRUMENTO", "")}</div>
                        <div class="planra-preview-act"><b>Contenidos:</b> {fila.get("CONTENIDOS", "")[:150]}{'...' if len(str(fila.get('CONTENIDOS', ''))) > 150 else ''}</div>
                    </div>
                    """, unsafe_allow_html=True)

                # 8. GENERACIÓN DEL ARCHIVO WORD
                buffer_docx = generar_documento_word(datos_json, datos_formulario)

                st.success(f"✅ ¡Matriz Curricular generada con éxito! ({len(matriz)} actividades)")
                
                if flags.get("reintento"):
                    st.info("ℹ️ La primera llamada se cortó por tokens; el reintento automático tuvo éxito.")
                    
                # FIX: Eliminado width="stretch" que lanzaba TypeError. Se utiliza use_container_width=True
                st.download_button(
                    label="📥 Descargar Matriz Oficial de Planificación (.docx)",
                    data=buffer_docx,
                    file_name=f"Matriz_Planificacion_RA_{sanear_nombre(modulo)}.docx",
                    mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                    type="primary",
                    use_container_width=True,
                )
            except Exception as e:
                # Capturamos todas las excepciones juntas para limpiar el código
                if hasattr(ia, "render_error_ia"):
                    ia.render_error_ia(e, texto_crudo)
                else:
                    st.error(f"Error procesando la solicitud: {str(e)}")