"""
auditor_matriz.py — Auditor de Matriz Curricular ETP (v2 · REGENERADO)
Auditoría de calidad curricular en 6 dimensiones con puntuación 0-100,
distribución Bloom, alertas críticas y plan de mejora.
• IA vía core/ia (solicitar_json, marcadores, reintento, auditoría).
• Entrada flexible: PDF, Excel o texto pegado.
• Super interfaz: hero, stepper, tarjetas de puntuación, visualización Bloom.
• Word profesional de auditoría con dashboard y firmas.
PAQUETE APLICADO:
• FIX CRÍTICO: add_page_number sin Run.add_run (AttributeError).
• DEPRECACIONES: use_container_width → width="stretch".
• 1) Historial persistente SQLite (auditorias_matriz) + pestaña Historial.
• 2) Gráfico de barras por dimensión en pantalla.
• 3) Auto-relleno de docente desde usuario_display_nombre.
• 4) Export CSV de dimensiones y plan de mejora.
• 5) Guardado del último puntaje por módulo (lee la Sala de Situación).
"""
import re
from datetime import datetime, date
from io import BytesIO
import sqlite3

import pandas as pd
import streamlit as st
from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement, parse_xml
from docx.oxml.ns import qn, nsdecls
from docx.shared import Inches, Pt, RGBColor

try:
    from pypdf import PdfReader
except Exception:
    try:
        from PyPDF2 import PdfReader
    except Exception:
        PdfReader = None

from core import ia

# ═══════════════════════════════════════════════════════════════════════════
# PERSISTENCIA SQLite (mejoras 1 y 5)
# ═══════════════════════════════════════════════════════════════════════════
DB_NAME = "gestion_etp.db"


def _conn():
    return sqlite3.connect(DB_NAME, check_same_thread=False)


def asegurar_tabla_auditorias():
    conn = _conn()
    conn.execute('''CREATE TABLE IF NOT EXISTS auditorias_matriz (
        id INTEGER PRIMARY KEY AUTOINCREMENT,
        fecha TEXT, modulo TEXT, tipo_matriz TEXT, docente TEXT, evaluador TEXT,
        puntuacion INTEGER, nivel TEXT, completitud INTEGER
    )''')
    conn.commit()
    conn.close()


def insertar_auditoria(fecha, modulo, tipo_matriz, docente, evaluador,
                       puntuacion, nivel, completitud):
    conn = _conn()
    conn.execute('''INSERT INTO auditorias_matriz
        (fecha, modulo, tipo_matriz, docente, evaluador, puntuacion, nivel, completitud)
        VALUES (?,?,?,?,?,?,?,?)''',
        (fecha, modulo, tipo_matriz, docente, evaluador, puntuacion, nivel, completitud))
    conn.commit()
    conn.close()


def listar_auditorias():
    conn = _conn()
    cur = conn.cursor()
    cur.execute("SELECT id, fecha, modulo, tipo_matriz, docente, evaluador, "
                "puntuacion, nivel, completitud FROM auditorias_matriz ORDER BY id DESC")
    cols = ["id", "fecha", "modulo", "tipo_matriz", "docente", "evaluador",
            "puntuacion", "nivel", "completitud"]
    rows = [dict(zip(cols, r)) for r in cur.fetchall()]
    conn.close()
    return rows


asegurar_tabla_auditorias()

# ═══════════════════════════════════════════════════════════════════════════
# DIMENSIONES DE AUDITORÍA
# ═══════════════════════════════════════════════════════════════════════════
DIMENSIONES_AUDITORIA = {
    "COMPLETITUD": {
        "nombre": "Completitud Curricular", "icono": "📋",
        "descripcion": "Verifica que todos los elementos requeridos estén presentes: RA, CE, EC, contenidos, actividades e instrumentos.",
        "color": "#3B82F6",
    },
    "ALINEACION": {
        "nombre": "Alineación RA-CE-EC", "icono": "🎯",
        "descripcion": "Evalúa la coherencia entre Resultados de Aprendizaje, Criterios de Evaluación y Elementos de Capacidad.",
        "color": "#8B5CF6",
    },
    "BLOOM": {
        "nombre": "Distribución Cognitiva Bloom", "icono": "🧠",
        "descripcion": "Analiza el equilibrio de niveles cognitivos según la Taxonomía de Bloom.",
        "color": "#EC4899",
    },
    "EVALUACION": {
        "nombre": "Coherencia Evaluativa", "icono": "📊",
        "descripcion": "Verifica que los instrumentos de evaluación sean pertinentes y estén alineados a los RA.",
        "color": "#10B981",
    },
    "TEMPORAL": {
        "nombre": "Dosificación Temporal", "icono": "⏱️",
        "descripcion": "Evalúa si la distribución de tiempo es realista y proporcional a la complejidad.",
        "color": "#F59E0B",
    },
    "PERTINENCIA": {
        "nombre": "Pertinencia Técnica", "icono": "🔧",
        "descripcion": "Verifica que los contenidos sean relevantes para el campo técnico-profesional ETP.",
        "color": "#EF4444",
    },
}

# ═══════════════════════════════════════════════════════════════════════════
# SUPER INTERFAZ — ESTILOS
# ═══════════════════════════════════════════════════════════════════════════
st.markdown("""
<style>
@import url('https://fonts.googleapis.com/css2?family=Inter:wght@300;400;500;600;700;800;900&display=swap');
html, body, [class*="css"] { font-family: 'Inter', sans-serif; background-color: #F0F4F8; color: #1E293B; }
.audit-hero { background: linear-gradient(135deg, #0F172A 0%, #065F46 40%, #059669 70%, #34D399 100%);
color: #fff; padding: 2.2rem; border-radius: 20px; margin-bottom: 1.5rem;
box-shadow: 0 25px 50px rgba(6,95,70,0.3); position: relative; overflow: hidden; }
.audit-hero::before { content: '🧾'; position: absolute; right: 2rem; top: 50%; transform: translateY(-50%);
font-size: 6rem; opacity: 0.15; }
.audit-hero-title { font-size: 2.4rem; font-weight: 900; letter-spacing: -0.03em; margin-bottom: 0.4rem; }
.audit-hero-sub { font-size: 1.05rem; opacity: 0.9; line-height: 1.5; }
.audit-hero-badge { display: inline-block; background: rgba(255,255,255,0.15); border: 1px solid rgba(255,255,255,0.25);
border-radius: 8px; padding: 4px 12px; font-size: 0.8rem; font-weight: 600; margin-top: 0.8rem; margin-right: 8px; }
.audit-stepper { display: flex; align-items: center; justify-content: center; gap: 0; margin: 1.5rem 0; }
.audit-step { display: flex; flex-direction: column; align-items: center; gap: 6px; }
.audit-step-circle { width: 48px; height: 48px; border-radius: 50%; display: flex; align-items: center;
justify-content: center; font-weight: 800; font-size: 1.2rem; transition: all 0.3s ease; }
.audit-step-circle.inactive { background: #E2E8F0; color: #94A3B8; border: 3px solid #CBD5E1; }
.audit-step-circle.active { background: linear-gradient(135deg, #059669, #34D399); color: #fff;
border: 3px solid #059669; box-shadow: 0 4px 15px rgba(5,150,105,0.4); }
.audit-step-circle.done { background: linear-gradient(135deg, #0284C7, #38BDF8); color: #fff;
border: 3px solid #0284C7; box-shadow: 0 4px 15px rgba(2,132,199,0.3); }
.audit-step-label { font-size: 0.72rem; font-weight: 700; color: #64748B; text-transform: uppercase; letter-spacing: 0.03em; }
.audit-step-line { width: 70px; height: 3px; background: #CBD5E1; margin: 0 4px; margin-bottom: 22px; }
.audit-step-line.done { background: linear-gradient(90deg, #0284C7, #38BDF8); }
.dim-card { background: #fff; border: 2px solid #E2E8F0; border-radius: 14px; padding: 1.2rem;
transition: all 0.25s ease; height: 100%; box-shadow: 0 2px 8px rgba(0,0,0,0.04); }
.dim-card:hover { transform: translateY(-3px); box-shadow: 0 8px 20px rgba(0,0,0,0.1); }
.dim-header { display: flex; align-items: center; gap: 10px; margin-bottom: 10px; }
.dim-icono { font-size: 1.8rem; }
.dim-nombre { font-weight: 700; font-size: 0.95rem; color: #0F172A; }
.dim-score-container { margin: 12px 0; }
.dim-score-bar { height: 10px; border-radius: 5px; background: #E2E8F0; overflow: hidden; }
.dim-score-fill { height: 100%; border-radius: 5px; transition: width 0.8s ease; }
.dim-score-value { font-size: 1.6rem; font-weight: 800; margin-top: 6px; }
.dim-nivel { display: inline-block; padding: 3px 10px; border-radius: 6px; font-size: 0.72rem;
font-weight: 700; text-transform: uppercase; letter-spacing: 0.03em; }
.score-global-card { background: linear-gradient(135deg, #0F172A, #1E293B); border-radius: 20px;
padding: 2rem; text-align: center; color: #fff; box-shadow: 0 20px 40px rgba(15,23,42,0.3); }
.score-global-value { font-size: 4rem; font-weight: 900; line-height: 1; }
.score-global-label { font-size: 0.9rem; opacity: 0.7; margin-top: 8px; text-transform: uppercase; letter-spacing: 0.05em; }
.bloom-bar-container { display: flex; align-items: center; gap: 10px; margin-bottom: 8px; }
.bloom-bar-label { width: 100px; font-size: 0.8rem; font-weight: 600; color: #475569; text-align: right; }
.bloom-bar-track { flex: 1; height: 24px; background: #E2E8F0; border-radius: 6px; overflow: hidden; }
.bloom-bar-fill { height: 100%; border-radius: 6px; display: flex; align-items: center; padding-left: 8px;
font-size: 0.72rem; font-weight: 700; color: #fff; transition: width 0.8s ease; }
.audit-section-title { color: #065F46; font-weight: 700; font-size: 1.12rem; border-bottom: 2px solid #D1FAE5;
padding-bottom: 8px; margin: 1.2rem 0 0.9rem 0; }
.audit-alert-critical { background: #FEF2F2; border-left: 5px solid #EF4444; border-radius: 8px;
padding: 12px 16px; margin-bottom: 8px; }
.audit-alert-text { color: #991B1B; font-size: 0.88rem; font-weight: 500; }
.audit-strength { background: #F0FDF4; border-left: 5px solid #22C55E; border-radius: 8px;
padding: 12px 16px; margin-bottom: 8px; }
.audit-strength-text { color: #166534; font-size: 0.88rem; font-weight: 500; }
</style>
""", unsafe_allow_html=True)

# ═══════════════════════════════════════════════════════════════════════════
# ESTADO DE SESIÓN
# ═══════════════════════════════════════════════════════════════════════════
def init_estado():
    if "auditoria_resultado" not in st.session_state:
        st.session_state.auditoria_resultado = None
    if "auditoria_fase" not in st.session_state:
        st.session_state.auditoria_fase = 0  # 0=config, 2=resultados


init_estado()

# ═══════════════════════════════════════════════════════════════════════════
# UTILIDADES DE EXTRACCIÓN
# ═══════════════════════════════════════════════════════════════════════════
def extraer_texto_pdf(archivo, max_caracteres=80000):
    if PdfReader is None:
        raise RuntimeError("No hay librería PDF. Instala pypdf o PyPDF2.")
    archivo.seek(0)
    reader = PdfReader(archivo)
    texto = "".join((p.extract_text() or "") for p in reader.pages)
    texto = re.sub(r"\s+", " ", texto).strip()
    return texto[:max_caracteres]


def extraer_texto_excel(archivo, max_caracteres=80000):
    archivo.seek(0)
    try:
        df = pd.read_excel(archivo, sheet_name=None)
        textos = []
        for nombre_hoja, hoja_df in df.items():
            textos.append(f"--- Hoja: {nombre_hoja} ---")
            textos.append(hoja_df.to_string(index=False))
        return "\n".join(textos)[:max_caracteres]
    except Exception as e:
        raise RuntimeError(f"Error leyendo Excel: {e}")


# ═══════════════════════════════════════════════════════════════════════════
# UTILIDADES WORD (FIX CRÍTICO en add_page_number)
# ═══════════════════════════════════════════════════════════════════════════
def shade_cell(cell, color):
    shd = parse_xml(r'<w:shd {} w:fill="{}"/>'.format(nsdecls("w"), color))
    cell._tc.get_or_add_tcPr().append(shd)


def set_cell_text(cell, text, bold=False, center=False, color=None):
    cell.text = ""
    p = cell.paragraphs[0]
    run = p.add_run(str(text))
    run.bold = bold
    if center:
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    if color:
        shade_cell(cell, color)


def add_page_number(doc):
    """Agrega 'Página X de Y' al pie de CADA sección.
    FIX: add_run() se invoca sobre el Paragraph (p), nunca sobre un Run."""
    for section in doc.sections:
        footer = section.footer
        footer.is_linked_to_previous = False
        p = footer.paragraphs[0] if footer.paragraphs else footer.add_paragraph()
        p.text = ""
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER

        def _campo(parrafo, instruccion):
            r = parrafo.add_run()
            i = OxmlElement("w:fldChar"); i.set(qn("w:fldCharType"), "begin")
            t = OxmlElement("w:instrText"); t.set(qn("xml:space"), "preserve"); t.text = instruccion
            f = OxmlElement("w:fldChar"); f.set(qn("w:fldCharType"), "end")
            r._r.append(i); r._r.append(t); r._r.append(f)
            return r

        r1 = _campo(p, "PAGE")
        r_sep = p.add_run(" / ")   # ← FIX: era run.add_run(...) → AttributeError
        r2 = _campo(p, "NUMPAGES")
        for r in (r1, r_sep, r2):
            r.font.size = Pt(8)
            r.font.color.rgb = RGBColor(0x64, 0x74, 0x8B)


def nivel_color(nivel):
    niveles = {
        "Excelente": ("D1FAE5", "065F46"),
        "Bueno": ("DBEAFE", "1E40AF"),
        "Aceptable": ("FEF3C7", "92400E"),
        "Deficiente": ("FEE2E2", "991B1B"),
    }
    return niveles.get(nivel, ("F1F5F9", "475569"))


def nivel_css(nivel):
    niveles = {
        "Excelente": ("#D1FAE5", "#065F46"),
        "Bueno": ("#DBEAFE", "#1E40AF"),
        "Aceptable": ("#FEF3C7", "#92400E"),
        "Deficiente": ("#FEE2E2", "#991B1B"),
    }
    return niveles.get(nivel, ("#F1F5F9", "#475569"))


# ═══════════════════════════════════════════════════════════════════════════
# PROMPT DE AUDITORÍA
# ═══════════════════════════════════════════════════════════════════════════
def construir_prompt_auditoria(matriz_texto, modulo, tipo_matriz, enfoque):
    dimensiones_txt = "\n".join([
        f"- {clave}: {d['nombre']} — {d['descripcion']}"
        for clave, d in DIMENSIONES_AUDITORIA.items()
    ])
    return f"""Actúa como un Auditor Curricular Senior del MINERD, Especialista en Educación Técnico Profesional (ETP) y Experto en Diseño Curricular por Competencias.
Tu tarea es auditar la siguiente matriz curricular y emitir un informe de calidad profesional.

MATRIZ CURRICULAR A AUDITAR:
Tipo: {tipo_matriz}
Módulo/Asignatura: {modulo}
Enfoque de auditoría: {enfoque}

CONTENIDO DE LA MATRIZ:
{matriz_texto}

DIMENSIONES DE AUDITORÍA (puntúa cada una de 0 a 100):
{dimensiones_txt}

REGLAS DE AUDITORÍA:
- Sé riguroso y específico. Cada hallazgo debe ser concreto y accionable.
- La puntuación debe reflejar la calidad real, no ser inflada.
- Identifica elementos faltantes, inconsistencias y áreas de mejora.
- Detecta la distribución de niveles cognitivos Bloom en las actividades/evaluaciones.
- Señala alertas críticas que requieran atención inmediata.
- Destaca las fortalezas genuinas de la matriz.
- El plan de mejora debe ser priorizado y realista.

NIVELES DE PUNTUACIÓN:
90-100: Excelente (cumple plenamente)
75-89: Bueno (cumple con observaciones menores)
60-74: Aceptable (requiere mejoras significativas)
0-59: Deficiente (requiere reestructuración)

Devuelve ÚNICAMENTE JSON válido con este formato exacto:
{{
  "RESUMEN_EJECUTIVO": "Resumen de 2-3 párrafos del estado general de la matriz...",
  "PUNTUACION_GLOBAL": 85,
  "NIVEL_GLOBAL": "Bueno",
  "DIMENSIONES": [
    {{
      "CLAVE": "COMPLETITUD",
      "PUNTUACION": 90,
      "NIVEL": "Excelente",
      "HALLAZGOS": ["Hallazgo específico 1", "Hallazgo específico 2"],
      "RECOMENDACIONES": ["Recomendación accionable 1", "Recomendación accionable 2"]
    }}
  ],
  "MATRIZ_ANALIZADA": {{
    "RA_DETECTADOS": 5, "CE_DETECTADOS": 12, "EC_DETECTADOS": 8,
    "CONTENIDOS_DETECTADOS": 25, "ACTIVIDADES_DETECTADAS": 15,
    "INSTRUMENTOS_DETECTADOS": 10, "SEMANAS_DETECTADAS": 38
  }},
  "DISTRIBUCION_BLOOM": {{
    "Recordar": 15, "Comprender": 25, "Aplicar": 30,
    "Analizar": 15, "Evaluar": 10, "Crear": 5
  }},
  "ALERTAS_CRITICAS": ["Alerta crítica 1", "Alerta crítica 2"],
  "FORTALEZAS": ["Fortaleza 1", "Fortaleza 2", "Fortaleza 3"],
  "PLAN_MEJORA": [
    {{ "PRIORIDAD": "Alta", "ACCION": "Acción específica a realizar", "PLAZO": "Corto plazo (1-2 semanas)" }}
  ]
}}
"""


# ═══════════════════════════════════════════════════════════════════════════
# GENERACIÓN WORD PROFESIONAL
# ═══════════════════════════════════════════════════════════════════════════
def build_audit_docx(resultado, meta):
    doc = Document()
    doc.styles["Normal"].font.name = "Calibri"
    doc.styles["Normal"].font.size = Pt(11)
    for section in doc.sections:
        section.left_margin = Inches(0.8)
        section.right_margin = Inches(0.8)

    # ─── PORTADA ───
    for _ in range(3):
        doc.add_paragraph()
    p_inst = doc.add_paragraph()
    p_inst.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run_inst = p_inst.add_run(meta.get("institucion", ""))
    run_inst.bold = True
    run_inst.font.size = Pt(14)
    run_inst.font.color.rgb = RGBColor(0x06, 0x5F, 0x46)
    doc.add_paragraph()
    p_titulo = doc.add_paragraph()
    p_titulo.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run_t = p_titulo.add_run("INFORME DE AUDITORÍA DE MATRIZ CURRICULAR")
    run_t.bold = True
    run_t.font.size = Pt(22)
    run_t.font.color.rgb = RGBColor(0x06, 0x5F, 0x46)
    p_sub = doc.add_paragraph()
    p_sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run_sub = p_sub.add_run(f"Módulo: {meta.get('modulo', '')}")
    run_sub.font.size = Pt(14)
    run_sub.font.color.rgb = RGBColor(0x64, 0x74, 0x8B)
    doc.add_paragraph()
    p_meta = doc.add_paragraph()
    p_meta.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p_meta.add_run(f"Auditor: {meta.get('docente', '')}").font.size = Pt(12)
    p_fecha = doc.add_paragraph()
    p_fecha.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run_fecha = p_fecha.add_run(f"Fecha: {datetime.now().strftime('%d/%m/%Y')}")
    run_fecha.font.size = Pt(11)
    run_fecha.font.color.rgb = RGBColor(0x94, 0xA3, 0xB8)
    doc.add_page_break()

    # ─── RESUMEN EJECUTIVO ───
    doc.add_heading("1. Resumen Ejecutivo", level=1)
    p_score = doc.add_paragraph()
    p_score.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run_score = p_score.add_run(f"Puntuación Global: {resultado.get('PUNTUACION_GLOBAL', 0)}/100")
    run_score.bold = True
    run_score.font.size = Pt(18)
    p_nivel = doc.add_paragraph()
    p_nivel.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run_nivel = p_nivel.add_run(f"Nivel: {resultado.get('NIVEL_GLOBAL', 'N/A')}")
    run_nivel.bold = True
    run_nivel.font.size = Pt(14)
    doc.add_paragraph()
    p_resumen = doc.add_paragraph(resultado.get("RESUMEN_EJECUTIVO", ""))
    p_resumen.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    doc.add_page_break()

    # ─── DASHBOARD DE DIMENSIONES ───
    doc.add_heading("2. Dashboard de Dimensiones", level=1)
    tabla_dim = doc.add_table(rows=1, cols=4)
    tabla_dim.style = "Table Grid"
    hdr = tabla_dim.rows[0].cells
    set_cell_text(hdr[0], "Dimensión", bold=True, color="D1FAE5")
    set_cell_text(hdr[1], "Puntuación", bold=True, center=True, color="D1FAE5")
    set_cell_text(hdr[2], "Nivel", bold=True, center=True, color="D1FAE5")
    set_cell_text(hdr[3], "Hallazgos Clave", bold=True, color="D1FAE5")
    for dim in resultado.get("DIMENSIONES", []):
        row = tabla_dim.add_row().cells
        clave = dim.get("CLAVE", "")
        dim_info = DIMENSIONES_AUDITORIA.get(clave, {"nombre": clave, "icono": "📋"})
        set_cell_text(row[0], f"{dim_info.get('icono', '')} {dim_info.get('nombre', clave)}", bold=True)
        set_cell_text(row[1], str(dim.get("PUNTUACION", 0)), center=True)
        nivel = dim.get("NIVEL", "N/A")
        bg_color, _ = nivel_color(nivel)
        set_cell_text(row[2], nivel, center=True, color=bg_color)
        hallazgos = dim.get("HALLAZGOS", [])
        set_cell_text(row[3], hallazgos[0] if hallazgos else "—")
    doc.add_page_break()

    # ─── MATRIZ ANALIZADA ───
    doc.add_heading("3. Elementos Detectados en la Matriz", level=1)
    matriz_info = resultado.get("MATRIZ_ANALIZADA", {})
    tabla_elem = doc.add_table(rows=1, cols=2)
    tabla_elem.style = "Table Grid"
    hdr_e = tabla_elem.rows[0].cells
    set_cell_text(hdr_e[0], "Elemento", bold=True, color="DBEAFE")
    set_cell_text(hdr_e[1], "Cantidad Detectada", bold=True, center=True, color="DBEAFE")
    elementos = [
        ("Resultados de Aprendizaje (RA)", matriz_info.get("RA_DETECTADOS", 0)),
        ("Criterios de Evaluación (CE)", matriz_info.get("CE_DETECTADOS", 0)),
        ("Elementos de Capacidad (EC)", matriz_info.get("EC_DETECTADOS", 0)),
        ("Contenidos", matriz_info.get("CONTENIDOS_DETECTADOS", 0)),
        ("Actividades", matriz_info.get("ACTIVIDADES_DETECTADAS", 0)),
        ("Instrumentos de Evaluación", matriz_info.get("INSTRUMENTOS_DETECTADOS", 0)),
        ("Semanas/Períodos", matriz_info.get("SEMANAS_DETECTADAS", 0)),
    ]
    for nombre, cantidad in elementos:
        row = tabla_elem.add_row().cells
        set_cell_text(row[0], nombre)
        set_cell_text(row[1], str(cantidad), center=True)
    doc.add_paragraph()

    # ─── DISTRIBUCIÓN BLOOM ───
    doc.add_heading("4. Distribución Cognitiva Bloom", level=1)
    bloom = resultado.get("DISTRIBUCION_BLOOM", {})
    tabla_bloom = doc.add_table(rows=1, cols=3)
    tabla_bloom.style = "Table Grid"
    hdr_b = tabla_bloom.rows[0].cells
    set_cell_text(hdr_b[0], "Nivel Cognitivo", bold=True, color="F3E8FF")
    set_cell_text(hdr_b[1], "Porcentaje", bold=True, center=True, color="F3E8FF")
    set_cell_text(hdr_b[2], "Distribución", bold=True, color="F3E8FF")
    bloom_colors = {"Recordar": "FEE2E2", "Comprender": "FEF3C7", "Aplicar": "D1FAE5",
                    "Analizar": "DBEAFE", "Evaluar": "E9D5FF", "Crear": "FCE7F3"}
    for nivel_bloom, pct in bloom.items():
        row = tabla_bloom.add_row().cells
        set_cell_text(row[0], nivel_bloom, bold=True, color=bloom_colors.get(nivel_bloom, "F1F5F9"))
        set_cell_text(row[1], f"{pct}%", center=True)
        barra = "█" * int(pct / 5) + "░" * max(0, 20 - int(pct / 5))
        set_cell_text(row[2], barra)
    doc.add_page_break()

    # ─── HALLAZGOS DETALLADOS ───
    doc.add_heading("5. Hallazgos y Recomendaciones por Dimensión", level=1)
    for dim in resultado.get("DIMENSIONES", []):
        clave = dim.get("CLAVE", "")
        dim_info = DIMENSIONES_AUDITORIA.get(clave, {"nombre": clave, "icono": "📋"})
        doc.add_heading(
            f"{dim_info.get('icono', '')} {dim_info.get('nombre', clave)} — "
            f"{dim.get('PUNTUACION', 0)}/100 ({dim.get('NIVEL', 'N/A')})", level=2)
        p_h = doc.add_paragraph()
        p_h.add_run("Hallazgos:").bold = True
        for hallazgo in dim.get("HALLAZGOS", []):
            doc.add_paragraph(hallazgo, style="List Bullet")
        p_r = doc.add_paragraph()
        p_r.add_run("Recomendaciones:").bold = True
        for rec in dim.get("RECOMENDACIONES", []):
            doc.add_paragraph(rec, style="List Bullet")
        doc.add_paragraph()
    doc.add_page_break()

    # ─── ALERTAS CRÍTICAS ───
    alertas = resultado.get("ALERTAS_CRITICAS", [])
    if alertas:
        doc.add_heading("6. Alertas Críticas", level=1)
        p_alerta_intro = doc.add_paragraph()
        p_alerta_intro.add_run("⚠️ Las siguientes alertas requieren atención inmediata:").bold = True
        for alerta in alertas:
            p_a = doc.add_paragraph(style="List Bullet")
            run_a = p_a.add_run(alerta)
            run_a.font.color.rgb = RGBColor(0xDC, 0x26, 0x26)
        doc.add_paragraph()

    # ─── FORTALEZAS ───
    doc.add_heading("7. Fortalezas Identificadas", level=1)
    for fortaleza in resultado.get("FORTALEZAS", []):
        p_f = doc.add_paragraph(style="List Bullet")
        run_f = p_f.add_run(fortaleza)
        run_f.font.color.rgb = RGBColor(0x05, 0x96, 0x69)
    doc.add_paragraph()

    # ─── PLAN DE MEJORA ───
    doc.add_heading("8. Plan de Mejora Priorizado", level=1)
    plan = resultado.get("PLAN_MEJORA", [])
    if plan:
        tabla_plan = doc.add_table(rows=1, cols=3)
        tabla_plan.style = "Table Grid"
        hdr_p = tabla_plan.rows[0].cells
        set_cell_text(hdr_p[0], "Prioridad", bold=True, center=True, color="FEF3C7")
        set_cell_text(hdr_p[1], "Acción", bold=True, color="FEF3C7")
        set_cell_text(hdr_p[2], "Plazo", bold=True, center=True, color="FEF3C7")
        for item in plan:
            row = tabla_plan.add_row().cells
            prioridad = item.get("PRIORIDAD", "Media")
            prio_color = {"Alta": "FEE2E2", "Media": "FEF3C7", "Baja": "D1FAE5"}.get(prioridad, "F1F5F9")
            set_cell_text(row[0], prioridad, center=True, color=prio_color)
            set_cell_text(row[1], item.get("ACCION", ""))
            set_cell_text(row[2], item.get("PLAZO", ""), center=True)
    doc.add_paragraph()

    # ─── FIRMAS ───
    doc.add_paragraph()
    tabla_firmas = doc.add_table(rows=2, cols=3)
    tabla_firmas.cell(0, 0).text = "_________________________"
    tabla_firmas.cell(0, 1).text = "_________________________"
    tabla_firmas.cell(0, 2).text = "_________________________"
    tabla_firmas.cell(1, 0).text = "Auditor Curricular"
    tabla_firmas.cell(1, 1).text = "Coordinador/a ETP"
    tabla_firmas.cell(1, 2).text = "Director/a Académico"
    for row in tabla_firmas.rows:
        for cell in row.cells:
            cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER

    add_page_number(doc)
    buffer = BytesIO()
    doc.save(buffer)
    buffer.seek(0)
    return buffer


# ═══════════════════════════════════════════════════════════════════════════
# RENDERIZADO DE STEPPER
# ═══════════════════════════════════════════════════════════════════════════
def render_stepper(fase_actual):
    pasos = [("1", "Cargar Matriz"), ("2", "Analizar"), ("3", "Resultados")]
    html = '<div class="audit-stepper">'
    for i, (num, label) in enumerate(pasos):
        estado = "done" if i + 1 < fase_actual else ("active" if i + 1 == fase_actual else "inactive")
        html += f'''
        <div class="audit-step">
            <div class="audit-step-circle {estado}">{"✓" if estado == "done" else num}</div>
            <div class="audit-step-label">{label}</div>
        </div>'''
        if i < len(pasos) - 1:
            line_class = "done" if i + 1 < fase_actual else ""
            html += f'<div class="audit-step-line {line_class}"></div>'
    html += '</div>'
    st.markdown(html, unsafe_allow_html=True)


# ═══════════════════════════════════════════════════════════════════════════
# HERO + TABS PRINCIPALES
# ═══════════════════════════════════════════════════════════════════════════
ia.panel_sidebar_ia("Auditor de Matriz Curricular")
st.markdown("""
<div class="audit-hero">
    <div class="audit-hero-title">🧾 Auditor de Matriz Curricular ETP</div>
    <div class="audit-hero-sub">Auditoría de calidad curricular en 6 dimensiones · Puntuación 0-100 · Distribución Bloom · Plan de mejora</div>
    <div>
        <span class="audit-hero-badge">📋 6 Dimensiones</span>
        <span class="audit-hero-badge">🧠 Taxonomía Bloom</span>
        <span class="audit-hero-badge">🤖 Asistido por IA</span>
        <span class="audit-hero-badge">📄 Word Profesional</span>
    </div>
</div>
""", unsafe_allow_html=True)

tab_audit, tab_hist = st.tabs(["🔍 Auditoría de Matriz", "📋 Historial"])

# ═══════════════════════════════════════════════════════════════════════════
# TAB 1: AUDITORÍA
# ═══════════════════════════════════════════════════════════════════════════
with tab_audit:
    render_stepper(st.session_state.auditoria_fase + 1)

    if st.session_state.auditoria_fase == 0:
        st.markdown('<div class="audit-section-title">📋 Cargar Matriz Curricular</div>', unsafe_allow_html=True)
        with st.form("form_auditoria"):
            col1, col2 = st.columns(2)
            with col1:
                modulo = st.text_input("Módulo / Asignatura", placeholder="Ej: MF_358_3 Impuestos al Consumo")
                # Mejora 3: auto-relleno desde sesión
                docente = st.text_input(
                    "Auditor / Docente",
                    value=st.session_state.get("usuario_display_nombre", "") or "Ing. Bernardo Antonio Hernández Batista",
                )
            with col2:
                institucion = st.text_input("Institución", value="Politécnico Salesiano Arquides Calderón")
                tipo_matriz = st.selectbox("Tipo de Matriz", [
                    "Matriz de Planificación Modular (RA)", "Matriz de Ponderación",
                    "Matriz de Plan Diario", "Matriz Curricular Completa", "Otra",
                ])
            enfoque = st.text_area("Enfoque de auditoría (opcional)", height=60,
                                   placeholder="Ej: Verificar alineación con el diseño curricular oficial, dosificación temporal...")
            st.markdown('<div class="audit-section-title">📄 Fuente de la Matriz</div>', unsafe_allow_html=True)
            tab_pdf, tab_excel, tab_texto = st.tabs(["📕 PDF", "📊 Excel", "✍️ Texto"])
            with tab_pdf:
                archivo_pdf = st.file_uploader("Subir PDF de la matriz", type=["pdf"])
            with tab_excel:
                archivo_excel = st.file_uploader("Subir Excel de la matriz", type=["xlsx", "xls"])
            with tab_texto:
                texto_directo = st.text_area("Pegar contenido de la matriz", height=200,
                                             placeholder="Pega aquí el contenido de la matriz curricular...")
            max_tokens, temperature = ia.control_avanzado(default_tokens=16384, tope=32000, default_temp=0.15)
            st.markdown("---")
            btn_auditar = st.form_submit_button("🔍 Iniciar Auditoría Curricular", type="primary", width="stretch")

        if btn_auditar:
            if not modulo:
                st.warning("⚠️ Indica al menos el módulo/asignatura.")
            else:
                try:
                    if archivo_pdf:
                        matriz_texto = extraer_texto_pdf(archivo_pdf)
                        st.info(f"📕 PDF cargado: {len(matriz_texto)} caracteres extraídos.")
                    elif archivo_excel:
                        matriz_texto = extraer_texto_excel(archivo_excel)
                        st.info(f"📊 Excel cargado: {len(matriz_texto)} caracteres extraídos.")
                    elif texto_directo:
                        matriz_texto = texto_directo
                        st.info(f"✍️ Texto directo: {len(matriz_texto)} caracteres.")
                    else:
                        st.warning("⚠️ Debes cargar un PDF, Excel o pegar el texto de la matriz.")
                        st.stop()
                except Exception as e:
                    st.error(f"❌ Error extrayendo contenido: {e}")
                    st.stop()

                with st.spinner("🧠 Auditando matriz curricular con IA..."):
                    try:
                        prompt = construir_prompt_auditoria(matriz_texto, modulo, tipo_matriz, enfoque)
                        resultado, flags = ia.solicitar_json(prompt, max_tokens=max_tokens,
                                                             temperature=temperature, modulo="auditor_matriz")
                        st.session_state.auditoria_resultado = {
                            "datos": resultado, "flags": flags,
                            "meta": {"modulo": modulo, "docente": docente, "institucion": institucion,
                                     "tipo_matriz": tipo_matriz, "fecha": datetime.now().strftime("%Y-%m-%d")},
                        }
                        # Mejoras 1 y 5: persistir auditoría
                        dims = resultado.get("DIMENSIONES", [])
                        promedio = round(sum(float(d.get("PUNTUACION", 0)) for d in dims) / len(dims)) if dims else 0
                        matriz_an = resultado.get("MATRIZ_ANALIZADA", {})
                        total_elem = len(matriz_an)
                        presentes = sum(1 for v in matriz_an.values() if isinstance(v, (int, float)) and v > 0)
                        completitud = round(presentes / total_elem * 100) if total_elem else 0
                        insertar_auditoria(
                            fecha=datetime.now().strftime("%Y-%m-%d"), modulo=modulo,
                            tipo_matriz=tipo_matriz, docente=docente, evaluador=docente,
                            puntuacion=int(resultado.get("PUNTUACION_GLOBAL", promedio)),
                            nivel=resultado.get("NIVEL_GLOBAL", "N/A"), completitud=completitud,
                        )
                        st.session_state.auditoria_fase = 2
                        st.toast("✅ Auditoría completada.", icon="🧾")
                        st.rerun()
                    except Exception as e:
                        ia.render_error_ia(e)

    elif st.session_state.auditoria_fase == 2:
        resultado = st.session_state.auditoria_resultado
        datos = resultado["datos"]
        meta = resultado["meta"]
        st.markdown('<div class="audit-section-title">🎯 Resultados de la Auditoría</div>', unsafe_allow_html=True)

        col_score, col_info = st.columns([1, 2])
        with col_score:
            puntuacion_global = datos.get("PUNTUACION_GLOBAL", 0)
            nivel_global = datos.get("NIVEL_GLOBAL", "N/A")
            score_color = {"Excelente": "#10B981", "Bueno": "#3B82F6",
                           "Aceptable": "#F59E0B", "Deficiente": "#EF4444"}.get(nivel_global, "#64748B")
            st.markdown(f"""
            <div class="score-global-card">
                <div class="score-global-value" style="color: {score_color};">{puntuacion_global}</div>
                <div class="score-global-label">Puntuación Global / 100</div>
                <div style="margin-top: 12px; font-size: 1.2rem; font-weight: 700; color: {score_color};">{nivel_global}</div>
            </div>""", unsafe_allow_html=True)
        with col_info:
            st.markdown(f"**Módulo:** {meta.get('modulo', '')}")
            st.markdown(f"**Tipo:** {meta.get('tipo_matriz', '')}")
            st.markdown(f"**Fecha:** {meta.get('fecha', '')}")
            st.markdown("---")
            with st.expander("📄 Resumen Ejecutivo", expanded=True):
                st.write(datos.get("RESUMEN_EJECUTIVO", ""))
        st.markdown("---")

        # ─── DIMENSIONES + gráfico (mejora 2) ───
        st.markdown('<div class="audit-section-title">📊 Dashboard de Dimensiones</div>', unsafe_allow_html=True)
        cols = st.columns(3)
        for idx, dim in enumerate(datos.get("DIMENSIONES", [])):
            with cols[idx % 3]:
                clave = dim.get("CLAVE", "")
                dim_info = DIMENSIONES_AUDITORIA.get(clave, {"nombre": clave, "icono": "📋", "color": "#64748B"})
                puntuacion = dim.get("PUNTUACION", 0)
                nivel = dim.get("NIVEL", "N/A")
                nivel_bg, nivel_text = nivel_css(nivel)
                st.markdown(f"""
                <div class="dim-card" style="border-left: 5px solid {dim_info.get('color', '#64748B')};">
                    <div class="dim-header">
                        <div class="dim-icono">{dim_info.get('icono', '📋')}</div>
                        <div class="dim-nombre">{dim_info.get('nombre', clave)}</div>
                    </div>
                    <div class="dim-score-container">
                        <div class="dim-score-bar">
                            <div class="dim-score-fill" style="width: {puntuacion}%; background: {dim_info.get('color', '#64748B')};"></div>
                        </div>
                        <div class="dim-score-value" style="color: {dim_info.get('color', '#64748B')};">{puntuacion}/100</div>
                    </div>
                    <div class="dim-nivel" style="background: {nivel_bg}; color: {nivel_text};">{nivel}</div>
                </div>""", unsafe_allow_html=True)

        # Mejora 2: gráfico de barras por dimensión
        df_dims = pd.DataFrame(datos.get("DIMENSIONES", []))
        if not df_dims.empty:
            df_dims["PUNTUACION"] = pd.to_numeric(df_dims["PUNTUACION"], errors="coerce").fillna(0)
            st.markdown("**Puntuación por dimensión:**")
            st.bar_chart(df_dims.set_index("CLAVE")["PUNTUACION"])
        st.markdown("---")

        # ─── BLOOM ───
        st.markdown('<div class="audit-section-title">🧠 Distribución Cognitiva Bloom</div>', unsafe_allow_html=True)
        bloom = datos.get("DISTRIBUCION_BLOOM", {})
        bloom_colors_css = {"Recordar": "#EF4444", "Comprender": "#F59E0B", "Aplicar": "#10B981",
                            "Analizar": "#3B82F6", "Evaluar": "#8B5CF6", "Crear": "#EC4899"}
        for nivel_bloom, pct in bloom.items():
            color = bloom_colors_css.get(nivel_bloom, "#64748B")
            st.markdown(f"""
            <div class="bloom-bar-container">
                <div class="bloom-bar-label">{nivel_bloom}</div>
                <div class="bloom-bar-track">
                    <div class="bloom-bar-fill" style="width: {pct}%; background: {color};">{pct}%</div>
                </div>
            </div>""", unsafe_allow_html=True)
        st.markdown("---")

        # ─── ALERTAS / FORTALEZAS ───
        alertas = datos.get("ALERTAS_CRITICAS", [])
        if alertas:
            st.markdown('<div class="audit-section-title">🚨 Alertas Críticas</div>', unsafe_allow_html=True)
            for alerta in alertas:
                st.markdown(f'<div class="audit-alert-critical"><div class="audit-alert-text">⚠️ {alerta}</div></div>',
                            unsafe_allow_html=True)
        fortalezas = datos.get("FORTALEZAS", [])
        if fortalezas:
            st.markdown('<div class="audit-section-title">💪 Fortalezas Identificadas</div>', unsafe_allow_html=True)
            for fortaleza in fortalezas:
                st.markdown(f'<div class="audit-strength"><div class="audit-strength-text">✅ {fortaleza}</div></div>',
                            unsafe_allow_html=True)
        st.markdown("---")

        # ─── HALLAZGOS ───
        st.markdown('<div class="audit-section-title">🔍 Hallazgos y Recomendaciones por Dimensión</div>', unsafe_allow_html=True)
        for dim in datos.get("DIMENSIONES", []):
            clave = dim.get("CLAVE", "")
            dim_info = DIMENSIONES_AUDITORIA.get(clave, {"nombre": clave, "icono": "📋"})
            with st.expander(f"{dim_info.get('icono', '📋')} {dim_info.get('nombre', clave)} — {dim.get('PUNTUACION', 0)}/100 ({dim.get('NIVEL', 'N/A')})"):
                st.markdown("**Hallazgos:**")
                for hallazgo in dim.get("HALLAZGOS", []):
                    st.markdown(f"- {hallazgo}")
                st.markdown("**Recomendaciones:**")
                for rec in dim.get("RECOMENDACIONES", []):
                    st.markdown(f"- 💡 {rec}")

        # ─── PLAN DE MEJORA ───
        plan = datos.get("PLAN_MEJORA", [])
        if plan:
            st.markdown('<div class="audit-section-title">📋 Plan de Mejora Priorizado</div>', unsafe_allow_html=True)
            df_plan = pd.DataFrame(plan)
            st.dataframe(df_plan, width="stretch", hide_index=True)   # ← deprecación corregida
        st.markdown("---")

        # ─── TABS VISTA / JSON / DEBUG ───
        tab_preview, tab_json, tab_debug = st.tabs(["👁️ Vista Previa", "🧾 JSON Completo", "🐛 Depuración"])
        with tab_preview:
            st.json(datos)
        with tab_json:
            st.json({"resultado": datos, "meta": meta, "flags": resultado.get("flags", {})})
        with tab_debug:
            st.write("Meta:", meta)
            st.write("Flags:", resultado.get("flags", {}))
        st.markdown("---")

        # ─── DESCARGAS (Word + CSV, mejora 4) ───
        col_dl1, col_dl2, col_dl3 = st.columns(3)
        with col_dl1:
            if st.button("📥 Generar Informe Word (.docx)", type="primary", width="stretch"):
                with st.spinner("📄 Construyendo informe de auditoría..."):
                    st.session_state.audit_buffer = build_audit_docx(datos, meta)
        if hasattr(st.session_state, "audit_buffer") and st.session_state.audit_buffer:
            nombre_archivo = ia.sanear_nombre_archivo(f"Auditoria_{meta.get('modulo', 'matriz')}")
            st.download_button(label="⬇️ Descargar Informe de Auditoría (.docx)",
                               data=st.session_state.audit_buffer,
                               file_name=f"{nombre_archivo}.docx",
                               mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                               type="primary", width="stretch")
        with col_dl2:
            if not df_dims.empty:
                st.download_button("📄 Dimensiones (.csv)",
                                   data=df_dims.to_csv(index=False).encode("utf-8-sig"),
                                   file_name=f"Dimensiones_{ia.sanear_nombre_archivo(meta.get('modulo', 'matriz'))}.csv",
                                   mime="text/csv", width="stretch")
        with col_dl3:
            if plan:
                st.download_button("📄 Plan de Mejora (.csv)",
                                   data=pd.DataFrame(plan).to_csv(index=False).encode("utf-8-sig"),
                                   file_name=f"PlanMejora_{ia.sanear_nombre_archivo(meta.get('modulo', 'matriz'))}.csv",
                                   mime="text/csv", width="stretch")
        st.markdown("---")
        if st.button("🔄 Nueva Auditoría", width="stretch"):
            st.session_state.auditoria_fase = 0
            st.session_state.auditoria_resultado = None
            if hasattr(st.session_state, "audit_buffer"):
                del st.session_state.audit_buffer
            st.rerun()

# ═══════════════════════════════════════════════════════════════════════════
# TAB 2: HISTORIAL (mejora 1)
# ═══════════════════════════════════════════════════════════════════════════
with tab_hist:
    st.markdown('<div class="audit-section-title">🗄️ Historial de Auditorías de Matriz</div>', unsafe_allow_html=True)
    hist = listar_auditorias()
    if not hist:
        st.info("Aún no hay auditorías registradas. Ejecuta la primera en la pestaña '🔍 Auditoría de Matriz'.")
    else:
        df_hist = pd.DataFrame(hist)
        col_f1, col_f2 = st.columns(2)
        with col_f1:
            f_mod = st.selectbox("Filtrar por módulo", ["Todos"] + sorted(df_hist["modulo"].unique().tolist()))
        with col_f2:
            f_niv = st.selectbox("Filtrar por nivel", ["Todos"] + sorted(df_hist["nivel"].unique().tolist()))
        df_f = df_hist.copy()
        if f_mod != "Todos":
            df_f = df_f[df_f["modulo"] == f_mod]
        if f_niv != "Todos":
            df_f = df_f[df_f["nivel"] == f_niv]

        m1, m2, m3 = st.columns(3)
        with m1:
            st.metric("Auditorías", len(df_f))
        with m2:
            st.metric("Puntaje promedio", round(df_f["puntuacion"].mean(), 1) if not df_f.empty else 0)
        with m3:
            st.metric("Completitud promedio", f"{round(df_f['completitud'].mean(), 1)}%" if not df_f.empty else "0%")

        st.markdown("#### 🏆 Ranking por módulo (puntaje promedio)")
        ranking = df_hist.groupby("modulo")["puntuacion"].mean().sort_values(ascending=False).round(1)
        st.bar_chart(ranking)

        st.dataframe(df_f.drop(columns=["id"]), width="stretch", hide_index=True)
        st.download_button("📄 Exportar historial (.csv)",
                           data=df_f.to_csv(index=False).encode("utf-8-sig"),
                           file_name="Historial_Auditorias_Matriz.csv",
                           mime="text/csv", width="stretch")