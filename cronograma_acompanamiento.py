"""
cronograma_acompanamiento.py — Cronograma de Acompañamiento Pedagógico ETP (NIVEL DIOS)
Gestión de visitas con Cascada Inteligente (Docente -> Módulo -> Sección),
persistencia SQLite, vista previa tipo calendario semanal y exportación.
"""
import streamlit as st
import pandas as pd
from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import nsdecls
from docx.oxml import parse_xml
from io import BytesIO
import datetime
import sqlite3

# ═══════════════════════════════════════════════════════════════════════════
# SECCIÓN 1 — Base de Datos SQLite (Conexión Unificada)
# ═══════════════════════════════════════════════════════════════════════════
DB_NAME = "gestion_etp.db"

def _conn():
    return sqlite3.connect(DB_NAME, check_same_thread=False)

def init_db():
    conn = _conn()
    cursor = conn.cursor()
    cursor.execute('''
    CREATE TABLE IF NOT EXISTS cronograma (
        id INTEGER PRIMARY KEY AUTOINCREMENT,
        centro TEXT, coordinador TEXT, anio TEXT, trimestre TEXT,
        docente TEXT, modulo TEXT, seccion TEXT, hora TEXT,
        dia TEXT, modalidad TEXT, observaciones TEXT
    )
    ''')
    # Asegurar tabla docentes para que la cascada no falle
    cursor.execute('''CREATE TABLE IF NOT EXISTS docentes (
        docente TEXT, modulo TEXT, seccion TEXT, password TEXT DEFAULT '1234', usuario TEXT
    )''')
    conn.commit()
    conn.close()

init_db()

def listar_docentes_bd():
    conn = _conn()
    cur = conn.cursor()
    try:
        cur.execute("SELECT DISTINCT docente FROM docentes WHERE docente IS NOT NULL AND docente != ''")
        rows = [r[0] for r in cur.fetchall()]
    except Exception:
        rows = []
    conn.close()
    return rows

def obtener_modulos_usuario_bd(docente):
    conn = _conn()
    cur = conn.cursor()
    try:
        cur.execute("SELECT modulo, seccion FROM docentes WHERE docente=? AND modulo IS NOT NULL AND modulo != ''", (docente,))
        rows = [list(r) for r in cur.fetchall()]
    except Exception:
        rows = []
    conn.close()
    return rows

def parse_time(t_str):
    if pd.isna(t_str) or not str(t_str).strip():
        return None
    try:
        return datetime.datetime.strptime(str(t_str), "%H:%M").time()
    except Exception:
        return None

def cargar_cronograma(anio, trimestre):
    conn = _conn()
    query = """
    SELECT docente as Docente, modulo as Módulo, seccion as Sección,
           hora as Hora, dia as Día, modalidad as Modalidad, observaciones as Observaciones
    FROM cronograma WHERE anio = ? AND trimestre = ?
    """
    df = pd.read_sql_query(query, conn, params=(anio, trimestre))
    conn.close()
    if not df.empty and 'Hora' in df.columns:
        df['Hora'] = df['Hora'].apply(parse_time)
    else:
        df = pd.DataFrame(columns=["Docente", "Módulo", "Sección", "Hora", "Día", "Modalidad", "Observaciones"])
    return df

def guardar_cronograma_completo(df, centro, coordinador, anio, trimestre):
    conn = _conn()
    cursor = conn.cursor()
    cursor.execute("DELETE FROM cronograma WHERE anio = ? AND trimestre = ?", (anio, trimestre))
    for _, row in df.iterrows():
        hora_obj = row.get('Hora')
        hora_str = hora_obj.strftime("%H:%M") if pd.notna(hora_obj) and isinstance(hora_obj, datetime.time) else ""
        cursor.execute('''
            INSERT INTO cronograma (centro, coordinador, anio, trimestre, docente, modulo, seccion, hora, dia, modalidad, observaciones)
            VALUES (?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?)
        ''', (centro, coordinador, anio, trimestre,
              str(row.get('Docente', '')), str(row.get('Módulo', '')), str(row.get('Sección', '')),
              hora_str, str(row.get('Día', '')), str(row.get('Modalidad', '')), str(row.get('Observaciones', ''))))
    conn.commit()
    conn.close()

def agregar_visita_unica(centro, coord, anio, trim, doc, mod, sec, hora, dia, mod_tipo, obs):
    conn = _conn()
    conn.execute('''
        INSERT INTO cronograma (centro, coordinador, anio, trimestre, docente, modulo, seccion, hora, dia, modalidad, observaciones)
        VALUES (?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?)
    ''', (centro, coord, anio, trim, doc, mod, sec, hora, dia, mod_tipo, obs))
    conn.commit()
    conn.close()

# ═══════════════════════════════════════════════════════════════════════════
# SECCIÓN 2 — Constantes y utilidades visuales
# ═══════════════════════════════════════════════════════════════════════════
DIAS_SEMANA = ["Lunes", "Martes", "Miércoles", "Jueves", "Viernes"]
MODALIDAD_COLORES = {
    "Presencial": ("#DBEAFE", "#1E40AF", "🏫"),
    "Virtual": ("#D1FAE5", "#065F46", "💻"),
    "Mixta": ("#EDE9FE", "#5B21B6", "🔄"),
}

def format_time_display(val):
    if pd.notna(val) and isinstance(val, datetime.time):
        return val.strftime("%I:%M %p")
    return str(val) if pd.notna(val) else "—"

# ═══════════════════════════════════════════════════════════════════════════
# SECCIÓN 3 — Exportación Word y Excel
# ═══════════════════════════════════════════════════════════════════════════
def shade_cell(cell, color):
    shd = parse_xml(r'<w:shd {} w:fill="{}"/>'.format(nsdecls('w'), color))
    cell._tc.get_or_add_tcPr().append(shd)

def generar_word_cronograma(df, info):
    doc = Document()
    doc.styles['Normal'].font.name = 'Calibri'
    doc.styles['Normal'].font.size = Pt(10)
    for section in doc.sections:
        section.left_margin = Inches(0.5)
        section.right_margin = Inches(0.5)
        section.top_margin = Inches(0.5)
        section.bottom_margin = Inches(0.5)

    p_titulo = doc.add_paragraph()
    p_titulo.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r1 = p_titulo.add_run("MINISTERIO DE EDUCACIÓN DE LA REPÚBLICA DOMINICANA\n")
    r1.bold = True; r1.font.size = Pt(12)
    r2 = p_titulo.add_run(f"{info['centro']}\n")
    r2.bold = True; r2.font.size = Pt(11)
    r3 = p_titulo.add_run("CRONOGRAMA DE ACOMPAÑAMIENTO PEDAGÓGICO ETP")
    r3.bold = True; r3.font.size = Pt(12); r3.font.color.rgb = RGBColor(37, 99, 235)

    p_datos = doc.add_paragraph()
    p_datos.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p_datos.add_run(
        f"Coordinador/a: {info['coordinador']}   |   Trimestre: {info['trimestre']}   |   Año Escolar: {info['anio']}"
    ).italic = True
    doc.add_paragraph()

    num_cols = len(df.columns) + 1
    tabla = doc.add_table(rows=1, cols=num_cols)
    tabla.style = 'Table Grid'
    tabla.alignment = WD_TABLE_ALIGNMENT.CENTER
    hdr_cells = tabla.rows[0].cells
    hdr_cells[0].text = "No."
    for i, col_name in enumerate(df.columns):
        hdr_cells[i+1].text = col_name
    for cell in hdr_cells:
        shade_cell(cell, "1E40AF")
        for paragraph in cell.paragraphs:
            for run in paragraph.runs:
                run.font.color.rgb = RGBColor(255, 255, 255)
                run.bold = True
                run.font.size = Pt(9)
            paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER

    for idx, row in df.iterrows():
        row_cells = tabla.add_row().cells
        row_cells[0].text = str(idx + 1)
        row_cells[0].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
        for i, val in enumerate(row):
            if pd.isna(val):
                val = ""
            elif isinstance(val, datetime.time):
                val = val.strftime("%I:%M %p")
            row_cells[i+1].text = str(val)
            for paragraph in row_cells[i+1].paragraphs:
                for run in paragraph.runs:
                    run.font.size = Pt(9)
        if idx % 2 == 1:
            for cell in row_cells:
                shade_cell(cell, "F1F5F9")

    anchos = [Inches(0.3), Inches(1.5), Inches(1.3), Inches(0.6), Inches(0.8), Inches(0.8), Inches(0.8), Inches(1.6)]
    for row in tabla.rows:
        for i, width in enumerate(anchos):
            if i < len(row.cells):
                row.cells[i].width = width

    doc.add_paragraph("\n\n")
    t_firmas = doc.add_table(rows=2, cols=3)
    t_firmas.cell(0, 0).text = "__________________________"
    t_firmas.cell(0, 1).text = "__________________________"
    t_firmas.cell(0, 2).text = "__________________________"
    t_firmas.cell(1, 0).text = "Director/a de Centro"
    t_firmas.cell(1, 1).text = "Coordinador/a Módulos Formativos ETP"
    t_firmas.cell(1, 2).text = "Docente ETP"
    for row in t_firmas.rows:
        for cell in row.cells:
            cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER

    buffer = BytesIO()
    doc.save(buffer)
    buffer.seek(0)
    return buffer

def generar_excel_cronograma(df):
    df_export = df.copy()
    if "Hora" in df_export.columns:
        df_export["Hora"] = df_export["Hora"].apply(
            lambda val: val.strftime("%I:%M %p") if pd.notna(val) and isinstance(val, datetime.time) else val
        )
    buffer = BytesIO()
    with pd.ExcelWriter(buffer, engine='openpyxl') as writer:
        df_export.to_excel(writer, index=False, sheet_name='Cronograma')
    buffer.seek(0)
    return buffer

# ═══════════════════════════════════════════════════════════════════════════
# SECCIÓN 4 — Estilos
# ═══════════════════════════════════════════════════════════════════════════
st.markdown("""
<style>
@import url('https://fonts.googleapis.com/css2?family=Inter:wght@300;400;500;600;700;800;900&display=swap');

html, body, [class*="css"] {
    font-family: 'Inter', sans-serif;
    background-color: #F0F4F8;
    color: #1E293B;
}

.crono-hero {
    background: linear-gradient(135deg, #0F172A 0%, #065F46 40%, #10B981 70%, #34D399 100%);
    color: #fff;
    padding: 2.2rem;
    border-radius: 20px;
    margin-bottom: 1.5rem;
    box-shadow: 0 25px 50px rgba(6, 95, 70, 0.3);
    position: relative;
    overflow: hidden;
}

.crono-hero::before {
    content: '';
    position: absolute;
    top: -50%; left: -50%; width: 200%; height: 200%;
    background: radial-gradient(circle, rgba(255,255,255,0.1) 0%, transparent 60%);
    animation: cronoPulse 6s ease-in-out infinite;
}

@keyframes cronoPulse {
    0%, 100% { transform: scale(1); opacity: 0.5; }
    50% { transform: scale(1.1); opacity: 0.8; }
}

.crono-hero-title { font-size: 2.4rem; font-weight: 900; letter-spacing: -0.03em; margin-bottom: 0.4rem; position: relative; }
.crono-hero-sub { font-size: 1.05rem; opacity: 0.9; line-height: 1.5; position: relative; }

.crono-hero-badge {
    display: inline-block;
    background: rgba(255,255,255,0.15);
    border: 1px solid rgba(255,255,255,0.25);
    border-radius: 8px;
    padding: 4px 12px;
    font-size: 0.8rem;
    font-weight: 600;
    margin-top: 0.8rem;
    margin-right: 8px;
    position: relative;
}

.crono-section-title {
    color: #065F46;
    font-weight: 700;
    font-size: 1.12rem;
    border-bottom: 2px solid #D1FAE5;
    padding-bottom: 8px;
    margin: 1.2rem 0 0.9rem 0;
}

.crono-kpi-card {
    background: #fff;
    border-radius: 14px;
    padding: 1.2rem;
    text-align: center;
    border: 1px solid #E2E8F0;
    box-shadow: 0 4px 12px rgba(0,0,0,0.06);
    transition: all 0.3s ease;
}

.crono-kpi-card:hover { transform: translateY(-4px); box-shadow: 0 12px 24px rgba(0,0,0,0.12); }
.crono-kpi-icon { font-size: 2rem; margin-bottom: 0.5rem; }
.crono-kpi-value { font-size: 2rem; font-weight: 800; color: #065F46; }
.crono-kpi-label { font-size: 0.75rem; font-weight: 700; color: #94A3B8; text-transform: uppercase; letter-spacing: 0.05em; margin-top: 0.3rem; }

/* Cascada Inteligente */
.box-cascada {
    background: #FFFFFF; 
    padding: 20px; 
    border-radius: 12px; 
    border: 1px solid #D1FAE5; 
    margin-bottom: 20px;
    box-shadow: 0 4px 12px rgba(0,0,0,0.03);
}

/* Calendario semanal */
.crono-day-header {
    background: linear-gradient(135deg, #065F46, #10B981);
    color: #fff;
    padding: 10px 14px;
    border-radius: 10px 10px 0 0;
    font-weight: 800;
    font-size: 0.95rem;
    text-align: center;
}

.crono-day-column {
    background: #F8FAFC;
    border: 1px solid #E2E8F0;
    border-radius: 0 0 10px 10px;
    padding: 10px;
    min-height: 200px;
}

.crono-visit-card {
    background: #fff;
    border-radius: 10px;
    padding: 10px 12px;
    margin-bottom: 10px;
    border-left: 4px solid #3B82F6;
    box-shadow: 0 2px 6px rgba(0,0,0,0.05);
    transition: all 0.2s ease;
}

.crono-visit-card:hover { transform: translateX(3px); box-shadow: 0 6px 16px rgba(0,0,0,0.1); }

.crono-visit-time { font-size: 0.72rem; font-weight: 800; color: #64748B; text-transform: uppercase; letter-spacing: 0.03em; }
.crono-visit-docente { font-weight: 700; font-size: 0.88rem; color: #0F172A; margin-top: 2px; }
.crono-visit-modulo { font-size: 0.78rem; color: #475569; margin-top: 2px; }
.crono-visit-meta { font-size: 0.7rem; color: #94A3B8; margin-top: 4px; }

.crono-empty {
    text-align: center;
    color: #CBD5E1;
    font-size: 0.82rem;
    padding: 2rem 0;
}
</style>
""", unsafe_allow_html=True)

# ═══════════════════════════════════════════════════════════════════════════
# SECCIÓN 5 — Hero y Filtros Globales
# ═══════════════════════════════════════════════════════════════════════════
st.markdown("""
<div class="crono-hero">
    <div class="crono-hero-title">🗓️ Cronograma de Acompañamiento Pedagógico</div>
    <div class="crono-hero-sub">
        Gestión Inteligente en Cascada · Vista de Calendario · Sincronizado con Docentes
    </div>
    <div>
        <span class="crono-hero-badge" style="display:inline-block; margin:0; box-shadow:none; border: 1px solid rgba(255,255,255,0.25);">📅 Calendario Semanal</span>
        <span class="crono-hero-badge" style="display:inline-block; margin:0; box-shadow:none; border: 1px solid rgba(255,255,255,0.25);">🧠 Selección Automática</span>
        <span class="crono-hero-badge" style="display:inline-block; margin:0; box-shadow:none; border: 1px solid rgba(255,255,255,0.25);">💾 Persistencia SQLite</span>
        <span class="crono-hero-badge" style="display:inline-block; margin:0; box-shadow:none; border: 1px solid rgba(255,255,255,0.25);">📥 Reportes Automáticos</span>
    </div>
</div>
""", unsafe_allow_html=True)

st.markdown('<div class="crono-section-title">📋 1. Definir Período Activo</div>', unsafe_allow_html=True)

col_f1, col_f2 = st.columns(2)
with col_f1:
    anio_escolar = st.selectbox("Año Escolar", ["2025-2026", "2026-2027", "2027-2028"])
with col_f2:
    trimestre_sel = st.selectbox("Período de Acompañamiento", ["Primer Trimestre", "Segundo Trimestre", "Tercer Trimestre", "Anual"])

df_cronograma_db = cargar_cronograma(anio_escolar, trimestre_sel)

# ═══════════════════════════════════════════════════════════════════════════
# SECCIÓN 6 — KPIs
# ═══════════════════════════════════════════════════════════════════════════
if not df_cronograma_db.empty:
    total_visitas = len(df_cronograma_db)
    docentes_unicos = df_cronograma_db["Docente"].nunique()
    modulos_unicos = df_cronograma_db["Módulo"].nunique()
    modalidad_top = df_cronograma_db["Modalidad"].mode()
    modalidad_top = modalidad_top.iloc[0] if not modalidad_top.empty else "—"

    col_k1, col_k2, col_k3, col_k4 = st.columns(4)
    with col_k1:
        st.markdown(f'<div class="crono-kpi-card"><div class="crono-kpi-icon">📅</div><div class="crono-kpi-value">{total_visitas}</div><div class="crono-kpi-label">Visitas Totales</div></div>', unsafe_allow_html=True)
    with col_k2:
        st.markdown(f'<div class="crono-kpi-card"><div class="crono-kpi-icon">🧑‍🏫</div><div class="crono-kpi-value">{docentes_unicos}</div><div class="crono-kpi-label">Docentes Únicos</div></div>', unsafe_allow_html=True)
    with col_k3:
        st.markdown(f'<div class="crono-kpi-card"><div class="crono-kpi-icon">📚</div><div class="crono-kpi-value">{modulos_unicos}</div><div class="crono-kpi-label">Módulos Programados</div></div>', unsafe_allow_html=True)
    with col_k4:
        st.markdown(f'<div class="crono-kpi-card"><div class="crono-kpi-icon">🎯</div><div class="crono-kpi-value" style="font-size:1.3rem;">{modalidad_top}</div><div class="crono-kpi-label">Modalidad Dominante</div></div>', unsafe_allow_html=True)
    st.markdown("---")

tab_editor, tab_preview, tab_analitica, tab_export = st.tabs([
    "📝 Editor", "📅 Vista Previa", "📊 Analítica", "📥 Exportar"
])

# ═══════════════════════════════════════════════════════════════════════════
# TAB 1: EDITOR (Cascada Inteligente + Edición Masiva)
# ═══════════════════════════════════════════════════════════════════════════
with tab_editor:
    centro_coordinador = st.session_state.get("crono_centro", "Politécnico Salesiano Arquides Calderón")
    nombre_coordinador = st.session_state.get("coordinador_nombre", "Ing. Bernardo Antonio Hernández Batista")

    # --- PARTE 1: AGREGAR NUEVA VISITA (CASCADA) ---
    st.markdown('<div class="crono-section-title">➕ Programar Nueva Visita (Cascada Automática)</div>', unsafe_allow_html=True)
    st.markdown('<div class="box-cascada">', unsafe_allow_html=True)
    
    docentes_disponibles = listar_docentes_bd()
    
    col_a, col_b, col_c = st.columns(3)
    with col_a:
        doc_sel = st.selectbox("👤 Docente", ["Seleccionar..."] + sorted(docentes_disponibles) + ["✍️ Escribir manualmente"])
        if doc_sel == "✍️ Escribir manualmente":
            doc_val = st.text_input("Nombre del docente")
            modulos_docente = []
        elif doc_sel != "Seleccionar...":
            doc_val = doc_sel
            modulos_docente = obtener_modulos_usuario_bd(doc_val)
        else:
            doc_val = ""
            modulos_docente = []

    with col_b:
        if modulos_docente:
            mods_unicos = sorted(list(set(m[0] for m in modulos_docente)))
            mod_sel = st.selectbox("📚 Módulo", ["Seleccionar..."] + mods_unicos + ["✍️ Escribir manualmente"])
            if mod_sel == "✍️ Escribir manualmente":
                mod_val = st.text_input("Escribir Módulo")
            elif mod_sel != "Seleccionar...":
                mod_val = mod_sel
            else:
                mod_val = ""
        else:
            mod_val = st.text_input("📚 Módulo Formativo")

    with col_c:
        if modulos_docente and mod_val and mod_val != "✍️ Escribir manualmente":
            secciones_vinculadas = sorted(list(set(m[1] for m in modulos_docente if m[0] == mod_val)))
            sec_sel = st.selectbox("🏫 Sección", secciones_vinculadas + ["✍️ Otra"])
            if sec_sel == "✍️ Otra":
                sec_val = st.text_input("Escribir sección")
            else:
                sec_val = sec_sel
        else:
            sec_val = st.text_input("🏫 Sección/Grado")

    col_d, col_e, col_f = st.columns([1, 1, 1])
    with col_d:
        dia_val = st.selectbox("📅 Día de la semana", DIAS_SEMANA)
    with col_e:
        hora_val = st.time_input("⏰ Hora", value=datetime.time(8, 0))
    with col_f:
        mod_tipo = st.selectbox("📍 Modalidad", ["Presencial", "Virtual", "Mixta"])
        
    obs_val = st.text_input("📝 Notas u Observaciones previas")

    if st.button("➕ Agregar al Cronograma", type="primary", use_container_width=True):
        if not doc_val or not mod_val:
            st.warning("⚠️ Debes seleccionar el Docente y el Módulo.")
        else:
            hora_str = hora_val.strftime("%H:%M")
            agregar_visita_unica(centro_coordinador, nombre_coordinador, anio_escolar, trimestre_sel,
                                 doc_val, mod_val, sec_val, hora_str, dia_val, mod_tipo, obs_val)
            st.toast(f"✅ ¡Visita programada para el {dia_val}!", icon="📅")
            st.rerun()

    st.markdown('</div>', unsafe_allow_html=True)
    
    # --- PARTE 2: EDICIÓN MASIVA ---
    st.markdown('<div class="crono-section-title">✏️ Gestión Masiva de Visitas Programadas</div>', unsafe_allow_html=True)
    st.info("💡 En esta tabla puedes ajustar rápidamente horas, días o borrar filas. Haz clic en **Guardar Cambios** cuando termines.")

    with st.form("form_cronograma_masivo"):
        column_config = {
            "Docente": st.column_config.SelectboxColumn("Docente", options=sorted(docentes_disponibles), width="medium", required=True),
            "Módulo": st.column_config.TextColumn("Módulo", width="medium"),
            "Sección": st.column_config.TextColumn("Sección", width="small"),
            "Hora": st.column_config.TimeColumn("Hora", format="hh:mm A", step=300, width="small"),
            "Día": st.column_config.SelectboxColumn("Día", options=DIAS_SEMANA, width="small"),
            "Modalidad": st.column_config.SelectboxColumn("Modalidad", options=["Presencial", "Virtual", "Mixta"], width="small"),
            "Observaciones": st.column_config.TextColumn("Observaciones", width="large"),
        }

        edited_df = st.data_editor(
            df_cronograma_db,
            column_config=column_config,
            num_rows="dynamic",
            use_container_width=True,
            hide_index=True,
        )

        if st.form_submit_button("💾 Guardar Cambios en la Tabla", type="primary", use_container_width=True):
            cleaned_df = edited_df.dropna(subset=['Docente'], how='any')
            guardar_cronograma_completo(cleaned_df, centro_coordinador, nombre_coordinador, anio_escolar, trimestre_sel)
            st.toast(f"✅ Cambios guardados ({len(cleaned_df)} registros)", icon="💾")
            st.rerun()

# ═══════════════════════════════════════════════════════════════════════════
# TAB 2: VISTA PREVIA (Calendario Semanal)
# ═══════════════════════════════════════════════════════════════════════════
with tab_preview:
    st.markdown('<div class="crono-section-title">📅 Vista Previa — Calendario Semanal</div>', unsafe_allow_html=True)

    busqueda = st.text_input("🔍 Buscar docente, módulo o sección", placeholder="Escribe para filtrar...")
    df_preview = df_cronograma_db.copy()
    if busqueda.strip():
        texto = busqueda.strip().lower()
        df_preview = df_preview[df_preview.apply(lambda row: texto in " ".join(str(v).lower() for v in row if pd.notna(v)), axis=1)]

    if df_preview.empty:
        st.info("📭 No hay visitas registradas para este período. Agrega visitas en la pestaña '📝 Editor'.")
    else:
        cols_dias = st.columns(5)
        for i, dia in enumerate(DIAS_SEMANA):
            with cols_dias[i]:
                st.markdown(f'<div class="crono-day-header">{dia}</div>', unsafe_allow_html=True)
                st.markdown('<div class="crono-day-column">', unsafe_allow_html=True)

                visitas_dia = df_preview[df_preview["Día"] == dia].copy()
                visitas_dia["_hora_sort"] = visitas_dia["Hora"].apply(lambda x: x if isinstance(x, datetime.time) else datetime.time(23, 59))
                visitas_dia = visitas_dia.sort_values("_hora_sort")

                if visitas_dia.empty:
                    st.markdown('<div class="crono-empty">Sin visitas</div>', unsafe_allow_html=True)
                else:
                    for _, visit in visitas_dia.iterrows():
                        modalidad = str(visit.get("Modalidad", "Presencial"))
                        color_bg, color_text, icono_mod = MODALIDAD_COLORES.get(modalidad, ("#DBEAFE", "#1E40AF", "🏫"))
                        hora_display = format_time_display(visit.get("Hora"))
                        seccion = str(visit.get("Sección", "")) if pd.notna(visit.get("Sección")) else ""
                        observaciones = str(visit.get("Observaciones", "")) if pd.notna(visit.get("Observaciones")) else ""
                        obs_html = f'<div class="crono-visit-meta">📝 {observaciones[:60]}</div>' if observaciones else ""

                        # Tarjeta limpia sin la barra de modalidad
                        st.markdown(f"""
                        <div class="crono-visit-card" style="border-left-color: {color_text};">
                            <div class="crono-visit-time">🕐 {hora_display}</div>
                            <div class="crono-visit-docente">{visit.get('Docente', '')}</div>
                            <div class="crono-visit-modulo">📚 {visit.get('Módulo', '')}</div>
                            <div class="crono-visit-meta">🏫 Sec: {seccion}</div>
                            {obs_html}
                        </div>
                        """, unsafe_allow_html=True)
                st.markdown('</div>', unsafe_allow_html=True)
        st.markdown("---")
        st.caption(f"📊 Mostrando {len(df_preview)} de {len(df_cronograma_db)} visitas")

# ═══════════════════════════════════════════════════════════════════════════
# TAB 3: ANALÍTICA
# ═══════════════════════════════════════════════════════════════════════════
with tab_analitica:
    st.markdown('<div class="crono-section-title">📊 Analítica del Cronograma</div>', unsafe_allow_html=True)
    if df_cronograma_db.empty:
        st.info("📭 No hay datos para analizar.")
    else:
        st.markdown("#### 📅 Distribución de Visitas por Día")
        df_dias = df_cronograma_db["Día"].value_counts().reindex(DIAS_SEMANA, fill_value=0)
        st.bar_chart(df_dias)
        st.markdown("---")
        col_m1, col_m2 = st.columns(2)
        with col_m1:
            st.markdown("#### 🎯 Distribución por Modalidad")
            df_modalidad = df_cronograma_db["Modalidad"].value_counts()
            st.bar_chart(df_modalidad)
        with col_m2:
            st.markdown("#### 🧑‍🏫 Visitas por Docente (Top 10)")
            df_docentes = df_cronograma_db["Docente"].value_counts().head(10)
            st.bar_chart(df_docentes)
        st.markdown("---")
        st.markdown("#### 📋 Resumen Detallado por Docente")
        df_resumen = df_cronograma_db.groupby("Docente").agg(
            Visitas=("Docente", "count"), Módulos=("Módulo", "nunique"), Días=("Día", "nunique"),
        ).reset_index().sort_values("Visitas", ascending=False)
        st.dataframe(df_resumen, use_container_width=True, hide_index=True)

# ═══════════════════════════════════════════════════════════════════════════
# TAB 4: EXPORTAR
# ═══════════════════════════════════════════════════════════════════════════
with tab_export:
    st.markdown('<div class="crono-section-title">📥 Exportar Documentos</div>', unsafe_allow_html=True)
    if df_cronograma_db.empty:
        st.info("📭 No hay datos para exportar.")
    else:
        centro_export = st.session_state.get("crono_centro", "Politécnico Salesiano Arquides Calderón")
        coordinador_export = st.session_state.get("coordinador_nombre", "Ing. Bernardo Antonio Hernández Batista")
        info_datos = {"centro": centro_export, "coordinador": coordinador_export, "trimestre": trimestre_sel, "anio": anio_escolar}

        col_e1, col_e2 = st.columns(2)
        with col_e1:
            st.markdown("#### 📄 Documento Word")
            st.caption("Tabla profesional con encabezados sombreados y firmas.")
            buffer_word = generar_word_cronograma(df_cronograma_db, info_datos)
            st.download_button(label="📥 Descargar Word (.docx)", data=buffer_word,
                               file_name=f"Cronograma_{trimestre_sel.replace(' ', '_')}.docx",
                               mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                               type="primary", use_container_width=True)
        with col_e2:
            st.markdown("#### 📊 Hoja Excel")
            st.caption("Datos completos en formato tabular.")
            buffer_excel = generar_excel_cronograma(df_cronograma_db)
            st.download_button(label="📥 Descargar Excel (.xlsx)", data=buffer_excel,
                               file_name=f"Cronograma_{trimestre_sel.replace(' ', '_')}.docx".replace(".docx", ".xlsx"),
                               mime="application/vnd.openxmlformats-officedocument.spreadsheetml.sheet",
                               use_container_width=True)
        st.markdown("---")
        st.markdown("#### 👁️ Vista Previa")
        st.dataframe(df_cronograma_db, use_container_width=True, hide_index=True)