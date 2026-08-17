"""
generador_guia.py — Generador de Libros y Guías Educativas ETP (REGENERADO · limpio)
Generación en 2 fases: estructura del libro → contenido por capítulos.
Exportación Word profesional con portada, índice TOC, paginación y formato editorial.
• IA vía core/ia (solicitar_json, marcadores, reintento, auditoría).
• Super interfaz: hero, stepper, cards por capítulo, métricas, toasts.
• Anclaje opcional a PDF curricular.
"""
import re
from datetime import datetime
from io import BytesIO

import streamlit as st
from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement, parse_xml
from docx.oxml.ns import qn, nsdecls
from docx.shared import Inches, Pt, RGBColor

try:
    from pypdf import PdfReader
except Exception:
    try:
        from PyPDF2 import PdfReader
    except Exception:
        PdfReader = None

from core import ia

# ═══════════════════════════════════════════════════════════════════════════
# SUPER INTERFAZ — ESTILOS (limpios)
# ═══════════════════════════════════════════════════════════════════════════
st.markdown("""
<style>
@import url('https://fonts.googleapis.com/css2?family=Inter:wght@300;400;500;600;700;800;900&display=swap');
html, body, [class*="css"] { font-family: 'Inter', sans-serif; background-color: #F0F4F8; color: #1E293B; }

.guia-hero {
    background: linear-gradient(135deg, #1E1B4B 0%, #4C1D95 40%, #7C3AED 70%, #A78BFA 100%);
    color: #fff; padding: 2.2rem; border-radius: 20px; margin-bottom: 1.5rem;
    box-shadow: 0 25px 50px rgba(76, 29, 149, 0.3); position: relative; overflow: hidden;
}
.guia-hero::before { content: '📖'; position: absolute; right: 2rem; top: 50%;
    transform: translateY(-50%); font-size: 6rem; opacity: 0.15; }
.guia-hero-title { font-size: 2.4rem; font-weight: 900; letter-spacing: -0.03em; margin-bottom: 0.4rem; }
.guia-hero-sub { font-size: 1.05rem; opacity: 0.9; line-height: 1.5; }
.guia-hero-badge { display: inline-block; background: rgba(255,255,255,0.15);
    border: 1px solid rgba(255,255,255,0.25); border-radius: 8px; padding: 4px 12px;
    font-size: 0.8rem; font-weight: 600; margin-top: 0.8rem; margin-right: 8px; }

.stepper-container { display: flex; align-items: center; justify-content: center; gap: 0; margin: 1.5rem 0; }
.stepper-step { display: flex; flex-direction: column; align-items: center; gap: 6px; }
.stepper-circle { width: 48px; height: 48px; border-radius: 50%; display: flex;
    align-items: center; justify-content: center; font-weight: 800; font-size: 1.2rem;
    transition: all 0.3s ease; }
.stepper-circle.inactive { background: #E2E8F0; color: #94A3B8; border: 3px solid #CBD5E1; }
.stepper-circle.active { background: linear-gradient(135deg, #7C3AED, #A78BFA); color: #fff;
    border: 3px solid #7C3AED; box-shadow: 0 4px 15px rgba(124, 58, 237, 0.4); }
.stepper-circle.done { background: linear-gradient(135deg, #059669, #34D399); color: #fff;
    border: 3px solid #059669; box-shadow: 0 4px 15px rgba(5, 150, 105, 0.3); }
.stepper-label { font-size: 0.75rem; font-weight: 700; color: #64748B;
    text-transform: uppercase; letter-spacing: 0.03em; }
.stepper-line { width: 80px; height: 3px; background: #CBD5E1; margin: 0 4px; margin-bottom: 22px; }
.stepper-line.done { background: linear-gradient(90deg, #059669, #34D399); }

.capitulo-card { background: #fff; border: 2px solid #E2E8F0; border-left: 5px solid #7C3AED;
    border-radius: 12px; padding: 1rem 1.2rem; margin-bottom: 0.8rem; transition: all 0.2s ease; }
.capitulo-card:hover { border-color: #7C3AED; transform: translateX(4px);
    box-shadow: 0 4px 15px rgba(124, 58, 237, 0.12); }
.capitulo-num { font-size: 0.7rem; font-weight: 800; color: #7C3AED;
    text-transform: uppercase; letter-spacing: 0.05em; }
.capitulo-titulo { font-weight: 700; font-size: 1rem; color: #0F172A; margin-top: 2px; }
.capitulo-desc { font-size: 0.82rem; color: #64748B; margin-top: 4px; line-height: 1.4; }

.guia-stat { background: #fff; border-radius: 12px; padding: 1rem; text-align: center;
    border: 1px solid #E2E8F0; box-shadow: 0 2px 6px rgba(0,0,0,0.04); }
.guia-stat-value { font-size: 1.8rem; font-weight: 800; color: #4C1D95; }
.guia-stat-label { font-size: 0.72rem; font-weight: 700; color: #94A3B8;
    text-transform: uppercase; letter-spacing: 0.05em; }

.guia-section-title { color: #4C1D95; font-weight: 700; font-size: 1.12rem;
    border-bottom: 2px solid #EDE9FE; padding-bottom: 8px; margin: 1.2rem 0 0.9rem 0; }
</style>
""", unsafe_allow_html=True)

# ═══════════════════════════════════════════════════════════════════════════
# ESTADO DE SESIÓN
# ═══════════════════════════════════════════════════════════════════════════
def init_estado():
    if "guia_fase" not in st.session_state:
        st.session_state.guia_fase = 0  # 0=config, 1=estructura, 2=contenido, 3=listo
    if "guia_estructura" not in st.session_state:
        st.session_state.guia_estructura = None
    if "guia_capitulos_contenido" not in st.session_state:
        st.session_state.guia_capitulos_contenido = []
    if "guia_meta" not in st.session_state:
        st.session_state.guia_meta = {}
    if "guia_buffer" not in st.session_state:
        st.session_state.guia_buffer = None

init_estado()

# ═══════════════════════════════════════════════════════════════════════════
# UTILIDADES
# ═══════════════════════════════════════════════════════════════════════════
def extraer_texto_pdf(archivo, max_caracteres=80000):
    if PdfReader is None:
        raise RuntimeError("No hay librería PDF. Instala pypdf o PyPDF2.")
    archivo.seek(0)
    reader = PdfReader(archivo)
    texto = "".join((p.extract_text() or "") for p in reader.pages)
    texto = re.sub(r"\s+", " ", texto).strip()
    return texto[:max_caracteres]


def shade_cell(cell, color):
    shd = parse_xml(r'<w:shd {} w:fill="{}"/>'.format(nsdecls("w"), color))
    cell._tc.get_or_add_tcPr().append(shd)


def set_cell_text(cell, text, bold=False, center=False, color=None):
    cell.text = ""
    p = cell.paragraphs[0]
    run = p.add_run(str(text))
    run.bold = bold
    if center:
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    if color:
        shade_cell(cell, color)


# ── PAGINACIÓN (limpia, sin duplicados) ──
def _agregar_campo_numero(parrafo, instruccion):
    """Inserta un campo dinámico (PAGE o NUMPAGES) en el párrafo dado."""
    run = parrafo.add_run()
    fld_begin = OxmlElement("w:fldChar")
    fld_begin.set(qn("w:fldCharType"), "begin")
    instr = OxmlElement("w:instrText")
    instr.set(qn("xml:space"), "preserve")
    instr.text = instruccion
    fld_end = OxmlElement("w:fldChar")
    fld_end.set(qn("w:fldCharType"), "end")
    run._r.append(fld_begin)
    run._r.append(instr)
    run._r.append(fld_end)
    return run


def add_page_number(doc):
    """Agrega 'Página X de Y' al pie de página de la primera sección."""
    pie = doc.sections[0].footer.paragraphs[0]
    pie.alignment = WD_ALIGN_PARAGRAPH.CENTER
    pie.text = ""
    r_texto = pie.add_run("Página ")
    _agregar_campo_numero(pie, "PAGE")
    r_sep = pie.add_run(" / ")
    _agregar_campo_numero(pie, "NUMPAGES")
    for r in (r_texto, r_sep):
        r.font.size = Pt(8)
        r.font.color.rgb = RGBColor(0x64, 0x74, 0x8B)


def add_toc(doc):
    """Inserta un índice (TOC) que se actualiza al abrir en Word."""
    p = doc.add_paragraph()
    run = p.add_run()
    fld = OxmlElement("w:fldSimple")
    fld.set(qn("w:instr"), r'TOC \o "1-3" \h \z \u')
    r = OxmlElement("w:r")
    t = OxmlElement("w:t")
    t.text = "Actualiza el índice: clic derecho → Actualizar campos"
    r.append(t)
    fld.append(r)
    run._r.append(fld)

# ═══════════════════════════════════════════════════════════════════════════
# PROMPTS (claves limpias, sin espacios)
# ═══════════════════════════════════════════════════════════════════════════
def prompt_fase1_estructura(titulo, audiencia, temas, num_capitulos, contexto_pdf):
    return f"""Actúa como un Autor Profesional de Libros Educativos y Diseñador Curricular ETP del MINERD.
Tu tarea es crear la ESTRUCTURA COMPLETA de un libro/guía educativa.

DATOS DEL LIBRO:
Título: {titulo}
Audiencia: {audiencia}
Temas principales a cubrir: {temas}
Cantidad de capítulos solicitada: {num_capitulos}

{f"CONTEXTO CURRICULAR (PDF cargado - úsalo como base): {contexto_pdf[:3000]}" if contexto_pdf else "Sin PDF curricular: genera contenido estándar de alta calidad."}

REGLAS:
- Genera EXACTAMENTE {num_capitulos} capítulos.
- Cada capítulo debe tener entre 3 y 6 secciones.
- Incluye un glosario de al menos 8 términos.
- La estructura debe ser progresiva: de lo básico a lo avanzado.
- Cada sección debe tener un título claro y descriptivo.

Devuelve ÚNICAMENTE JSON válido con este formato exacto:
{{
  "TITULO": "{titulo}",
  "SUBTITULO": "Subtítulo del libro...",
  "INTRODUCCION": "Introducción general del libro (2-3 párrafos)...",
  "OBJETIVOS": ["Objetivo 1", "Objetivo 2", "Objetivo 3"],
  "CAPITULOS": [
    {{
      "NUMERO": 1,
      "TITULO": "Título del Capítulo 1",
      "RESUMEN": "Resumen breve de este capítulo (1-2 oraciones)",
      "SECCIONES": [
        {{
          "TITULO": "Título de la sección",
          "TIPO": "teoria|practica|ejemplo|actividad",
          "DESCRIPCION": "Qué cubrirá esta sección"
        }}
      ]
    }}
  ],
  "GLOSARIO": [
    {{ "TERMINO": "...", "DEFINICION": "...", "EJEMPLO": "..." }}
  ],
  "BIBLIOGRAFIA": ["Referencia 1", "Referencia 2"]
}}
"""


def prompt_fase2_capitulo(capitulo_estructura, titulo_libro, audiencia, contexto_pdf):
    secciones_txt = "\n".join([
        f"- {s.get('TITULO', '')} ({s.get('TIPO', 'teoria')}): {s.get('DESCRIPCION', '')}"
        for s in capitulo_estructura.get("SECCIONES", [])
    ])
    return f"""Actúa como un Autor Profesional de Libros Educativos ETP.
Escribe el CONTENIDO COMPLETO del Capítulo {capitulo_estructura.get('NUMERO', 1)} del libro "{titulo_libro}".

DATOS DEL CAPÍTULO:
Título: {capitulo_estructura.get('TITULO', '')}
Resumen: {capitulo_estructura.get('RESUMEN', '')}
Secciones a desarrollar:
{secciones_txt}

Audiencia: {audiencia}
{f"CONTEXTO CURRICULAR: {contexto_pdf[:2000]}" if contexto_pdf else ""}

REGLAS:
- Escribe contenido completo, detallado y profesional para CADA sección.
- Usa lenguaje apropiado para la audiencia indicada.
- Incluye ejemplos prácticos relevantes para la ETP.
- Cada sección debe tener entre 150 y 400 palabras.
- Si el tipo es "actividad", incluye instrucciones paso a paso.
- Si el tipo es "ejemplo", incluye un caso concreto resuelto.
- Mantén coherencia con el resto del libro.

Devuelve ÚNICAMENTE JSON válido:
{{
  "NUMERO_CAPITULO": {capitulo_estructura.get('NUMERO', 1)},
  "TITULO_CAPITULO": "{capitulo_estructura.get('TITULO', '')}",
  "INTRODUCCION_CAPITULO": "Introducción del capítulo (1 párrafo)...",
  "SECCIONES_CONTENIDO": [
    {{
      "TITULO": "Título de la sección",
      "CONTENIDO": "Contenido completo de la sección...",
      "TIPO": "teoria|practica|ejemplo|actividad"
    }}
  ],
  "RESUMEN_CAPITULO": "Resumen del capítulo (1 párrafo)...",
  "PREGUNTAS_REVISION": ["Pregunta 1", "Pregunta 2", "Pregunta 3"]
}}
"""

# ═══════════════════════════════════════════════════════════════════════════
# GENERACIÓN WORD PROFESIONAL
# ═══════════════════════════════════════════════════════════════════════════
def build_guia_docx(estructura, capitulos_contenido, meta):
    doc = Document()
    doc.styles["Normal"].font.name = "Calibri"
    doc.styles["Normal"].font.size = Pt(11)
    for section in doc.sections:
        section.left_margin = Inches(0.9)
        section.right_margin = Inches(0.9)
        section.top_margin = Inches(0.8)
        section.bottom_margin = Inches(0.8)

    # ─── PORTADA ───
    for _ in range(4):
        doc.add_paragraph()
    p_inst = doc.add_paragraph()
    p_inst.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run_inst = p_inst.add_run(meta.get("institucion", ""))
    run_inst.bold = True
    run_inst.font.size = Pt(14)
    run_inst.font.color.rgb = RGBColor(0x1E, 0x1B, 0x4B)
    doc.add_paragraph()

    p_titulo = doc.add_paragraph()
    p_titulo.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run_t = p_titulo.add_run(estructura.get("TITULO", meta.get("titulo", "Guía Educativa")))
    run_t.bold = True
    run_t.font.size = Pt(28)
    run_t.font.color.rgb = RGBColor(0x4C, 0x1D, 0x95)

    if estructura.get("SUBTITULO"):
        p_sub = doc.add_paragraph()
        p_sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run_sub = p_sub.add_run(estructura["SUBTITULO"])
        run_sub.font.size = Pt(14)
        run_sub.font.color.rgb = RGBColor(0x64, 0x74, 0x8B)

    doc.add_paragraph()
    doc.add_paragraph()

    p_autor = doc.add_paragraph()
    p_autor.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run_autor = p_autor.add_run(f"Autor: {meta.get('docente', '')}")
    run_autor.font.size = Pt(12)

    p_fecha = doc.add_paragraph()
    p_fecha.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run_fecha = p_fecha.add_run(f"{meta.get('audience', '')} · {datetime.now().strftime('%Y')}")
    run_fecha.font.size = Pt(11)
    run_fecha.font.color.rgb = RGBColor(0x94, 0xA3, 0xB8)
    doc.add_page_break()

    # ─── ÍNDICE ───
    doc.add_heading("Índice", level=1)
    add_toc(doc)
    doc.add_page_break()

    # ─── INTRODUCCIÓN ───
    if estructura.get("INTRODUCCION"):
        doc.add_heading("Introducción", level=1)
        p_intro = doc.add_paragraph(estructura["INTRODUCCION"])
        p_intro.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        doc.add_paragraph()
    if estructura.get("OBJETIVOS"):
        doc.add_heading("Objetivos del Libro", level=2)
        for obj in estructura["OBJETIVOS"]:
            p_obj = doc.add_paragraph(obj, style="List Bullet")
            p_obj.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        doc.add_paragraph()
    doc.add_page_break()

    # ─── CAPÍTULOS ───
    for cap_content in capitulos_contenido:
        num_cap = cap_content.get("NUMERO_CAPITULO", 0)
        titulo_cap = cap_content.get("TITULO_CAPITULO", f"Capítulo {num_cap}")
        doc.add_heading(f"Capítulo {num_cap}: {titulo_cap}", level=1)
        if cap_content.get("INTRODUCCION_CAPITULO"):
            p_int = doc.add_paragraph(cap_content["INTRODUCCION_CAPITULO"])
            p_int.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
            doc.add_paragraph()
        for sec in cap_content.get("SECCIONES_CONTENIDO", []):
            tipo = sec.get("TIPO", "teoria")
            titulo_sec = sec.get("TITULO", "")
            contenido_sec = sec.get("CONTENIDO", "")
            if tipo == "actividad":
                doc.add_heading(f"📋 Actividad: {titulo_sec}", level=2)
            elif tipo == "ejemplo":
                doc.add_heading(f"💡 Ejemplo: {titulo_sec}", level=2)
            elif tipo == "practica":
                doc.add_heading(f"🔧 Práctica: {titulo_sec}", level=2)
            else:
                doc.add_heading(titulo_sec, level=2)
            p_cont = doc.add_paragraph(contenido_sec)
            p_cont.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
            doc.add_paragraph()
        if cap_content.get("RESUMEN_CAPITULO"):
            doc.add_heading("Resumen del Capítulo", level=2)
            p_res = doc.add_paragraph(cap_content["RESUMEN_CAPITULO"])
            p_res.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
            doc.add_paragraph()
        if cap_content.get("PREGUNTAS_REVISION"):
            doc.add_heading("Preguntas de Revisión", level=2)
            for i, preg in enumerate(cap_content["PREGUNTAS_REVISION"], 1):
                doc.add_paragraph(f"{i}. {preg}", style="List Number")
        doc.add_page_break()

    # ─── GLOSARIO ───
    if estructura.get("GLOSARIO"):
        doc.add_heading("Glosario", level=1)
        tabla_g = doc.add_table(rows=1, cols=3)
        tabla_g.style = "Table Grid"
        hdr = tabla_g.rows[0].cells
        set_cell_text(hdr[0], "Término", bold=True, color="EDE9FE")
        set_cell_text(hdr[1], "Definición", bold=True, color="EDE9FE")
        set_cell_text(hdr[2], "Ejemplo", bold=True, color="EDE9FE")
        for item in estructura["GLOSARIO"]:
            row = tabla_g.add_row().cells
            set_cell_text(row[0], item.get("TERMINO", ""), bold=True)
            set_cell_text(row[1], item.get("DEFINICION", ""))
            set_cell_text(row[2], item.get("EJEMPLO", ""))
        doc.add_page_break()

    # ─── BIBLIOGRAFÍA ───
    if estructura.get("BIBLIOGRAFIA"):
        doc.add_heading("Bibliografía y Referencias", level=1)
        for ref in estructura["BIBLIOGRAFIA"]:
            doc.add_paragraph(ref, style="List Bullet")

    # ─── PAGINACIÓN ───
    add_page_number(doc)

    buffer = BytesIO()
    doc.save(buffer)
    buffer.seek(0)
    return buffer

# ═══════════════════════════════════════════════════════════════════════════
# RENDERIZADO DE STEPPER
# ═══════════════════════════════════════════════════════════════════════════
def render_stepper(fase_actual):
    pasos = [
        ("1", "Configurar"),
        ("2", "Estructura"),
        ("3", "Capítulos"),
        ("4", "Descargar"),
    ]
    html = '<div class="stepper-container">'
    for i, (num, label) in enumerate(pasos):
        if i + 1 < fase_actual:
            estado = "done"
        elif i + 1 == fase_actual:
            estado = "active"
        else:
            estado = "inactive"
        html += f'''
        <div class="stepper-step">
            <div class="stepper-circle {estado}">{"✓" if estado == "done" else num}</div>
            <div class="stepper-label">{label}</div>
        </div>
        '''
        if i < len(pasos) - 1:
            line_class = "done" if i + 1 < fase_actual else ""
            html += f'<div class="stepper-line {line_class}"></div>'
    html += '</div>'
    st.markdown(html, unsafe_allow_html=True)

# ═══════════════════════════════════════════════════════════════════════════
# HERO Y STEPPER
# ═══════════════════════════════════════════════════════════════════════════
ia.panel_sidebar_ia("Generador de Libros/Guías")

st.markdown("""
<div class="guia-hero">
    <div class="guia-hero-title">📖 Generador de Libros y Guías Educativas</div>
    <div class="guia-hero-sub">
        Creación profesional de libros y guías ETP en 2 fases: estructura → contenido por capítulos
    </div>
    <div>
        <span class="guia-hero-badge">📚 Generación en 2 Fases</span>
        <span class="guia-hero-badge">🤖 Asistido por IA</span>
        <span class="guia-hero-badge">📄 Word Profesional</span>
        <span class="guia-hero-badge">📑 Índice TOC</span>
    </div>
</div>
""", unsafe_allow_html=True)

render_stepper(st.session_state.guia_fase + 1)

# ═══════════════════════════════════════════════════════════════════════════
# FASE 1: CONFIGURACIÓN
# ═══════════════════════════════════════════════════════════════════════════
if st.session_state.guia_fase == 0:
    st.markdown('<div class="guia-section-title">📋 Configuración del Libro</div>', unsafe_allow_html=True)
    with st.form("form_guia_config"):
        col1, col2 = st.columns(2)
        with col1:
            titulo = st.text_input("Título del Libro/Guía", placeholder="Ej: Guía de Impuestos al Consumo")
            audiencia = st.text_input("Audiencia", placeholder="Ej: Estudiantes de 5to Bachillerato Técnico")
            docente = st.text_input("Autor/Docente", value="Ing. Bernardo Antonio Hernández Batista")
        with col2:
            institucion = st.text_input("Institución", value="Politécnico Salesiano Arquides Calderón")
            num_capitulos = st.slider("Número de capítulos", 3, 15, 6)
        temas = st.text_area(
            "Temas principales (separados por coma)",
            height=80,
            placeholder="Ej: Impuestos al consumo, ITBIS, ISC, importación de vehículos, matrícula, placas",
        )
        st.markdown('<div class="guia-section-title">📄 Fuente Curricular (Opcional)</div>', unsafe_allow_html=True)
        archivo_pdf = st.file_uploader("Subir PDF curricular para anclar el contenido", type=["pdf"])
        max_tokens, temperature = ia.control_avanzado(default_tokens=16384, tope=32000, default_temp=0.2)
        btn_generar = st.form_submit_button(
            "⚡ Fase 1: Generar Estructura del Libro",
            type="primary",
            width="stretch",
        )

    if btn_generar:
        if not titulo or not temas:
            st.warning("⚠️ Completa al menos el título y los temas.")
        else:
            contexto_pdf = ""
            if archivo_pdf:
                try:
                    contexto_pdf = extraer_texto_pdf(archivo_pdf)
                    st.info(f"📄 PDF cargado: {len(contexto_pdf)} caracteres extraídos.")
                except Exception as e:
                    st.warning(f"⚠️ No se pudo leer el PDF: {e}")
            with st.spinner("🧠 Generando estructura del libro con IA..."):
                try:
                    prompt = prompt_fase1_estructura(titulo, audiencia, temas, num_capitulos, contexto_pdf)
                    resultado, flags = ia.solicitar_json(
                        prompt, max_tokens=max_tokens, temperature=temperature, modulo="generador_guia"
                    )
                    st.session_state.guia_estructura = resultado
                    st.session_state.guia_meta = {
                        "titulo": titulo,
                        "audience": audiencia,
                        "docente": docente,
                        "institucion": institucion,
                        "temas": temas,
                        "contexto_pdf": contexto_pdf[:2000],
                        "max_tokens": max_tokens,
                        "temperature": temperature,
                    }
                    st.session_state.guia_fase = 1
                    st.session_state.guia_capitulos_contenido = []
                    st.session_state.guia_buffer = None
                    st.toast("✅ Estructura del libro generada.", icon="📖")
                    st.rerun()
                except Exception as e:
                    ia.render_error_ia(e)

# ═══════════════════════════════════════════════════════════════════════════
# FASE 2: ESTRUCTURA GENERADA → GENERAR CAPÍTULOS
# ═══════════════════════════════════════════════════════════════════════════
elif st.session_state.guia_fase == 1:
    estructura = st.session_state.guia_estructura
    st.markdown('<div class="guia-section-title">📖 Estructura del Libro Generada</div>', unsafe_allow_html=True)

    cols = st.columns(4)
    with cols[0]:
        st.markdown(f'<div class="guia-stat"><div class="guia-stat-value">{len(estructura.get("CAPITULOS", []))}</div><div class="guia-stat-label">Capítulos</div></div>', unsafe_allow_html=True)
    with cols[1]:
        total_secciones = sum(len(c.get("SECCIONES", [])) for c in estructura.get("CAPITULOS", []))
        st.markdown(f'<div class="guia-stat"><div class="guia-stat-value">{total_secciones}</div><div class="guia-stat-label">Secciones</div></div>', unsafe_allow_html=True)
    with cols[2]:
        st.markdown(f'<div class="guia-stat"><div class="guia-stat-value">{len(estructura.get("GLOSARIO", []))}</div><div class="guia-stat-label">Glosario</div></div>', unsafe_allow_html=True)
    with cols[3]:
        st.markdown(f'<div class="guia-stat"><div class="guia-stat-value">{len(estructura.get("BIBLIOGRAFIA", []))}</div><div class="guia-stat-label">Referencias</div></div>', unsafe_allow_html=True)

    st.markdown(f"#### {estructura.get('TITULO', '')}")
    if estructura.get("SUBTITULO"):
        st.caption(estructura["SUBTITULO"])
    st.markdown("---")

    for cap in estructura.get("CAPITULOS", []):
        st.markdown(f"""
        <div class="capitulo-card">
            <div class="capitulo-num">Capítulo {cap.get('NUMERO', '')}</div>
            <div class="capitulo-titulo">{cap.get('TITULO', '')}</div>
            <div class="capitulo-desc">{cap.get('RESUMEN', '')}</div>
            <div class="capitulo-desc" style="margin-top:6px; font-size:0.75rem; color:#7C3AED;">
                📑 {len(cap.get('SECCIONES', []))} secciones
            </div>
        </div>
        """, unsafe_allow_html=True)

    st.markdown("---")
    col_btn1, col_btn2 = st.columns(2)
    with col_btn1:
        if st.button("⚡ Fase 2: Generar Contenido de Todos los Capítulos", type="primary", width="stretch"):
            st.session_state.guia_fase = 2
            st.session_state.guia_capitulos_contenido = []
            st.rerun()
    with col_btn2:
        if st.button("🔄 Regenerar Estructura", width="stretch"):
            st.session_state.guia_fase = 0
            st.session_state.guia_estructura = None
            st.session_state.guia_buffer = None
            st.rerun()

# ═══════════════════════════════════════════════════════════════════════════
# FASE 3: GENERANDO CAPÍTULOS
# ═══════════════════════════════════════════════════════════════════════════
elif st.session_state.guia_fase == 2:
    estructura = st.session_state.guia_estructura
    meta = st.session_state.guia_meta
    capitulos = estructura.get("CAPITULOS", [])

    st.markdown('<div class="guia-section-title">✍️ Generando Contenido de los Capítulos</div>', unsafe_allow_html=True)

    progress_placeholder = st.empty()
    status_placeholder = st.empty()

    for idx, cap in enumerate(capitulos):
        progress_placeholder.progress((idx) / max(len(capitulos), 1), text=f"Capítulo {idx+1} de {len(capitulos)}")
        status_placeholder.info(f"✍️ Escribiendo: {cap.get('TITULO', f'Capítulo {idx+1}')}...")
        try:
            prompt = prompt_fase2_capitulo(
                cap, meta.get("titulo", ""), meta.get("audience", ""), meta.get("contexto_pdf", "")
            )
            resultado_cap, flags = ia.solicitar_json(
                prompt,
                max_tokens=meta.get("max_tokens", 16384),
                temperature=meta.get("temperature", 0.2),
                modulo="generador_guia_cap",
            )
            st.session_state.guia_capitulos_contenido.append(resultado_cap)
        except Exception as e:
            st.warning(f"⚠️ Error generando capítulo {idx+1}: {e}")
            st.session_state.guia_capitulos_contenido.append({
                "NUMERO_CAPITULO": cap.get("NUMERO", idx + 1),
                "TITULO_CAPITULO": cap.get("TITULO", f"Capítulo {idx+1}"),
                "INTRODUCCION_CAPITULO": f"[Error al generar: {str(e)[:100]}]",
                "SECCIONES_CONTENIDO": [],
                "RESUMEN_CAPITULO": "",
                "PREGUNTAS_REVISION": [],
            })

    progress_placeholder.progress(1.0, text="¡Completado!")
    status_placeholder.success("✅ Todos los capítulos generados.")
    st.session_state.guia_fase = 3
    st.toast("📖 ¡Libro completo generado!", icon="🎉")
    st.rerun()

# ═══════════════════════════════════════════════════════════════════════════
# FASE 4: RESULTADO FINAL Y DESCARGA
# ═══════════════════════════════════════════════════════════════════════════
elif st.session_state.guia_fase == 3:
    estructura = st.session_state.guia_estructura
    capitulos_contenido = st.session_state.guia_capitulos_contenido
    meta = st.session_state.guia_meta

    st.markdown('<div class="guia-section-title">🎉 Libro Generado Completamente</div>', unsafe_allow_html=True)

    total_palabras = sum(
        len(sec.get("CONTENIDO", "").split())
        for cap in capitulos_contenido
        for sec in cap.get("SECCIONES_CONTENIDO", [])
    )

    cols = st.columns(4)
    with cols[0]:
        st.markdown(f'<div class="guia-stat"><div class="guia-stat-value">{len(capitulos_contenido)}</div><div class="guia-stat-label">Capítulos</div></div>', unsafe_allow_html=True)
    with cols[1]:
        st.markdown(f'<div class="guia-stat"><div class="guia-stat-value">{total_palabras:,}</div><div class="guia-stat-label">Palabras</div></div>', unsafe_allow_html=True)
    with cols[2]:
        st.markdown(f'<div class="guia-stat"><div class="guia-stat-value">{len(estructura.get("GLOSARIO", []))}</div><div class="guia-stat-label">Glosario</div></div>', unsafe_allow_html=True)
    with cols[3]:
        tiempo_lectura = max(1, total_palabras // 200)
        st.markdown(f'<div class="guia-stat"><div class="guia-stat-value">{tiempo_lectura} min</div><div class="guia-stat-label">Lectura Est.</div></div>', unsafe_allow_html=True)

    st.markdown("---")

    tab_preview, tab_json, tab_debug = st.tabs(["👁️ Vista Previa", "🧾 JSON Completo", "🐛 Depuración"])
    with tab_preview:
        st.markdown(f"## {estructura.get('TITULO', '')}")
        if estructura.get("INTRODUCCION"):
            with st.expander("📄 Introducción"):
                st.write(estructura["INTRODUCCION"])
        for cap in capitulos_contenido:
            with st.expander(f"📖 Capítulo {cap.get('NUMERO_CAPITULO', '')}: {cap.get('TITULO_CAPITULO', '')}"):
                if cap.get("INTRODUCCION_CAPITULO"):
                    st.write(cap["INTRODUCCION_CAPITULO"])
                for sec in cap.get("SECCIONES_CONTENIDO", []):
                    st.markdown(f"**{sec.get('TITULO', '')}**")
                    contenido = sec.get("CONTENIDO", "")
                    st.write(contenido[:500] + "..." if len(contenido) > 500 else contenido)
    with tab_json:
        st.json({"estructura": estructura, "capitulos": capitulos_contenido})
    with tab_debug:
        st.write("Meta:", meta)
        st.write(f"Capítulos generados: {len(capitulos_contenido)}")

    st.markdown("---")

    col_dl1, col_dl2 = st.columns(2)
    with col_dl1:
        if st.button("📥 Generar Documento Word (.docx)", type="primary", width="stretch"):
            with st.spinner("📄 Construyendo documento Word profesional..."):
                st.session_state.guia_buffer = build_guia_docx(estructura, capitulos_contenido, meta)

    if st.session_state.guia_buffer is not None:
        nombre_archivo = ia.sanear_nombre_archivo(meta.get("titulo", "guia"))
        st.download_button(
            label="⬇️ Descargar Libro/Guía (.docx)",
            data=st.session_state.guia_buffer,
            file_name=f"{nombre_archivo}.docx",
            mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
            type="primary",
            width="stretch",
        )

    with col_dl2:
        if st.button("🔄 Crear Nuevo Libro", width="stretch"):
            st.session_state.guia_fase = 0
            st.session_state.guia_estructura = None
            st.session_state.guia_capitulos_contenido = []
            st.session_state.guia_meta = {}
            st.session_state.guia_buffer = None
            st.rerun()