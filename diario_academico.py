"""
diario_academico.py — Planificación de Clase Diaria Académica (Paso 26 v2 · Nivel Dios)
Identificación inteligente de malla curricular y secuencias didácticas MINERD.
• Selector de 9 materias académicas + talleres optativos (sin malla).
• Extracción automática de secuencias didácticas desde la malla curricular.
• Auto-identificación de competencias, contenidos e indicadores.
• IA vía core/ia + Word profesional MINERD académico.
"""
import re
from datetime import datetime, date
from io import BytesIO
from typing import Any, Dict, List, Optional

import pandas as pd
import streamlit as st
from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement, parse_xml
from docx.oxml.ns import qn, nsdecls
from docx.shared import Inches, Pt, RGBColor

try:
    from pypdf import PdfReader
except Exception:
    try:
        from PyPDF2 import PdfReader
    except Exception:
        PdfReader = None

from core import ia

# ═══════════════════════════════════════════════════════════════════════════
# CATÁLOGOS MINERD
# ═══════════════════════════════════════════════════════════════════════════
AREAS_ACADEMICAS = [
    {"nombre": "Matemáticas", "icono": "🔢", "color": "#3B82F6"},
    {"nombre": "Lengua Española", "icono": "📖", "color": "#EF4444"},
    {"nombre": "Ciencias Sociales", "icono": "🌍", "color": "#F59E0B"},
    {"nombre": "Ciencias de la Naturaleza", "icono": "🔬", "color": "#10B981"},
    {"nombre": "Inglés", "icono": "🇬🇧", "color": "#8B5CF6"},
    {"nombre": "Francés", "icono": "🇫🇷", "color": "#06B6D4"},
    {"nombre": "Educación Física", "icono": "⚽", "color": "#84CC16"},
    {"nombre": "Educación Artística", "icono": "🎨", "color": "#EC4899"},
    {"nombre": "Formación Integral Humana y Religiosa", "icono": "🕊️", "color": "#A78BFA"},
]

TALLERES_OPTATIVOS = [
    "Emprendimiento",
    "Formación y Orientación Laboral (FOL)",
    "Tecnologías Digitales / Ofimática",
    "Formación en Centros de Trabajo (FCT)",
    "Robótica y Programación",
    "Inglés Técnico",
    "Educación Financiera",
    "Liderazgo y Trabajo en Equipo",
]

COMPETENCIAS_FUNDAMENTALES = [
    "Ética y Ciudadana",
    "Comunicativa",
    "Pensamiento Lógico, Creativo y Crítico",
    "Resolución de Problemas",
    "Científica y Tecnológica",
    "Ambiental y de la Salud",
    "Desarrollo Personal y Espiritual",
]

ESTRATEGIAS_EA = [
    "Indagación dialógica",
    "Estudio de Caso",
    "Aprendizaje basado en problemas",
    "Aprendizaje colaborativo",
    "Investigación dirigida",
    "Juego de roles",
    "Debate",
    "Taller práctico",
    "Aprendizaje por descubrimiento",
    "Exposición dialogada",
    "Método de proyectos",
]

MODOS_PLANIFICACION = {
    "secuencia_didactica": {
        "nombre": "Secuencia Didáctica",
        "icono": "📅",
        "descripcion": "Planificación dentro de una secuencia didáctica de la malla curricular.",
        "color": "#0D9488",
    },
    "situacion_aprendizaje": {
        "nombre": "Situación de Aprendizaje",
        "icono": "🎯",
        "descripcion": "Planificación contextualizada en una situación real o simulada.",
        "color": "#0891B2",
    },
}

# ═══════════════════════════════════════════════════════════════════════════
# SUPER INTERFAZ — ESTILOS
# ═══════════════════════════════════════════════════════════════════════════
st.markdown("""
<style>
@import url('https://fonts.googleapis.com/css2?family=Inter:wght@300;400;500;600;700;800;900&display=swap');

html, body, [class*="css"] {
    font-family: 'Inter', sans-serif;
    background-color: #F0F4F8;
    color: #1E293B;
}

.diario-hero {
    background: linear-gradient(135deg, #042F2E 0%, #0F766E 40%, #14B8A6 70%, #5EEAD4 100%);
    color: #fff;
    padding: 2.2rem;
    border-radius: 20px;
    margin-bottom: 1.5rem;
    box-shadow: 0 25px 50px rgba(15, 118, 110, 0.3);
    position: relative;
    overflow: hidden;
}

.diario-hero::before {
    content: '🗓️';
    position: absolute;
    right: 2rem;
    top: 50%;
    transform: translateY(-50%);
    font-size: 6rem;
    opacity: 0.15;
}

.diario-hero-title { font-size: 2.4rem; font-weight: 900; letter-spacing: -0.03em; margin-bottom: 0.4rem; }
.diario-hero-sub { font-size: 1.05rem; opacity: 0.9; line-height: 1.5; }

.diario-hero-badge {
    display: inline-block;
    background: rgba(255,255,255,0.15);
    border: 1px solid rgba(255,255,255,0.25);
    border-radius: 8px;
    padding: 4px 12px;
    font-size: 0.8rem;
    font-weight: 600;
    margin-top: 0.8rem;
    margin-right: 8px;
}

.diario-stepper {
    display: flex;
    align-items: center;
    justify-content: center;
    gap: 0;
    margin: 1.5rem 0;
}

.diario-step { display: flex; flex-direction: column; align-items: center; gap: 6px; }

.diario-step-circle {
    width: 48px; height: 48px; border-radius: 50%;
    display: flex; align-items: center; justify-content: center;
    font-weight: 800; font-size: 1.2rem;
    transition: all 0.3s ease;
}

.diario-step-circle.inactive { background: #E2E8F0; color: #94A3B8; border: 3px solid #CBD5E1; }
.diario-step-circle.active {
    background: linear-gradient(135deg, #0F766E, #14B8A6);
    color: #fff; border: 3px solid #0F766E;
    box-shadow: 0 4px 15px rgba(15, 118, 110, 0.4);
}
.diario-step-circle.done {
    background: linear-gradient(135deg, #059669, #34D399);
    color: #fff; border: 3px solid #059669;
    box-shadow: 0 4px 15px rgba(5, 150, 105, 0.3);
}

.diario-step-label { font-size: 0.72rem; font-weight: 700; color: #64748B; text-transform: uppercase; letter-spacing: 0.03em; }
.diario-step-line { width: 70px; height: 3px; background: #CBD5E1; margin: 0 4px; margin-bottom: 22px; }
.diario-step-line.done { background: linear-gradient(90deg, #059669, #34D399); }

.diario-area-card {
    background: #fff;
    border: 3px solid #E2E8F0;
    border-radius: 14px;
    padding: 1.2rem;
    transition: all 0.25s ease;
    cursor: pointer;
    text-align: center;
    box-shadow: 0 2px 8px rgba(0,0,0,0.04);
}

.diario-area-card:hover { transform: translateY(-4px); box-shadow: 0 10px 25px rgba(0,0,0,0.1); }
.diario-area-card.selected { border-color: #14B8A6; background: #F0FDFA; box-shadow: 0 6px 20px rgba(20, 184, 166, 0.2); }

.diario-area-icono { font-size: 2.5rem; margin-bottom: 0.6rem; }
.diario-area-nombre { font-weight: 700; font-size: 0.92rem; color: #0F172A; }

.diario-sec-card {
    background: #fff;
    border: 2px solid #E2E8F0;
    border-left: 5px solid #14B8A6;
    border-radius: 12px;
    padding: 1rem 1.2rem;
    margin-bottom: 0.8rem;
    transition: all 0.2s ease;
}

.diario-sec-card:hover { border-color: #14B8A6; transform: translateX(4px); box-shadow: 0 4px 15px rgba(20, 184, 166, 0.12); }
.diario-sec-card.selected { border-color: #0F766E; background: #F0FDFA; }

.diario-sec-nombre { font-weight: 700; font-size: 1rem; color: #0F172A; }
.diario-sec-meta { font-size: 0.82rem; color: #64748B; margin-top: 4px; }

.diario-section-title {
    color: #0F766E;
    font-weight: 700;
    font-size: 1.12rem;
    border-bottom: 2px solid #CCFBF1;
    padding-bottom: 8px;
    margin: 1.2rem 0 0.9rem 0;
}

.diario-stat {
    background: #fff;
    border-radius: 12px;
    padding: 1rem;
    text-align: center;
    border: 1px solid #E2E8F0;
    box-shadow: 0 2px 6px rgba(0,0,0,0.04);
}

.diario-stat-value { font-size: 1.8rem; font-weight: 800; color: #0F766E; }
.diario-stat-label { font-size: 0.72rem; font-weight: 700; color: #94A3B8; text-transform: uppercase; letter-spacing: 0.05em; }
</style>
""", unsafe_allow_html=True)

# ═══════════════════════════════════════════════════════════════════════════
# ESTADO DE SESIÓN
# ═══════════════════════════════════════════════════════════════════════════
def init_estado():
    if "diario_fase" not in st.session_state:
        st.session_state.diario_fase = 0  # 0=materia+malla, 1=secuencia, 2=plan
    if "diario_tipo" not in st.session_state:
        st.session_state.diario_tipo = None  # "academica" | "taller"
    if "diario_area" not in st.session_state:
        st.session_state.diario_area = None
    if "diario_secuencias" not in st.session_state:
        st.session_state.diario_secuencias = None
    if "diario_secuencia_sel" not in st.session_state:
        st.session_state.diario_secuencia_sel = None
    if "diario_resultado" not in st.session_state:
        st.session_state.diario_resultado = None
    if "diario_modo" not in st.session_state:
        st.session_state.diario_modo = None

init_estado()

# ═══════════════════════════════════════════════════════════════════════════
# UTILIDADES
# ═══════════════════════════════════════════════════════════════════════════
def extraer_texto_pdf(archivo, max_caracteres=100000):
    if PdfReader is None:
        raise RuntimeError("No hay librería PDF. Instala pypdf o PyPDF2.")
    archivo.seek(0)
    reader = PdfReader(archivo)
    texto = "".join((p.extract_text() or "") for p in reader.pages)
    texto = re.sub(r"\s+", " ", texto).strip()
    return texto[:max_caracteres]

def shade_cell(cell, color):
    shd = parse_xml(r'<w:shd {} w:fill="{}"/>'.format(nsdecls("w"), color))
    cell._tc.get_or_add_tcPr().append(shd)

def set_cell_text(cell, text, bold=False, center=False, color=None):
    cell.text = ""
    p = cell.paragraphs[0]
    run = p.add_run(str(text))
    run.bold = bold
    if center:
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    if color:
        shade_cell(cell, color)

def add_page_number(doc):
    for section in doc.sections:
        footer = section.footer
        footer.is_linked_to_previous = False
        p = footer.paragraphs[0] if footer.paragraphs else footer.add_paragraph()
        p.text = ""
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run()
        fld1 = OxmlElement("w:fldChar"); fld1.set(qn("w:fldCharType"), "begin")
        instr = OxmlElement("w:instrText"); instr.set(qn("xml:space"), "preserve"); instr.text = "PAGE"
        fld2 = OxmlElement("w:fldChar"); fld2.set(qn("w:fldCharType"), "end")
        run._r.append(fld1); run._r.append(instr); run._r.append(fld2)
        run.add_run(" / ")
        run2 = p.add_run()
        fld3 = OxmlElement("w:fldChar"); fld3.set(qn("w:fldCharType"), "begin")
        instr2 = OxmlElement("w:instrText"); instr2.set(qn("xml:space"), "preserve"); instr2.text = "NUMPAGES"
        fld4 = OxmlElement("w:fldChar"); fld4.set(qn("w:fldCharType"), "end")
        run2._r.append(fld3); run2._r.append(instr2); run2._r.append(fld4)
        run.font.size = Pt(8)
        run.font.color.rgb = RGBColor(0x64, 0x74, 0x8B)
        run2.font.size = Pt(8)
        run2.font.color.rgb = RGBColor(0x64, 0x74, 0x8B)

# ═══════════════════════════════════════════════════════════════════════════
# PROMPT 1: EXTRAER SECUENCIAS DIDÁCTICAS DE LA MALLA
# ═══════════════════════════════════════════════════════════════════════════
def prompt_extraer_secuencias(area, grado, texto_malla):
    return f"""Actúa como un Especialista Curricular del MINERD, experto en Diseño Curricular del Nivel Secundario.

Tu tarea es extraer TODAS las Secuencias Didácticas (o Situaciones de Aprendizaje) correspondientes al área de "{area}" para el grado "{grado}" a partir de la malla curricular proporcionada.

MALLA CURRICULAR:
{texto_malla}

REGLAS:
- Identifica CADA secuencia didáctica o situación de aprendizaje del grado indicado.
- Para cada secuencia, extrae: nombre, semanas, actividades, competencias específicas, contenidos (conceptos, procedimientos, actitudes, indicadores de logro) y situación de aprendizaje (si existe).
- Si la malla usa "Situación de Aprendizaje" en lugar de "Secuencia Didáctica", trátalas como equivalentes.
- Sé exhaustivo: no omitas ninguna secuencia del grado.

Devuelve ÚNICAMENTE JSON válido:
{{
  "AREA": "{area}",
  "GRADO": "{grado}",
  "SECUENCIAS": [
    {{
      "NUMERO": 1,
      "NOMBRE": "Nombre de la secuencia didáctica...",
      "SEMANAS": "1-2",
      "SITUACION_APRENDIZAJE": "Descripción de la situación de aprendizaje (si existe)...",
      "COMPETENCIAS_ESPECIFICAS": ["CE-XXX-...", "CE-YYY-..."],
      "CONTENIDOS": {{
        "CONCEPTOS": "Conceptos...",
        "PROCEDIMIENTOS": "Procedimientos...",
        "ACTITUDES": "Actitudes y valores...",
        "INDICADORES_LOGRO": "IL-X-..., IL-Y-..."
      }},
      "ACTIVIDADES": [
        {{"NUMERO": 1, "NOMBRE": "Nombre de la actividad..."}}
      ]
    }}
  ]
}}
"""

# ═══════════════════════════════════════════════════════════════════════════
# PROMPT 2: GENERAR PLAN DIARIO
# ═══════════════════════════════════════════════════════════════════════════
def prompt_generar_plan(datos_form, secuencia, modo):
    if secuencia:
        contexto_curricular = f"""
DATOS CURRICULARES IDENTIFICADOS AUTOMÁTICAMENTE DE LA MALLA:
Secuencia Didáctica: {secuencia.get('NOMBRE', '')}
Semanas: {secuencia.get('SEMANAS', '')}
Situación de Aprendizaje: {secuencia.get('SITUACION_APRENDIZAJE', 'No especificada')}
Competencias Específicas: {', '.join(secuencia.get('COMPETENCIAS_ESPECIFICAS', []))}
Conceptos: {secuencia.get('CONTENIDOS', {}).get('CONCEPTOS', '')}
Procedimientos: {secuencia.get('CONTENIDOS', {}).get('PROCEDIMIENTOS', '')}
Actitudes y Valores: {secuencia.get('CONTENIDOS', {}).get('ACTITUDES', '')}
Indicadores de Logro: {secuencia.get('CONTENIDOS', {}).get('INDICADORES_LOGRO', '')}
"""
    else:
        contexto_curricular = f"""
DATOS DEL TALLER OPTATIVO:
Taller: {datos_form.get('taller', '')}
Tema de la sesión: {datos_form.get('tema_taller', '')}
Nota: Los talleres optativos no requieren malla curricular. Genera el plan basado en el tema indicado.
"""

    competencias_fund_txt = "\n".join([f"- {c}" for c in datos_form.get("competencias_fundamentales", [])])
    estrategias_txt = "\n".join([f"- {e}" for e in datos_form.get("estrategias", [])])

    return f"""Actúa como un Especialista Curricular del MINERD, Experto en Planificación de Clase Diaria para el Nivel Secundario.

Tu tarea es diseñar una Planificación de Clase Diaria completa y profesional siguiendo el formato oficial MINERD.

DATOS GENERALES:
Centro Educativo: {datos_form.get('centro', '')}
Docente: {datos_form.get('docente', '')}
Área: {datos_form.get('area', '')}
Asignatura: {datos_form.get('asignatura', '')}
Grado y Sección: {datos_form.get('grado', '')}
Nivel: Secundario | Ciclo: {datos_form.get('ciclo', '')}
Duración: {datos_form.get('duracion', '50 minutos')}
Fecha: {datos_form.get('fecha', datetime.now().strftime('%d/%m/%Y'))}

{contexto_curricular}

COMPETENCIAS FUNDAMENTALES A DESARROLLAR:
{competencias_fund_txt}

ESTRATEGIAS DE ENSEÑANZA-APRENDIZAJE:
{estrategias_txt}

REGLAS DE DISEÑO:
- Los MOMENTOS PEDAGÓGICOS deben tener 3 fases cada uno: Inicio (10 min), Desarrollo (30 min), Cierre (10 min).
- Cada fase debe tener actividades concretas, recursos y tiempo estimado.
- El Inicio debe incluir: motivación, recuperación de saberes previos y presentación de la intención pedagógica.
- El Desarrollo debe incluir: construcción del aprendizaje con actividades prácticas y colaborativas.
- El Cierre debe incluir: metacognición, consolidación y preguntas de reflexión.
- La LISTA DE COTEJO debe tener entre 4 y 6 criterios alineados a los indicadores de logro.
- Las ADAPTACIONES NEAE deben ser específicas y aplicables.
- Las EVIDENCIAS deben ser tangibles y verificables.

Devuelve ÚNICAMENTE JSON válido con este formato exacto:
{{
  "INTENCION_PEDAGOGICA": "Descripción de la intención pedagógica del día...",
  "MOMENTOS_PEDAGOGICOS": {{
    "INICIO": {{
      "TIEMPO": "10 minutos",
      "FASES": [
        {{"FASE": "Motivación y activación", "ACTIVIDADES": "Descripción...", "RECURSOS": "Recursos..."}},
        {{"FASE": "Recuperación de saberes", "ACTIVIDADES": "Descripción...", "RECURSOS": "Recursos..."}},
        {{"FASE": "Presentación de intención", "ACTIVIDADES": "Descripción...", "RECURSOS": "Recursos..."}}
      ]
    }},
    "DESARROLLO": {{
      "TIEMPO": "30 minutos",
      "FASES": [
        {{"FASE": "Fase 1", "ACTIVIDADES": "Descripción...", "RECURSOS": "Recursos..."}},
        {{"FASE": "Fase 2", "ACTIVIDADES": "Descripción...", "RECURSOS": "Recursos..."}},
        {{"FASE": "Fase 3", "ACTIVIDADES": "Descripción...", "RECURSOS": "Recursos..."}}
      ]
    }},
    "CIERRE": {{
      "TIEMPO": "10 minutos",
      "FASES": [
        {{"FASE": "Actividad de cierre", "ACTIVIDADES": "Descripción...", "RECURSOS": "Recursos..."}},
        {{"FASE": "Metacognición", "ACTIVIDADES": "Descripción...", "RECURSOS": "Recursos..."}},
        {{"FASE": "Preguntas de reflexión", "ACTIVIDADES": "Descripción...", "RECURSOS": "Recursos..."}}
      ]
    }}
  }},
  "EVIDENCIAS": "Descripción de las evidencias de aprendizaje...",
  "EVALUACION": {{
    "TIPO": "Diagnóstica (Inicio), Formativa (Desarrollo y Cierre)",
    "AGENTE": "Heteroevaluación, Coevaluación, Autoevaluación",
    "TECNICAS": ["Observación directa", "Pregunta y respuesta"],
    "INSTRUMENTOS": ["Lista de cotejo"]
  }},
  "LISTA_COTEJO": [
    {{"NUMERO": 1, "CRITERIO": "Criterio de evaluación 1..."}},
    {{"NUMERO": 2, "CRITERIO": "Criterio de evaluación 2..."}}
  ],
  "ADAPTACIONES_NEAE": "Descripción de las adaptaciones para estudiantes con NEAE...",
  "OBSERVACIONES": "Observaciones adicionales..."
}}
"""

# ═══════════════════════════════════════════════════════════════════════════
# GENERACIÓN WORD PROFESIONAL
# ═══════════════════════════════════════════════════════════════════════════
def build_diario_docx(datos, meta, secuencia):
    doc = Document()
    doc.styles["Normal"].font.name = "Calibri"
    doc.styles["Normal"].font.size = Pt(10)

    for section in doc.sections:
        section.left_margin = Inches(0.6)
        section.right_margin = Inches(0.6)
        section.top_margin = Inches(0.5)
        section.bottom_margin = Inches(0.5)

    # ─── ENCABEZADO MINERD ───
    p_minerd = doc.add_paragraph()
    p_minerd.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run_m = p_minerd.add_run("MINISTERIO DE EDUCACIÓN DE LA REPÚBLICA DOMINICANA")
    run_m.bold = True
    run_m.font.size = Pt(11)

    p_titulo = doc.add_paragraph()
    p_titulo.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run_t = p_titulo.add_run("PLANIFICACIÓN DE CLASE DIARIA")
    run_t.bold = True
    run_t.font.size = Pt(14)
    run_t.font.color.rgb = RGBColor(0x0F, 0x76, 0x6E)

    p_sub = doc.add_paragraph()
    p_sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run_sub = p_sub.add_run("Nivel Secundario · Modalidad Académica")
    run_sub.font.size = Pt(10)
    run_sub.font.color.rgb = RGBColor(0x64, 0x74, 0x8B)

    doc.add_paragraph()

    # ─── DATOS GENERALES ───
    p_dg = doc.add_paragraph()
    run_dg = p_dg.add_run("DATOS GENERALES")
    run_dg.bold = True
    run_dg.font.size = Pt(11)
    run_dg.font.color.rgb = RGBColor(0x0F, 0x76, 0x6E)

    tabla_datos = doc.add_table(rows=4, cols=4)
    tabla_datos.style = "Table Grid"

    datos_gen = [
        ("Nombre completo", meta.get("docente", ""), "Cédula", meta.get("cedula", "")),
        ("Regional", meta.get("regional", ""), "Distrito", meta.get("distrito", "")),
        ("Centro Educativo", meta.get("centro", ""), "Código", meta.get("codigo", "")),
        ("Nivel/Sub-Sistema", "Secundario", "Ciclo", meta.get("ciclo", "")),
        ("Grado y Sección", meta.get("grado", ""), "Duración", meta.get("duracion", "50 minutos")),
        ("Área", meta.get("area", ""), "Asignatura", meta.get("asignatura", "")),
        ("Modalidad", "Académica", "Fecha", meta.get("fecha", "")),
    ]

    for i, (l1, v1, l2, v2) in enumerate(datos_gen):
        row = tabla_datos.rows[i].cells
        set_cell_text(row[0], l1, bold=True, color="CCFBF1")
        set_cell_text(row[1], v1)
        set_cell_text(row[2], l2, bold=True, color="CCFBF1")
        set_cell_text(row[3], v2)

    doc.add_paragraph()

    # ─── SECUENCIA DIDÁCTICA ───
    if secuencia:
        p_sec = doc.add_paragraph()
        run_sec = p_sec.add_run("SECUENCIA DIDÁCTICA")
        run_sec.bold = True
        run_sec.font.size = Pt(11)
        run_sec.font.color.rgb = RGBColor(0x0F, 0x76, 0x6E)

        tabla_sec = doc.add_table(rows=1, cols=4)
        tabla_sec.style = "Table Grid"
        row_sec = tabla_sec.rows[0].cells
        set_cell_text(row_sec[0], "Secuencia Didáctica", bold=True, color="CCFBF1")
        set_cell_text(row_sec[1], secuencia.get("NOMBRE", ""))
        set_cell_text(row_sec[2], f"Semanas: {secuencia.get('SEMANAS', '')}", bold=True, color="CCFBF1")
        set_cell_text(row_sec[3], f"Actividad: {meta.get('actividad', '')}")

        doc.add_paragraph()

        # ─── COMPETENCIAS ESPECÍFICAS ───
        p_ce = doc.add_paragraph()
        run_ce = p_ce.add_run("COMPETENCIAS ESPECÍFICAS")
        run_ce.bold = True
        run_ce.font.size = Pt(11)
        run_ce.font.color.rgb = RGBColor(0x0F, 0x76, 0x6E)

        for comp in secuencia.get("COMPETENCIAS_ESPECIFICAS", []):
            doc.add_paragraph(comp, style="List Bullet")

        doc.add_paragraph()

        # ─── CONTENIDOS ───
        p_cont = doc.add_paragraph()
        run_cont = p_cont.add_run("CONTENIDOS")
        run_cont.bold = True
        run_cont.font.size = Pt(11)
        run_cont.font.color.rgb = RGBColor(0x0F, 0x76, 0x6E)

        tabla_cont = doc.add_table(rows=2, cols=4)
        tabla_cont.style = "Table Grid"

        hdr_cont = tabla_cont.rows[0].cells
        set_cell_text(hdr_cont[0], "Conceptos", bold=True, center=True, color="CCFBF1")
        set_cell_text(hdr_cont[1], "Procedimientos", bold=True, center=True, color="CCFBF1")
        set_cell_text(hdr_cont[2], "Actitudes y Valores", bold=True, center=True, color="CCFBF1")
        set_cell_text(hdr_cont[3], "Indicadores de Logro", bold=True, center=True, color="CCFBF1")

        contenidos = secuencia.get("CONTENIDOS", {})
        row_cont = tabla_cont.rows[1].cells
        set_cell_text(row_cont[0], contenidos.get("CONCEPTOS", ""))
        set_cell_text(row_cont[1], contenidos.get("PROCEDIMIENTOS", ""))
        set_cell_text(row_cont[2], contenidos.get("ACTITUDES", ""))
        set_cell_text(row_cont[3], contenidos.get("INDICADORES_LOGRO", ""))

        doc.add_paragraph()

    # ─── COMPETENCIAS FUNDAMENTALES ───
    p_cf = doc.add_paragraph()
    run_cf = p_cf.add_run("COMPETENCIAS FUNDAMENTALES")
    run_cf.bold = True
    run_cf.font.size = Pt(11)
    run_cf.font.color.rgb = RGBColor(0x0F, 0x76, 0x6E)

    competencias_sel = meta.get("competencias_fundamentales", [])
    competencias_txt = "  ".join([
        f"{'☑' if c in competencias_sel else '☐'} {c}"
        for c in COMPETENCIAS_FUNDAMENTALES
    ])
    doc.add_paragraph(competencias_txt)
    doc.add_paragraph()

    # ─── ESTRATEGIAS DE ENSEÑANZA-APRENDIZAJE ───
    p_estr = doc.add_paragraph()
    run_estr = p_estr.add_run("ESTRATEGIAS DE ENSEÑANZA-APRENDIZAJE")
    run_estr.bold = True
    run_estr.font.size = Pt(11)
    run_estr.font.color.rgb = RGBColor(0x0F, 0x76, 0x6E)

    p_int = doc.add_paragraph()
    p_int.add_run("Intención pedagógica del día: ").bold = True
    p_int.add_run(datos.get("INTENCION_PEDAGOGICA", ""))

    p_estr_met = doc.add_paragraph()
    p_estr_met.add_run("Estrategia/Metodología: ").bold = True
    p_estr_met.add_run(", ".join(meta.get("estrategias", [])))

    doc.add_paragraph()

    # ─── MOMENTOS PEDAGÓGICOS ───
    p_mom = doc.add_paragraph()
    run_mom = p_mom.add_run("MOMENTOS PEDAGÓGICOS")
    run_mom.bold = True
    run_mom.font.size = Pt(11)
    run_mom.font.color.rgb = RGBColor(0x0F, 0x76, 0x6E)

    momentos = datos.get("MOMENTOS_PEDAGOGICOS", {})

    for nombre_momento, clave in [("INICIO", "INICIO"), ("DESARROLLO", "DESARROLLO"), ("CIERRE", "CIERRE")]:
        momento = momentos.get(clave, {})
        tiempo = momento.get("TIEMPO", "")

        p_mom_header = doc.add_paragraph()
        run_mh = p_mom_header.add_run(f"{nombre_momento} ({tiempo})")
        run_mh.bold = True
        run_mh.font.size = Pt(10)
        run_mh.font.color.rgb = RGBColor(0x08, 0x91, 0xB2)

        for fase in momento.get("FASES", []):
            p_fase = doc.add_paragraph()
            p_fase.add_run(f"{fase.get('FASE', '')}: ").bold = True
            p_fase.add_run(fase.get("ACTIVIDADES", ""))

            p_rec = doc.add_paragraph()
            p_rec.add_run("Recursos: ").bold = True
            p_rec.add_run(fase.get("RECURSOS", ""))
            p_rec.runs[0].font.size = Pt(9)

        doc.add_paragraph()

    # ─── EVIDENCIAS ───
    p_evid = doc.add_paragraph()
    p_evid.add_run("EVIDENCIAS: ").bold = True
    p_evid.add_run(datos.get("EVIDENCIAS", ""))

    doc.add_paragraph()

    # ─── EVALUACIÓN ───
    p_eval = doc.add_paragraph()
    run_eval = p_eval.add_run("EVALUACIÓN")
    run_eval.bold = True
    run_eval.font.size = Pt(11)
    run_eval.font.color.rgb = RGBColor(0x0F, 0x76, 0x6E)

    evaluacion = datos.get("EVALUACION", {})

    tabla_eval = doc.add_table(rows=2, cols=4)
    tabla_eval.style = "Table Grid"

    hdr_eval = tabla_eval.rows[0].cells
    set_cell_text(hdr_eval[0], "Tipo", bold=True, center=True, color="CCFBF1")
    set_cell_text(hdr_eval[1], "Agente", bold=True, center=True, color="CCFBF1")
    set_cell_text(hdr_eval[2], "Técnicas", bold=True, center=True, color="CCFBF1")
    set_cell_text(hdr_eval[3], "Instrumentos", bold=True, center=True, color="CCFBF1")

    row_eval = tabla_eval.rows[1].cells
    set_cell_text(row_eval[0], evaluacion.get("TIPO", ""))
    set_cell_text(row_eval[1], evaluacion.get("AGENTE", ""))
    set_cell_text(row_eval[2], ", ".join(evaluacion.get("TECNICAS", [])))
    set_cell_text(row_eval[3], ", ".join(evaluacion.get("INSTRUMENTOS", [])))

    doc.add_paragraph()

    # ─── ADAPTACIONES NEAE ───
    p_neae = doc.add_paragraph()
    p_neae.add_run("ADAPTACIONES (Si aplica, para estudiantes con NEAE): ").bold = True
    p_neae.add_run(datos.get("ADAPTACIONES_NEAE", "No aplica"))

    doc.add_paragraph()

    # ─── OBSERVACIONES ───
    p_obs = doc.add_paragraph()
    p_obs.add_run("OBSERVACIONES: ").bold = True
    p_obs.add_run(datos.get("OBSERVACIONES", ""))

    doc.add_page_break()

    # ─── LISTA DE COTEJO ───
    p_lc = doc.add_paragraph()
    run_lc = p_lc.add_run("LISTA DE COTEJO")
    run_lc.bold = True
    run_lc.font.size = Pt(12)
    run_lc.font.color.rgb = RGBColor(0x0F, 0x76, 0x6E)

    p_lc_tema = doc.add_paragraph()
    p_lc_tema.add_run(f"Tema: {meta.get('asignatura', '')} - {meta.get('grado', '')}")

    tabla_lc = doc.add_table(rows=1, cols=5)
    tabla_lc.style = "Table Grid"

    hdr_lc = tabla_lc.rows[0].cells
    set_cell_text(hdr_lc[0], "No.", bold=True, center=True, color="CCFBF1")
    set_cell_text(hdr_lc[1], "Criterios de Evaluación", bold=True, color="CCFBF1")
    set_cell_text(hdr_lc[2], "Sí", bold=True, center=True, color="CCFBF1")
    set_cell_text(hdr_lc[3], "No", bold=True, center=True, color="CCFBF1")
    set_cell_text(hdr_lc[4], "Observación", bold=True, color="CCFBF1")

    for item in datos.get("LISTA_COTEJO", []):
        row_lc = tabla_lc.add_row().cells
        set_cell_text(row_lc[0], str(item.get("NUMERO", "")), center=True)
        set_cell_text(row_lc[1], item.get("CRITERIO", ""))
        set_cell_text(row_lc[2], "", center=True)
        set_cell_text(row_lc[3], "", center=True)
        set_cell_text(row_lc[4], "")

    doc.add_paragraph()

    # ─── ESCALA DE VALORACIÓN ───
    p_ev = doc.add_paragraph()
    run_ev = p_ev.add_run("ESCALA DE VALORACIÓN")
    run_ev.bold = True
    run_ev.font.size = Pt(11)
    run_ev.font.color.rgb = RGBColor(0x0F, 0x76, 0x6E)

    tabla_ev = doc.add_table(rows=4, cols=2)
    tabla_ev.style = "Table Grid"

    hdr_ev = tabla_ev.rows[0].cells
    set_cell_text(hdr_ev[0], "Sigla", bold=True, center=True, color="CCFBF1")
    set_cell_text(hdr_ev[1], "Descripción", bold=True, color="CCFBF1")

    escala = [
        ("L", "Logrado (4 o 5, Máximo 2 indicadores en EP y los demás en L)"),
        ("EP", "En proceso (3 o 4, Máximo 2 indicadores en L, pero 3 en EP)"),
        ("NA", "Necesita apoyo (3 o más indicadores en NA, con 1 o 2 EP y sin apenas ninguna L)"),
    ]

    for i, (sigla, desc) in enumerate(escala, 1):
        row_ev = tabla_ev.rows[i].cells
        set_cell_text(row_ev[0], sigla, center=True)
        set_cell_text(row_ev[1], desc)

    doc.add_paragraph()

    # ─── FIRMAS ───
    doc.add_paragraph()
    tabla_firmas = doc.add_table(rows=2, cols=3)
    tabla_firmas.cell(0, 0).text = "_________________________"
    tabla_firmas.cell(0, 1).text = "_________________________"
    tabla_firmas.cell(0, 2).text = "_________________________"
    tabla_firmas.cell(1, 0).text = "Director/a de Centro Educativo"
    tabla_firmas.cell(1, 1).text = "Coordinador/a Académico"
    tabla_firmas.cell(1, 2).text = "Docente"

    for row in tabla_firmas.rows:
        for cell in row.cells:
            cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER

    add_page_number(doc)

    buffer = BytesIO()
    doc.save(buffer)
    buffer.seek(0)
    return buffer

# ═══════════════════════════════════════════════════════════════════════════
# RENDERIZADO DE STEPPER
# ═══════════════════════════════════════════════════════════════════════════
def render_stepper(fase_actual):
    pasos = [
        ("1", "Materia y Malla"),
        ("2", "Secuencia"),
        ("3", "Plan Diario"),
    ]

    html = '<div class="diario-stepper">'
    for i, (num, label) in enumerate(pasos):
        if i + 1 < fase_actual:
            estado = "done"
        elif i + 1 == fase_actual:
            estado = "active"
        else:
            estado = "inactive"

        html += f'''
        <div class="diario-step">
            <div class="diario-step-circle {estado}">{"✓" if estado == "done" else num}</div>
            <div class="diario-step-label">{label}</div>
        </div>
        '''
        if i < len(pasos) - 1:
            line_class = "done" if i + 1 < fase_actual else ""
            html += f'<div class="diario-step-line {line_class}"></div>'

    html += '</div>'
    st.markdown(html, unsafe_allow_html=True)

# ═══════════════════════════════════════════════════════════════════════════
# HERO Y STEPPER
# ═══════════════════════════════════════════════════════════════════════════
ia.panel_sidebar_ia("Plan Diario Académico")

st.markdown("""
<div class="diario-hero">
    <div class="diario-hero-title">🗓️ Planificación de Clase Diaria Académica</div>
    <div class="diario-hero-sub">
        Identificación inteligente de malla curricular · Secuencias Didácticas · Formato MINERD
    </div>
    <div>
        <span class="diario-hero-badge">📚 9 Materias Académicas</span>
        <span class="diario-hero-badge">🔧 Talleres Optativos</span>
        <span class="diario-hero-badge">🤖 Asistido por IA</span>
        <span class="diario-hero-badge">📄 Word MINERD</span>
    </div>
</div>
""", unsafe_allow_html=True)

render_stepper(st.session_state.diario_fase + 1)

# ═══════════════════════════════════════════════════════════════════════════
# FASE 0: MATERIA Y MALLA
# ═══════════════════════════════════════════════════════════════════════════
if st.session_state.diario_fase == 0:
    st.markdown('<div class="diario-section-title">📚 Selecciona el Tipo de Planificación</div>', unsafe_allow_html=True)

    col_tipo1, col_tipo2 = st.columns(2)

    with col_tipo1:
        selected_acad = st.session_state.diario_tipo == "academica"
        css_acad = "diario-area-card selected" if selected_acad else "diario-area-card"
        st.markdown(f"""
        <div class="{css_acad}">
            <div class="diario-area-icono">📚</div>
            <div class="diario-area-nombre">Materia Académica</div>
            <div style="font-size: 0.82rem; color: #64748B; margin-top: 0.4rem;">
                Requiere malla curricular. El módulo identificará automáticamente secuencias didácticas,
                competencias, contenidos e indicadores.
            </div>
        </div>
        """, unsafe_allow_html=True)
        if st.button("Seleccionar Materia Académica", key="tipo_academica", use_container_width=True,
                     type="primary" if selected_acad else "secondary"):
            st.session_state.diario_tipo = "academica"
            st.session_state.diario_area = None
            st.session_state.diario_secuencias = None
            st.rerun()

    with col_tipo2:
        selected_taller = st.session_state.diario_tipo == "taller"
        css_taller = "diario-area-card selected" if selected_taller else "diario-area-card"
        st.markdown(f"""
        <div class="{css_taller}">
            <div class="diario-area-icono">🔧</div>
            <div class="diario-area-nombre">Taller Optativo</div>
            <div style="font-size: 0.82rem; color: #64748B; margin-top: 0.4rem;">
                No requiere malla curricular. Planificación directa basada en el tema del taller.
            </div>
        </div>
        """, unsafe_allow_html=True)
        if st.button("Seleccionar Taller Optativo", key="tipo_taller", use_container_width=True,
                     type="primary" if selected_taller else "secondary"):
            st.session_state.diario_tipo = "taller"
            st.session_state.diario_area = None
            st.session_state.diario_secuencias = None
            st.rerun()

    # ─── MATERIA ACADÉMICA ───
    if st.session_state.diario_tipo == "academica":
        st.markdown('<hr style="border: none; height: 2px; background: linear-gradient(90deg, #14B8A6, #5EEAD4); margin: 1.5rem 0;">', unsafe_allow_html=True)
        st.markdown('<div class="diario-section-title">📖 Selecciona la Materia Académica</div>', unsafe_allow_html=True)

        cols_areas = st.columns(3)
        for idx, area_info in enumerate(AREAS_ACADEMICAS):
            col = cols_areas[idx % 3]
            with col:
                selected = st.session_state.diario_area == area_info["nombre"]
                css_class = "diario-area-card selected" if selected else "diario-area-card"
                st.markdown(f"""
                <div class="{css_class}" style="border-color: {area_info['color'] if selected else '#E2E8F0'};">
                    <div class="diario-area-icono">{area_info['icono']}</div>
                    <div class="diario-area-nombre">{area_info['nombre']}</div>
                </div>
                """, unsafe_allow_html=True)
                if st.button(f"Seleccionar", key=f"area_{idx}", use_container_width=True,
                             type="primary" if selected else "secondary"):
                    st.session_state.diario_area = area_info["nombre"]
                    st.session_state.diario_secuencias = None
                    st.rerun()

        if st.session_state.diario_area:
            st.markdown(f"**Materia seleccionada:** {st.session_state.diario_area}")
            st.markdown('<div class="diario-section-title">📄 Cargar Malla Curricular</div>', unsafe_allow_html=True)

            with st.form("form_malla"):
                col_g1, col_g2 = st.columns(2)
                with col_g1:
                    grado = st.text_input("Grado y Sección", placeholder="Ej: 4to B")
                with col_g2:
                    centro = st.text_input("Centro Educativo", value="Centro Educativo Ejemplo")

                tab_pdf, tab_texto = st.tabs(["📕 PDF de la Malla", "✍️ Pegar Texto de la Malla"])

                texto_malla = ""
                with tab_pdf:
                    archivo_malla = st.file_uploader("Subir PDF de la malla curricular", type=["pdf"])

                with tab_texto:
                    texto_malla_directo = st.text_area(
                        "Pegar el contenido de la malla curricular",
                        height=200,
                        placeholder="Pega aquí el contenido de la malla curricular del área..."
                    )

                btn_extraer = st.form_submit_button(
                    "🔍 Extraer Secuencias Didácticas con IA",
                    type="primary",
                    use_container_width=True
                )

            if btn_extraer:
                if not grado:
                    st.warning("⚠️ Indica el grado y sección.")
                else:
                    try:
                        if archivo_malla:
                            texto_malla = extraer_texto_pdf(archivo_malla)
                        elif texto_malla_directo:
                            texto_malla = texto_malla_directo
                        else:
                            st.warning("⚠️ Debes cargar un PDF o pegar el texto de la malla.")
                            st.stop()
                    except Exception as e:
                        st.error(f"❌ Error extrayendo la malla: {e}")
                        st.stop()

                    with st.spinner("🧠 Extrayendo secuencias didácticas de la malla curricular..."):
                        try:
                            prompt = prompt_extraer_secuencias(st.session_state.diario_area, grado, texto_malla)
                            resultado_malla, flags = ia.solicitar_json(
                                prompt, max_tokens=16384, temperature=0.1, modulo="diario_academico_malla"
                            )

                            st.session_state.diario_secuencias = resultado_malla.get("SECUENCIAS", [])
                            st.session_state.diario_grado = grado
                            st.session_state.diario_centro = centro
                            st.session_state.diario_fase = 1
                            st.toast(f"✅ {len(st.session_state.diario_secuencias)} secuencias didácticas extraídas.", icon="📚")
                            st.rerun()

                        except Exception as e:
                            ia.render_error_ia(e)

    # ─── TALLER OPTATIVO ───
    elif st.session_state.diario_tipo == "taller":
        st.markdown('<hr style="border: none; height: 2px; background: linear-gradient(90deg, #14B8A6, #5EEAD4); margin: 1.5rem 0;">', unsafe_allow_html=True)
        st.markdown('<div class="diario-section-title">🔧 Selecciona el Taller Optativo</div>', unsafe_allow_html=True)

        with st.form("form_taller"):
            col_t1, col_t2 = st.columns(2)
            with col_t1:
                taller_sel = st.selectbox("Taller Optativo", TALLERES_OPTATIVOS + ["Otro (escribir)"])
                if taller_sel == "Otro (escribir)":
                    taller_custom = st.text_input("Nombre del taller", placeholder="Escribe el nombre del taller")
                else:
                    taller_custom = taller_sel
            with col_t2:
                grado = st.text_input("Grado y Sección", placeholder="Ej: 5to A")
                centro = st.text_input("Centro Educativo", value="Centro Educativo Ejemplo")

            btn_taller = st.form_submit_button(
                "➡️ Continuar a Configuración del Plan",
                type="primary",
                use_container_width=True
            )

        if btn_taller:
            if not grado:
                st.warning("⚠️ Indica el grado y sección.")
            else:
                st.session_state.diario_area = taller_custom
                st.session_state.diario_grado = grado
                st.session_state.diario_centro = centro
                st.session_state.diario_secuencias = None
                st.session_state.diario_fase = 2  # Ir directo a generación (sin secuencia)
                st.rerun()

# ═══════════════════════════════════════════════════════════════════════════
# FASE 1: SELECCIÓN DE SECUENCIA
# ═══════════════════════════════════════════════════════════════════════════
elif st.session_state.diario_fase == 1:
    secuencias = st.session_state.diario_secuencias or []

    st.markdown(f'<div class="diario-section-title">📅 Secuencias Didácticas de {st.session_state.diario_area} - {st.session_state.diario_grado}</div>', unsafe_allow_html=True)

    if not secuencias:
        st.warning("⚠️ No se encontraron secuencias didácticas en la malla. Intenta con otro formato de malla.")
        if st.button("🔄 Volver a cargar malla", use_container_width=True):
            st.session_state.diario_fase = 0
            st.session_state.diario_secuencias = None
            st.rerun()
    else:
        st.info(f"📚 Se encontraron **{len(secuencias)}** secuencias didácticas. Selecciona una para generar el plan diario.")

        for idx, sec in enumerate(secuencias):
            selected = st.session_state.diario_secuencia_sel == idx
            css_class = "diario-sec-card selected" if selected else "diario-sec-card"

            competencias = sec.get("COMPETENCIAS_ESPECIFICAS", [])
            contenidos = sec.get("CONTENIDOS", {})
            actividades = sec.get("ACTIVIDADES", [])

            st.markdown(f"""
            <div class="{css_class}">
                <div class="diario-sec-nombre">📅 Secuencia {sec.get('NUMERO', idx+1)}: {sec.get('NOMBRE', '')}</div>
                <div class="diario-sec-meta">
                    Semanas: {sec.get('SEMANAS', '')} · 
                    Competencias: {len(competencias)} · 
                    Actividades: {len(actividades)}
                </div>
            </div>
            """, unsafe_allow_html=True)

            with st.expander(f"Ver detalles de la Secuencia {sec.get('NUMERO', idx+1)}", expanded=False):
                st.markdown("**Competencias Específicas:**")
                for comp in competencias:
                    st.markdown(f"- {comp}")
                st.markdown("**Contenidos:**")
                st.markdown(f"- Conceptos: {contenidos.get('CONCEPTOS', '')[:200]}")
                st.markdown(f"- Procedimientos: {contenidos.get('PROCEDIMIENTOS', '')[:200]}")
                st.markdown(f"- Indicadores: {contenidos.get('INDICADORES_LOGRO', '')[:200]}")

            if st.button(f"Seleccionar Secuencia {sec.get('NUMERO', idx+1)}", key=f"sec_{idx}",
                         use_container_width=True, type="primary" if selected else "secondary"):
                st.session_state.diario_secuencia_sel = idx
                st.session_state.diario_fase = 2
                st.rerun()

        st.markdown("---")
        if st.button("🔄 Volver a cargar malla", use_container_width=True):
            st.session_state.diario_fase = 0
            st.session_state.diario_secuencias = None
            st.session_state.diario_secuencia_sel = None
            st.rerun()

# ═══════════════════════════════════════════════════════════════════════════
# FASE 2: GENERACIÓN DEL PLAN DIARIO
# ═══════════════════════════════════════════════════════════════════════════
elif st.session_state.diario_fase == 2:
    secuencia = None
    if st.session_state.diario_secuencias and st.session_state.diario_secuencia_sel is not None:
        secuencia = st.session_state.diario_secuencias[st.session_state.diario_secuencia_sel]

    st.markdown('<div class="diario-section-title">🗓️ Configuración del Plan Diario</div>', unsafe_allow_html=True)

    if secuencia:
        st.success(f"✅ Secuencia seleccionada: **{secuencia.get('NOMBRE', '')}** (Semanas {secuencia.get('SEMANAS', '')})")
        with st.expander("📋 Datos curriculares identificados automáticamente"):
            st.json(secuencia)
    else:
        st.info(f"🔧 Taller optativo: **{st.session_state.diario_area}** (sin malla curricular)")

    with st.form("form_plan_diario"):
        col1, col2 = st.columns(2)
        with col1:
            docente = st.text_input("Docente", value="Ing. Bernardo Antonio Hernández Batista")
            cedula = st.text_input("Cédula", placeholder="Ej: 000-0000-0000A")
            regional = st.text_input("Regional", value="06")
            distrito = st.text_input("Distrito", value="06")
            codigo = st.text_input("Código del Centro", value="00000")
        with col2:
            centro = st.text_input("Centro Educativo", value=st.session_state.diario_centro)
            area = st.text_input("Área", value=st.session_state.diario_area)
            asignatura = st.text_input("Asignatura", placeholder="Ej: Geometría, Álgebra")
            grado = st.text_input("Grado y Sección", value=st.session_state.diario_grado)
            ciclo = st.text_input("Ciclo", value="2do ciclo")
            duracion = st.text_input("Duración", value="50 minutos")
            fecha = st.date_input("Fecha", value=date.today())

        if secuencia:
            actividad = st.selectbox(
                "Actividad de la Secuencia",
                [f"Actividad {a.get('NUMERO', i+1)}: {a.get('NOMBRE', '')}" for i, a in enumerate(secuencia.get("ACTIVIDADES", []))] or ["Actividad única"]
            )
        else:
            tema_taller = st.text_area("Tema de la sesión del taller", height=70, placeholder="Describe el tema de la sesión...")

        competencias_fundamentales = st.multiselect(
            "Competencias Fundamentales a desarrollar",
            COMPETENCIAS_FUNDAMENTALES,
            default=["Pensamiento Lógico, Creativo y Crítico", "Resolución de Problemas"]
        )

        estrategias = st.multiselect(
            "Estrategias de Enseñanza-Aprendizaje",
            ESTRATEGIAS_EA,
            default=["Indagación dialógica", "Estudio de Caso"]
        )

        max_tokens, temperature = ia.control_avanzado(default_tokens=16384, tope=32000, default_temp=0.2)

        st.markdown("---")
        btn_generar = st.form_submit_button(
            "⚡ Generar Plan Diario con IA",
            type="primary",
            use_container_width=True
        )

    if btn_generar:
        with st.spinner("🧠 Generando plan diario con IA..."):
            try:
                datos_form = {
                    "centro": centro,
                    "docente": docente,
                    "cedula": cedula,
                    "regional": regional,
                    "distrito": distrito,
                    "codigo": codigo,
                    "area": area,
                    "asignatura": asignatura,
                    "grado": grado,
                    "ciclo": ciclo,
                    "duracion": duracion,
                    "fecha": fecha.strftime("%d/%m/%Y"),
                    "competencias_fundamentales": competencias_fundamentales,
                    "estrategias": estrategias,
                    "actividad": actividad if secuencia else "",
                    "taller": st.session_state.diario_area if not secuencia else "",
                    "tema_taller": tema_taller if not secuencia else "",
                }

                prompt = prompt_generar_plan(datos_form, secuencia, "secuencia_didactica" if secuencia else "taller")

                resultado, flags = ia.solicitar_json(
                    prompt, max_tokens=max_tokens, temperature=temperature, modulo="diario_academico"
                )

                st.session_state.diario_resultado = {
                    "datos": resultado,
                    "flags": flags,
                    "meta": datos_form,
                    "secuencia": secuencia,
                }
                st.session_state.diario_fase = 3
                st.toast("✅ Plan diario generado.", icon="🗓️")
                st.rerun()

            except Exception as e:
                ia.render_error_ia(e)

# ═══════════════════════════════════════════════════════════════════════════
# FASE 3: RESULTADOS
# ═══════════════════════════════════════════════════════════════════════════
elif st.session_state.diario_fase == 3:
    resultado = st.session_state.diario_resultado
    datos = resultado["datos"]
    meta = resultado["meta"]
    secuencia = resultado.get("secuencia")

    st.markdown('<div class="diario-section-title">🎉 Plan Diario Generado</div>', unsafe_allow_html=True)

    # ─── MÉTRICAS ───
    col_s1, col_s2, col_s3, col_s4 = st.columns(4)
    with col_s1:
        st.markdown(f'<div class="diario-stat"><div class="diario-stat-value">{len(meta.get("competencias_fundamentales", []))}</div><div class="diario-stat-label">Comp. Fundamentales</div></div>', unsafe_allow_html=True)
    with col_s2:
        momentos = datos.get("MOMENTOS_PEDAGOGICOS", {})
        total_fases = sum(len(m.get("FASES", [])) for m in momentos.values())
        st.markdown(f'<div class="diario-stat"><div class="diario-stat-value">{total_fases}</div><div class="diario-stat-label">Fases Pedagógicas</div></div>', unsafe_allow_html=True)
    with col_s3:
        st.markdown(f'<div class="diario-stat"><div class="diario-stat-value">{len(datos.get("LISTA_COTEJO", []))}</div><div class="diario-stat-label">Criterios Cotejo</div></div>', unsafe_allow_html=True)
    with col_s4:
        st.markdown(f'<div class="diario-stat"><div class="diario-stat-value">{len(meta.get("estrategias", []))}</div><div class="diario-stat-label">Estrategias</div></div>', unsafe_allow_html=True)

    st.markdown("---")

    # ─── TÍTULO ───
    st.markdown(f"## 🗓️ Plan Diario: {meta.get('asignatura', '')} - {meta.get('grado', '')}")
    st.caption(f"Docente: {meta.get('docente', '')} | Fecha: {meta.get('fecha', '')}")

    st.markdown("---")

    # ─── TABS ───
    tab_preview, tab_json, tab_debug = st.tabs(["👁️ Vista Previa", "🧾 JSON Completo", "🐛 Depuración"])

    with tab_preview:
        with st.expander("🎯 Intención Pedagógica", expanded=True):
            st.write(datos.get("INTENCION_PEDAGOGICA", ""))

        with st.expander("📅 Momentos Pedagógicos"):
            momentos = datos.get("MOMENTOS_PEDAGOGICOS", {})
            for nombre, clave in [("INICIO", "INICIO"), ("DESARROLLO", "DESARROLLO"), ("CIERRE", "CIERRE")]:
                momento = momentos.get(clave, {})
                st.markdown(f"**{nombre}** ({momento.get('TIEMPO', '')})")
                for fase in momento.get("FASES", []):
                    st.markdown(f"- **{fase.get('FASE', '')}**: {fase.get('ACTIVIDADES', '')[:150]}...")
                st.markdown("---")

        with st.expander("📋 Lista de Cotejo"):
            for item in datos.get("LISTA_COTEJO", []):
                st.markdown(f"{item.get('NUMERO', '')}. {item.get('CRITERIO', '')}")

        with st.expander("📊 Evaluación"):
            evaluacion = datos.get("EVALUACION", {})
            st.markdown(f"**Tipo:** {evaluacion.get('TIPO', '')}")
            st.markdown(f"**Agente:** {evaluacion.get('AGENTE', '')}")
            st.markdown(f"**Técnicas:** {', '.join(evaluacion.get('TECNICAS', []))}")
            st.markdown(f"**Instrumentos:** {', '.join(evaluacion.get('INSTRUMENTOS', []))}")

    with tab_json:
        st.json(datos)

    with tab_debug:
        st.write("Meta:", meta)
        st.write("Flags:", resultado.get("flags", {}))

    st.markdown("---")

    # ─── DESCARGA WORD ───
    col_dl1, col_dl2 = st.columns(2)

    with col_dl1:
        if st.button("📥 Generar Word (.docx)", type="primary", use_container_width=True):
            with st.spinner("📄 Construyendo documento Word..."):
                buffer = build_diario_docx(datos, meta, secuencia)
                st.session_state.diario_buffer = buffer

    if hasattr(st.session_state, "diario_buffer") and st.session_state.diario_buffer:
        nombre_archivo = ia.sanear_nombre_archivo(f"Plan_Diario_{meta.get('asignatura', 'academico')}_{meta.get('grado', '')}")
        st.download_button(
            label="⬇️ Descargar Plan Diario (.docx)",
            data=st.session_state.diario_buffer,
            file_name=f"{nombre_archivo}.docx",
            mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
            type="primary",
            use_container_width=True
        )

    with col_dl2:
        if st.button("🔄 Nuevo Plan Diario", use_container_width=True):
            st.session_state.diario_fase = 0
            st.session_state.diario_tipo = None
            st.session_state.diario_area = None
            st.session_state.diario_secuencias = None
            st.session_state.diario_secuencia_sel = None
            st.session_state.diario_resultado = None
            if hasattr(st.session_state, "diario_buffer"):
                del st.session_state.diario_buffer
            st.rerun()