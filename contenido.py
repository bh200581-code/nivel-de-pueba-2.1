"""
contenido.py — Generador de Contenidos y Actividades ETP (MIGRADO · Paso 17 v2)
Mejorado SOBRE EL ORIGINAL: conserva su riqueza (18 secciones, normalización de
claves, validación con advertencias, anclaje a PDF, rúbrica L/EP/NA, solucionario)
y lo migra a core/ia con super-interfaz y Word paginado.
"""
import copy
import json
import re
import unicodedata
from datetime import date, datetime
from io import BytesIO
from typing import Any, Dict, List, Tuple

import streamlit as st
from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml import parse_xml
from docx.oxml.ns import nsdecls, qn
from docx.shared import Inches, Pt

try:
    from pypdf import PdfReader
except Exception:
    try:
        from PyPDF2 import PdfReader
    except Exception:
        PdfReader = None

from core import ia

TOPE_MAX_TOKENS = 32000
CLAVES_LISTA = [
    "CONOCIMIENTOS_PREVIOS", "ERRORES_COMUNES", "CONTENIDO_TEORICO", "GLOSARIO",
    "CONEXIONES_INTERCURRICULARES", "TICKET_SALIDA", "SIMULADORES_RECURSOS",
    "WEBGRAFIA", "CRITERIOS_EVALUACION", "AUTOEVALUACION", "GUIA_RESPUESTAS",
]
CLAVES_DICT = ["ACTIVIDAD_PRINCIPAL", "ACTIVIDAD_REFUERZO", "ACTIVIDAD_EXTENSION", "TAREA_INDEPENDIENTE"]

# ═══════════════════════════════════════════════════════════════════════════
# UTILIDADES (conservadas del original)
# ═══════════════════════════════════════════════════════════════════════════
def limpiar_texto(valor: Any) -> str:
    if valor is None:
        return ""
    return re.sub(r"\s+", " ", str(valor)).strip()

def quitar_acentos(texto: str) -> str:
    return "".join(c for c in unicodedata.normalize("NFD", texto) if unicodedata.category(c) != "Mn")

def normalizar_clave(clave: Any) -> str:
    clave = limpiar_texto(clave)
    clave = quitar_acentos(clave).upper()
    return re.sub(r"\s+", "_", clave)

def normalizar_claves(objeto: Any) -> Any:
    if isinstance(objeto, dict):
        return {normalizar_clave(k): normalizar_claves(v) for k, v in objeto.items()}
    if isinstance(objeto, list):
        return [normalizar_claves(item) for item in objeto]
    return objeto

def convertir_puntos(valor: Any) -> int:
    if isinstance(valor, (int, float)):
        return int(valor)
    m = re.search(r"\d+", str(valor))
    return int(m.group()) if m else 0

def validar_y_normalizar_datos(datos, cantidad_criterios, valor_puntos) -> Tuple[Dict, List[str]]:
    advertencias: List[str] = []
    if not isinstance(datos, dict):
        return {}, ["La respuesta no tiene formato de objeto JSON."]
    datos = normalizar_claves(datos)
    base = {
        "RESUMEN_EJECUTIVO": "", "CONOCIMIENTOS_PREVIOS": [], "ERRORES_COMUNES": [],
        "CONTENIDO_TEORICO": [], "GLOSARIO": [], "CONEXIONES_INTERCURRICULARES": [],
        "ACTIVIDAD_PRINCIPAL": {}, "ACTIVIDAD_REFUERZO": {}, "ACTIVIDAD_EXTENSION": {},
        "ADAPTACIONES_NEAE": "", "TICKET_SALIDA": [], "TAREA_INDEPENDIENTE": {},
        "SIMULADORES_RECURSOS": [], "WEBGRAFIA": [], "CRITERIOS_EVALUACION": [],
        "AUTOEVALUACION": [], "GUIA_RESPUESTAS": [],
    }
    for clave, default in base.items():
        if clave not in datos:
            datos[clave] = copy.deepcopy(default)
            advertencias.append(f"Falta la sección '{clave}'. Se creó vacía.")
    for clave in CLAVES_LISTA:
        if not isinstance(datos.get(clave), list):
            datos[clave] = []
            advertencias.append(f"La sección '{clave}' no era una lista y se normalizó.")
    for clave in CLAVES_DICT:
        if not isinstance(datos.get(clave), dict):
            datos[clave] = {}
            advertencias.append(f"La sección '{clave}' no era un objeto y se normalizó.")
    criterios = datos.get("CRITERIOS_EVALUACION", [])
    if len(criterios) != cantidad_criterios:
        advertencias.append(f"Se pidieron {cantidad_criterios} criterios y la IA devolvió {len(criterios)}.")
    suma = sum(convertir_puntos(c.get("PUNTOS", 0)) if isinstance(c, dict) else 0 for c in criterios)
    if suma != valor_puntos:
        advertencias.append(f"La suma de puntos es {suma}, pero se solicitaron {valor_puntos}.")
    return datos, advertencias

def extraer_texto_pdf(archivo, max_caracteres=80000) -> str:
    if PdfReader is None:
        raise RuntimeError("No hay librería PDF disponible. Instala pypdf o PyPDF2.")
    archivo.seek(0)
    reader = PdfReader(archivo)
    texto = "".join((p.extract_text() or "") for p in reader.pages)
    texto = re.sub(r"\s+", " ", texto).strip()
    return texto[:max_caracteres]

# ═══════════════════════════════════════════════════════════════════════════
# PROMPT MAESTRO (con marcadores seguros)
# ═══════════════════════════════════════════════════════════════════════════
def construir_prompt_maestro(f: Dict) -> str:
    contexto = (
        f"DOCUMENTO CURRICULAR OFICIAL CARGADO (BASE ÚNICA Y OBLIGATORIA):\n{f['texto_pdf']}\n"
        "REGLA DE ANCLAJE: Todo el contenido teórico, glosario y actividades deben derivarse "
        "EXCLUSIVAMENTE de este documento. PROHIBIDO conocimiento externo no presente en el PDF."
        if f["texto_pdf"] else
        "REGLA DE CONTENIDO: Genera contenido con rigor técnico y estándares reales de la industria."
    )
    return f"""Actúa como Catedrático Universitario de Alto Nivel, Especialista Curricular ETP del MINERD y Experto en Diseño Instruccional con Taxonomía de Bloom.
INSUMOS:
Contenido a desarrollar: {f['contenido']}
Actividad propuesta: {f['actividad']}
Instrumento seleccionado: {f['instrumento']}
Valor total: {f['valor_puntos']} puntos
Cantidad EXACTA de criterios de evaluación: {f['cant_criterios']}
Duración de la sesión: {f['duracion_sesion']}
Nivel de Bloom objetivo: {f['nivel_bloom_objetivo']}
Características del grupo: {f['caracteristicas_grupo']}
Resultado de Aprendizaje (RA): {f['ra']}
Criterio de Evaluación / EC: {f['ce_ec']}
{contexto}
REGLAS DE GENERACIÓN:
RESUMEN EJECUTIVO: síntesis de 3 líneas. CONOCIMIENTOS PREVIOS: 3-4 prerrequisitos.
ERRORES COMUNES: 3-4 errores con corrección. CONTENIDO TEÓRICO: mínimo 3 secciones con nivel Bloom, título, contenido y ayuda visual.
GLOSARIO: mínimo 5 términos con definición y ejemplo. CONEXIONES: 2-3 vínculos.
ACTIVIDAD PRINCIPAL: título, indicador y pasos con tiempo/modalidad/recurso. REFUERZO y EXTENSIÓN diferenciadas.
ADAPTACIONES NEAE o "Sin adaptaciones adicionales requeridas". TICKET: 3 preguntas. TAREA con entregable.
SIMULADORES: 3 recursos con nombre/descripción/URL/Bloom. WEBGRAFÍA pertinente.
CRITERIOS: exactamente {f['cant_criterios']} con descriptores L/EP/NA y puntos que sumen {f['valor_puntos']}.
AUTOEVALUACIÓN: 3-4 ítems. GUÍA DE RESPUESTAS: resultado esperado por paso.
CODIFICACIÓN OBLIGATORIA: salto de línea → {ia.MARKER_NL} · comilla doble → {ia.MARKER_DQ} · tabulación → {ia.MARKER_TAB}.
Nunca uses saltos de línea ni comillas dobles literales dentro de los valores.
FORMATO DE SALIDA ESTRICTO (JSON NATIVO):
{{
 "RESUMEN_EJECUTIVO": "...", "CONOCIMIENTOS_PREVIOS": ["..."],
 "ERRORES_COMUNES": [{{"ERROR": "...", "CORRECCION": "..."}}],
 "CONTENIDO_TEORICO": [{{"NIVEL_BLOOM": "...", "TITULO_SECCION": "...", "CONTENIDO": "...", "AYUDA_VISUAL": "..."}}],
 "GLOSARIO": [{{"TERMINO": "...", "DEFINICION": "...", "EJEMPLO_APLICACION": "..."}}],
 "CONEXIONES_INTERCURRICULARES": ["..."],
 "ACTIVIDAD_PRINCIPAL": {{"TITULO": "...", "INDICADOR_LOGRO": "...", "PASOS": [{{"PASO": "...", "TIEMPO": "...", "MODALIDAD": "...", "RECURSO": "..."}}]}},
 "ACTIVIDAD_REFUERZO": {{"TITULO": "...", "PASOS": ["..."]}},
 "ACTIVIDAD_EXTENSION": {{"TITULO": "...", "DESCRIPCION": "...", "ENTREGABLE": "..."}},
 "ADAPTACIONES_NEAE": "...", "TICKET_SALIDA": ["..."],
 "TAREA_INDEPENDIENTE": {{"DESCRIPCION": "...", "ENTREGABLE": "..."}},
 "SIMULADORES_RECURSOS": [{{"TIPO": "...", "NOMBRE": "...", "DESCRIPCION": "...", "URL": "...", "NIVEL_BLOOM": "..."}}],
 "WEBGRAFIA": ["..."],
 "CRITERIOS_EVALUACION": [{{"CRITERIO": "...", "INDICADOR": "...", "LOGRADO": "...", "EN_PROCESO": "...", "NECESITA_APOYO": "...", "PUNTOS": 20}}],
 "AUTOEVALUACION": ["..."], "GUIA_RESPUESTAS": [{{"PASO": "...", "RESPUESTA_ESPERADA": "..."}}]
}}
"""

# ═══════════════════════════════════════════════════════════════════════════
# WORD (mejorado: paginación + pie institucional)
# ═══════════════════════════════════════════════════════════════════════════
def shade_cell(cell, color):
    shd = parse_xml(r'<w:shd {} w:fill="{}"/>'.format(nsdecls("w"), color))
    cell._tc.get_or_add_tcPr().append(shd)

def set_cell_text(cell, text, bold=False, center=False, color=None):
    cell.text = ""
    p = cell.paragraphs[0]
    run = p.add_run(str(text))
    run.bold = bold
    if center:
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    if color:
        shade_cell(cell, color)

def add_bullet(doc, text):
    try:
        p = doc.add_paragraph(text, style="List Bullet")
    except Exception:
        p = doc.add_paragraph(f"• {text}")
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

def agregar_numeracion_pagina(doc):
    pie = doc.sections[0].footer.paragraphs[0]
    pie.alignment = WD_ALIGN_PARAGRAPH.CENTER
    pie.text = ""
    run = pie.add_run()
    def _campo(instr):
        i = OxmlElement("w:fldChar"); i.set(qn("w:fldCharType"), "begin")
        t = OxmlElement("w:instrText"); t.set(qn("xml:space"), "preserve"); t.text = instr
        f = OxmlElement("w:fldChar"); f.set(qn("w:fldCharType"), "end")
        run._r.append(i); run._r.append(t); run._r.append(f)
    run.add_text("Página "); _campo("PAGE"); run.add_text(" de "); _campo("NUMPAGES")
    run.font.size = Pt(8); run.font.color.rgb = RGBColor(0x64, 0x74, 0x8B) if False else None

def build_docx(datos: Dict, meta: Dict) -> BytesIO:
    from docx.shared import RGBColor
    doc = Document()
    doc.styles["Normal"].font.name = "Calibri"
    doc.styles["Normal"].font.size = Pt(11)
    for s in doc.sections:
        s.left_margin = Inches(0.75); s.right_margin = Inches(0.75)
    p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.add_run(meta.get("politecnico", "")).bold = True
    p2 = doc.add_paragraph(); p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p2.add_run("Material Didáctico Integral — Generador de Contenidos ETP")
    r.bold = True; r.font.size = Pt(13); r.font.color.rgb = RGBColor(0x16, 0xA3, 0x4A)
    fecha = meta.get("fecha") or date.today()
    if isinstance(fecha, datetime):
        fecha = fecha.date()
    doc.add_paragraph(f"Docente: {meta.get('docente','')} | Módulo: {meta.get('asignatura','')} | Fecha: {fecha.strftime('%d/%m/%Y')}")
    doc.add_paragraph(f"Duración: {meta.get('duracion_sesion','')} | Nivel Bloom: {meta.get('nivel_bloom_objetivo','')}")
    doc.add_paragraph("Estudiante: ______________________________________ | Sección: _______")
    doc.add_paragraph("_" * 70)

    if limpiar_texto(datos.get("RESUMEN_EJECUTIVO")):
        doc.add_heading("I. Resumen Ejecutivo (Para el Docente)", level=1)
        pr = doc.add_paragraph(limpiar_texto(datos["RESUMEN_EJECUTIVO"])); pr.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    if meta.get("ra") or meta.get("ce_ec"):
        doc.add_heading("II. Alineación Curricular", level=1)
        if meta.get("ra"):
            pa = doc.add_paragraph(); pa.add_run("Resultado de Aprendizaje (RA): ").bold = True; pa.add_run(meta["ra"])
        if meta.get("ce_ec"):
            pc = doc.add_paragraph(); pc.add_run("Criterio de Evaluación / EC: ").bold = True; pc.add_run(meta["ce_ec"])
    if datos.get("CONOCIMIENTOS_PREVIOS"):
        doc.add_heading("III. Conocimientos Previos Requeridos", level=1)
        for it in datos["CONOCIMIENTOS_PREVIOS"]: add_bullet(doc, limpiar_texto(it))
    if datos.get("ERRORES_COMUNES"):
        doc.add_heading("IV. Errores Comunes y Concepciones Previas", level=1)
        t = doc.add_table(rows=1, cols=2); t.style = "Table Grid"
        set_cell_text(t.rows[0].cells[0], "⚠️ Error / Concepción Errónea", bold=True, color="FEE2E2")
        set_cell_text(t.rows[0].cells[1], "✅ Corrección Pedagógica", bold=True, color="FEE2E2")
        for it in datos["ERRORES_COMUNES"]:
            row = t.add_row().cells
            if isinstance(it, dict):
                set_cell_text(row[0], limpiar_texto(it.get("ERROR", "")))
                set_cell_text(row[1], limpiar_texto(it.get("CORRECCION", "")))
            else:
                set_cell_text(row[0], limpiar_texto(it)); set_cell_text(row[1], "")
    if datos.get("CONTENIDO_TEORICO"):
        doc.add_heading("V. Desarrollo de Contenido Teórico", level=1)
        for sec in datos["CONTENIDO_TEORICO"]:
            if isinstance(sec, dict):
                ps = doc.add_paragraph()
                ps.add_run(f"[Bloom: {limpiar_texto(sec.get('NIVEL_BLOOM',''))}] ").bold = True
                ps.add_run(limpiar_texto(sec.get("TITULO_SECCION",""))).bold = True
                pc = doc.add_paragraph(limpiar_texto(sec.get("CONTENIDO",""))); pc.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
                if sec.get("AYUDA_VISUAL"):
                    pv = doc.add_paragraph(); pv.add_run("💡 Ayuda visual sugerida: ").bold = True
                    pv.add_run(limpiar_texto(sec["AYUDA_VISUAL"])).italic = True
    if datos.get("GLOSARIO"):
        doc.add_heading("VI. Glosario Técnico", level=1)
        t = doc.add_table(rows=1, cols=3); t.style = "Table Grid"
        for i, h in enumerate(["Término", "Definición", "Ejemplo de Aplicación"]):
            set_cell_text(t.rows[0].cells[i], h, bold=True, color="E2E8F0")
        for it in datos["GLOSARIO"]:
            if isinstance(it, dict):
                row = t.add_row().cells
                set_cell_text(row[0], limpiar_texto(it.get("TERMINO","")), bold=True)
                set_cell_text(row[1], limpiar_texto(it.get("DEFINICION","")))
                set_cell_text(row[2], limpiar_texto(it.get("EJEMPLO_APLICACION","")))
    if datos.get("CONEXIONES_INTERCURRICULARES"):
        doc.add_heading("VII. Conexiones Intercurriculares", level=1)
        for it in datos["CONEXIONES_INTERCURRICULARES"]: add_bullet(doc, limpiar_texto(it))
    ap = datos.get("ACTIVIDAD_PRINCIPAL", {})
    if isinstance(ap, dict) and ap:
        doc.add_heading(f"VIII. Actividad Principal: {limpiar_texto(ap.get('TITULO',''))}", level=1)
        if ap.get("INDICADOR_LOGRO"):
            pi = doc.add_paragraph(); pi.add_run("Indicador de Logro: ").bold = True; pi.add_run(limpiar_texto(ap["INDICADOR_LOGRO"]))
        pasos = ap.get("PASOS", [])
        if pasos:
            t = doc.add_table(rows=1, cols=4); t.style = "Table Grid"
            for i, h in enumerate(["Paso", "Tiempo", "Modalidad", "Recurso"]):
                set_cell_text(t.rows[0].cells[i], h, bold=True, center=True, color="DBEAFE")
            for paso in pasos:
                row = t.add_row().cells
                if isinstance(paso, dict):
                    set_cell_text(row[0], limpiar_texto(paso.get("PASO","")))
                    set_cell_text(row[1], limpiar_texto(paso.get("TIEMPO","")), center=True)
                    set_cell_text(row[2], limpiar_texto(paso.get("MODALIDAD","")), center=True)
                    set_cell_text(row[3], limpiar_texto(paso.get("RECURSO","")))
    ar = datos.get("ACTIVIDAD_REFUERZO", {})
    if isinstance(ar, dict) and ar:
        doc.add_heading(f"IX. Actividad de Refuerzo: {limpiar_texto(ar.get('TITULO',''))}", level=1)
        for paso in ar.get("PASOS", []): add_bullet(doc, limpiar_texto(paso))
    ae = datos.get("ACTIVIDAD_EXTENSION", {})
    if isinstance(ae, dict) and ae:
        doc.add_heading(f"X. Actividad de Extensión: {limpiar_texto(ae.get('TITULO',''))}", level=1)
        pd_ = doc.add_paragraph(); pd_.add_run("Descripción: ").bold = True; pd_.add_run(limpiar_texto(ae.get("DESCRIPCION","")))
        pe = doc.add_paragraph(); pe.add_run("Entregable: ").bold = True; pe.add_run(limpiar_texto(ae.get("ENTREGABLE","")))
    doc.add_heading("XI. Adaptaciones para NEAE", level=1)
    doc.add_paragraph(limpiar_texto(datos.get("ADAPTACIONES_NEAE")) or "Sin adaptaciones adicionales requeridas.")
    if datos.get("TICKET_SALIDA"):
        doc.add_heading("XII. Ticket de Salida", level=1)
        for i, preg in enumerate(datos["TICKET_SALIDA"], 1):
            pt = doc.add_paragraph(); pt.add_run(f"{i}. ").bold = True; pt.add_run(limpiar_texto(preg))
            doc.add_paragraph("R: _______________________________________________")
    ti = datos.get("TAREA_INDEPENDIENTE", {})
    if isinstance(ti, dict) and ti:
        doc.add_heading("XIII. Tarea Independiente", level=1)
        pd_ = doc.add_paragraph(); pd_.add_run("Descripción: ").bold = True; pd_.add_run(limpiar_texto(ti.get("DESCRIPCION","")))
        pe = doc.add_paragraph(); pe.add_run("Entregable: ").bold = True; pe.add_run(limpiar_texto(ti.get("ENTREGABLE","")))
    if datos.get("SIMULADORES_RECURSOS"):
        doc.add_heading("XIV. Repositorio de Simuladores y Recursos", level=1)
        for sim in datos["SIMULADORES_RECURSOS"]:
            if isinstance(sim, dict):
                ps = doc.add_paragraph(style="List Bullet")
                ps.add_run(f"[{limpiar_texto(sim.get('TIPO',''))}] ").bold = True
                ps.add_run(f"{limpiar_texto(sim.get('NOMBRE',''))}: ").bold = True
                ps.add_run(f"{limpiar_texto(sim.get('DESCRIPCION',''))} ")
                ps.add_run(f"[Bloom: {limpiar_texto(sim.get('NIVEL_BLOOM',''))}] ").italic = True
                if sim.get("URL"): ps.add_run(f"🔗 {limpiar_texto(sim['URL'])}").italic = True
    if datos.get("WEBGRAFIA"):
        doc.add_heading("XV. Fuentes y Referencias", level=1)
        for ref in datos["WEBGRAFIA"]: add_bullet(doc, limpiar_texto(ref))
    # ── Sección docente ──
    doc.add_page_break()
    pdt = doc.add_paragraph(); pdt.alignment = WD_ALIGN_PARAGRAPH.CENTER
    rdt = pdt.add_run("🔑 SECCIÓN PARA EL DOCENTE\n"); rdt.bold = True; rdt.font.size = Pt(14)
    pdt.add_run("Rúbrica Multinivel · Guía de Respuestas · Autoevaluación")
    doc.add_paragraph("_" * 70)
    criterios = datos.get("CRITERIOS_EVALUACION", [])
    if criterios:
        doc.add_heading("Rúbrica de Evaluación", level=1)
        t = doc.add_table(rows=1, cols=6); t.style = "Table Grid"
        hdr = ["No.", "Criterio e Indicador", "L (Logrado)", "EP (En Proceso)", "NA (Necesita Apoyo)", "Pts"]
        for i, h in enumerate(hdr):
            set_cell_text(t.rows[0].cells[i], h, bold=True, center=(i != 1), color="E2E8F0")
        for idx, crit in enumerate(criterios, 1):
            if isinstance(crit, dict):
                row = t.add_row().cells
                set_cell_text(row[0], str(idx), center=True)
                set_cell_text(row[1], f"• {limpiar_texto(crit.get('CRITERIO',''))}\n• Ind: {limpiar_texto(crit.get('INDICADOR',''))}")
                set_cell_text(row[2], limpiar_texto(crit.get("LOGRADO","")), color="D1FAE5")
                set_cell_text(row[3], limpiar_texto(crit.get("EN_PROCESO","")), color="FEF3C7")
                set_cell_text(row[4], limpiar_texto(crit.get("NECESITA_APOYO","")), color="FEE2E2")
                set_cell_text(row[5], str(convertir_puntos(crit.get("PUNTOS", 0))), center=True)
        row = t.add_row().cells
        set_cell_text(row[1], "TOTAL", bold=True)
        set_cell_text(row[5], str(meta.get("valor_puntos", 0)), bold=True, center=True)
    if datos.get("AUTOEVALUACION"):
        doc.add_heading("Autoevaluación del Estudiante", level=2)
        t = doc.add_table(rows=len(datos["AUTOEVALUACION"]) + 1, cols=4); t.style = "Table Grid"
        for i, h in enumerate(["Reflexión", "Sí", "Parcialmente", "No"]):
            set_cell_text(t.rows[0].cells[i], h, bold=True, center=(i != 0), color="F1F5F9")
        for idx, it in enumerate(datos["AUTOEVALUACION"], 1):
            set_cell_text(t.cell(idx, 0), limpiar_texto(it))
            for j in range(1, 4): set_cell_text(t.cell(idx, j), "☐", center=True)
    if datos.get("GUIA_RESPUESTAS"):
        doc.add_heading("Guía de Respuestas (Solucionario)", level=2)
        t = doc.add_table(rows=1, cols=2); t.style = "Table Grid"
        set_cell_text(t.rows[0].cells[0], "Paso", bold=True, color="DBEAFE")
        set_cell_text(t.rows[0].cells[1], "Respuesta / Resultado Esperado", bold=True, color="DBEAFE")
        for it in datos["GUIA_RESPUESTAS"]:
            if isinstance(it, dict):
                row = t.add_row().cells
                set_cell_text(row[0], limpiar_texto(it.get("PASO","")), bold=True)
                set_cell_text(row[1], limpiar_texto(it.get("RESPUESTA_ESPERADA","")))
    agregar_numeracion_pagina(doc)
    buffer = BytesIO(); doc.save(buffer); buffer.seek(0)
    return buffer

# ═══════════════════════════════════════════════════════════════════════════
# INTERFAZ (SUPER UI)
# ═══════════════════════════════════════════════════════════════════════════
st.markdown("""
<style>
@import url('https://fonts.googleapis.com/css2?family=Inter:wght@300;400;500;600;700;800&display=swap');
html, body, [class*="css"] { font-family: 'Inter', sans-serif; background-color: #F8FAFC; color: #1E293B; }
.cont-hero { background: linear-gradient(135deg, #0F172A 0%, #14532D 55%, #16A34A 100%); color: #fff;
    padding: 1.8rem; border-radius: 18px; margin-bottom: 1.25rem; box-shadow: 0 18px 35px rgba(15,23,42,.18); }
.cont-title { font-size: 2rem; font-weight: 800; letter-spacing: -0.02em; }
.cont-sub { opacity: .88; font-size: 1rem; margin-top: .3rem; }
.section-title { color: #1D4ED8; font-weight: 700; font-size: 1.12rem; border-bottom: 2px solid #DBEAFE;
    padding-bottom: 8px; margin: 1.2rem 0 .9rem 0; }
</style>
""", unsafe_allow_html=True)

ia.panel_sidebar_ia("Generador de Contenidos")

st.markdown("""
<div class="cont-hero">
    <div class="cont-title">📚 Generador de Contenidos y Actividades ETP</div>
    <div class="cont-sub">Contenido anclado al currículo · Progresión Bloom · Diferenciación · Rúbrica multinivel · Solucionario</div>
</div>
""", unsafe_allow_html=True)

if "contenido_resultado" not in st.session_state:
    st.session_state.contenido_resultado = None

with st.form("form_contenido", clear_on_submit=False):
    st.markdown('<div class="section-title">📄 1. Fuente Curricular (Opcional)</div>', unsafe_allow_html=True)
    archivo_pdf = st.file_uploader("Cargue el PDF del módulo para anclar el contenido", type=["pdf"],
                                   help="Si subes el PDF, la IA generará contenido fiel al currículo oficial.")
    st.markdown('<div class="section-title">🏫 2. Datos Institucionales y Curriculares</div>', unsafe_allow_html=True)
    c1, c2 = st.columns(2)
    with c1:
        docente = st.text_input("Nombre del Docente", value="Ing. Bernardo Antonio Hernández Batista")
        asignatura = st.text_input("Módulo / Asignatura", placeholder="Ej: Ofimática, Redes LAN")
    with c2:
        politecnico = st.text_input("Centro Educativo", value="Politécnico Salesiano Arquides Calderón")
        fecha = st.date_input("Fecha de Aplicación", value=date.today())
    c_ra, c_ce = st.columns(2)
    with c_ra:
        ra = st.text_area("Resultado de Aprendizaje (RA)", height=70, placeholder="Pega el RA completo")
    with c_ce:
        ce_ec = st.text_area("Criterio de Evaluación (CE) / Elemento de Capacidad (EC)", height=70)
    st.markdown('<div class="section-title">📚 3. Base Pedagógica y Contenido</div>', unsafe_allow_html=True)
    contenido = st.text_area("Contenido a Desarrollar", height=90, placeholder="Ej: Configuración de subredes IP y enrutamiento estático.")
    actividad = st.text_area("Actividad Práctica de Clase", height=90, placeholder="Ej: Los estudiantes simularán una red LAN...")
    st.markdown('<div class="section-title">👥 4. Perfil del Grupo y Condiciones</div>', unsafe_allow_html=True)
    cp1, cp2, cp3 = st.columns([2, 1, 1])
    with cp1:
        caracteristicas_grupo = st.text_area("Características del grupo / NEAE", height=68)
    with cp2:
        duracion_sesion = st.text_input("Duración de la sesión", value="50 minutos")
    with cp3:
        nivel_bloom_objetivo = st.selectbox("Nivel Bloom objetivo",
            ["Recordar / Comprender", "Comprender / Aplicar", "Aplicar / Analizar", "Analizar / Evaluar", "Evaluar / Crear"], index=2)
    st.markdown('<div class="section-title">📋 5. Estrategia de Evaluación</div>', unsafe_allow_html=True)
    ce1, ce2, ce3 = st.columns([3, 1, 1])
    with ce1:
        instrumento = st.selectbox("Técnica / Instrumento de Evaluación", [
            "Rúbrica Analítica de Competencias ETP (con niveles L / EP / NA)",
            "Lista de Cotejo Avanzada (Indicadores de Logro)",
            "Escala Estimativa con Descriptores de Desempeño",
            "Guía de Observación Metodológica",
            "Registro de Desempeño Técnico"])
    with ce2:
        valor_puntos = st.number_input("Valor (Puntos)", min_value=1, max_value=100, value=100)
    with ce3:
        cant_criterios = st.number_input("Criterios", min_value=3, max_value=8, value=5)
    max_tokens, temperature = ia.control_avanzado(default_tokens=16384, tope=TOPE_MAX_TOKENS, default_temp=0.1)
    modo_debug = st.checkbox("🐛 Modo depuración", value=False)
    st.markdown("<br>", unsafe_allow_html=True)
    submit_button = st.form_submit_button("⚙️ Generar Material Didáctico Integral (Word)", type="primary", width="stretch")

if submit_button:
    cfg = ia.config_ia()
    if not cfg["api_key"]:
        st.error("🔒 Debes configurar tu API Key en la barra lateral.")
    elif not asignatura or not contenido:
        st.warning("📝 Completa la asignatura y el contenido a desarrollar.")
    else:
        with st.spinner(f"🧠 Generando contenido integral con {cfg['modelo']}..."):
            texto_crudo = None
            try:
                texto_pdf = extraer_texto_pdf(archivo_pdf) if archivo_pdf else ""
                prompt = construir_prompt_maestro({
                    "contenido": contenido, "actividad": actividad, "instrumento": instrumento,
                    "valor_puntos": int(valor_puntos), "cant_criterios": int(cant_criterios),
                    "duracion_sesion": duracion_sesion, "nivel_bloom_objetivo": nivel_bloom_objetivo,
                    "caracteristicas_grupo": caracteristicas_grupo, "ra": ra, "ce_ec": ce_ec,
                    "texto_pdf": texto_pdf,
                })
                texto_crudo, flags = ia.solicitar_ia(
                    prompt, modo="json", max_tokens=max_tokens, temperature=temperature, modulo="contenido")
                datos_brutos = ia.decodificar_marcadores(ia.parsear_json_robusto(texto_crudo))
                datos, advertencias = validar_y_normalizar_datos(datos_brutos, int(cant_criterios), int(valor_puntos))
                st.session_state.contenido_resultado = {
                    "datos": datos, "raw": texto_crudo, "advertencias": advertencias,
                    "flags": flags, "pdf_anclado": bool(texto_pdf),
                    "meta": {"docente": docente, "asignatura": asignatura, "politecnico": politecnico,
                             "fecha": fecha, "ra": ra, "ce_ec": ce_ec, "duracion_sesion": duracion_sesion,
                             "nivel_bloom_objetivo": nivel_bloom_objetivo, "instrumento": instrumento,
                             "valor_puntos": int(valor_puntos)},
                }
                st.toast("✅ Material Didáctico Integral generado.", icon="📚")
            except ValueError as ve:
                ia.render_error_ia(ve, texto_crudo)
            except Exception as e:
                ia.render_error_ia(e, texto_crudo)

resultado = st.session_state.contenido_resultado
if resultado:
    st.markdown('<div class="section-title">📦 Resultado generado</div>', unsafe_allow_html=True)
    d = resultado["datos"]
    m1, m2, m3, m4 = st.columns(4)
    with m1: st.metric("📄 Secciones", sum(1 for k in CLAVES_LISTA + ["RESUMEN_EJECUTIVO", "ADAPTACIONES_NEAE"] if d.get(k)))
    with m2: st.metric("📚 Secciones Bloom", len(d.get("CONTENIDO_TEORICO", [])))
    with m3: st.metric("🎯 Criterios", len(d.get("CRITERIOS_EVALUACION", [])))
    with m4: st.metric("🔗 Simuladores", len(d.get("SIMULADORES_RECURSOS", [])))
    if resultado.get("pdf_anclado"):
        st.info("📌 Contenido anclado al PDF curricular.")
    else:
        st.warning("⚠️ Sin PDF: contenido generado desde descripción libre.")
    if resultado["flags"].get("reintento"):
        st.info("ℹ️ La primera llamada se cortó por tokens; el reintento automático tuvo éxito.")
    if resultado.get("advertencias"):
        with st.expander(f"⚠️ Advertencias de validación ({len(resultado['advertencias'])})"):
            for a in resultado["advertencias"]: st.markdown(f"- {a}")
    tab_prev, tab_json, tab_debug = st.tabs(["👁️ Vista previa", "🧾 JSON", "🐛 Depuración"])
    with tab_prev:
        st.markdown("#### 🧭 Resumen ejecutivo")
        st.write(limpiar_texto(d.get("RESUMEN_EJECUTIVO")) or "Sin resumen.")
        with st.expander("📘 Contenido teórico"):
            if d.get("CONTENIDO_TEORICO"):
                for sec in d["CONTENIDO_TEORICO"]:
                    if isinstance(sec, dict):
                        st.markdown(f"**[{sec.get('NIVEL_BLOOM','')}] {sec.get('TITULO_SECCION','')}**")
                        st.caption(limpiar_texto(sec.get("CONTENIDO",""))[:400])
            else: st.info("Sin contenido teórico.")
        with st.expander("🎯 Actividad principal"):
            st.json(d.get("ACTIVIDAD_PRINCIPAL", {}))
        with st.expander("🧾 Criterios de evaluación"):
            if d.get("CRITERIOS_EVALUACION"):
                import pandas as pd
                st.dataframe(pd.DataFrame(d["CRITERIOS_EVALUACION"]), width="stretch", hide_index=True)
    with tab_json:
        st.json(d)
    with tab_debug:
        st.text_area("Respuesta cruda", resultado.get("raw", ""), height=220)
        st.write(resultado["flags"])
    buffer = build_docx(d, resultado["meta"])
    nombre = f"Material_Integral_{ia.sanear_nombre_archivo(resultado['meta']['asignatura'])}_{resultado['meta']['fecha'].strftime('%Y%m%d')}.docx"
    st.download_button("📥 Descargar Material Didáctico Integral (.docx)", data=buffer, file_name=nombre,
                       mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                       type="primary", width="stretch")
    if st.button("🗑️ Limpiar resultado generado", width="stretch"):
        st.session_state.contenido_resultado = None
        st.rerun()