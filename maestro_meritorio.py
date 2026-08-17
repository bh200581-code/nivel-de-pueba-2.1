"""
maestro_meritorio.py — Selección Maestro Meritorio ETP (Fígital 100% Nativo)
• Conexión con base de datos para lista desplegable automática de Docentes.
• Generación nativa de Links y Códigos QR para evaluar online.
• Sistema de evaluación en 4 fases con instrumentos imprimibles en Word.
• GUARDADO PERMANENTE: Las fases se guardan por partes en la base de datos.
• Corrector IA de fichas escaneadas e informe integrado final.
"""
import re
import sqlite3
import json
import socket
import random
import string
from datetime import datetime
from io import BytesIO
from typing import Any, Dict, Optional
from urllib.parse import quote

import pandas as pd
import streamlit as st
from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement, parse_xml
from docx.oxml.ns import qn, nsdecls
from docx.shared import Inches, Pt, RGBColor

try:
    from pypdf import PdfReader
except Exception:
    try:
        from PyPDF2 import PdfReader
    except Exception:
        PdfReader = None

from core import ia

DB_NAME = "gestion_etp.db"

# ═══════════════════════════════════════════════════════════════════════════
# CATÁLOGOS DE CRITERIOS POR FASE (alineados al MINERD)
# ═══════════════════════════════════════════════════════════════════════════
CRITERIOS_AUTOEVALUACION = [
    {"id": "AE1", "criterio": "Dominio y actualización del contenido técnico de mi módulo formativo", "dimension": "Dominio Técnico"},
    {"id": "AE2", "criterio": "Alineación de mi planificación diaria con los RA, CE y EC del diseño curricular", "dimension": "Planificación"},
    {"id": "AE3", "criterio": "Aplico estrategias de enseñanza activas y pertinentes a la modalidad ETP", "dimension": "Metodología"},
    {"id": "AE4", "criterio": "Aplico evaluación formativa y sumativa coherente con los resultados de aprendizaje", "dimension": "Evaluación"},
    {"id": "AE5", "criterio": "Atiendo la diversidad y aplico adaptaciones para estudiantes con NEAE", "dimension": "Diversidad"},
    {"id": "AE6", "criterio": "Integro recursos TIC y herramientas tecnológicas en mi práctica docente", "dimension": "Recursos y TIC"},
    {"id": "AE7", "criterio": "Promuevo un ambiente de aprendizaje respetuoso y motivador", "dimension": "Ambiente"},
    {"id": "AE8", "criterio": "Me comprometo con mi desarrollo profesional continuo y la innovación", "dimension": "Desarrollo Profesional"},
]

CRITERIOS_ESTUDIANTES = [
    {"id": "EST1", "criterio": "El docente explica los temas de forma clara y comprensible", "dimension": "Comunicación"},
    {"id": "EST2", "criterio": "El docente trata a todos los estudiantes con respeto y amabilidad", "dimension": "Convivencia"},
    {"id": "EST3", "criterio": "El docente organiza bien la clase y aprovecha el tiempo", "dimension": "Organización"},
    {"id": "EST4", "criterio": "El docente me ayuda cuando tengo dificultades para aprender", "dimension": "Apoyo"},
    {"id": "EST5", "criterio": "El docente me motiva a aprender y dar lo mejor de mí", "dimension": "Motivación"},
    {"id": "EST6", "criterio": "El docente evalúa mis aprendizajes de forma justa", "dimension": "Equidad"},
    {"id": "EST7", "criterio": "El docente relaciona los temas con situaciones reales del trabajo", "dimension": "Pertinencia"},
    {"id": "EST8", "criterio": "El docente usa recursos y tecnologías que facilitan mi aprendizaje", "dimension": "Recursos"},
]

CRITERIOS_COORDINACION = [
    {"id": "CO1", "criterio": "Dominio técnico-profesional del módulo formativo que imparte", "dimension": "Dominio Técnico"},
    {"id": "CO2", "criterio": "Alineación de la planificación con RA, CE y EC del diseño curricular", "dimension": "Planificación"},
    {"id": "CO3", "criterio": "Aplicación de metodologías activas pertinentes a la ETP", "dimension": "Metodología"},
    {"id": "CO4", "criterio": "Evaluación formativa y retroalimentación oportuna a los estudiantes", "dimension": "Evaluación"},
    {"id": "CO5", "criterio": "Gestión del ambiente de aprendizaje y convivencia en aula/taller", "dimension": "Ambiente"},
    {"id": "CO6", "criterio": "Vinculación de la enseñanza con situaciones reales del sector productivo", "dimension": "Pertinencia ETP"},
    {"id": "CO7", "criterio": "Uso de recursos, equipos y tecnologías pertinentes al módulo", "dimension": "Recursos"},
    {"id": "CO8", "criterio": "Atención a la diversidad y adecuaciones curriculares para NEAE", "dimension": "Diversidad"},
    {"id": "CO9", "criterio": "Documentación del proceso de enseñanza (planificaciones, registros, evidencias)", "dimension": "Documentación"},
    {"id": "CO10", "criterio": "Participación en capacitación continua y desarrollo profesional", "dimension": "Desarrollo Profesional"},
    {"id": "CO11", "criterio": "Colaboración con el equipo docente y la coordinación", "dimension": "Trabajo Colaborativo"},
    {"id": "CO12", "criterio": "Mejora continua a partir de los resultados de evaluación", "dimension": "Mejora Continua"},
]

CRITERIOS_DIRECCION = [
    {"id": "DI1", "criterio": "Puntualidad y asistencia a sus responsabilidades docentes", "dimension": "Responsabilidad"},
    {"id": "DI2", "criterio": "Ética, responsabilidad y compromiso profesional", "dimension": "Ética"},
    {"id": "DI3", "criterio": "Alineación con la misión institucional y el modelo educativo", "dimension": "Identidad Institucional"},
    {"id": "DI4", "criterio": "Trabajo en equipo y clima escolar positivo", "dimension": "Clima Escolar"},
    {"id": "DI5", "criterio": "Comunicación efectiva con la comunidad educativa", "dimension": "Comunicación"},
    {"id": "DI6", "criterio": "Participación en actividades y proyectos institucionales", "dimension": "Participación"},
    {"id": "DI7", "criterio": "Innovación y mejora continua de su práctica docente", "dimension": "Innovación"},
    {"id": "DI8", "criterio": "Relación con las familias y la comunidad", "dimension": "Comunidad"},
    {"id": "DI9", "criterio": "Cuidado de los espacios, equipos y recursos institucionales", "dimension": "Recursos"},
    {"id": "DI10", "criterio": "Contribución a los indicadores de calidad institucional", "dimension": "Calidad"},
]

FASES_EVALUACION = {
    "autoevaluacion": {
        "nombre": "Autoevaluación Docente",
        "icono": "🪞",
        "color": "#8B5CF6",
        "peso": 0.15,
        "evaluador": "El propio docente",
        "criterios": CRITERIOS_AUTOEVALUACION,
        "instrucciones": "Evalúe su propia práctica pedagógica con honestidad y reflexión crítica. Seleccione el nivel que mejor describe su desempeño.",
    },
    "estudiantes": {
        "nombre": "Evaluación por Estudiantes",
        "icono": "🎓",
        "color": "#10B981",
        "peso": 0.25,
        "evaluador": "Los estudiantes del módulo",
        "criterios": CRITERIOS_ESTUDIANTES,
        "instrucciones": "Evalúa a tu docente seleccionando el nivel que mejor describe su desempeño. Tu opinión es 100% anónima y confidencial.",
    },
    "coordinacion": {
        "nombre": "Evaluación por Coordinación ETP",
        "icono": "📋",
        "color": "#3B82F6",
        "peso": 0.35,
        "evaluador": "Coordinación Técnico-Pedagógica ETP",
        "criterios": CRITERIOS_COORDINACION,
        "instrucciones": "Evalúa el desempeño docente a partir de la observación de aula, revisión de planificaciones y evidencias.",
    },
    "direccion": {
        "nombre": "Evaluación por Dirección",
        "icono": "🏛️",
        "color": "#F59E0B",
        "peso": 0.25,
        "evaluador": "Dirección del Centro Educativo",
        "criterios": CRITERIOS_DIRECCION,
        "instrucciones": "Evalúa el desempeño institucional considerando responsabilidad, ética, participación y clima escolar.",
    },
}

ESCALA_VALORACION = {
    4: ("Excelente", "Demuestra SIEMPRE el criterio con alta calidad y consistencia"),
    3: ("Bueno", "Demuestra FRECUENTEMENTE el criterio con calidad"),
    2: ("Aceptable", "Demuestra OCASIONALMENTE el criterio"),
    1: ("Requiere Mejora", "Demuestra RARA VEZ o NUNCA el criterio"),
}

# ═══════════════════════════════════════════════════════════════════════════
# ESTILOS COMPARTIDOS
# ═══════════════════════════════════════════════════════════════════════════
st.markdown("""
<style>
@import url('https://fonts.googleapis.com/css2?family=Inter:wght@300;400;500;600;700;800;900&display=swap');
html, body, [class*="css"] { font-family: 'Inter', sans-serif; background-color: #F0F4F8; color: #1E293B; }

.meritorio-hero {
    background: linear-gradient(135deg, #451A03 0%, #92400E 30%, #D97706 60%, #F59E0B 100%);
    color: #fff; padding: 2.2rem; border-radius: 20px; margin-bottom: 1.5rem;
    box-shadow: 0 25px 50px rgba(146, 64, 14, 0.3); position: relative; overflow: hidden;
}
.meritorio-hero::before { content: '🏆'; position: absolute; right: 2rem; top: 50%; transform: translateY(-50%); font-size: 6rem; opacity: 0.15; }
.meritorio-hero-title { font-size: 2.4rem; font-weight: 900; letter-spacing: -0.03em; margin-bottom: 0.4rem; }
.meritorio-hero-sub { font-size: 1.05rem; opacity: 0.9; line-height: 1.5; }
.meritorio-hero-badge { display: inline-block; background: rgba(255,255,255,0.15); border: 1px solid rgba(255,255,255,0.25); border-radius: 8px; padding: 4px 12px; font-size: 0.8rem; font-weight: 600; margin-top: 0.8rem; margin-right: 8px; }

.meritorio-stepper { display: flex; align-items: center; justify-content: center; gap: 0; margin: 1.5rem 0; }
.meritorio-step { display: flex; flex-direction: column; align-items: center; gap: 6px; }
.meritorio-step-circle { width: 52px; height: 52px; border-radius: 50%; display: flex; align-items: center; justify-content: center; font-weight: 800; font-size: 1.3rem; }
.meritorio-step-circle.inactive { background: #E2E8F0; color: #94A3B8; border: 3px solid #CBD5E1; }
.meritorio-step-circle.active { background: linear-gradient(135deg, #D97706, #F59E0B); color: #fff; border: 3px solid #D97706; box-shadow: 0 4px 15px rgba(217, 119, 6, 0.4); }
.meritorio-step-circle.done { background: linear-gradient(135deg, #059669, #34D399); color: #fff; border: 3px solid #059669; box-shadow: 0 4px 15px rgba(5, 150, 105, 0.3); }
.meritorio-step-label { font-size: 0.72rem; font-weight: 700; color: #64748B; text-transform: uppercase; }
.meritorio-step-line { width: 60px; height: 3px; background: #CBD5E1; margin: 0 4px; margin-bottom: 22px; }
.meritorio-step-line.done { background: linear-gradient(90deg, #059669, #34D399); }

.meritorio-section-title { color: #92400E; font-weight: 700; font-size: 1.12rem; border-bottom: 2px solid #FEF3C7; padding-bottom: 8px; margin: 1.2rem 0 0.9rem 0; }
.meritorio-stat { background: #fff; border-radius: 12px; padding: 1rem; text-align: center; border: 1px solid #E2E8F0; box-shadow: 0 2px 6px rgba(0,0,0,0.04); }
.meritorio-stat-value { font-size: 1.8rem; font-weight: 800; color: #92400E; }
.meritorio-stat-label { font-size: 0.72rem; font-weight: 700; color: #94A3B8; text-transform: uppercase; }
.meritorio-fase-card { background: #fff; border: 2px solid #E2E8F0; border-radius: 14px; padding: 1.2rem; box-shadow: 0 2px 8px rgba(0,0,0,0.04); }
.meritorio-progress-bar { height: 8px; border-radius: 4px; background: #E2E8F0; overflow: hidden; margin-top: 8px; }
.meritorio-progress-fill { height: 100%; border-radius: 4px; }
.pd-linkbox { background:#F5F3FF; border:2px dashed #A78BFA; border-radius:12px; padding:1rem; margin-top:0.8rem; font-family:monospace; font-size:0.95rem; word-break:break-all; }
.pd-qcard { background:#fff; border:2px solid #E2E8F0; border-radius:14px; padding:1.2rem; margin-bottom:0.6rem; box-shadow:0 2px 8px rgba(0,0,0,0.04); }
</style>
""", unsafe_allow_html=True)

# ═══════════════════════════════════════════════════════════════════════════
# BASE DE DATOS Y ESTADO CON GUARDADO PERMANENTE
# ═══════════════════════════════════════════════════════════════════════════
def _conn():
    return sqlite3.connect(DB_NAME, check_same_thread=False)

def init_db():
    conn = _conn()
    cursor = conn.cursor()
    cursor.execute('''CREATE TABLE IF NOT EXISTS mm_links (
        codigo TEXT PRIMARY KEY,
        docente TEXT,
        periodo TEXT,
        fase TEXT,
        fecha TEXT
    )''')
    cursor.execute('''CREATE TABLE IF NOT EXISTS mm_respuestas (
        id INTEGER PRIMARY KEY AUTOINCREMENT,
        codigo TEXT,
        evaluador TEXT,
        scores_json TEXT,
        comentarios TEXT,
        fecha TEXT
    )''')
    # ── TABLA PARA GUARDADO PARCIAL ──
    cursor.execute('''CREATE TABLE IF NOT EXISTS mm_consolidado (
        docente TEXT,
        periodo TEXT,
        fase TEXT,
        scores_json TEXT,
        comentarios TEXT,
        PRIMARY KEY (docente, periodo, fase)
    )''')
    conn.commit()
    conn.close()

init_db()

def init_estado():
    if "meritorio_docente" not in st.session_state: st.session_state.meritorio_docente = None
    if "meritorio_periodo" not in st.session_state: st.session_state.meritorio_periodo = "2025-2026 - Primer Trimestre"
    if "meritorio_scores" not in st.session_state:
        st.session_state.meritorio_scores = {"autoevaluacion": {}, "estudiantes": {}, "coordinacion": {}, "direccion": {}}
    if "meritorio_comentarios" not in st.session_state:
        st.session_state.meritorio_comentarios = {"autoevaluacion": "", "estudiantes": "", "coordinacion": "", "direccion": ""}
    if "meritorio_integrado" not in st.session_state: st.session_state.meritorio_integrado = None

init_estado()

def cargar_consolidado(docente, periodo):
    """Carga los resultados parciales de la Base de Datos para que no se pierdan y se unan al final."""
    conn = _conn()
    cur = conn.cursor()
    cur.execute("SELECT fase, scores_json, comentarios FROM mm_consolidado WHERE docente=? AND periodo=?", (docente, periodo))
    rows = cur.fetchall()
    conn.close()
    
    # Limpiamos para no arrastrar datos de otro docente
    st.session_state.meritorio_scores = {"autoevaluacion": {}, "estudiantes": {}, "coordinacion": {}, "direccion": {}}
    st.session_state.meritorio_comentarios = {"autoevaluacion": "", "estudiantes": "", "coordinacion": "", "direccion": ""}
    
    for fase, scores_json, comentarios in rows:
        if fase in st.session_state.meritorio_scores:
            try:
                st.session_state.meritorio_scores[fase] = json.loads(scores_json)
                st.session_state.meritorio_comentarios[fase] = comentarios if comentarios else ""
            except Exception: pass

def guardar_consolidado(docente, periodo, fase, scores, comentarios):
    """Guarda una fase completada en la Base de Datos de forma permanente."""
    conn = _conn()
    conn.execute('''REPLACE INTO mm_consolidado (docente, periodo, fase, scores_json, comentarios) 
                    VALUES (?, ?, ?, ?, ?)''', 
                 (docente, periodo, fase, json.dumps(scores), comentarios))
    conn.commit()
    conn.close()

# ═══════════════════════════════════════════════════════════════════════════
# UTILIDADES WORD, PDF Y LINKS
# ═══════════════════════════════════════════════════════════════════════════
def shade_cell(cell, color):
    shd = parse_xml(r'<w:shd {} w:fill="{}"/>'.format(nsdecls("w"), color))
    cell._tc.get_or_add_tcPr().append(shd)

def set_cell_text(cell, text, bold=False, center=False, color=None):
    cell.text = ""
    p = cell.paragraphs[0]
    run = p.add_run(str(text))
    run.bold = bold
    if center: p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    if color: shade_cell(cell, color)

def add_page_number(doc):
    for section in doc.sections:
        footer = section.footer
        footer.is_linked_to_previous = False
        p = footer.paragraphs[0] if footer.paragraphs else footer.add_paragraph()
        p.text = ""
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        run = p.add_run()
        fld1 = OxmlElement("w:fldChar"); fld1.set(qn("w:fldCharType"), "begin")
        instr = OxmlElement("w:instrText"); instr.set(qn("xml:space"), "preserve"); instr.text = "PAGE"
        fld2 = OxmlElement("w:fldChar"); fld2.set(qn("w:fldCharType"), "end")
        run._r.append(fld1); run._r.append(instr); run._r.append(fld2)
        
        run_sep = p.add_run(" / ")
        
        run2 = p.add_run()
        fld3 = OxmlElement("w:fldChar"); fld3.set(qn("w:fldCharType"), "begin")
        instr2 = OxmlElement("w:instrText"); instr2.set(qn("xml:space"), "preserve"); instr2.text = "NUMPAGES"
        fld4 = OxmlElement("w:fldChar"); fld4.set(qn("w:fldCharType"), "end")
        run2._r.append(fld3); run2._r.append(instr2); run2._r.append(fld4)
        
        for r in (run, run_sep, run2):
            r.font.size = Pt(8)
            r.font.color.rgb = RGBColor(0x64, 0x74, 0x8B)

def extraer_texto_pdf(archivo):
    if PdfReader is None: raise RuntimeError("Instala pypdf o PyPDF2.")
    archivo.seek(0)
    reader = PdfReader(archivo)
    texto = "".join((p.extract_text() or "") for p in reader.pages)
    return re.sub(r"\s+", " ", texto).strip()[:40000]

def detectar_ip_local():
    try:
        s = socket.socket(socket.AF_INET, socket.SOCK_DGRAM)
        s.connect(("8.8.8.8", 80))
        ip = s.getsockname()[0]
        s.close()
        return ip
    except Exception: return "127.0.0.1"

def construir_link_eval(codigo):
    base = st.secrets.get("url_publica", "") if hasattr(st, "secrets") else ""
    if not base:
        import os
        base = os.environ.get("STREAMLIT_URL", "")
    if base: return f"{base.rstrip('/')}/?eval={codigo}"
    return f"http://{detectar_ip_local()}:8501/?eval={codigo}"

def qr_url(link):
    return f"https://api.qrserver.com/v1/create-qr-code/?size=200x200&data={quote(link)}"

def obtener_docentes_bd() -> list:
    conn = _conn()
    cur = conn.cursor()
    try:
        cur.execute("SELECT DISTINCT docente FROM docentes WHERE docente IS NOT NULL AND docente != '' ORDER BY docente")
        rows = [r[0] for r in cur.fetchall()]
    except Exception: rows = []
    conn.close()
    return rows

def promediar_respuestas(codigo):
    conn = _conn()
    cur = conn.cursor()
    cur.execute("SELECT scores_json, comentarios FROM mm_respuestas WHERE codigo=?", (codigo,))
    rows = cur.fetchall()
    conn.close()
    
    if not rows: return {}, ""
    
    sum_scores = {}
    counts = {}
    comentarios_list = []
    
    for r in rows:
        try:
            scores = json.loads(r[0])
        except Exception: continue
        
        coment = r[1]
        if coment and coment.strip(): comentarios_list.append(coment.strip())
        
        for k, v in scores.items():
            sum_scores[k] = sum_scores.get(k, 0) + v
            counts[k] = counts.get(k, 0) + 1
            
    avg_scores = {}
    for k in sum_scores:
        avg_scores[k] = int(round(sum_scores[k] / counts[k]))
        
    comentarios_unidos = "\n---\n".join(comentarios_list)
    return avg_scores, comentarios_unidos

# ═══════════════════════════════════════════════════════════════════════════
# VISTA ONLINE PARA LLENAR (Sin Login)
# ═══════════════════════════════════════════════════════════════════════════
def render_vista_online(codigo):
    st.markdown("<style>[data-testid='stSidebar'] {display:none !important;}</style>", unsafe_allow_html=True)
    
    conn = _conn()
    cur = conn.cursor()
    cur.execute("SELECT docente, periodo, fase FROM mm_links WHERE codigo=?", (codigo,))
    row = cur.fetchone()
    conn.close()
    
    if not row:
        st.error("❌ Enlace de evaluación no válido o expirado.")
        st.stop()
        
    docente, periodo, fase_key = row
    fase = FASES_EVALUACION.get(fase_key)
    
    if st.session_state.get(f"eval_enviada_{codigo}"):
        st.markdown(f"""
        <div class="meritorio-hero" style="background: linear-gradient(135deg, #065F46, #10B981);">
            <div class="meritorio-hero-title">🎉 ¡Evaluación Completada!</div>
            <div class="meritorio-hero-sub">Tus respuestas han sido registradas con éxito. Gracias por tu tiempo.</div>
        </div>
        """, unsafe_allow_html=True)
        st.stop()
        
    st.markdown(f"""
    <div class="meritorio-hero">
        <div class="meritorio-hero-title">{fase['icono']} {fase['nombre']}</div>
        <div class="meritorio-hero-sub">Docente evaluado: <b>{docente}</b> | Período: {periodo}</div>
    </div>
    """, unsafe_allow_html=True)
    
    st.info(fase["instrucciones"])
    
    with st.form("form_eval_online"):
        st.markdown('<div class="meritorio-section-title">📋 Criterios de Evaluación</div>', unsafe_allow_html=True)
        respuestas = {}
        for i, crit in enumerate(fase["criterios"], 1):
            st.markdown(f'<div class="pd-qcard"><b>{i}. {crit["criterio"]}</b><br><small style="color:#64748B;">Dimensión: {crit.get("dimension", "")}</small></div>', unsafe_allow_html=True)
            opciones = ["Seleccione...", "4 - Excelente", "3 - Bueno", "2 - Aceptable", "1 - Requiere Mejora"]
            resp = st.selectbox(f"Valoración para {crit['id']}", opciones, key=f"ol_{crit['id']}", label_visibility="collapsed")
            respuestas[crit["id"]] = resp
            
        st.markdown('<div class="meritorio-section-title">💬 Comentarios Opcionales</div>', unsafe_allow_html=True)
        comentarios = st.text_area("¿Tienes alguna observación adicional que justifique tu evaluación?", height=100)
        
        submit = st.form_submit_button("✅ Enviar Evaluación", type="primary", use_container_width=True)
        
    if submit:
        faltan = [k for k, v in respuestas.items() if v == "Seleccione..."]
        if faltan:
            st.warning("⚠️ Por favor, califica todos los criterios antes de enviar.")
        else:
            scores_json = {}
            for k, v in respuestas.items():
                scores_json[k] = int(v.split(" - ")[0])
            
            conn = _conn()
            conn.execute("INSERT INTO mm_respuestas (codigo, evaluador, scores_json, comentarios, fecha) VALUES (?, ?, ?, ?, ?)",
                         (codigo, "Anónimo", json.dumps(scores_json), comentarios, datetime.now().strftime("%Y-%m-%d %H:%M")))
            conn.commit()
            conn.close()
            
            st.session_state[f"eval_enviada_{codigo}"] = True
            st.rerun()

# ═══════════════════════════════════════════════════════════════════════════
# DETECCIÓN DE MODO (ONLINE VS GESTIÓN)
# ═══════════════════════════════════════════════════════════════════════════
_codigo_qs = st.query_params.get("eval")
if _codigo_qs:
    render_vista_online(_codigo_qs)
    st.stop()

# ═══════════════════════════════════════════════════════════════════════════
# CORE LÓGICO Y CÁLCULOS
# ═══════════════════════════════════════════════════════════════════════════
def calcular_promedio_fase(scores: Dict[str, int]) -> Optional[float]:
    if not scores: return None
    valores = [v for v in scores.values() if v is not None and v > 0]
    if not valores: return None
    return round(sum(valores) / len(valores), 2)

def calcular_integrado() -> Optional[Dict[str, Any]]:
    promedios = {}
    pesos_total = 0
    suma_ponderada = 0
    for fase_key, fase_info in FASES_EVALUACION.items():
        scores = st.session_state.meritorio_scores.get(fase_key, {})
        promedio = calcular_promedio_fase(scores)
        promedios[fase_key] = promedio
        if promedio is not None:
            suma_ponderada += promedio * fase_info["peso"]
            pesos_total += fase_info["peso"]
    if pesos_total == 0: return None
    promedio_final = round(suma_ponderada / pesos_total, 2)

    if promedio_final >= 3.60:
        dictamen = "Maestro Meritorio"; desc = "Candidato a reconocimiento institucional"; color = "#10B981"
    elif promedio_final >= 3.00:
        dictamen = "Maestro Destacado"; desc = "Desempeño sobresaliente"; color = "#3B82F6"
    elif promedio_final >= 2.50:
        dictamen = "En Proceso de Desarrollo"; desc = "Con plan de mejora específico"; color = "#F59E0B"
    else:
        dictamen = "Requiere Plan de Mejora"; desc = "Con acompañamiento de coordinación"; color = "#EF4444"
    return {"promedios": promedios, "promedio_final": promedio_final, "dictamen": dictamen, "descripcion": desc, "color": color}

def progreso_fase(fase_key: str) -> float:
    fase = FASES_EVALUACION[fase_key]
    scores = st.session_state.meritorio_scores.get(fase_key, {})
    total = len(fase["criterios"])
    completados = sum(1 for v in scores.values() if v is not None and v > 0)
    return (completados / total * 100) if total > 0 else 0

# ═══════════════════════════════════════════════════════════════════════════
# PROMPTS IA Y GENERADORES WORD
# ═══════════════════════════════════════════════════════════════════════════
def prompt_analisis_integrado(docente, integrado, comentarios):
    promedios_txt = "\n".join([f"- {FASES_EVALUACION[k]['nombre']}: {v if v else 'Sin datos'} (peso {FASES_EVALUACION[k]['peso']*100:.0f}%)" for k, v in integrado["promedios"].items()])
    comentarios_txt = "\n".join([f"- {FASES_EVALUACION[k]['nombre']}: {v if v else 'Sin comentarios'}" for k, v in comentarios.items()])
    return f"""Actúa como Especialista en Desarrollo Docente y Evaluación del Desempeño del MINERD.
Genera un ANÁLISIS CUALITATIVO INTEGRADO del docente evaluado.
DOCENTE: {docente} | PROMEDIO: {integrado['promedio_final']}/4.00 | DICTAMEN: {integrado['dictamen']}
RESULTADOS POR FASE:\n{promedios_txt}\nCOMENTARIOS:\n{comentarios_txt}
Devuelve ÚNICAMENTE JSON válido:
{{
 "ANALISIS_GENERAL": "...",
 "FORTALEZAS": ["..."],
 "AREAS_MEJORA": ["..."],
 "RECOMENDACIONES": ["..."],
 "DICTAMEN_JUSTIFICADO": "..."
}}"""

def corregir_instrumento_ia(texto_escaneado, fase_key):
    fase = FASES_EVALUACION[fase_key]
    criterios_str = json.dumps(fase["criterios"], ensure_ascii=False)
    prompt = f"""Actúa como Coordinador Evaluador del MINERD.
Aquí tienes los CRITERIOS DE LA FASE ({fase['nombre']}):
{criterios_str}
TEXTO EXTRAÍDO DEL INSTRUMENTO ESCANEADO:
{texto_escaneado}
TAREA:
1. Interpreta las puntuaciones (escala 1 al 4) asignadas a cada ID de criterio.
2. Extrae observaciones o comentarios generales.
Devuelve ÚNICAMENTE JSON válido:
{{
 "SCORES": {{"AE1": 4, "AE2": 3}},
 "COMENTARIOS": "..."
}}"""
    datos, _ = ia.solicitar_json(prompt, max_tokens=2000, temperature=0.1, modulo="corrector_meritorio")
    return ia.decodificar_marcadores(datos)

def build_instrumento_docx(fase_key: str) -> BytesIO:
    fase = FASES_EVALUACION[fase_key]
    criterios = fase["criterios"]
    docente = st.session_state.meritorio_docente or "_________________________"
    periodo = st.session_state.meritorio_periodo or "_________________________"

    doc = Document()
    doc.styles["Normal"].font.name = "Calibri"; doc.styles["Normal"].font.size = Pt(10)
    for section in doc.sections: section.left_margin = section.right_margin = Inches(0.6); section.top_margin = section.bottom_margin = Inches(0.5)

    p_minerd = doc.add_paragraph(); p_minerd.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p_minerd.add_run("MINISTERIO DE EDUCACIÓN DE LA REPÚBLICA DOMINICANA\n").bold = True
    p_centro = doc.add_paragraph(); p_centro.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r_c = p_centro.add_run("Politécnico Salesiano Arquides Calderón\n"); r_c.bold = True; r_c.font.size = Pt(13); r_c.font.color.rgb = RGBColor(0x92, 0x40, 0x0E)
    r_t = p_centro.add_run("SELECCIÓN MAESTRO MERITORIO ETP\n"); r_t.bold = True; r_t.font.size = Pt(15); r_t.font.color.rgb = RGBColor(0x92, 0x40, 0x0E)
    r_f = p_centro.add_run(f"{fase['icono']} {fase['nombre'].upper()}"); r_f.bold = True; r_f.font.size = Pt(13)

    tabla_datos = doc.add_table(rows=2, cols=4); tabla_datos.style = "Table Grid"
    row1 = tabla_datos.rows[0].cells
    set_cell_text(row1[0], "Docente", bold=True, color="FEF3C7"); set_cell_text(row1[1], docente)
    set_cell_text(row1[2], "Período", bold=True, color="FEF3C7"); set_cell_text(row1[3], periodo)
    row2 = tabla_datos.rows[1].cells
    set_cell_text(row2[0], "Evaluador", bold=True, color="FEF3C7"); set_cell_text(row2[1], fase["evaluador"])
    set_cell_text(row2[2], "Peso", bold=True, color="FEF3C7"); set_cell_text(row2[3], f"{fase['peso']*100:.0f}%")
    doc.add_paragraph()

    p_instr = doc.add_paragraph(); p_instr.add_run("INSTRUCCIONES").bold = True
    doc.add_paragraph(fase["instrucciones"])

    p_escala = doc.add_paragraph(); p_escala.add_run("ESCALA DE VALORACIÓN (1 al 4)").bold = True
    tabla_escala = doc.add_table(rows=5, cols=3); tabla_escala.style = "Table Grid"
    hdr_esc = tabla_escala.rows[0].cells
    set_cell_text(hdr_esc[0], "Valor", bold=True, center=True, color="FEF3C7")
    set_cell_text(hdr_esc[1], "Nivel", bold=True, center=True, color="FEF3C7")
    set_cell_text(hdr_esc[2], "Descriptor", bold=True, center=True, color="FEF3C7")
    for i, (valor, (nivel, descriptor)) in enumerate(ESCALA_VALORACION.items(), 1):
        row = tabla_escala.rows[i].cells
        set_cell_text(row[0], str(valor), center=True, bold=True)
        set_cell_text(row[1], nivel, bold=True)
        set_cell_text(row[2], descriptor)
    doc.add_paragraph()

    tabla_crit = doc.add_table(rows=1, cols=7); tabla_crit.style = "Table Grid"
    hdr = tabla_crit.rows[0].cells
    set_cell_text(hdr[0], "No.", bold=True, center=True, color="FEF3C7")
    set_cell_text(hdr[1], "Dimensión", bold=True, color="FEF3C7")
    set_cell_text(hdr[2], "Criterio de Evaluación", bold=True, color="FEF3C7")
    for val, col_idx in [("4", 3), ("3", 4), ("2", 5), ("1", 6)]: set_cell_text(hdr[col_idx], val, bold=True, center=True)

    for idx, crit in enumerate(criterios, 1):
        row = tabla_crit.add_row().cells
        set_cell_text(row[0], str(idx), center=True)
        set_cell_text(row[1], crit.get("dimension", ""))
        set_cell_text(row[2], crit["criterio"])
        for col_idx in [3, 4, 5, 6]: set_cell_text(row[col_idx], "☐", center=True)

    doc.add_paragraph(); p_com = doc.add_paragraph(); p_com.add_run("COMENTARIOS Y OBSERVACIONES").bold = True
    doc.add_paragraph("_" * 70); doc.add_paragraph("_" * 70); doc.add_paragraph()

    tabla_firmas = doc.add_table(rows=2, cols=3)
    tabla_firmas.cell(0, 0).text = "_________________________"; tabla_firmas.cell(0, 1).text = "_________________________"; tabla_firmas.cell(0, 2).text = "_________________________"
    tabla_firmas.cell(1, 0).text = f"{fase['evaluador']}"; tabla_firmas.cell(1, 1).text = "Docente Evaluado"; tabla_firmas.cell(1, 2).text = "Coordinación ETP"
    for row in tabla_firmas.rows:
        for cell in row.cells: cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER

    add_page_number(doc)
    buffer = BytesIO(); doc.save(buffer); buffer.seek(0)
    return buffer

def build_informe_integrado_docx(integrado: Dict, analisis_ia: Optional[Dict]) -> BytesIO:
    docente = st.session_state.meritorio_docente or "_________________________"
    periodo = st.session_state.meritorio_periodo or "_________________________"
    doc = Document()
    doc.styles["Normal"].font.name = "Calibri"; doc.styles["Normal"].font.size = Pt(11)
    for section in doc.sections: section.left_margin = section.right_margin = Inches(0.7)

    p_minerd = doc.add_paragraph(); p_minerd.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p_minerd.add_run("MINISTERIO DE EDUCACIÓN DE LA REPÚBLICA DOMINICANA\n").bold = True
    p_centro = doc.add_paragraph(); p_centro.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r_c = p_centro.add_run("Politécnico Salesiano Arquides Calderón\n"); r_c.bold = True; r_c.font.size = Pt(14); r_c.font.color.rgb = RGBColor(0x92, 0x40, 0x0E)
    r_t = p_centro.add_run("INFORME INTEGRADOR — MAESTRO MERITORIO ETP\n"); r_t.bold = True; r_t.font.size = Pt(18); r_t.font.color.rgb = RGBColor(0x92, 0x40, 0x0E)
    
    doc.add_paragraph(f"Docente: {docente}   |   Período: {periodo}   |   Fecha: {datetime.now().strftime('%d/%m/%Y')}")
    
    p_dict = doc.add_paragraph(); p_dict.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p_dict.add_run(f"Puntaje Final: {integrado['promedio_final']}/4.00 · Dictamen: {integrado['dictamen']}").bold = True

    doc.add_heading("Resultados por Fase", level=2)
    tabla_fases = doc.add_table(rows=1, cols=5); tabla_fases.style = "Table Grid"
    hdr = tabla_fases.rows[0].cells
    for i, h in enumerate(["Fase", "Evaluador", "Peso", "Puntaje", "Ponderado"]): set_cell_text(hdr[i], h, bold=True, color="FEF3C7")
    for fase_key, fase_info in FASES_EVALUACION.items():
        row = tabla_fases.add_row().cells
        promedio = integrado["promedios"].get(fase_key)
        ponderado = round(promedio * fase_info["peso"], 2) if promedio else 0
        set_cell_text(row[0], f"{fase_info['icono']} {fase_info['nombre']}", bold=True)
        set_cell_text(row[1], fase_info["evaluador"])
        set_cell_text(row[2], f"{fase_info['peso']*100:.0f}%", center=True)
        set_cell_text(row[3], f"{promedio if promedio else 'N/A'}/4.00", center=True)
        set_cell_text(row[4], str(ponderado), center=True)

    if analisis_ia:
        doc.add_heading("Análisis Cualitativo (IA)", level=2)
        doc.add_paragraph(analisis_ia.get("ANALISIS_GENERAL", ""))
        doc.add_heading("Fortalezas", level=3)
        for f in analisis_ia.get("FORTALEZAS", []): doc.add_paragraph(f, style="List Bullet")
        doc.add_heading("Áreas de Mejora", level=3)
        for a in analisis_ia.get("AREAS_MEJORA", []): doc.add_paragraph(a, style="List Bullet")

    add_page_number(doc)
    buffer = BytesIO(); doc.save(buffer); buffer.seek(0)
    return buffer

def render_stepper():
    fases_keys = list(FASES_EVALUACION.keys())
    partes = ['<div class="meritorio-stepper">']
    for i, fase_key in enumerate(fases_keys):
        fase = FASES_EVALUACION[fase_key]
        progreso = progreso_fase(fase_key)
        estado = "done" if progreso >= 100 else ("active" if progreso > 0 else "inactive")
        partes.append(f'<div class="meritorio-step"><div class="meritorio-step-circle {estado}">{fase["icono"]}</div><div class="meritorio-step-label">{fase["nombre"].split()[0]}</div></div>')
        if i < len(fases_keys) - 1:
            line_class = "done" if progreso >= 100 else ""
            partes.append(f'<div class="meritorio-step-line {line_class}"></div>')
    partes.append('<div class="meritorio-step-line"></div>')
    integrado = calcular_integrado()
    estado_final = "done" if integrado else "inactive"
    partes.append(f'<div class="meritorio-step"><div class="meritorio-step-circle {estado_final}">🏆</div><div class="meritorio-step-label">Resultado</div></div></div>')
    st.markdown("".join(partes), unsafe_allow_html=True)

# ═══════════════════════════════════════════════════════════════════════════
# RENDERIZADO DE FASE (COORDINADOR)
# ═══════════════════════════════════════════════════════════════════════════
def render_fase(fase_key: str):
    fase = FASES_EVALUACION[fase_key]
    criterios = fase["criterios"]
    scores = st.session_state.meritorio_scores[fase_key]
    progreso = progreso_fase(fase_key)

    st.markdown(f"""
    <div class="meritorio-fase-card" style="border-left: 5px solid {fase['color']}; margin-bottom: 20px;">
        <div style="display: flex; align-items: center; gap: 12px; margin-bottom: 12px;">
            <div style="font-size: 2.5rem;">{fase['icono']}</div>
            <div>
                <div style="font-weight: 800; font-size: 1.3rem; color: #0F172A;">{fase['nombre']}</div>
                <div style="font-size: 0.88rem; color: #64748B;">Evaluador: {fase['evaluador']} · Peso: {fase['peso']*100:.0f}%</div>
            </div>
        </div>
        <div class="meritorio-progress-bar"><div class="meritorio-progress-fill" style="width: {progreso}%; background: {fase['color']};"></div></div>
        <div style="font-size: 0.78rem; color: #64748B; margin-top: 4px;">{progreso:.0f}% completado</div>
    </div>
    """, unsafe_allow_html=True)

    tab_fisico, tab_digital, tab_escaner, tab_manual = st.tabs(["🖨️ Físico (Word)", "🔗 Digital (QR y Link)", "📸 Escáner IA", "✍️ Manual"])

    with tab_fisico:
        st.markdown('<div class="meritorio-section-title">📄 Generar Instrumento Físico</div>', unsafe_allow_html=True)
        if not st.session_state.meritorio_docente:
            st.warning("⚠️ Primero selecciona un docente en la pestaña '🎓 Docente'.")
        else:
            buf = build_instrumento_docx(fase_key)
            nombre_doc = f"Instrumento_{fase_key}_{st.session_state.meritorio_docente.replace(' ', '_')}.docx"
            st.download_button(
                label=f"🖨️ Descargar Instrumento {fase['nombre']} (.docx)", 
                data=buf,
                file_name=nombre_doc,
                mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                use_container_width=True, 
                key=f"dl_docx_{fase_key}"
            )

    with tab_digital:
        st.markdown('<div class="meritorio-section-title">🔗 Evaluación 100% en Línea</div>', unsafe_allow_html=True)
        st.info("Genera un enlace interactivo y un código QR para que los evaluadores completen la rúbrica desde cualquier celular o PC.")
        
        if not st.session_state.meritorio_docente:
            st.warning("⚠️ Primero selecciona un docente en la pestaña '🎓 Docente'.")
        else:
            if st.button("🔗 Generar Enlace y QR", type="primary", use_container_width=True, key=f"gen_lnk_{fase_key}"):
                codigo = "EV-" + "".join(random.choices(string.ascii_uppercase + string.digits, k=6))
                conn = _conn()
                conn.execute("INSERT INTO mm_links (codigo, docente, periodo, fase, fecha) VALUES (?,?,?,?,?)",
                             (codigo, st.session_state.meritorio_docente, st.session_state.meritorio_periodo, fase_key, datetime.now().strftime("%Y-%m-%d")))
                conn.commit()
                conn.close()
                st.session_state[f"link_activo_{fase_key}"] = codigo
                st.success("✅ Enlace generado correctamente.")

            if st.session_state.get(f"link_activo_{fase_key}"):
                codigo = st.session_state[f"link_activo_{fase_key}"]
                link = construir_link_eval(codigo)
                
                c_qr, c_txt = st.columns([1, 2])
                with c_qr: st.image(qr_url(link), width=180)
                with c_txt:
                    st.markdown(f"**Código de acceso:** `{codigo}`")
                    st.markdown(f'<div class="pd-linkbox">{link}</div>', unsafe_allow_html=True)
                    st.caption("Comparte este QR/Link en WhatsApp, Classroom o proyéctalo en la pantalla.")
                
                conn = _conn()
                cur = conn.cursor()
                cur.execute("SELECT COUNT(*) FROM mm_respuestas WHERE codigo=?", (codigo,))
                count = cur.fetchone()[0]
                conn.close()
                
                st.markdown("---")
                st.markdown(f"#### 📥 Respuestas Recibidas: **{count}**")
                
                if count > 0:
                    if st.button("🔄 Importar y Promediar Respuestas", type="primary", use_container_width=True, key=f"sync_{fase_key}"):
                        avg_scores, com_unidos = promediar_respuestas(codigo)
                        st.session_state.meritorio_scores[fase_key] = avg_scores
                        st.session_state.meritorio_comentarios[fase_key] = com_unidos
                        guardar_consolidado(st.session_state.meritorio_docente, st.session_state.meritorio_periodo, fase_key, avg_scores, com_unidos)
                        st.toast("✅ Puntuaciones importadas y guardadas permanentemente", icon="📥")
                        st.rerun()

    with tab_escaner:
        st.markdown('<div class="meritorio-section-title">🤖 Corrector IA de Instrumentos Escaneados</div>', unsafe_allow_html=True)
        st.info("Sube el PDF del instrumento aplicado y llenado a mano. La IA interpretará las notas (1 al 4) y registrará los puntajes automáticamente.")
        
        archivo_escaneado = st.file_uploader(f"Sube {fase['nombre']} escaneado (PDF)", type=["pdf"], key=f"up_{fase_key}")
        if st.button(f"🤖 Procesar Escaneo de {fase['nombre']}", type="primary", use_container_width=True, disabled=not archivo_escaneado):
            if not st.session_state.get("api_key_global", ""): st.error("🔒 Configura tu API Key.")
            else:
                with st.spinner("Leyendo instrumento y extrayendo puntuaciones con IA..."):
                    try:
                        texto_pdf = extraer_texto_pdf(archivo_escaneado)
                        res_ia = corregir_instrumento_ia(texto_pdf, fase_key)
                        
                        scores_ia = res_ia.get("SCORES", {})
                        coms_ia = res_ia.get("COMENTARIOS", "")
                        
                        for cid, val in scores_ia.items():
                            st.session_state.meritorio_scores[fase_key][cid] = int(val)
                        if coms_ia: st.session_state.meritorio_comentarios[fase_key] = coms_ia
                        
                        guardar_consolidado(st.session_state.meritorio_docente, st.session_state.meritorio_periodo, fase_key, st.session_state.meritorio_scores[fase_key], st.session_state.meritorio_comentarios[fase_key])
                        st.toast(f"✅ Puntuaciones de {fase['nombre']} asignadas y guardadas por IA", icon="🎯")
                        st.rerun()
                    except Exception as e: ia.render_error_ia(e)

    with tab_manual:
        st.markdown('<div class="meritorio-section-title">✍️ Registro Manual de Puntuaciones</div>', unsafe_allow_html=True)
        with st.form(f"form_scores_{fase_key}"):
            for crit in criterios:
                col_c1, col_c2 = st.columns([3, 1])
                with col_c1: st.markdown(f"**{crit['id']}** ({crit.get('dimension','')}) — {crit['criterio']}")
                with col_c2:
                    val_act = scores.get(crit["id"], 0)
                    scores[crit["id"]] = st.selectbox("Punt", [0, 4, 3, 2, 1], index=[0, 4, 3, 2, 1].index(val_act) if val_act in [0, 4, 3, 2, 1] else 0, format_func=lambda x: "—" if x==0 else f"{x} - {ESCALA_VALORACION[x][0]}", key=f"s_{fase_key}_{crit['id']}")

            st.session_state.meritorio_comentarios[fase_key] = st.text_area("Comentarios:", value=st.session_state.meritorio_comentarios.get(fase_key, ""), height=80, key=f"c_{fase_key}")
            if st.form_submit_button("💾 Guardar Puntuaciones", type="primary", use_container_width=True):
                guardar_consolidado(st.session_state.meritorio_docente, st.session_state.meritorio_periodo, fase_key, scores, st.session_state.meritorio_comentarios[fase_key])
                st.toast(f"✅ Puntuaciones guardadas permanentemente.", icon=fase['icono'])
                st.rerun()

    promedio = calcular_promedio_fase(scores)
    if promedio is not None:
        st.markdown("---")
        c_p1, c_p2, c_p3 = st.columns(3)
        c_p1.metric("Promedio de la Fase", f"{promedio}/4.00")
        c_p2.metric("Peso", f"{fase['peso']*100:.0f}%")
        c_p3.metric("Ponderado", f"{round(promedio * fase['peso'], 2)}")

# ═══════════════════════════════════════════════════════════════════════════
# HERO E INTERFAZ PRINCIPAL
# ═══════════════════════════════════════════════════════════════════════════
ia.panel_sidebar_ia("Selección Maestro Meritorio")

st.markdown("""
<div class="meritorio-hero">
    <div class="meritorio-hero-title">🏆 Selección Maestro Meritorio ETP (Fígital Nativo)</div>
    <div class="meritorio-hero-sub">Evaluación en 4 fases · Generación de Enlaces y QRs locales · Guardado en BD · Corrector IA de escaneos</div>
    <div>
        <span class="meritorio-hero-badge">🪞 Autoevaluación</span>
        <span class="meritorio-hero-badge">🎓 Estudiantes</span>
        <span class="meritorio-hero-badge">📋 Coordinación</span>
        <span class="meritorio-hero-badge">🏛️ Dirección</span>
    </div>
</div>
""", unsafe_allow_html=True)

render_stepper()

tab_config, tab_f1, tab_f2, tab_f3, tab_f4, tab_integrado = st.tabs([
    "🎓 Docente",
    f"{FASES_EVALUACION['autoevaluacion']['icono']} Autoevaluación",
    f"{FASES_EVALUACION['estudiantes']['icono']} Estudiantes",
    f"{FASES_EVALUACION['coordinacion']['icono']} Coordinación",
    f"{FASES_EVALUACION['direccion']['icono']} Dirección",
    "🏆 Resultado Integrado",
])

with tab_config:
    st.markdown('<div class="meritorio-section-title">🎓 Configuración de la Evaluación</div>', unsafe_allow_html=True)
    docentes_bd = obtener_docentes_bd()
    with st.form("form_config_meritorio"):
        col_cfg1, col_cfg2 = st.columns(2)
        with col_cfg1:
            if docentes_bd:
                docente_sel = st.selectbox("👤 Seleccionar Docente", ["— Seleccionar del directorio —"] + docentes_bd + ["✍️ Escribir manualmente"])
                if docente_sel == "✍️ Escribir manualmente": docente_manual = st.text_input("Nombre completo del docente")
                elif docente_sel != "— Seleccionar del directorio —": docente_manual = docente_sel
                else: docente_manual = ""
            else:
                docente_manual = st.text_input("Nombre completo del docente", placeholder="Escribe el nombre...")
        with col_cfg2:
            periodos_disponibles = ["2025-2026 - Primer Trimestre", "2025-2026 - Segundo Trimestre", "2025-2026 - Tercer Trimestre", "2026-2027 - Primer Trimestre", "2026-2027 - Segundo Trimestre", "2026-2027 - Tercer Trimestre", "Evaluación Anual Completa"]
            periodo_sel = st.selectbox("📅 Período de Evaluación", periodos_disponibles)
        if st.form_submit_button("💾 Configurar Evaluación", type="primary", use_container_width=True):
            if not docente_manual.strip(): st.warning("⚠️ Debes seleccionar o escribir el nombre del docente.")
            else:
                st.session_state.meritorio_docente = docente_manual.strip()
                st.session_state.meritorio_periodo = periodo_sel
                cargar_consolidado(st.session_state.meritorio_docente, st.session_state.meritorio_periodo)
                st.toast(f"✅ Configurado para {docente_manual.strip()}. Datos cargados.", icon="🎓")
                st.rerun()

    if st.session_state.meritorio_docente:
        st.success(f"✅ Evaluando actualmente a: **{st.session_state.meritorio_docente}** | Período: {st.session_state.meritorio_periodo}")

with tab_f1: render_fase("autoevaluacion")
with tab_f2: render_fase("estudiantes")
with tab_f3: render_fase("coordinacion")
with tab_f4: render_fase("direccion")

with tab_integrado:
    st.markdown('<div class="meritorio-section-title">🏆 Resultado Integrado Final</div>', unsafe_allow_html=True)
    if not st.session_state.meritorio_docente: 
        st.warning("⚠️ Configura el docente primero.")
    else:
        # Aseguramos cargar SIEMPRE los datos más frescos al entrar a esta pestaña
        cargar_consolidado(st.session_state.meritorio_docente, st.session_state.meritorio_periodo)
        
        integrado = calcular_integrado()
        if not integrado: 
            st.info("📭 Completa y guarda al menos una fase de evaluación para ver el resultado.")
        else:
            c_score, c_info = st.columns([1, 2])
            with c_score:
                st.markdown(f"""
                <div style="background: #1E293B; border-radius: 20px; padding: 2rem; text-align: center; color: #fff;">
                    <div style="font-size: 4rem; font-weight: 900; color: {integrado['color']};">{integrado['promedio_final']}</div>
                    <div style="font-size: 0.9rem; opacity: 0.7;">Puntuación Final / 4.00</div>
                    <div style="margin-top: 12px; font-size: 1.2rem; font-weight: 700; color: {integrado['color']};">{integrado['dictamen']}</div>
                </div>
                """, unsafe_allow_html=True)
            with c_info:
                st.markdown(f"**Docente:** {st.session_state.meritorio_docente}")
                st.markdown(f"**Período:** {st.session_state.meritorio_periodo}")
                st.markdown(f"**Descripción:** {integrado['descripcion']}")
                st.markdown("---")
                for fase_key, fase_info in FASES_EVALUACION.items():
                    p = integrado["promedios"].get(fase_key)
                    st.markdown(f"{fase_info['icono']} **{fase_info['nombre']}** ({fase_info['peso']*100:.0f}%): {p if p else 'N/A'}/4.00")

            st.markdown("---")
            if st.button("🧠 Generar Análisis Cualitativo con IA", type="primary", use_container_width=True):
                with st.spinner("Analizando desempeño..."):
                    try:
                        prompt = prompt_analisis_integrado(st.session_state.meritorio_docente, integrado, st.session_state.meritorio_comentarios)
                        res_ia, _ = ia.solicitar_json(prompt, max_tokens=8192, temperature=0.3, modulo="maestro_meritorio")
                        st.session_state.meritorio_integrado = {"analisis_ia": res_ia, "integrado": integrado}
                        st.toast("✅ Análisis generado", icon="🤖")
                    except Exception as e: 
                        ia.render_error_ia(e)

            if st.session_state.meritorio_integrado:
                ai = st.session_state.meritorio_integrado.get("analisis_ia", {})
                st.info(ai.get("ANALISIS_GENERAL", ""))
                with st.expander("💪 Fortalezas y Áreas de Mejora", expanded=True):
                    for f in ai.get("FORTALEZAS", []): st.write(f"• {f}")
                    for a in ai.get("AREAS_MEJORA", []): st.write(f"• {a}")

            st.markdown("---")
            # El botón de descarga ahora está visible directamente, lo que garantiza que nunca falle la descarga.
            ai_dat = st.session_state.meritorio_integrado.get("analisis_ia") if st.session_state.meritorio_integrado else None
            buf_inf = build_informe_integrado_docx(integrado, ai_dat)
            
            st.download_button(
                label="📄 Descargar Informe Integrado (.docx)", 
                data=buf_inf,
                file_name=f"Informe_Meritorio_{st.session_state.meritorio_docente.replace(' ', '_')}.docx",
                mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document", 
                type="primary", 
                use_container_width=True
            )