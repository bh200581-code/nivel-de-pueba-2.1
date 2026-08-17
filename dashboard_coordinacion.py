"""
Dashboard de Coordinación ETP — Sala de Situación Profesional
Panel gerencial de alto nivel para coordinación académico-formativa.
NUEVO: Integración con prueba_diagnostica.py para monitorear resultados y promedios.
"""
from __future__ import annotations

import html
import logging
import sqlite3
from contextlib import closing
from dataclasses import dataclass
from datetime import date, datetime, timedelta
from io import BytesIO
from pathlib import Path

import pandas as pd
import streamlit as st
from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Inches, Pt, RGBColor

# ---------------------------------------------------------------------------
# Configuración básica
# ---------------------------------------------------------------------------
logging.basicConfig(
    level=logging.INFO,
    format="%(asctime)s | %(levelname)s | %(name)s | %(message)s",
)
logger = logging.getLogger("dashboard_coordinacion")

try:
    st.set_page_config(
        page_title="Sala de Situación ETP",
        page_icon="📊",
        layout="wide",
        initial_sidebar_state="expanded",
    )
except Exception:
    pass

# ---------------------------------------------------------------------------
# Constantes de dominio y configuración
# ---------------------------------------------------------------------------
APP_TITLE = "Sala de Situación ETP"
INSTITUCION = "Politécnico Arquides Calderón"
COORDINADOR_DEFAULT = "Ing. Bernardo Hernández"
INCIDENCIAS_TABLE = "incidencias"
ACOMPANAMIENTOS_TABLE = "evidencias_acompanamiento"
CRONOGRAMA_TABLE = "cronograma"
RESPUESTAS_DIAG_TABLE = "respuestas_diagnosticas"

ALLOWED_TABLES = {
    "acompanamientos",
    "evidencias_acompanamiento",
    "alertas",
    "acuerdos",
    "incidencias",
    "cronograma",
    "pruebas_diagnosticas",
    "respuestas_diagnosticas"
}
GRAVEDADES_CRITICAS = ("Grave", "Muy Grave")
CACHE_TTL_SECONDS = 30
MESES_ES = [
    "enero", "febrero", "marzo", "abril", "mayo", "junio",
    "julio", "agosto", "septiembre", "octubre", "noviembre", "diciembre",
]
DIAS_SEMANA_ES = ["Lunes", "Martes", "Miércoles", "Jueves", "Viernes", "Sábado", "Domingo"]

def resolver_ruta_bd() -> Path:
    rutas_posibles = [
        Path("gestion_etp.db"),
        Path(__file__).resolve().parent / "gestion_etp.db",
        Path.cwd() / "gestion_etp.db",
    ]
    for ruta in rutas_posibles:
        if ruta.exists():
            return ruta
    return Path("gestion_etp.db")

DB_PATH = resolver_ruta_bd()

# ---------------------------------------------------------------------------
# Estilos CSS
# ---------------------------------------------------------------------------
CSS = """
<style>
:root {
  --bg: #F8FAFC; --card: #FFFFFF; --border: #E2E8F0;
  --text-main: #0F172A; --text-soft: #475569; --text-muted: #94A3B8;
  --primary: #2563EB; --success: #10B981; --warning: #F59E0B;
  --danger: #EF4444; --info: #3B82F6;
  --shadow: 0 10px 30px rgba(15, 23, 42, 0.08);
}
html, body, [class*="css"] {
  font-family: "Segoe UI", "Inter", "Helvetica Neue", Arial, sans-serif;
  background-color: var(--bg); color: var(--text-main);
}
.dashboard-hero {
  background: linear-gradient(135deg, #0F172A 0%, #1E3A8A 55%, #2563EB 100%);
  color: white; padding: 2rem; border-radius: 20px; margin-bottom: 1.5rem;
  box-shadow: var(--shadow); display: flex; justify-content: space-between;
  gap: 1rem; flex-wrap: wrap;
}
.hero-greeting { font-size: 1rem; font-weight: 600; opacity: 0.92; margin: 0 0 0.35rem 0; }
.dashboard-title { font-size: 2.35rem; font-weight: 800; letter-spacing: -0.03em; margin: 0; line-height: 1.05; }
.dashboard-subtitle { margin-top: 0.55rem; font-size: 1rem; opacity: 0.82; }
.hero-badges { display: flex; flex-direction: column; gap: 0.5rem; align-items: flex-end; }
.badge { display: inline-block; padding: 0.45rem 0.8rem; border-radius: 999px; font-size: 0.82rem;
  font-weight: 700; border: 1px solid rgba(255,255,255,0.22); background: rgba(255,255,255,0.12); }
.badge-info { background: rgba(59,130,246,0.18); }
.badge-success { background: rgba(16,185,129,0.18); }
.kpi-card { background: var(--card); border: 1px solid var(--border); border-radius: 18px;
  padding: 1.25rem; box-shadow: var(--shadow); min-height: 148px; display: flex;
  flex-direction: column; justify-content: space-between;
  transition: transform 0.18s ease, box-shadow 0.18s ease; }
.kpi-card:hover { transform: translateY(-2px); box-shadow: 0 14px 34px rgba(15,23,42,0.12); }
.tone-info { border-top: 5px solid var(--info); }
.tone-alert { border-top: 5px solid var(--danger); }
.tone-warning { border-top: 5px solid var(--warning); }
.tone-success { border-top: 5px solid var(--success); }
.tone-primary { border-top: 5px solid var(--primary); }
.kpi-header { display: flex; justify-content: space-between; align-items: flex-start; gap: 0.75rem; }
.kpi-label { font-size: 0.84rem; font-weight: 700; color: var(--text-soft);
  text-transform: uppercase; letter-spacing: 0.045em; line-height: 1.25; }
.kpi-icon { font-size: 1.35rem; line-height: 1; }
.kpi-value { font-size: 2.45rem; font-weight: 800; color: var(--text-main); line-height: 1;
  margin: 0.7rem 0 0.35rem 0; }
.kpi-caption { color: var(--text-muted); font-size: 0.88rem; line-height: 1.3; }
.action-card { background: var(--card); border: 1px solid var(--border); border-radius: 16px;
  padding: 1rem 1.1rem; box-shadow: var(--shadow); margin-bottom: 0.65rem; }
.action-title { font-size: 1.02rem; font-weight: 700; color: var(--text-main); margin-bottom: 0.35rem; }
.action-desc { color: var(--text-soft); font-size: 0.9rem; line-height: 1.35; }
</style>
"""
st.markdown(CSS, unsafe_allow_html=True)

# ---------------------------------------------------------------------------
# Modelos
# ---------------------------------------------------------------------------
@dataclass(frozen=True)
class Kpi:
    label: str
    value: int
    caption: str
    tone: str
    icon: str

@dataclass(frozen=True)
class Accion:
    titulo: str
    descripcion: str
    pagina: str
    emoji: str

# ---------------------------------------------------------------------------
# Utilidades
# ---------------------------------------------------------------------------
def obtener_saludo(hora: int) -> str:
    if hora < 12: return "Buenos días"
    if hora < 19: return "Buenas tardes"
    return "Buenas noches"

def formatear_fecha_larga(fecha: date) -> str:
    return f"{fecha.day} de {MESES_ES[fecha.month - 1]} de {fecha.year}"

def formatear_gravedad(valor: object) -> str:
    if not isinstance(valor, str): return ""
    mapa = {"Leve": "🟢 Leve", "Moderada": "🟡 Moderada", "Grave": "🟠 Grave", "Muy Grave": "🔴 Muy Grave"}
    return mapa.get(valor.strip(), valor.strip())

def navegar_a_pagina(pagina: str) -> None:
    try:
        st.switch_page(pagina)
    except Exception:
        st.error(f"No fue posible abrir la página `{pagina}`.")

# ---------------------------------------------------------------------------
# Acceso a datos
# ---------------------------------------------------------------------------
def tabla_existe(conn: sqlite3.Connection, tabla: str) -> bool:
    with closing(conn.cursor()) as cursor:
        cursor.execute("SELECT name FROM sqlite_master WHERE type = 'table' AND name = ?", (tabla,))
        return cursor.fetchone() is not None

def obtener_columnas(conn: sqlite3.Connection, tabla: str) -> list[str]:
    with closing(conn.cursor()) as cursor:
        cursor.execute(f"PRAGMA table_info({tabla})")
        return [row[1] for row in cursor.fetchall()]

def obtener_conteo(tabla: str, where: str = "", params: tuple = ()) -> int:
    if not DB_PATH.exists(): return 0
    try:
        with closing(sqlite3.connect(DB_PATH)) as conn:
            if not tabla_existe(conn, tabla): return 0
            sql = f"SELECT COUNT(*) FROM {tabla}"
            if where: sql += f" {where.strip()}"
            with closing(conn.cursor()) as cursor:
                cursor.execute(sql, params)
                row = cursor.fetchone()
                return int(row[0]) if row else 0
    except sqlite3.Error:
        return 0

@st.cache_data(ttl=CACHE_TTL_SECONDS, show_spinner=False)
def cargar_incidencias() -> pd.DataFrame:
    expected = ["id", "fecha", "estudiante", "grado", "gravedad", "falta"]
    if not DB_PATH.exists(): return pd.DataFrame(columns=expected)
    try:
        with closing(sqlite3.connect(DB_PATH)) as conn:
            if not tabla_existe(conn, INCIDENCIAS_TABLE): return pd.DataFrame(columns=expected)
            cols = obtener_columnas(conn, INCIDENCIAS_TABLE)
            if not cols: return pd.DataFrame(columns=expected)
            sel = [c for c in expected if c in cols] or ["*"]
            ord_c = "id" if "id" in cols else "rowid"
            df = pd.read_sql_query(f"SELECT {', '.join(sel)} FROM {INCIDENCIAS_TABLE} ORDER BY {ord_c} DESC", conn)
    except sqlite3.Error:
        return pd.DataFrame(columns=expected)
    df = df.reindex(columns=expected, fill_value="")
    df["fecha"] = pd.to_datetime(df["fecha"], errors="coerce")
    return df

@st.cache_data(ttl=CACHE_TTL_SECONDS, show_spinner=False)
def cargar_acompanamientos() -> pd.DataFrame:
    expected = ["id", "fecha", "docente", "modulo", "seccion", "tipo", "area", "puntuacion", "nivel"]
    if not DB_PATH.exists(): return pd.DataFrame(columns=expected)
    try:
        with closing(sqlite3.connect(DB_PATH)) as conn:
            if not tabla_existe(conn, ACOMPANAMIENTOS_TABLE): return pd.DataFrame(columns=expected)
            cols = obtener_columnas(conn, ACOMPANAMIENTOS_TABLE)
            sel = [c for c in expected if c in cols] or ["*"]
            df = pd.read_sql_query(f"SELECT {', '.join(sel)} FROM {ACOMPANAMIENTOS_TABLE} ORDER BY id DESC", conn)
    except sqlite3.Error:
        return pd.DataFrame(columns=expected)
    df = df.reindex(columns=expected, fill_value="")
    df["fecha"] = pd.to_datetime(df["fecha"], errors="coerce")
    df["puntuacion"] = pd.to_numeric(df["puntuacion"], errors="coerce").fillna(0)
    return df

@st.cache_data(ttl=CACHE_TTL_SECONDS, show_spinner=False)
def cargar_cronograma() -> pd.DataFrame:
    expected = ["id", "centro", "coordinador", "anio", "trimestre", "docente", "modulo", "seccion", "hora", "dia", "modalidad", "observaciones"]
    if not DB_PATH.exists(): return pd.DataFrame(columns=expected)
    try:
        with closing(sqlite3.connect(DB_PATH)) as conn:
            if not tabla_existe(conn, CRONOGRAMA_TABLE): return pd.DataFrame(columns=expected)
            cols = obtener_columnas(conn, CRONOGRAMA_TABLE)
            sel = [c for c in expected if c in cols] or ["*"]
            df = pd.read_sql_query(f"SELECT {', '.join(sel)} FROM {CRONOGRAMA_TABLE}", conn)
            return df
    except sqlite3.Error:
        return pd.DataFrame(columns=expected)

@st.cache_data(ttl=CACHE_TTL_SECONDS, show_spinner=False)
def cargar_respuestas_diagnosticas() -> pd.DataFrame:
    """Extrae las calificaciones de las pruebas diagnósticas cruzadas con su información base."""
    expected = ["id", "prueba_codigo", "estudiante", "puntaje", "nivel", "fecha", "titulo", "modulo", "docente"]
    if not DB_PATH.exists(): return pd.DataFrame(columns=expected)
    try:
        with closing(sqlite3.connect(DB_PATH)) as conn:
            if not tabla_existe(conn, RESPUESTAS_DIAG_TABLE): return pd.DataFrame(columns=expected)
            
            # Unimos las respuestas con la tabla de pruebas para tener el contexto (Módulo, Docente, etc.)
            query = f"""
            SELECT r.id, r.prueba_codigo, r.estudiante, r.puntaje, r.nivel, r.fecha,
                   p.titulo, p.modulo, p.docente
            FROM {RESPUESTAS_DIAG_TABLE} r
            LEFT JOIN pruebas_diagnosticas p ON r.prueba_codigo = p.codigo
            ORDER BY r.id DESC
            """
            df = pd.read_sql_query(query, conn)
            return df
    except sqlite3.Error:
        return pd.DataFrame(columns=expected)

@st.cache_data(ttl=60, show_spinner=False)
def listar_tablas_existentes() -> set[str]:
    if not DB_PATH.exists(): return set()
    try:
        with closing(sqlite3.connect(DB_PATH)) as conn:
            with closing(conn.cursor()) as cursor:
                cursor.execute("SELECT name FROM sqlite_master WHERE type = 'table' AND name NOT LIKE 'sqlite_%'")
                return {row[0] for row in cursor.fetchall()}
    except sqlite3.Error:
        return set()

# ---------------------------------------------------------------------------
# Lógica de negocio / KPIs
# ---------------------------------------------------------------------------
def calcular_kpis(df_incidencias: pd.DataFrame, df_acompanamientos: pd.DataFrame, 
                  df_cronograma: pd.DataFrame, df_respuestas: pd.DataFrame, coordinador: str) -> list[Kpi]:
    hoy = date.today()
    ayer = hoy - timedelta(days=1)
    
    total_acompanamientos = len(df_acompanamientos)
    alertas_rojas = obtener_conteo("alertas", "WHERE estado IN (?, ?)", ("Crítico", "Roja"))
    acuerdos_recientes = obtener_conteo("acuerdos", "WHERE fecha >= ?", (ayer.isoformat(),))
    
    visitas_programadas = 0
    if not df_cronograma.empty:
        df_mi_crono = df_cronograma[df_cronograma["coordinador"] == coordinador]
        visitas_programadas = len(df_mi_crono)

    pruebas_evaluadas = len(df_respuestas) if not df_respuestas.empty else 0

    incidencias_criticas_mes = 0
    if not df_incidencias.empty:
        df_criticas = df_incidencias[df_incidencias["gravedad"].isin(GRAVEDADES_CRITICAS)]
        if not df_criticas.empty:
            fechas_validas = df_criticas["fecha"].notna()
            if fechas_validas.any():
                mask_mes = (fechas_validas & (df_criticas["fecha"].dt.year == hoy.year) & (df_criticas["fecha"].dt.month == hoy.month))
                incidencias_criticas_mes = int(mask_mes.sum())
            else:
                incidencias_criticas_mes = int(len(df_criticas))
                
    return [
        Kpi("Prog. Cronograma", visitas_programadas, "Visitas planificadas", "primary", "🗓️"),
        Kpi("Acompañamientos", total_acompanamientos, "Fichas evaluadas", "info", "📋"),
        Kpi("Pruebas Diag.", pruebas_evaluadas, "Estudiantes evaluados", "success", "🩺"),
        Kpi("Alertas rojas", alertas_rojas, "Riesgo priorizado", "alert", "🚨"),
        Kpi("Acuerdos", acuerdos_recientes, "En 48 hrs", "success", "🤝"),
        Kpi("Incidencias", incidencias_criticas_mes, "Críticas del mes", "warning", "⚠️"),
    ]

# ---------------------------------------------------------------------------
# Componentes de UI
# ---------------------------------------------------------------------------
def render_encabezado() -> None:
    ahora = datetime.now()
    saludo = obtener_saludo(ahora.hour)
    fecha_larga = formatear_fecha_larga(ahora.date())
    coordinador = st.session_state.get("coordinador_nombre", COORDINADOR_DEFAULT)
    st.markdown(
        f"""
        <section class="dashboard-hero" role="banner">
            <div>
                <p class="hero-greeting">{saludo}, {html.escape(coordinador)} 👋</p>
                <h1 class="dashboard-title">{html.escape(APP_TITLE)}</h1>
                <p class="dashboard-subtitle">{html.escape(INSTITUCION)} · {fecha_larga}</p>
            </div>
            <div class="hero-badges">
                <span class="badge badge-info">Panel gerencial</span>
                <span class="badge badge-success">Modo coordinación</span>
            </div>
        </section>
        """, unsafe_allow_html=True
    )

def render_barra_herramientas() -> None:
    col_btn, col_state = st.columns([1, 3])
    with col_btn:
        if st.button("🔄 Actualizar datos", use_container_width=True, type="primary"):
            st.cache_data.clear()
            try: st.toast("Datos actualizados correctamente", icon="✅")
            except AttributeError: st.success("Datos actualizados correctamente")
            st.rerun()
    with col_state:
        estado_bd = "✅ Base de datos conectada" if DB_PATH.exists() else "⚠️ Sin conexión a BD"
        hora_actual = datetime.now().strftime("%H:%M:%S")
        st.caption(f"{estado_bd} · Última actualización: {hora_actual}")

def render_alertas_cronograma(df_cronograma: pd.DataFrame, coordinador: str):
    """Muestra alerta en pantalla para las visitas de hoy según el cronograma."""
    if df_cronograma.empty: return
    
    df_mi_crono = df_cronograma[df_cronograma["coordinador"] == coordinador]
    if df_mi_crono.empty: return
    
    dia_hoy = DIAS_SEMANA_ES[date.today().weekday()]
    df_hoy = df_mi_crono[df_mi_crono["dia"] == dia_hoy]
    
    if not df_hoy.empty:
        st.warning(f"🔔 **¡Tienes {len(df_hoy)} acompañamiento(s) programado(s) en el cronograma para hoy ({dia_hoy})!**")
        for _, row in df_hoy.sort_values("hora").iterrows():
            st.markdown(f"- ⏰ **{row['hora']}** | 🧑‍🏫 **{row['docente']}** | 📚 {row['modulo']} ({row['seccion']}) | 📍 {row['modalidad']}")
        st.markdown("---")

def render_kpi_card(kpi: Kpi) -> None:
    st.markdown(
        f"""
        <article class="kpi-card tone-{kpi.tone}">
            <div class="kpi-header">
                <span class="kpi-label">{html.escape(kpi.label)}</span>
                <span class="kpi-icon">{kpi.icon}</span>
            </div>
            <div class="kpi-value">{kpi.value}</div>
            <div class="kpi-caption">{html.escape(kpi.caption)}</div>
        </article>
        """, unsafe_allow_html=True
    )

def render_tab_acompanamientos(df: pd.DataFrame) -> None:
    st.markdown("### Historial de Acompañamientos Docentes")
    if df.empty:
        st.info("No hay acompañamientos registrados en la base de datos.")
        return

    col1, col2, col3 = st.columns(3)
    with col1:
        docentes_lista = [str(d).strip() for d in df["docente"].dropna().unique().tolist() if str(d).strip()]
        docente_sel = st.selectbox("Filtrar por Docente", ["Todos"] + sorted(docentes_lista))
    with col2:
        areas_lista = [str(a).strip() for a in df["area"].dropna().unique().tolist() if str(a).strip()]
        area_sel = st.selectbox("Filtrar por Área", ["Todas"] + sorted(areas_lista))
    with col3:
        nivel_sel = st.selectbox("Filtrar por Nivel", ["Todos", "Excelente", "Bueno", "Aceptable con Mejoras", "No Aceptable"])

    df_filtrado = df.copy()
    if docente_sel != "Todos": df_filtrado = df_filtrado[df_filtrado["docente"] == docente_sel]
    if area_sel != "Todas": df_filtrado = df_filtrado[df_filtrado["area"] == area_sel]
    if nivel_sel != "Todos": df_filtrado = df_filtrado[df_filtrado["nivel"] == nivel_sel]

    m1, m2, m3 = st.columns(3)
    with m1: st.metric("Total Acompañamientos", len(df_filtrado))
    with m2: 
        promedio = df_filtrado["puntuacion"].mean() if not df_filtrado.empty else 0
        st.metric("Puntaje Promedio General", f"{promedio:.1f} / 100")
    with m3:
        excelentes = len(df_filtrado[df_filtrado["nivel"] == "Excelente"])
        st.metric("Nivel Excelente", excelentes)

    if not df_filtrado.empty:
        df_mostrar = df_filtrado[["fecha", "docente", "modulo", "area", "tipo", "puntuacion", "nivel"]].copy()
        df_mostrar["fecha"] = pd.to_datetime(df_mostrar["fecha"]).dt.strftime("%d/%m/%Y").fillna("")
        st.dataframe(df_mostrar, use_container_width=True, hide_index=True)

        st.markdown("---")
        st.markdown("#### 🏆 Ranking y Desempeño Analítico")
        
        c_chart, c_rank = st.columns([1, 1.5])
        with c_chart:
            st.markdown("**Distribución de Niveles**")
            niveles = df_filtrado["nivel"].value_counts()
            st.bar_chart(niveles)
            
            st.markdown("**Promedio por Área Educativa**")
            df_area = df_filtrado.groupby("area").agg(Acompañamientos=("id", "count"), Promedio=("puntuacion", "mean")).reset_index()
            df_area["Promedio"] = df_area["Promedio"].round(1)
            st.dataframe(df_area, use_container_width=True, hide_index=True)
            
        with c_rank:
            df_ranking = df_filtrado.groupby("docente").agg(Acompañamientos=("id", "count"), Promedio=("puntuacion", "mean")).reset_index()
            df_ranking["Promedio"] = df_ranking["Promedio"].round(1)
            df_ranking = df_ranking.sort_values("Promedio", ascending=False)
            
            st.success("🌟 Top Docentes (Mejor Promedio)")
            st.dataframe(df_ranking.head(5), use_container_width=True, hide_index=True)
            
            st.error("⚠️ Atención Prioritaria (Menor Promedio)")
            st.dataframe(df_ranking.tail(5).sort_values("Promedio", ascending=True), use_container_width=True, hide_index=True)

def render_tab_pruebas(df: pd.DataFrame) -> None:
    st.markdown("### Historial de Pruebas Diagnósticas y Evaluaciones")
    if df.empty:
        st.info("No hay respuestas a pruebas diagnósticas registradas en la base de datos.")
        return

    col1, col2 = st.columns(2)
    with col1:
        docentes_lista = [str(d).strip() for d in df["docente"].dropna().unique().tolist() if str(d).strip()]
        docente_sel = st.selectbox("Filtrar por Docente Evaluador", ["Todos"] + sorted(docentes_lista), key="filt_doc_prueba")
    with col2:
        niveles_lista = [str(n).strip() for n in df["nivel"].dropna().unique().tolist() if str(n).strip()]
        nivel_sel = st.selectbox("Filtrar por Nivel de Logro", ["Todos"] + sorted(niveles_lista), key="filt_niv_prueba")

    df_filtrado = df.copy()
    if docente_sel != "Todos": df_filtrado = df_filtrado[df_filtrado["docente"] == docente_sel]
    if nivel_sel != "Todos": df_filtrado = df_filtrado[df_filtrado["nivel"] == nivel_sel]

    m1, m2, m3 = st.columns(3)
    with m1:
        st.metric("Estudiantes Evaluados", len(df_filtrado))
    with m2: 
        promedio = df_filtrado["puntaje"].mean() if not df_filtrado.empty else 0
        st.metric("Puntaje Promedio (0-100)", f"{promedio:.1f} / 100")
    with m3:
        logrados = len(df_filtrado[df_filtrado["nivel"].isin(["Logrado", "Excelente", "Muy Bueno"])])
        st.metric("Nivel Logrado / Excelente", logrados)

    if not df_filtrado.empty:
        df_mostrar = df_filtrado[["fecha", "estudiante", "modulo", "titulo", "puntaje", "nivel"]].copy()
        st.dataframe(df_mostrar, use_container_width=True, hide_index=True)

        st.markdown("---")
        c_chart, c_mod = st.columns(2)
        with c_chart:
            st.markdown("**Distribución de Niveles de Logro**")
            st.bar_chart(df_filtrado["nivel"].value_counts())
        with c_mod:
            st.markdown("**Promedio de Calificación por Módulo / Asignatura**")
            df_mod = df_filtrado.groupby("modulo")["puntaje"].mean().reset_index()
            df_mod["puntaje"] = df_mod["puntaje"].round(1)
            st.dataframe(df_mod.rename(columns={"modulo": "Módulo", "puntaje": "Promedio"}), use_container_width=True, hide_index=True)

def render_incidencias_criticas(df: pd.DataFrame) -> None:
    if df.empty: return
    df_crit = df[df["gravedad"].isin(GRAVEDADES_CRITICAS)]
    if df_crit.empty: return
    df_crit = df_crit.sort_values("fecha", ascending=False, na_position="last").head(5)
    st.markdown("#### 🚨 Incidencias críticas recientes (acción inmediata)")
    for _, row in df_crit.iterrows():
        fecha = row["fecha"].strftime("%d/%m/%Y") if pd.notna(row["fecha"]) else "Sin fecha"
        st.markdown(f"- **{fecha}** · {row['estudiante']} · {row['grado']} · {formatear_gravedad(row['gravedad'])} — {row['falta']}")

def generar_reporte_ejecutivo(kpis: list[Kpi], df: pd.DataFrame) -> BytesIO:
    doc = Document()
    doc.styles["Normal"].font.name = "Calibri"
    doc.styles["Normal"].font.size = Pt(11)
    for section in doc.sections:
        section.left_margin = Inches(0.9); section.right_margin = Inches(0.9)

    p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run("MINISTERIO DE EDUCACIÓN DE LA REPÚBLICA DOMINICANA\n"); r.bold = True; r.font.size = Pt(12)
    p.add_run(f"{INSTITUCION}\n").bold = True
    p.add_run("REPORTE EJECUTIVO — SALA DE SITUACIÓN ETP\n").bold = True
    p.add_run(f"Generado: {datetime.now().strftime('%d/%m/%Y %H:%M')}").italic = True

    doc.add_heading("I. Indicadores Prioritarios", level=2)
    for kpi in kpis: doc.add_paragraph(f"{kpi.icon} {kpi.label}: {kpi.value} — {kpi.caption}", style="List Bullet")

    doc.add_heading("II. Incidencias Críticas Recientes", level=2)
    df_crit = df[df["gravedad"].isin(GRAVEDADES_CRITICAS)] if not df.empty else df
    if df_crit.empty: doc.add_paragraph("No hay incidencias graves / muy graves registradas.")
    else:
        tabla = doc.add_table(rows=1, cols=4); tabla.style = "Table Grid"
        hdr = tabla.rows[0].cells
        for i, t in enumerate(["Fecha", "Estudiante", "Grado", "Falta"]):
            hdr[i].text = t; hdr[i].paragraphs[0].runs[0].bold = True
        for _, row in df_crit.sort_values("fecha", ascending=False, na_position="last").head(10).iterrows():
            c = tabla.add_row().cells
            c[0].text = row["fecha"].strftime("%d/%m/%Y") if pd.notna(row["fecha"]) else "Sin fecha"
            c[1].text = str(row["estudiante"]); c[2].text = str(row["grado"]); c[3].text = str(row["falta"])

    buffer = BytesIO(); doc.save(buffer); buffer.seek(0)
    return buffer

def render_tab_incidencias(df: pd.DataFrame) -> None:
    st.markdown("### Seguimiento de incidencias")
    if df.empty:
        st.info("No hay incidencias registradas.")
        return
    st.dataframe(df, use_container_width=True, hide_index=True, height=380)

def render_tab_acciones() -> None:
    st.markdown("### Acciones rápidas")
    acciones = [
        Accion("Cronograma de Acompañamiento", "Revisa tus visitas planificadas de la semana.", "cronograma_acompanamiento.py", "🗓️"),
        Accion("Acompañamiento Docente", "Evalúa y digitaliza fichas de aula (Fígital).", "acompanamiento.py", "📋"),
        Accion("Pruebas Diagnósticas", "Genera, aplica y corrige evaluaciones fígitales.", "prueba_diagnostica.py", "🩺"),
        Accion("Emitir Alerta Temprana", "Activa seguimiento inmediato a estudiantes en riesgo.", "alerta.py", "🚨"),
        Accion("Redactar Nuevo Acuerdo", "Genera acuerdos formativos con representantes.", "acuerdos.py", "📝"),
        Accion("Registrar Incidencia", "Documenta eventos disciplinares o académicos.", "incidencias.py", "⚠️"),
    ]
    col_izq, col_der = st.columns(2)
    for idx, accion in enumerate(acciones):
        target_col = col_izq if idx % 2 == 0 else col_der
        with target_col:
            st.markdown(f'<div class="action-card"><div class="action-title">{accion.emoji} {html.escape(accion.titulo)}</div><div class="action-desc">{html.escape(accion.descripcion)}</div></div>', unsafe_allow_html=True)
            if st.button(f"Abrir: {accion.titulo}", key=f"accion_{idx}", use_container_width=True):
                navegar_a_pagina(accion.pagina)

def render_tab_sistema() -> None:
    st.markdown("### Estado del sistema")
    tablas = listar_tablas_existentes()
    info = {
        "Ruta base de datos": str(DB_PATH),
        "Archivo BD existe": "Sí" if DB_PATH.exists() else "No",
        "Tabla cronograma": "Disponible" if CRONOGRAMA_TABLE in tablas else "No disponible",
        "Tabla acompañamientos": "Disponible" if ACOMPANAMIENTOS_TABLE in tablas else "No disponible",
        "Tabla pruebas diagnósticas": "Disponible" if RESPUESTAS_DIAG_TABLE in tablas else "No disponible",
    }
    st.table(pd.DataFrame(info.items(), columns=["Parámetro", "Valor"]))

# ---------------------------------------------------------------------------
# Flujo principal
# ---------------------------------------------------------------------------
def main() -> None:
    render_encabezado()
    render_barra_herramientas()
    
    df_incidencias = cargar_incidencias()
    df_acompanamientos = cargar_acompanamientos()
    df_cronograma = cargar_cronograma()
    df_respuestas_diag = cargar_respuestas_diagnosticas()
    coordinador = st.session_state.get("coordinador_nombre", COORDINADOR_DEFAULT)
    
    kpis = calcular_kpis(df_incidencias, df_acompanamientos, df_cronograma, df_respuestas_diag, coordinador)
    
    render_alertas_cronograma(df_cronograma, coordinador)
    
    tab_resumen, tab_incidencias, tab_acompanamientos, tab_pruebas, tab_acciones, tab_sistema = st.tabs(
        ["📊 Resumen ejecutivo", "🚨 Incidencias", "📋 Acompañamientos", "🩺 Pruebas Diagnósticas", "⚡ Acciones", "🧩 Sistema"]
    )
    
    with tab_resumen:
        st.markdown("### Indicadores prioritarios")
        cols = st.columns(len(kpis))
        for col, kpi in zip(cols, kpis):
            with col:
                render_kpi_card(kpi)
        st.divider()
        render_incidencias_criticas(df_incidencias)
        st.divider()
        if st.button("📄 Exportar Reporte Ejecutivo (.docx)", type="primary"):
            buffer = generar_reporte_ejecutivo(kpis, df_incidencias)
            st.download_button(
                label="⬇️ Descargar Reporte Ejecutivo (.docx)",
                data=buffer, file_name=f"Reporte_Ejecutivo_{date.today().strftime('%Y%m%d')}.docx",
                mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                type="primary",
            )
            
    with tab_incidencias:
        render_tab_incidencias(df_incidencias)
        
    with tab_acompanamientos:
        render_tab_acompanamientos(df_acompanamientos)
        
    with tab_pruebas:
        render_tab_pruebas(df_respuestas_diag)
        
    with tab_acciones:
        render_tab_acciones()
        
    with tab_sistema:
        render_tab_sistema()

main()