"""
alerta.py — Sistema de Alerta Temprana y Reforzamiento (PAQUETE 1-6)
• Abierto a Coordinación Y Docentes; docente auto-llenado.
• Plan completo persistido (plan_json) + ver/re-descargar en Historial.
• Descarga persistente vía session_state.
• Normalización única de severidad (BD + Word + badges).
• Export CSV/Excel en Dashboard e Historial.
• Contador de estudiantes mejorado.
"""
import datetime
import json
import re
import sqlite3
from io import BytesIO

import pandas as pd
import streamlit as st
from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import nsdecls, qn
from docx.oxml import parse_xml, OxmlElement
from docx.shared import Inches, Pt, RGBColor

from core import ia

# ═══════════════════════════════════════════════════════════════════════════
# BASE DE DATOS (con migración de plan_json)
# ═══════════════════════════════════════════════════════════════════════════
DB_NAME = "gestion_etp.db"


def init_db():
    conn = sqlite3.connect(DB_NAME, check_same_thread=False)
    cursor = conn.cursor()
    cursor.execute('''
    CREATE TABLE IF NOT EXISTS alertas (
        id INTEGER PRIMARY KEY AUTOINCREMENT,
        fecha TEXT, docente TEXT, asignatura TEXT, seccion TEXT,
        competencia TEXT, total_estudiantes INTEGER, estado TEXT,
        plan_json TEXT
    )
    ''')
    # Migración: si la tabla ya existía sin plan_json, agregarla
    try:
        cursor.execute("ALTER TABLE alertas ADD COLUMN plan_json TEXT")
        conn.commit()
    except Exception:
        pass
    conn.commit()
    return conn


conn = init_db()


def get_todas_alertas():
    cursor = conn.cursor()
    cursor.execute("""SELECT id, fecha, docente, asignatura, seccion, competencia,
                      total_estudiantes, estado FROM alertas""")
    rows = cursor.fetchall()
    columnas = ["id", "Fecha", "Docente", "Módulo/Asignatura", "Sección",
                "Competencia/RA", "Total Estudiantes", "Estado"]
    return [dict(zip(columnas, row)) for row in rows]


def get_plan_alerta(id_alerta):
    cursor = conn.cursor()
    cursor.execute("SELECT plan_json FROM alertas WHERE id = ?", (id_alerta,))
    row = cursor.fetchone()
    if row and row[0]:
        try:
            return json.loads(row[0])
        except Exception:
            return None
    return None


def insertar_alerta(fecha, docente, asignatura, seccion, competencia,
                    total_estudiantes, estado, plan_json=None):
    cursor = conn.cursor()
    cursor.execute('''
    INSERT INTO alertas (fecha, docente, asignatura, seccion, competencia,
                         total_estudiantes, estado, plan_json)
    VALUES (?, ?, ?, ?, ?, ?, ?, ?)
    ''', (fecha, docente, asignatura, seccion, competencia,
          total_estudiantes, estado, plan_json))
    conn.commit()


def eliminar_alerta(id_alerta):
    cursor = conn.cursor()
    cursor.execute("DELETE FROM alertas WHERE id = ?", (id_alerta,))
    conn.commit()


# ═══════════════════════════════════════════════════════════════════════════
# NORMALIZACIÓN ÚNICA DE SEVERIDAD (mejora 4)
# ═══════════════════════════════════════════════════════════════════════════
def normalizar_severidad(nivel):
    n = str(nivel or "").upper()
    if "ROJO" in n or "CRIT" in n:
        return "Roja"
    if "AMAR" in n:
        return "Amarilla"
    if "VERDE" in n:
        return "Verde"
    return "Indefinida"


SEVERIDAD_META = {
    "Roja":       {"icono": "🔴", "css": "alerta-severity-roja"},
    "Amarilla":   {"icono": "🟡", "css": "alerta-severity-amarilla"},
    "Verde":      {"icono": "🟢", "css": "alerta-severity-verde"},
    "Indefinida": {"icono": "⚪", "css": "alerta-severity-indefinida"},
}

MAPA_ALERTA_WORD = {
    "Roja":       ("🔴 ALERTA ROJA — Intervención URGENTE", "FECACA"),
    "Amarilla":   ("🟡 ALERTA AMARILLA — Requiere atención focalizada", "FEF3C7"),
    "Verde":      ("🟢 ALERTA VERDE — Reforzamiento leve", "D1FAE5"),
    "Indefinida": ("⚪ Sin clasificar", "F1F5F9"),
}


def _sev(estado):
    return SEVERIDAD_META.get(estado, SEVERIDAD_META["Indefinida"])


def badge_severidad(estado):
    s = _sev(estado)
    return f'<span class="{s["css"]}">{s["icono"]} {estado}</span>'


# ═══════════════════════════════════════════════════════════════════════════
# CONTADOR DE ESTUDIANTES MEJORADO (mejora 6)
# ═══════════════════════════════════════════════════════════════════════════
def contar_estudiantes(texto):
    if not texto or not texto.strip():
        return 0
    # 1) Ítems numerados (1. / 1) / 1-)
    nums = re.findall(r'(?:^|\n)\s*\d+\s*[\.\)\-]', texto)
    if nums:
        return len(nums)
    # 2) Varias líneas
    lineas = [l for l in texto.strip().split("\n") if l.strip()]
    if len(lineas) > 1:
        return len(lineas)
    # 3) Una línea: separadores coma / punto y coma
    return max(1, len(re.split(r'[,;]', texto.strip())))


# ═══════════════════════════════════════════════════════════════════════════
# ESTILOS
# ═══════════════════════════════════════════════════════════════════════════
st.markdown("""
<style>
@import url('https://fonts.googleapis.com/css2?family=Inter:wght@300;400;500;600;700;800;900&display=swap');
html, body, [class*="css"] { font-family: 'Inter', sans-serif; background-color: #F0F4F8; color: #1E293B; }
.alerta-hero { background: linear-gradient(135deg, #450A0A 0%, #991B1B 45%, #DC2626 75%, #F59E0B 100%);
color: #fff; padding: 2.2rem; border-radius: 20px; margin-bottom: 1.5rem;
box-shadow: 0 25px 50px rgba(153,27,27,0.3); position: relative; overflow: hidden; }
.alerta-hero::before { content: ''; position: absolute; top: -50%; left: -50%; width: 200%; height: 200%;
background: radial-gradient(circle, rgba(255,255,255,0.12) 0%, transparent 60%); animation: alertaPulse 6s ease-in-out infinite; }
@keyframes alertaPulse { 0%,100%{transform:scale(1);opacity:0.5} 50%{transform:scale(1.1);opacity:0.8} }
.alerta-hero-title { font-size: 2.4rem; font-weight: 900; letter-spacing: -0.03em; margin-bottom: 0.4rem; position: relative; }
.alerta-hero-sub { font-size: 1.05rem; opacity: 0.9; line-height: 1.5; position: relative; }
.alerta-hero-badge { display: inline-block; background: rgba(255,255,255,0.15); border: 1px solid rgba(255,255,255,0.25);
border-radius: 8px; padding: 4px 12px; font-size: 0.8rem; font-weight: 600; margin-top: 0.8rem; margin-right: 8px; position: relative; }
.alerta-section-title { color: #DC2626; font-weight: 700; font-size: 1.12rem; border-bottom: 2px solid #FEE2E2;
padding-bottom: 8px; margin: 1.2rem 0 0.9rem 0; }
.alerta-kpi-card { background: #fff; border-radius: 14px; padding: 1.2rem; text-align: center;
border: 1px solid #E2E8F0; box-shadow: 0 4px 12px rgba(0,0,0,0.06); transition: all 0.3s ease; }
.alerta-kpi-card:hover { transform: translateY(-4px); box-shadow: 0 12px 24px rgba(0,0,0,0.12); }
.alerta-kpi-icon { font-size: 2rem; margin-bottom: 0.5rem; }
.alerta-kpi-value { font-size: 2rem; font-weight: 800; color: #991B1B; }
.alerta-kpi-label { font-size: 0.75rem; font-weight: 700; color: #94A3B8; text-transform: uppercase; letter-spacing: 0.05em; margin-top: 0.3rem; }
.alerta-severity-roja { background: #FEE2E2; color: #991B1B; padding: 3px 10px; border-radius: 6px; font-size: 0.75rem; font-weight: 700; }
.alerta-severity-amarilla { background: #FEF3C7; color: #92400E; padding: 3px 10px; border-radius: 6px; font-size: 0.75rem; font-weight: 700; }
.alerta-severity-verde { background: #D1FAE5; color: #065F46; padding: 3px 10px; border-radius: 6px; font-size: 0.75rem; font-weight: 700; }
.alerta-severity-indefinida { background: #F1F5F9; color: #64748B; padding: 3px 10px; border-radius: 6px; font-size: 0.75rem; font-weight: 700; }
.alerta-live-box { background: #FFF7ED; border: 1px solid #FED7AA; border-radius: 12px; padding: 1rem 1.2rem; margin-top: 0.8rem; }
.alerta-live-count { font-size: 1.6rem; font-weight: 800; color: #C2410C; }
</style>
""", unsafe_allow_html=True)

# ═══════════════════════════════════════════════════════════════════════════
# GUARDIA (mejora 1: abierto a coordinación Y docentes)
# ═══════════════════════════════════════════════════════════════════════════
if not (st.session_state.get("coordinador_autenticado", False)
        or st.session_state.get("docente_autenticado", False)):
    st.error("🔒 Debes iniciar sesión para usar el Sistema de Alerta Temprana.")
    st.stop()

ia.panel_sidebar_ia("Alerta Temprana v4.0")

st.markdown("""
<div class="alerta-hero">
    <div class="alerta-hero-title">🚨 Sistema de Alerta Temprana y Reforzamiento</div>
    <div class="alerta-hero-sub">Diagnóstico inteligente de brechas · Plan de recuperación multinivel · Seguimiento permanente</div>
    <div>
        <span class="alerta-hero-badge">🎯 Severidad Automática</span>
        <span class="alerta-hero-badge">🧠 Estrategias A/B</span>
        <span class="alerta-hero-badge">📊 Dashboard Analítico</span>
        <span class="alerta-hero-badge">📄 Word Profesional</span>
    </div>
</div>
""", unsafe_allow_html=True)

# ═══════════════════════════════════════════════════════════════════════════
# KPIs GLOBALES
# ═══════════════════════════════════════════════════════════════════════════
db_alertas_kpi = get_todas_alertas()
df_kpi = pd.DataFrame(db_alertas_kpi)
if not df_kpi.empty:
    total_alertas = len(df_kpi)
    rojas = int((df_kpi["Estado"] == "Roja").sum())
    amarillas = int((df_kpi["Estado"] == "Amarilla").sum())
    verdes = int((df_kpi["Estado"] == "Verde").sum())
    total_estudiantes_rec = int(df_kpi["Total Estudiantes"].fillna(0).sum())
else:
    total_alertas = rojas = amarillas = verdes = total_estudiantes_rec = 0

k1, k2, k3, k4, k5 = st.columns(5)
with k1:
    st.markdown(f'<div class="alerta-kpi-card"><div class="alerta-kpi-icon">🚨</div><div class="alerta-kpi-value">{total_alertas}</div><div class="alerta-kpi-label">Alertas Totales</div></div>', unsafe_allow_html=True)
with k2:
    st.markdown(f'<div class="alerta-kpi-card"><div class="alerta-kpi-icon">🔴</div><div class="alerta-kpi-value" style="color:#DC2626;">{rojas}</div><div class="alerta-kpi-label">Críticas</div></div>', unsafe_allow_html=True)
with k3:
    st.markdown(f'<div class="alerta-kpi-card"><div class="alerta-kpi-icon">🟡</div><div class="alerta-kpi-value" style="color:#F59E0B;">{amarillas}</div><div class="alerta-kpi-label">Moderadas</div></div>', unsafe_allow_html=True)
with k4:
    st.markdown(f'<div class="alerta-kpi-card"><div class="alerta-kpi-icon">🟢</div><div class="alerta-kpi-value" style="color:#10B981;">{verdes}</div><div class="alerta-kpi-label">Leves</div></div>', unsafe_allow_html=True)
with k5:
    st.markdown(f'<div class="alerta-kpi-card"><div class="alerta-kpi-icon">🧑‍🎓</div><div class="alerta-kpi-value">{total_estudiantes_rec}</div><div class="alerta-kpi-label">En Recuperación</div></div>', unsafe_allow_html=True)

tab1, tab2, tab3 = st.tabs(["🚨 Generar Alerta Temprana", "📊 Dashboard de Alertas", "📋 Historial Institucional"])


# ═══════════════════════════════════════════════════════════════════════════
# WORD (usa severidad normalizada — mejora 4)
# ═══════════════════════════════════════════════════════════════════════════
def _shade_cell(cell, color):
    shd = parse_xml(r'<w:shd {} w:fill="{}"/>'.format(nsdecls('w'), color))
    cell._tc.get_or_add_tcPr().append(shd)


def _add_page_number(doc, centro):
    for section in doc.sections:
        footer = section.footer
        footer.is_linked_to_previous = False
        p = footer.paragraphs[0] if footer.paragraphs else footer.add_paragraph()
        p.text = ""
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run(f"{centro} · Sistema de Alerta Temprana ETP · Página ")
        run.font.size = Pt(7)
        run.font.color.rgb = RGBColor(0x64, 0x74, 0x8B)
        fld1 = OxmlElement("w:fldChar"); fld1.set(qn("w:fldCharType"), "begin")
        instr = OxmlElement("w:instrText"); instr.set(qn("xml:space"), "preserve"); instr.text = "PAGE"
        fld2 = OxmlElement("w:fldChar"); fld2.set(qn("w:fldCharType"), "end")
        run2 = p.add_run()
        run2._r.append(fld1); run2._r.append(instr); run2._r.append(fld2)
        run2.font.size = Pt(7)
        run2.font.color.rgb = RGBColor(0x64, 0x74, 0x8B)


def build_alerta_docx(datos, meta):
    estado = normalizar_severidad(datos.get("NIVEL_ALERTA_GLOBAL", ""))
    plan_estudiantes = datos.get("PLAN_ACCION_ESTUDIANTES", [])
    actividad_recu = datos.get("ACTIVIDAD_RECUPERACION_GRUPAL", {})
    indicadores = datos.get("INDICADORES_PROGRESO", [])

    doc = Document()
    doc.styles['Normal'].font.name = 'Calibri'
    doc.styles['Normal'].font.size = Pt(11)
    for section in doc.sections:
        section.left_margin = Inches(0.75)
        section.right_margin = Inches(0.75)

    p_enc = doc.add_paragraph()
    p_enc.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p_enc.add_run(f"{meta.get('politecnico', '')}\n").bold = True
    p_enc.add_run("Sistema de Alerta Temprana y Plan de Recuperación Integral (ETP)\n").bold = True
    doc.add_paragraph(f"Docente: {meta.get('docente', '')} | Módulo: {meta.get('asignatura', '')} | Sección: {meta.get('seccion', '')}")
    doc.add_paragraph(f"Competencia / R.A. Analizado: {meta.get('competencia', '')}")
    doc.add_paragraph(f"Tiempo Disponible: {meta.get('tiempo', '')}")
    doc.add_paragraph(f"Fecha: {datetime.date.today().strftime('%d/%m/%Y')}")
    doc.add_paragraph("_" * 70)

    doc.add_heading("📊 1. Diagnóstico de Alerta Temprana", level=1)
    texto_alerta, color_alerta = MAPA_ALERTA_WORD.get(estado, MAPA_ALERTA_WORD["Indefinida"])
    t_alerta = doc.add_table(rows=1, cols=1)
    t_alerta.style = 'Table Grid'
    cell_a = t_alerta.rows[0].cells[0]
    _shade_cell(cell_a, color_alerta)
    p_a = cell_a.paragraphs[0]
    p_a.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run_a = p_a.add_run(texto_alerta)
    run_a.bold = True
    run_a.font.size = Pt(12)
    doc.add_paragraph()
    p_diag = doc.add_paragraph(datos.get("DIAGNOSTICO_GENERAL", ""))
    p_diag.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    doc.add_paragraph()

    doc.add_heading("🎯 2. Matriz de Intervención Personalizada", level=1)
    tabla_resumen = doc.add_table(rows=1, cols=5)
    tabla_resumen.style = 'Table Grid'
    hdr = tabla_resumen.rows[0].cells
    for i, h_text in enumerate(["Estudiante", "Severidad", "Brecha Detectada", "Tiempo Est.", "Estrategia Principal"]):
        hdr[i].paragraphs[0].add_run(h_text).bold = True
        hdr[i].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
        _shade_cell(hdr[i], "E2E8F0")
    mapa_sev_color = {"CRÍTICO": "FEE2E2", "MODERADO": "FEF3C7", "LEVE": "D1FAE5"}
    for est in plan_estudiantes:
        rc = tabla_resumen.add_row().cells
        rc[0].text = str(est.get("ESTUDIANTE", ""))
        rc[0].paragraphs[0].runs[0].bold = True
        sev = str(est.get("SEVERIDAD", ""))
        rc[1].text = sev
        if sev in mapa_sev_color:
            _shade_cell(rc[1], mapa_sev_color[sev])
        rc[1].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
        rc[2].text = str(est.get("BRECHA_DETECTADA", ""))
        rc[3].text = str(est.get("TIEMPO_ESTIMADO", ""))
        rc[3].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
        rc[4].text = str(est.get("ESTRATEGIA_A", ""))
    doc.add_paragraph()

    doc.add_heading("📋 3. Desglose Detallado de Intervención por Estudiante", level=1)
    for idx, est in enumerate(plan_estudiantes):
        nombre = str(est.get("ESTUDIANTE", f"Estudiante {idx+1}"))
        sev = str(est.get("SEVERIDAD", ""))
        icono_sev = {"CRÍTICO": "🔴", "MODERADO": "🟡", "LEVE": "🟢"}.get(sev, "⚪")
        doc.add_heading(f"{icono_sev} {nombre} — Severidad: {sev}", level=3)
        t_det = doc.add_table(rows=7, cols=2)
        t_det.style = 'Table Grid'
        campos = [
            ("Brecha Detectada", str(est.get("BRECHA_DETECTADA", ""))),
            ("Tiempo Estimado de Recuperación", str(est.get("TIEMPO_ESTIMADO", ""))),
            ("Estrategia A (Principal)", str(est.get("ESTRATEGIA_A", ""))),
            ("Estrategia B (Alternativa)", str(est.get("ESTRATEGIA_B", ""))),
            ("Recursos Específicos", str(est.get("RECURSOS", ""))),
            ("Adaptaciones NEAE", str(est.get("ADAPTACION_NEAE", ""))),
            ("Acción del Docente", str(est.get("ACCION_DOCENTE", ""))),
        ]
        for r, (label, valor) in enumerate(campos):
            t_det.cell(r, 0).text = label
            t_det.cell(r, 0).paragraphs[0].runs[0].bold = True
            _shade_cell(t_det.cell(r, 0), "F1F5F9")
            t_det.cell(r, 1).text = valor
            t_det.cell(r, 1).paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        doc.add_paragraph()
        p_roles = doc.add_paragraph()
        p_roles.add_run("👤 Acción del Estudiante: ").bold = True
        p_roles.add_run(str(est.get("ACCION_ESTUDIANTE", "")))
        p_fam = doc.add_paragraph()
        p_fam.add_run("🏠 Acción de la Familia: ").bold = True
        p_fam.add_run(str(est.get("ACCION_FAMILIA", "")))
        doc.add_paragraph("_" * 50)
        doc.add_paragraph()

    doc.add_heading(f"⚙️ 4. Actividad de Recuperación: {actividad_recu.get('TITULO', 'Nivelación')}", level=1)
    p_act = doc.add_paragraph(str(actividad_recu.get("DESCRIPCION", "")))
    p_act.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    doc.add_heading("Pasos Metodológicos de la Tutoría:", level=3)
    for paso in actividad_recu.get("PASOS", []):
        p_paso = doc.add_paragraph(str(paso), style='List Bullet')
        p_paso.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    doc.add_paragraph()
    p_rec = doc.add_paragraph()
    p_rec.add_run("📦 Recursos de la Actividad: ").bold = True
    p_rec.add_run(str(actividad_recu.get("RECURSOS", "")))
    p_tiempo = doc.add_paragraph()
    p_tiempo.add_run("⏱️ Tiempo Estimado: ").bold = True
    p_tiempo.add_run(str(actividad_recu.get("TIEMPO_ESTIMADO", "")))
    doc.add_paragraph()

    doc.add_heading("📈 5. Indicadores de Progreso Intermedio", level=1)
    doc.add_paragraph("Utilice estos indicadores para monitorear la mejora del estudiante ANTES de la reevaluación final.")
    if indicadores:
        t_ind = doc.add_table(rows=len(indicadores) + 1, cols=4)
        t_ind.style = 'Table Grid'
        hdr_ind = t_ind.rows[0].cells
        for i, txt in enumerate(["No.", "Indicador de Progreso", "¿Evidencia?", "Fecha"]):
            hdr_ind[i].text = txt
            hdr_ind[i].paragraphs[0].runs[0].bold = True
            hdr_ind[i].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
            _shade_cell(hdr_ind[i], "E2E8F0")
        for i, ind in enumerate(indicadores):
            t_ind.cell(i+1, 0).text = str(i+1)
            t_ind.cell(i+1, 0).paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
            t_ind.cell(i+1, 1).text = str(ind)
            t_ind.cell(i+1, 2).text = "☐ Sí  ☐ No"
            t_ind.cell(i+1, 3).text = "___/___/____"
    doc.add_paragraph()

    doc.add_heading("✅ 6. Criterio de Cierre y Reevaluación", level=1)
    p_reval = doc.add_paragraph(datos.get("CRITERIO_REVALUACION", ""))
    p_reval.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    doc.add_paragraph()

    doc.add_heading("🆘 7. Plan de Contingencia", level=1)
    p_cont = doc.add_paragraph(datos.get("PLAN_CONTINGENCIA", ""))
    p_cont.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    doc.add_paragraph()

    doc.add_heading("💌 8. Guía de Comunicación con la Familia", level=1)
    doc.add_paragraph("Borrador sugerido para comunicar la situación del estudiante a su familia:")
    p_com = doc.add_paragraph()
    p_com.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    p_com.add_run(datos.get("COMUNICADO_FAMILIA", ""))
    doc.add_paragraph()

    doc.add_heading("📝 9. Registro de Seguimiento Semanal", level=1)
    doc.add_paragraph("Complete este registro cada semana para documentar la evolución de cada estudiante.")
    if plan_estudiantes:
        n_semanas = 4
        t_seg = doc.add_table(rows=len(plan_estudiantes) + 1, cols=n_semanas + 2)
        t_seg.style = 'Table Grid'
        hdr_seg = t_seg.rows[0].cells
        hdr_seg[0].text = "Estudiante"
        hdr_seg[0].paragraphs[0].runs[0].bold = True
        _shade_cell(hdr_seg[0], "E2E8F0")
        hdr_seg[1].text = "Severidad Inicial"
        hdr_seg[1].paragraphs[0].runs[0].bold = True
        _shade_cell(hdr_seg[1], "E2E8F0")
        for s in range(n_semanas):
            hdr_seg[s+2].text = f"Semana {s+1}"
            hdr_seg[s+2].paragraphs[0].runs[0].bold = True
            hdr_seg[s+2].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
            _shade_cell(hdr_seg[s+2], "F8FAFC")
        for i, est in enumerate(plan_estudiantes):
            t_seg.cell(i+1, 0).text = str(est.get("ESTUDIANTE", ""))
            t_seg.cell(i+1, 0).paragraphs[0].runs[0].bold = True
            sev_icono = {"CRÍTICO": "🔴", "MODERADO": "🟡", "LEVE": "🟢"}.get(str(est.get("SEVERIDAD", "")), "⚪")
            t_seg.cell(i+1, 1).text = f"{sev_icono} {str(est.get('SEVERIDAD', ''))}"
            t_seg.cell(i+1, 1).paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
    doc.add_paragraph()
    p_ley = doc.add_paragraph()
    run_ley = p_ley.add_run("Leyenda: L = Logrado | EP = En Proceso | NA = Necesita Apoyo | NP = No Presentó")
    run_ley.italic = True
    run_ley.font.size = Pt(9)
    doc.add_paragraph("\n\n")

    t_firmas = doc.add_table(rows=2, cols=3)
    t_firmas.cell(0, 0).text = "_________________________"
    t_firmas.cell(0, 1).text = "_________________________"
    t_firmas.cell(0, 2).text = "_________________________"
    t_firmas.cell(1, 0).text = "Director/a de Centro"
    t_firmas.cell(1, 1).text = "Coordinador/a ETP"
    t_firmas.cell(1, 2).text = "Docente"
    for row in t_firmas.rows:
        for cell in row.cells:
            cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER

    _add_page_number(doc, meta.get("politecnico", ""))
    buffer = BytesIO()
    doc.save(buffer)
    buffer.seek(0)
    return buffer


# ═══════════════════════════════════════════════════════════════════════════
# TAB 1: GENERAR ALERTA
# ═══════════════════════════════════════════════════════════════════════════
with tab1:
    # Mejora 1: auto-llenar docente con el usuario logueado
    _nombre_default = st.session_state.get("usuario_display_nombre") or "Ing. Bernardo Antonio Hernández Batista"

    with st.form("form_alerta", clear_on_submit=False):
        st.markdown('<div class="alerta-section-title">🏫 1. Datos del Contexto e Identificación</div>', unsafe_allow_html=True)
        col1, col2 = st.columns(2)
        with col1:
            docente = st.text_input("Nombre del Docente", value=_nombre_default)
            asignatura = st.text_input("Módulo / Asignatura", placeholder="Ej: Sistemas Operativos / Redes LAN")
        with col2:
            politecnico = st.text_input("Centro Educativo", value="Politécnico Salesiano Arquides Calderón")
            seccion = st.text_input("Sección / Grado", placeholder="Ej: 6to de Informática")

        st.markdown('<div class="alerta-section-title">⚠️ 2. Registro de Estudiantes y Brechas Detectadas</div>', unsafe_allow_html=True)
        competencia_evaluada = st.text_input(
            "Resultado de Aprendizaje (R.A.) o Competencia Evaluada",
            placeholder="Ej: R.A.1 Configurar los parámetros de red local...",
        )
        estudiantes_apoyo = st.text_area(
            "Estudiantes identificados 'En Proceso' o 'Necesitan Apoyo' (uno por línea)",
            height=120,
            placeholder="Ej:\n1. Carlos Pérez - Dificultad en el direccionamiento IP estático.\n2. María Gómez - Confusión en la identificación de topologías físicas.",
        )
        n_est = contar_estudiantes(estudiantes_apoyo)
        st.markdown(f"""
        <div class="alerta-live-box">
            <span class="alerta-live-count">{n_est}</span>
            <span style="color:#9A3412; font-weight:600;"> estudiante(s) detectado(s) para el plan de recuperación</span>
            <div style="font-size:0.8rem; color:#C2410C; margin-top:4px;">
                💡 Escribe un estudiante por línea (o numerados) para un conteo preciso.
            </div>
        </div>
        """, unsafe_allow_html=True)

        st.markdown('<div class="alerta-section-title">👥 3. Perfil del Grupo y Condiciones de Recuperación</div>', unsafe_allow_html=True)
        col_p1, col_p2 = st.columns(2)
        with col_p1:
            caracteristicas_grupo = st.text_area(
                "Características del grupo y NEAE", height=80,
                placeholder="Ej: Grupo visual-kinestésico. 2 estudiantes con dislexia, 1 con TDAH.",
            )
        with col_p2:
            tiempo_disponible = st.text_input(
                "Tiempo disponible para recuperación",
                placeholder="Ej: 3 semanas / 6 sesiones de tutoría",
            )

        max_tokens, temperature = ia.control_avanzado(default_tokens=16384, tope=32000, default_temp=0.2)
        st.markdown("<br>", unsafe_allow_html=True)
        submit_button = st.form_submit_button("⚙️ Generar Plan de Recuperación Integral y Guardar", type="primary", width="stretch")

    if submit_button:
        cfg = ia.config_ia()
        if not cfg["api_key"]:
            st.error("🔒 Debes ingresar tu API Key en la página de Inicio (barra lateral).")
        elif not asignatura or not estudiantes_apoyo or not competencia_evaluada:
            st.warning("📝 Por favor, completa la asignatura, la competencia y el listado de estudiantes.")
        else:
            with st.spinner(f'🧠 Diseñando plan de recuperación multinivel con {cfg["modelo"]}...'):
                texto_crudo = None
                try:
                    prompt_maestro = f"""Actúa como un Coordinador Pedagógico de Alto Nivel, Especialista en Educación Técnico Profesional (ETP) del MINERD y Experto en Pedagogía Diferenciada e Intervención Psicopedagógica.
CONTEXTO:
Competencia / R.A. Evaluado: {competencia_evaluada}
Estudiantes con brechas detectadas:
{estudiantes_apoyo}
Características del grupo y NEAE: {caracteristicas_grupo}
Tiempo disponible para recuperación: {tiempo_disponible}
OBJETIVO:
Diseñar un Sistema de Alerta Temprana INTEGRAL que incluya: diagnóstico con niveles de severidad, plan de recuperación multinivel con estrategias diferenciadas (A y B), recursos específicos, separación de responsabilidades (docente/estudiante/familia), indicadores de progreso intermedios, plan de contingencia y guía de comunicación familiar.
REGLAS ESTRICTAS:
CLASIFICACIÓN DE SEVERIDAD: Para cada estudiante, determina el nivel de severidad de su brecha:
CRÍTICO (🔴): Brecha fundamental que impide avanzar en el módulo. Requiere intervención inmediata.
MODERADO (🟡): Brecha significativa que dificulta el desempeño, pero no lo bloquea completamente.
LEVE (🟢): Brecha menor o laguna de conocimiento puntual. Se resuelve con práctica guiada.
ESTRATEGIA DIFERENCIADA: Para cada estudiante proporciona DOS estrategias:
ESTRATEGIA A (principal): La intervención más efectiva según el tipo de brecha y el perfil del estudiante.
ESTRATEGIA B (alternativa): Si la Estrategia A no produce resultados, qué enfoque diferente se puede intentar.
RECURSOS ESPECÍFICOS: Para cada estudiante sugiere al menos 2 recursos concretos (video tutorial, simulador, ejercicio práctico, lectura guiada, infografía, etc.).
ADAPTACIONES NEAE: Si el estudiante tiene necesidades educativas especiales, propone adaptaciones específicas. Si no aplica, indica "Sin adaptaciones adicionales requeridas".
SEPARACIÓN DE ROLES: Para cada estudiante indica qué debe hacer EL DOCENTE, EL ESTUDIANTE y LA FAMILIA.
INDICADORES DE PROGRESO: Define 3-4 indicadores intermedios observables.
PLAN DE CONTINGENCIA: Si después del tiempo estimado no hay mejora significativa, ¿qué acción se recomienda?
COMUNICACIÓN FAMILIAR: Redacta un borrador de comunicado respetuoso y profesional.
CODIFICACIÓN OBLIGATORIA (evita errores de formato JSON):
Si cualquier valor de texto necesitara un salto de línea, tabulación o comilla doble,
usa en su lugar estos marcadores: salto de línea → {ia.MARKER_NL} · comilla doble → {ia.MARKER_DQ} ·
tabulación → {ia.MARKER_TAB}. Nunca uses saltos de línea literales dentro de un valor JSON.
FORMATO DE SALIDA ESTRICTO (JSON NATIVO OBLIGATORIO):
{{
 "DIAGNOSTICO_GENERAL": "Análisis panorámico de las brechas detectadas...",
 "NIVEL_ALERTA_GLOBAL": "ROJO / AMARILLO / VERDE",
 "PLAN_ACCION_ESTUDIANTES": [
   {{
     "ESTUDIANTE": "Nombre del estudiante",
     "BRECHA_DETECTADA": "Resumen técnico de la dificultad",
     "SEVERIDAD": "CRÍTICO / MODERADO / LEVE",
     "TIEMPO_ESTIMADO": "Ej: 2 semanas / 4 sesiones",
     "ESTRATEGIA_A": "Intervención principal recomendada",
     "ESTRATEGIA_B": "Intervención alternativa si A no funciona",
     "RECURSOS": "Recurso 1: [nombre y tipo]. Recurso 2: [nombre y tipo].",
     "ADAPTACION_NEAE": "Adaptación específica o 'Sin adaptaciones adicionales requeridas'",
     "ACCION_DOCENTE": "Qué debe hacer específicamente el docente",
     "ACCION_ESTUDIANTE": "Qué debe hacer específicamente el estudiante",
     "ACCION_FAMILIA": "Cómo puede la familia apoyar desde casa"
   }}
 ],
 "ACTIVIDAD_RECUPERACION_GRUPAL": {{
   "TITULO": "Título de la actividad de nivelación",
   "DESCRIPCION": "Descripción metodológica completa",
   "PASOS": ["Paso 1...", "Paso 2...", "Paso 3...", "Paso 4..."],
   "RECURSOS": "Materiales y herramientas necesarios",
   "TIEMPO_ESTIMADO": "Duración de la actividad"
 }},
 "INDICADORES_PROGRESO": ["Indicador 1...", "Indicador 2...", "Indicador 3...", "Indicador 4..."],
 "CRITERIO_REVALUACION": "Cómo se comprobará que el estudiante superó la brecha",
 "PLAN_CONTINGENCIA": "Acción a tomar si no hay mejora...",
 "COMUNICADO_FAMILIA": "Borrador de comunicado respetuoso para la familia..."
}}
"""
                    texto_crudo, flags = ia.solicitar_ia(
                        prompt_maestro, modo="json", max_tokens=max_tokens,
                        temperature=temperature, modulo="alerta_temprana",
                    )
                    datos = ia.parsear_json_robusto(texto_crudo)
                    datos = ia.decodificar_marcadores(datos)

                    plan_estudiantes = datos.get("PLAN_ACCION_ESTUDIANTES", [])
                    estado_db = normalizar_severidad(datos.get("NIVEL_ALERTA_GLOBAL", ""))

                    meta_doc = {
                        "politecnico": politecnico, "docente": docente, "asignatura": asignatura,
                        "seccion": seccion, "competencia": competencia_evaluada,
                        "tiempo": tiempo_disponible,
                    }
                    # Mejora 2: persistir plan completo (datos + meta)
                    plan_json = json.dumps({"datos": datos, "meta": meta_doc}, ensure_ascii=False)

                    insertar_alerta(
                        fecha=datetime.date.today().strftime('%Y-%m-%d'),
                        docente=docente, asignatura=asignatura, seccion=seccion,
                        competencia=competencia_evaluada,
                        total_estudiantes=len(plan_estudiantes), estado=estado_db,
                        plan_json=plan_json,
                    )

                    buffer = build_alerta_docx(datos, meta_doc)
                    # Mejora 3: guardar bytes en session_state (descarga persistente)
                    st.session_state.alerta_buffer = buffer.getvalue()
                    st.session_state.alerta_filename = f"Plan_Recuperacion_Integral_{asignatura[:10]}.docx"

                    criticos = sum(1 for e in plan_estudiantes if str(e.get("SEVERIDAD", "")).upper() == "CRÍTICO")
                    moderados = sum(1 for e in plan_estudiantes if str(e.get("SEVERIDAD", "")).upper() == "MODERADO")
                    leves = sum(1 for e in plan_estudiantes if str(e.get("SEVERIDAD", "")).upper() == "LEVE")

                    st.markdown('<div class="alerta-section-title">📦 Resultado del Diagnóstico</div>', unsafe_allow_html=True)
                    s = _sev(estado_db)
                    st.markdown(f"""
                    <div class="alerta-live-box" style="text-align:center;">
                        <div style="font-size:1.3rem; font-weight:800; color:#991B1B;">
                            {s['icono']} Nivel de Alerta Global: {estado_db.upper()}
                        </div>
                    </div>
                    """, unsafe_allow_html=True)

                    col_r1, col_r2, col_r3, col_r4 = st.columns(4)
                    with col_r1:
                        st.markdown(f'<div class="alerta-kpi-card"><div class="alerta-kpi-icon">🔴</div><div class="alerta-kpi-value" style="color:#DC2626;">{criticos}</div><div class="alerta-kpi-label">Críticos</div></div>', unsafe_allow_html=True)
                    with col_r2:
                        st.markdown(f'<div class="alerta-kpi-card"><div class="alerta-kpi-icon">🟡</div><div class="alerta-kpi-value" style="color:#F59E0B;">{moderados}</div><div class="alerta-kpi-label">Moderados</div></div>', unsafe_allow_html=True)
                    with col_r3:
                        st.markdown(f'<div class="alerta-kpi-card"><div class="alerta-kpi-icon">🟢</div><div class="alerta-kpi-value" style="color:#10B981;">{leves}</div><div class="alerta-kpi-label">Leves</div></div>', unsafe_allow_html=True)
                    with col_r4:
                        st.markdown(f'<div class="alerta-kpi-card"><div class="alerta-kpi-icon">🧑‍🎓</div><div class="alerta-kpi-value">{len(plan_estudiantes)}</div><div class="alerta-kpi-label">Total</div></div>', unsafe_allow_html=True)

                    st.toast("✅ Plan de Recuperación generado y guardado.", icon="🚨")
                    if flags.get("reintento"):
                        st.info("ℹ️ La primera llamada se cortó por tokens; el reintento automático tuvo éxito.")
                except ValueError as ve:
                    ia.render_error_ia(ve, texto_crudo)
                except Exception as e:
                    ia.render_error_ia(e, texto_crudo)

    # Mejora 3: descarga persistente (sobrevive reruns)
    if st.session_state.get("alerta_buffer"):
        st.download_button(
            label="📥 Descargar Plan de Recuperación Integral (.docx)",
            data=st.session_state.alerta_buffer,
            file_name=st.session_state.get("alerta_filename", "Plan_Recuperacion.docx"),
            mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
            type="primary",
            width="stretch",
        )

# ═══════════════════════════════════════════════════════════════════════════
# TAB 2: DASHBOARD (con export CSV/Excel — mejora 5)
# ═══════════════════════════════════════════════════════════════════════════
with tab2:
    st.markdown('<div class="alerta-section-title">📊 Analítica de Alertas Tempranas</div>', unsafe_allow_html=True)
    db_alertas_dash = get_todas_alertas()
    if not db_alertas_dash:
        st.info("📭 Aún no hay alertas registradas. Genera la primera en la pestaña '🚨 Generar Alerta Temprana'.")
    else:
        df_dash = pd.DataFrame(db_alertas_dash)
        df_dash["Fecha_dt"] = pd.to_datetime(df_dash["Fecha"], errors="coerce")

        col_g1, col_g2 = st.columns(2)
        with col_g1:
            st.markdown("##### 🎯 Distribución por Severidad")
            st.bar_chart(df_dash["Estado"].value_counts())
        with col_g2:
            st.markdown("##### 📚 Alertas por Módulo/Asignatura")
            st.bar_chart(df_dash["Módulo/Asignatura"].value_counts().head(8))

        col_g3, col_g4 = st.columns(2)
        with col_g3:
            st.markdown("##### 🧑‍🏫 Alertas por Docente")
            st.bar_chart(df_dash["Docente"].value_counts().head(8))
        with col_g4:
            st.markdown("##### 📅 Alertas por Mes")
            df_dash["Mes"] = df_dash["Fecha_dt"].dt.strftime("%Y-%m")
            st.bar_chart(df_dash["Mes"].value_counts().sort_index())

        st.markdown("##### 📋 Detalle de Alertas")
        df_show = df_dash.drop(columns=["id", "Fecha_dt", "Mes"]).copy()
        df_show["Estado"] = df_show["Estado"].apply(lambda e: f"{_sev(e)['icono']} {e}")
        st.dataframe(df_show, width="stretch", hide_index=True)

        # Mejora 5: export CSV / Excel
        st.markdown("##### 📥 Exportar datos")
        c_csv, c_xls = st.columns(2)
        df_export = df_dash.drop(columns=["id", "Fecha_dt", "Mes"])
        with c_csv:
            st.download_button(
                "📄 Descargar CSV",
                data=df_export.to_csv(index=False).encode("utf-8-sig"),
                file_name="alertas_tempranas.csv",
                mime="text/csv",
                width="stretch",
            )
        with c_xls:
            buf_x = BytesIO()
            df_export.to_excel(buf_x, index=False, engine="openpyxl")
            buf_x.seek(0)
            st.download_button(
                "📊 Descargar Excel",
                data=buf_x,
                file_name="alertas_tempranas.xlsx",
                mime="application/vnd.openxmlformats-officedocument.spreadsheetml.sheet",
                width="stretch",
            )

# ═══════════════════════════════════════════════════════════════════════════
# TAB 3: HISTORIAL (ver/re-descargar plan — mejora 2 + export mejora 5)
# ═══════════════════════════════════════════════════════════════════════════
with tab3:
    st.markdown('<div class="alerta-section-title">🗄️ Archivo Institucional de Alertas</div>', unsafe_allow_html=True)
    db_alertas_hist = get_todas_alertas()
    if not db_alertas_hist:
        st.info("Aún no hay alertas de recuperación registradas en la base de datos.")
    else:
        df_hist = pd.DataFrame(db_alertas_hist)

        col_b1, col_b2, col_b3 = st.columns([2, 1, 1])
        with col_b1:
            busqueda = st.text_input("🔍 Buscar", placeholder="Docente, asignatura, competencia o sección...")
        with col_b2:
            filtro_sev = st.selectbox("Severidad", ["Todas", "Roja", "Amarilla", "Verde", "Indefinida"])
        with col_b3:
            filtro_doc = st.selectbox("Docente", ["Todos"] + sorted(df_hist["Docente"].dropna().unique().tolist()))

        df_filtrado = df_hist.copy()
        if busqueda.strip():
            texto = busqueda.strip().lower()
            df_filtrado = df_filtrado[
                df_filtrado.apply(lambda row: texto in " ".join(str(v).lower() for v in row if pd.notna(v)), axis=1)
            ]
        if filtro_sev != "Todas":
            df_filtrado = df_filtrado[df_filtrado["Estado"] == filtro_sev]
        if filtro_doc != "Todos":
            df_filtrado = df_filtrado[df_filtrado["Docente"] == filtro_doc]

        st.markdown(f"**{len(df_filtrado)}** alerta(s) encontrada(s)")
        df_show = df_filtrado.drop(columns=["id"]).copy()
        df_show["Estado"] = df_show["Estado"].apply(lambda e: f"{_sev(e)['icono']} {e}")
        st.dataframe(df_show, width="stretch", hide_index=True)

        # Mejora 2: ver / re-descargar plan guardado
        st.markdown("##### 👁️ Ver / 📥 Re-descargar Plan guardado")
        opciones_plan = {
            f"ID {row['id']}: {row['Fecha']} - {row['Docente']} ({row['Módulo/Asignatura']})": row['id']
            for row in db_alertas_hist
        }
        sel_plan = st.selectbox("Selecciona una alerta para ver su plan", list(opciones_plan.keys()),
                                index=None, placeholder="Elige...")
        if sel_plan:
            plan = get_plan_alerta(opciones_plan[sel_plan])
            if plan:
                datos_plan = plan.get("datos", {})
                meta_plan = plan.get("meta", {})
                with st.expander("📄 Resumen del plan", expanded=True):
                    st.markdown(f"**Nivel global:** {normalizar_severidad(datos_plan.get('NIVEL_ALERTA_GLOBAL', ''))}")
                    st.markdown(f"**Estudiantes:** {len(datos_plan.get('PLAN_ACCION_ESTUDIANTES', []))}")
                    st.markdown(f"**Diagnóstico:** {datos_plan.get('DIAGNOSTICO_GENERAL', '')[:400]}")
                buf_re = build_alerta_docx(datos_plan, meta_plan)
                st.download_button(
                    "📥 Re-descargar Word (.docx)",
                    data=buf_re,
                    file_name=f"Plan_Re_{meta_plan.get('asignatura', 'alerta')[:10]}.docx",
                    mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                    width="stretch",
                )
            else:
                st.info("Esta alerta fue creada antes de la función de plan guardado; no hay plan almacenado.")

        st.markdown("---")
        with st.expander("🗑️ Gestionar Historial (Eliminar Registro)"):
            st.warning("⚠️ Precaución: Eliminar un registro borrará la constancia de la alerta en el sistema.")
            opciones_del = {
                f"ID {row['id']}: {row['Fecha']} - {row['Docente']} ({row['Módulo/Asignatura']})": row['id']
                for row in db_alertas_hist
            }
            seleccion_del = st.selectbox("Selecciona la alerta a eliminar", list(opciones_del.keys()),
                                         index=None, placeholder="Elige...", key="sel_del")
            confirmar = st.checkbox("Confirmo eliminar esta alerta", key="conf_del_alerta")
            if st.button("Eliminar permanentemente", type="primary",
                         disabled=not (seleccion_del and confirmar)):
                eliminar_alerta(opciones_del[seleccion_del])
                st.toast("Alerta eliminada de la base de datos institucional.", icon="🗑️")
                st.rerun()