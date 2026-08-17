"""
acompanamiento.py — Acompañamiento Docente ETP (Flujo Unificado Fígital + Matrices IA)
• Cascada real: Docente → Módulo (asignados) → Sección (auto-actualizada).
• NUEVO: Creación dinámica de indicadores. La IA diseña una rúbrica única
  basada en el Área, Tipo de Acompañamiento y Módulo alineada a MINERD.
• Digitalización y corrección universal de la ficha generada.
"""
import base64
import re
import sqlite3
from datetime import datetime, date
from io import BytesIO

import pandas as pd
import streamlit as st
from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement, parse_xml
from docx.oxml.ns import qn, nsdecls
from docx.shared import Inches, Pt, RGBColor

try:
    from pypdf import PdfReader
except Exception:
    try:
        from PyPDF2 import PdfReader
    except Exception:
        PdfReader = None

from core import ia

DB_NAME = "gestion_etp.db"

# ═══════════════════════════════════════════════════════════════════════════
# ÁREAS Y TIPOS OFICIALES (MINERD)
# ═══════════════════════════════════════════════════════════════════════════
AREAS_Y_TIPOS = {
    "🔧 Área Técnica": [
        "Observación de Aula en Taller / Laboratorio",
        "Revisión de Planificación por Competencias",
        "Acompañamiento en Prácticas Profesionales (FCT)",
        "Evaluación de Desempeño Técnico y Retroalimentación",
        "Seguimiento a Plan de Mejora Técnico",
        "Verificación de Normas de Seguridad e Higiene"
    ],
    "📚 Área Académica": [
        "Observación de Aula Pedagógica",
        "Revisión de Planificación Diaria / Unidad",
        "Acompañamiento en Estrategias Metodológicas",
        "Análisis de Rendimiento y Retroalimentación",
        "Seguimiento a Plan de Mejora Académico",
        "Seguimiento de Adaptaciones NEAE"
    ]
}

# ═══════════════════════════════════════════════════════════════════════════
# BASE DE DATOS LOCAL
# ═══════════════════════════════════════════════════════════════════════════
def _conn():
    return sqlite3.connect(DB_NAME, check_same_thread=False)

def asegurar_tablas():
    conn = _conn()
    conn.execute('''CREATE TABLE IF NOT EXISTS evidencias_acompanamiento (
        id INTEGER PRIMARY KEY AUTOINCREMENT,
        fecha TEXT, docente TEXT, modulo TEXT, seccion TEXT, tipo TEXT, area TEXT,
        coordinador TEXT, puntuacion INTEGER, nivel TEXT,
        archivo_nombre TEXT, pdf_base64 TEXT, texto_extraido TEXT
    )''')
    conn.commit()
    conn.close()

def listar_docentes_bd():
    conn = _conn()
    cur = conn.cursor()
    cur.execute("SELECT DISTINCT docente FROM docentes WHERE docente IS NOT NULL AND docente != ''")
    rows = [r[0] for r in cur.fetchall()]
    conn.close()
    return rows

def obtener_modulos_usuario_bd(docente):
    conn = _conn()
    cur = conn.cursor()
    cur.execute("SELECT modulo, seccion FROM docentes WHERE docente=? AND modulo IS NOT NULL AND modulo != ''", (docente,))
    rows = [list(r) for r in cur.fetchall()]
    conn.close()
    return rows

def guardar_evidencia(fecha, docente, modulo, seccion, tipo, area, coordinador, puntuacion, nivel, archivo_nombre, pdf_bytes, texto):
    conn = _conn()
    conn.execute('''INSERT INTO evidencias_acompanamiento
        (fecha, docente, modulo, seccion, tipo, area, coordinador, puntuacion, nivel, archivo_nombre, pdf_base64, texto_extraido)
        VALUES (?,?,?,?,?,?,?,?,?,?,?,?)''',
        (fecha, docente, modulo, seccion, tipo, area, coordinador, puntuacion, nivel, archivo_nombre, base64.b64encode(pdf_bytes).decode(), texto))
    conn.commit()
    conn.close()

def listar_evidencias():
    conn = _conn()
    cur = conn.cursor()
    cur.execute("SELECT id, fecha, docente, modulo, seccion, tipo, area, puntuacion, nivel, archivo_nombre FROM evidencias_acompanamiento ORDER BY id DESC")
    cols = ["id", "fecha", "docente", "modulo", "seccion", "tipo", "area", "puntuacion", "nivel", "archivo_nombre"]
    rows = [dict(zip(cols, r)) for r in cur.fetchall()]
    conn.close()
    return rows

def obtener_evidencia(id_ev):
    conn = _conn()
    cur = conn.cursor()
    cur.execute("SELECT pdf_base64, texto_extraido FROM evidencias_acompanamiento WHERE id=?", (id_ev,))
    row = cur.fetchone()
    conn.close()
    return row if row else (None, None)

def eliminar_evidencia(id_ev):
    conn = _conn()
    conn.execute("DELETE FROM evidencias_acompanamiento WHERE id=?", (id_ev,))
    conn.commit()
    conn.close()

asegurar_tablas()

# ═══════════════════════════════════════════════════════════════════════════
# INTELIGENCIA ARTIFICIAL (Indicadores y Corrección)
# ═══════════════════════════════════════════════════════════════════════════
def generar_indicadores_ia(area, tipo, modulo):
    """Crea una matriz de evaluación única y apegada al contexto dominicano."""
    prompt = f"""Actúa como Técnico Nacional del MINERD (República Dominicana).
Crea una matriz de evaluación de exactamente 15 indicadores observables (evaluables con Sí/No).
Contexto del Acompañamiento:
- Área: {area}
- Tipo de Acompañamiento: {tipo}
- Módulo / Asignatura a evaluar: {modulo}

Los indicadores deben ser altamente específicos a esta combinación, profesionales, orientados a competencias y apegados al currículo dominicano.
CODIFICACIÓN: comilla doble → {ia.MARKER_DQ}
Devuelve ÚNICAMENTE un JSON válido con este formato:
{{
    "INDICADORES": [
        "1. [Texto del indicador observable]",
        "2. [Texto del indicador observable]",
        ...hasta 15
    ]
}}"""
    datos, _ = ia.solicitar_json(prompt, max_tokens=1500, temperature=0.4, modulo="acompanamiento_ind")
    return datos.get("INDICADORES", [
        "1. El docente demuestra dominio en el área evaluada.",
        "2. Se evidencia correspondencia con la ordenanza curricular."
    ])

def analizar_ficha_ia(texto, meta):
    """Audita el escaneo extrayendo los indicadores directamente de la imagen."""
    prompt = f"""Actúa como un Coordinador Pedagógico ETP del MINERD que EVALÚA una Ficha de
Acompañamiento llenada a mano (ahora digitalizada).
DATOS: Docente: {meta.get('docente','')} | Módulo: {meta.get('modulo','')} | Área: {meta.get('area','')} | Tipo: {meta.get('tipo','')}

TEXTO EXTRAÍDO DE LA FICHA ESCANEADA:
{texto}

TAREA OBLIGATORIA:
1. Extrae los indicadores de evaluación que aparecen explícitamente en el documento escaneado.
2. Interpreta si cada indicador fue logrado (Sí) o no (No) según las marcas del evaluador.
3. Calcula el PUNTUJE GLOBAL (0 a 100) en base a la proporción de Sí vs No.
4. Asigna un NIVEL GLOBAL: Excelente (90-100), Bueno (75-89), Aceptable con Mejoras (60-74), No Aceptable (0-59).
5. Resume fortalezas, áreas de mejora y cualquier observación adicional encontrada en el texto.

CODIFICACIÓN: salto de línea → {ia.MARKER_NL} · comilla doble → {ia.MARKER_DQ} · tabulación → {ia.MARKER_TAB}.
Devuelve ÚNICAMENTE JSON válido:
{{
 "INDICADORES": [ {{"CRITERIO": "...", "VALOR": "SI o NO", "OBS": "..."}} ],
 "PUNTUACION_GLOBAL": 85,
 "NIVEL_GLOBAL": "Bueno",
 "SI": 12, "NO": 3,
 "FORTALEZAS": ["..."], "AREAS_MEJORA": ["..."],
 "RESUMEN": "..."
}}"""
    datos, _ = ia.solicitar_json(prompt, max_tokens=3000, temperature=0.1, modulo="acompanamiento_eval")
    return ia.decodificar_marcadores(datos)

def extraer_texto_pdf(archivo, max_caracteres=60000):
    if PdfReader is None: raise RuntimeError("Instala pypdf o PyPDF2.")
    archivo.seek(0)
    reader = PdfReader(archivo)
    texto = "".join((p.extract_text() or "") for p in reader.pages)
    return re.sub(r"\s+", " ", texto).strip()[:max_caracteres]

# ═══════════════════════════════════════════════════════════════════════════
# GENERADORES DE WORD
# ═══════════════════════════════════════════════════════════════════════════
def shade_cell(cell, color):
    cell._tc.get_or_add_tcPr().append(parse_xml(r'<w:shd {} w:fill="{}"/>'.format(nsdecls("w"), color)))

def set_cell_text(cell, text, bold=False, center=False, color=None, size=None):
    cell.text = ""
    run = cell.paragraphs[0].add_run(str(text))
    run.bold = bold
    if size: run.font.size = Pt(size)
    if center: cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
    if color: shade_cell(cell, color)

def add_page_number(doc):
    for section in doc.sections:
        footer = section.footer
        footer.is_linked_to_previous = False
        p = footer.paragraphs[0] if footer.paragraphs else footer.add_paragraph()
        p.text = ""
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        def _campo(parrafo, instr):
            r = parrafo.add_run()
            i = OxmlElement("w:fldChar"); i.set(qn("w:fldCharType"), "begin")
            t = OxmlElement("w:instrText"); t.set(qn("xml:space"), "preserve"); t.text = instr
            f = OxmlElement("w:fldChar"); f.set(qn("w:fldCharType"), "end")
            r._r.append(i); r._r.append(t); r._r.append(f)
            return r
        r1 = _campo(p, "PAGE"); r_sep = p.add_run(" / "); r2 = _campo(p, "NUMPAGES")
        for r in (r1, r_sep, r2):
            r.font.size = Pt(8)
            r.font.color.rgb = RGBColor(0x64, 0x74, 0x8B)

def _portada_institucional(doc, titulo, subtitulo):
    banda = doc.add_table(rows=1, cols=1)
    banda.style = "Table Grid"
    celda = banda.cell(0, 0)
    shade_cell(celda, "1E40AF")
    p = celda.paragraphs[0]; p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run("POLITÉCNICO SALESIANO ARQUIDES CALDERÓN\n")
    r.bold = True; r.font.size = Pt(14); r.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
    r2 = p.add_run("“Formando Honrados Ciudadanos y Buenos Cristianos”\n")
    r2.italic = True; r2.font.size = Pt(9); r2.font.color.rgb = RGBColor(0xDB, 0xEA, 0xFE)
    r3 = p.add_run("Coordinación Técnica Pedagógica")
    r3.bold = True; r3.font.size = Pt(10); r3.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
    doc.add_paragraph()
    t = doc.add_paragraph(); t.alignment = WD_ALIGN_PARAGRAPH.CENTER
    rt = t.add_run(titulo); rt.bold = True; rt.font.size = Pt(16); rt.font.color.rgb = RGBColor(0x1E, 0x40, 0xAF)
    s = doc.add_paragraph(); s.alignment = WD_ALIGN_PARAGRAPH.CENTER
    rs = s.add_run(subtitulo); rs.font.size = Pt(11); rs.font.color.rgb = RGBColor(0x64, 0x74, 0x8B)
    doc.add_paragraph()

def build_ficha_vacia_docx(meta, indicadores):
    doc = Document()
    doc.styles["Normal"].font.name = "Calibri"
    doc.styles["Normal"].font.size = Pt(9)
    for s in doc.sections:
        s.left_margin = s.right_margin = s.top_margin = s.bottom_margin = Inches(0.6)

    _portada_institucional(doc, "INSTRUMENTO DE ACOMPAÑAMIENTO DOCENTE", f"Especializado para: {meta['area'].upper()}")
    
    doc.add_paragraph(f"Objetivo: Acompañar y evaluar la práctica formativa.   Fecha: {meta['fecha']}")
    doc.add_heading("DATOS GENERALES", level=2)
    t = doc.add_table(rows=2, cols=2); t.style = "Table Grid"
    t.cell(0, 0).text = f"Docente Observado: {meta['docente']}"; t.cell(0, 1).text = f"Acompañante: {meta['coordinador']}"
    t.cell(1, 0).text = f"Módulo/Asignatura: {meta['modulo']}"; t.cell(1, 1).text = f"Sección: {meta['seccion']}"
    doc.add_paragraph(f"\nTipo de Acompañamiento Específico: {meta['tipo']}\n")

    doc.add_heading("MATRIZ DE INDICADORES (Diseñada por IA según contexto MINERD)", level=2)
    t_ind = doc.add_table(rows=1, cols=4); t_ind.style = "Table Grid"
    t_ind.columns[0].width = Inches(4.0); t_ind.columns[1].width = Inches(0.4); t_ind.columns[2].width = Inches(0.4); t_ind.columns[3].width = Inches(2.0)
    set_cell_text(t_ind.cell(0, 0), "Indicador de Evaluación", bold=True, color="DBEAFE")
    set_cell_text(t_ind.cell(0, 1), "SÍ", bold=True, center=True, color="DBEAFE")
    set_cell_text(t_ind.cell(0, 2), "NO", bold=True, center=True, color="DBEAFE")
    set_cell_text(t_ind.cell(0, 3), "Observaciones", bold=True, color="DBEAFE")
    
    for ind in indicadores:
        row = t_ind.add_row().cells
        row[0].text = ind; row[0].paragraphs[0].runs[0].font.size = Pt(9)
        row[1].text = ""; row[2].text = ""; row[3].text = ""

    doc.add_paragraph()
    doc.add_heading("OBSERVACIONES CUALITATIVAS Y DIFICULTADES", level=2)
    t2 = doc.add_table(rows=4, cols=1); t2.style = "Table Grid"
    for r in t2.rows: r.cells[0].text = " "
    
    doc.add_paragraph("\n")
    t3 = doc.add_table(rows=2, cols=2)
    t3.cell(0, 0).text = "_________________________"; t3.cell(0, 1).text = "_________________________"
    t3.cell(1, 0).text = "Firma del Docente"; t3.cell(1, 1).text = "Firma del Acompañante"
    for r in t3.rows:
        for c in r.cells: c.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER

    add_page_number(doc)
    buf = BytesIO(); doc.save(buf); buf.seek(0)
    return buf

def build_evidencia_docx(datos, meta, puntuacion, nivel, si, no):
    doc = Document()
    doc.styles["Normal"].font.name = "Calibri"; doc.styles["Normal"].font.size = Pt(10)
    for s in doc.sections: s.left_margin = s.right_margin = Inches(0.7)
    
    _portada_institucional(doc, "REPORTE DE EVIDENCIA DIGITALIZADA Y CORRECCIÓN IA", f"{meta['area']} · {meta['tipo']}")
    
    t = doc.add_table(rows=4, cols=4); t.style = "Table Grid"
    t.cell(0, 0).text = "Docente:"; t.cell(0, 1).text = meta["docente"]
    t.cell(0, 2).text = "Módulo:"; t.cell(0, 3).text = meta["modulo"]
    t.cell(1, 0).text = "Sección:"; t.cell(1, 1).text = meta["seccion"]
    t.cell(1, 2).text = "Coordinador:"; t.cell(1, 3).text = meta["coordinador"]
    t.cell(2, 0).text = "Fecha:"; t.cell(2, 1).text = meta["fecha"]
    t.cell(2, 2).text = "Tipo:"; t.cell(2, 3).text = meta["tipo"]
    t.cell(3, 0).text = "Puntaje:"; t.cell(3, 1).text = f"{puntuacion}/100"
    t.cell(3, 2).text = "Nivel:"; t.cell(3, 3).text = nivel
    for row in t.rows:
        row.cells[0].paragraphs[0].runs[0].bold = True
        row.cells[2].paragraphs[0].runs[0].bold = True
        shade_cell(row.cells[0], "F1F5F9"); shade_cell(row.cells[2], "F1F5F9")
    
    color_fondo = {"Excelente": "D1FAE5", "Bueno": "DBEAFE", "Aceptable con Mejoras": "FEF3C7", "No Aceptable": "FEE2E2"}.get(nivel, "F1F5F9")
    doc.add_paragraph()
    t_res = doc.add_table(rows=1, cols=1); t_res.style = "Table Grid"
    shade_cell(t_res.cell(0, 0), color_fondo)
    p = t_res.cell(0, 0).paragraphs[0]; p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.add_run(f"RESULTADO: {nivel} · {puntuacion}/100 · Indicadores Sí: {si} · No: {no}").bold = True
    
    doc.add_paragraph()
    doc.add_heading("Matriz de Evaluación Extraída", level=2)
    t_ind = doc.add_table(rows=1, cols=3); t_ind.style = "Table Grid"
    set_cell_text(t_ind.cell(0, 0), "Indicador", bold=True, color="DBEAFE")
    set_cell_text(t_ind.cell(0, 1), "Valor", bold=True, center=True, color="DBEAFE")
    set_cell_text(t_ind.cell(0, 2), "Observación", bold=True, color="DBEAFE")
    for item in datos.get("INDICADORES", []) or []:
        row = t_ind.add_row().cells
        set_cell_text(row[0], item.get('CRITERIO',''), size=9)
        v = str(item.get("VALOR", "")).upper()
        set_cell_text(row[1], v, center=True, color={"SI": "D1FAE5", "NO": "FEE2E2"}.get(v, "F1F5F9"))
        set_cell_text(row[2], item.get("OBS", ""), size=9)
    
    doc.add_paragraph()
    doc.add_heading("Análisis Cualitativo de la IA", level=2)
    doc.add_paragraph(datos.get("RESUMEN", "")).alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    doc.add_heading("Fortalezas", level=3)
    for f in datos.get("FORTALEZAS", []): doc.add_paragraph(f, style="List Bullet")
    doc.add_heading("Áreas de Mejora", level=3)
    for m in datos.get("AREAS_MEJORA", []): doc.add_paragraph(m, style="List Bullet")
    
    add_page_number(doc)
    buf = BytesIO(); doc.save(buf); buf.seek(0)
    return buf

# ═══════════════════════════════════════════════════════════════════════════
# INTERFAZ (UI)
# ═══════════════════════════════════════════════════════════════════════════
st.markdown("""
<style>
@import url('https://fonts.googleapis.com/css2?family=Inter:wght@300;400;500;600;700;800&display=swap');
html, body, [class*="css"] { font-family: 'Inter', sans-serif; background-color: #F8FAFC; color: #1E293B; }
.acomp-hero { background: linear-gradient(135deg, #0F172A 0%, #1E40AF 40%, #3B82F6 100%);
color: #fff; padding: 2rem; border-radius: 16px; margin-bottom: 1.5rem; box-shadow: 0 10px 30px rgba(30,64,175,0.2); }
.acomp-hero-title { font-size: 2.2rem; font-weight: 800; letter-spacing: -0.02em; }
.acomp-hero-sub { font-size: 1rem; opacity: .9; }
.section-title { color: #1E40AF; font-weight: 700; font-size: 1.2rem; border-bottom: 2px solid #DBEAFE; padding-bottom: 8px; margin: 1.2rem 0 1rem 0; }
.form-box { background: #FFFFFF; padding: 20px; border-radius: 12px; border: 1px solid #E2E8F0; margin-bottom: 15px; }
.score-card { background: #1E293B; color: white; border-radius: 16px; padding: 25px; text-align: center; }
.score-number { font-size: 4rem; font-weight: 900; line-height: 1; color: #10B981; }
</style>
""", unsafe_allow_html=True)

if not st.session_state.get("coordinador_autenticado", False):
    st.error("🔒 Esta página es exclusiva de Coordinación."); st.stop()

ia.panel_sidebar_ia("Acompañamiento Fígital")

st.markdown("""
<div class="acomp-hero">
    <div class="acomp-hero-title">📋 Acompañamiento Docente Fígital (IA)</div>
    <div class="acomp-hero-sub">1. Diseña una ficha con IA según el área/módulo → 2. Imprime y evalúa en aula → 3. Escanea y deja que la IA corrija</div>
</div>
""", unsafe_allow_html=True)

# ── CONFIGURACIÓN GLOBAL (CASCADA) ──
st.markdown('<div class="section-title">⚙️ Contexto del Acompañamiento</div>', unsafe_allow_html=True)
st.markdown('<div class="form-box">', unsafe_allow_html=True)

docentes_bd = listar_docentes_bd()
c1, c2, c3 = st.columns(3)

with c1:
    doc_sel = st.selectbox("👤 Docente", ["Seleccionar..."] + sorted(docentes_bd) + ["✍️ Escribir manualmente"])
    if doc_sel == "✍️ Escribir manualmente":
        docente = st.text_input("Nombre del docente")
        modulos_del_docente = []
    elif doc_sel != "Seleccionar...":
        docente = doc_sel
        modulos_del_docente = obtener_modulos_usuario_bd(docente)
    else:
        docente = ""
        modulos_del_docente = []

with c2:
    if modulos_del_docente:
        mods_unicos = sorted(list(set(m[0] for m in modulos_del_docente)))
        mod_sel = st.selectbox("📚 Módulo / Asignatura", ["Seleccionar..."] + mods_unicos + ["✍️ Otro"])
        if mod_sel == "✍️ Otro":
            modulo = st.text_input("Escribir módulo")
        elif mod_sel != "Seleccionar...":
            modulo = mod_sel
        else:
            modulo = ""
    else:
        modulo = st.text_input("📚 Módulo / Asignatura")

with c3:
    if modulos_del_docente and modulo and modulo != "✍️ Otro":
        secs_vinculadas = sorted(list(set(m[1] for m in modulos_del_docente if m[0] == modulo)))
        sec_sel = st.selectbox("🏫 Sección", secs_vinculadas + ["✍️ Otra"])
        if sec_sel == "✍️ Otra":
            seccion = st.text_input("Escribir sección")
        else:
            seccion = sec_sel
    else:
        seccion = st.text_input("🏫 Sección")

c4, c5, c6 = st.columns([1, 1.5, 1])
with c4:
    area = st.selectbox("🏷️ Área Educativa", list(AREAS_Y_TIPOS.keys()))
with c5:
    tipo = st.selectbox("🎯 Enfoque / Tipo", AREAS_Y_TIPOS[area])
with c6:
    fecha = st.date_input("📅 Fecha", value=date.today())

coordinador = st.session_state.get("coordinador_nombre", "Coordinador Pedagógico")
st.markdown('</div>', unsafe_allow_html=True)

datos_completos = bool(docente.strip() and modulo.strip())
meta = {
    "docente": docente, "modulo": modulo, "seccion": seccion, "area": area,
    "tipo": tipo, "fecha": fecha.strftime("%d/%m/%Y"), "coordinador": coordinador
}

tab_print, tab_scan, tab_repo = st.tabs(["🖨️ 1. Diseñar Ficha (IA)", "📸 2. Evaluar Escaneo", "🗂️ 3. Historial"])

# ═══ TAB 1: DISEÑAR E IMPRIMIR ═══
with tab_print:
    st.markdown('<div class="section-title">🧠 Diseñar Instrumento A Medida</div>', unsafe_allow_html=True)
    st.info("La Inteligencia Artificial redactará 15 indicadores observables basados en el área, materia y tipo de acompañamiento para alinearse al diseño curricular dominicano.")
    
    if st.button("🤖 1. Generar Indicadores Especializados", type="primary", use_container_width=True, disabled=not datos_completos):
        if not st.session_state.get("api_key_global", ""):
            st.error("🔒 Configura tu API Key en la pantalla de inicio.")
        else:
            with st.spinner(f"Diseñando rúbrica especializada para {modulo} ({tipo})..."):
                try:
                    indicadores_ia = generar_indicadores_ia(area, tipo, modulo)
                    st.session_state.ficha_buffer = build_ficha_vacia_docx(meta, indicadores_ia)
                    st.success("✅ Ficha diseñada correctamente. Lista para descargar.")
                except Exception as e:
                    ia.render_error_ia(e)

    if st.session_state.get("ficha_buffer"):
        st.download_button("⬇️ 2. Descargar Ficha en Blanco (.docx)", data=st.session_state.ficha_buffer, 
                           file_name=f"Ficha_{ia.sanear_nombre_archivo(docente)}.docx",
                           mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                           use_container_width=True)

# ═══ TAB 2: DIGITALIZAR Y EVALUAR ═══
with tab_scan:
    st.markdown('<div class="section-title">📤 Procesar Ficha Llenada</div>', unsafe_allow_html=True)
    archivo_pdf = st.file_uploader("Sube la ficha completada a mano (formato PDF)", type=["pdf"])
    
    if st.button("🧠 Evaluar Documento con IA", type="primary", use_container_width=True, disabled=not (datos_completos and archivo_pdf)):
        if not st.session_state.get("api_key_global", ""):
            st.error("🔒 Configura tu API Key.")
        else:
            with st.spinner("Leyendo y evaluando documento..."):
                try:
                    pdf_bytes = archivo_pdf.getvalue()
                    texto = extraer_texto_pdf(archivo_pdf)
                    
                    if len(texto) < 30:
                        st.warning("⚠️ Nota: El PDF es una imagen plana. La IA intentará interpretarlo por OCR.")
                    
                    res, _ = analizar_ficha_ia(texto, meta)
                    puntos = int(res.get("PUNTUACION_GLOBAL", 0))
                    nivel = str(res.get("NIVEL_GLOBAL", "Evaluado"))
                    
                    guardar_evidencia(meta["fecha"], docente, modulo, seccion, tipo, area, coordinador, puntos, nivel, archivo_pdf.name, pdf_bytes, texto)
                    
                    st.session_state.eval_actual = res
                    st.success("✅ Evaluación completada y guardada en el repositorio.")
                except Exception as e:
                    ia.render_error_ia(e)

    if st.session_state.get("eval_actual"):
        r = st.session_state.eval_actual
        c_p, c_t = st.columns([1, 2])
        with c_p:
            color = {"Excelente": "#10B981", "Bueno": "#3B82F6", "Aceptable con Mejoras": "#F59E0B", "No Aceptable": "#EF4444"}.get(r.get("NIVEL_GLOBAL"), "#475569")
            st.markdown(f"""
            <div class="score-card" style="border-bottom: 5px solid {color};">
                <div>PUNTAJE OBTENIDO</div>
                <div class="score-number" style="color:{color};">{r.get('PUNTUACION_GLOBAL', 0)}</div>
                <div style="font-size:1.2rem; margin-top:5px;">{r.get('NIVEL_GLOBAL', '')}</div>
                <div style="font-size:0.9rem; margin-top:10px;">✅ {r.get('SI', 0)} | ❌ {r.get('NO', 0)}</div>
            </div>
            """, unsafe_allow_html=True)
        with c_t:
            st.info(f"**Análisis:** {r.get('RESUMEN', '')}")
            st.write("**✨ Fortalezas:** " + ", ".join(r.get("FORTALEZAS", [])))
            st.write("**🔧 Áreas de Mejora:** " + ", ".join(r.get("AREAS_MEJORA", [])))
        
        buf_ev = build_evidencia_docx(r, meta, r.get('PUNTUACION_GLOBAL',0), r.get('NIVEL_GLOBAL',''), r.get('SI',0), r.get('NO',0))
        st.download_button("⬇️ Descargar Reporte de Evidencia (.docx)", data=buf_ev,
                           file_name=f"Evidencia_{ia.sanear_nombre_archivo(docente)}.docx",
                           mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                           use_container_width=True)

# ═══ TAB 3: EVIDENCIAS ═══
with tab_repo:
    st.markdown('<div class="section-title">🗂️ Registro de Acompañamientos</div>', unsafe_allow_html=True)
    evs = listar_evidencias()
    if not evs:
        st.info("Aún no hay evidencias almacenadas.")
    else:
        df_ev = pd.DataFrame(evs)
        st.dataframe(df_ev[["fecha", "docente", "modulo", "area", "puntuacion", "nivel"]], use_container_width=True, hide_index=True)
        for ev in evs:
            with st.expander(f"📄 {ev['fecha']} | {ev['docente']} | Puntaje: {ev['puntuacion']} ({ev['nivel']})"):
                pdf_b64, texto = obtener_evidencia(ev["id"])
                if pdf_b64:
                    st.download_button("⬇️ Descargar PDF Escaneado", data=base64.b64decode(pdf_b64),
                                       file_name=ev["archivo_nombre"], mime="application/pdf", key=f"dl_pdf_{ev['id']}")
                if st.button("🗑️ Eliminar Registro", key=f"del_{ev['id']}"):
                    eliminar_evidencia(ev["id"])
                    st.rerun()