"""
bancoitems.py — Banco de Ítems ETP PRO (Paso 21 · Nivel Dios)
12 tipos de ítems profesionales alineados a la Taxonomía de Bloom y al MINERD.
Generación asistida por IA, interfaz premium y exportación Word institucional.
"""
import copy
import re
import uuid
from datetime import datetime
from io import BytesIO
from typing import Any, Dict, List, Optional

import pandas as pd
import streamlit as st
from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import parse_xml
from docx.oxml.ns import nsdecls
from docx.shared import Inches, Pt, RGBColor

from core import ia

# ═══════════════════════════════════════════════════════════════════════════
# CATÁLOGO DE TIPOS DE ÍTEMS (12 tipos profesionales)
# ═══════════════════════════════════════════════════════════════════════════
TIPOS_ITEM = {
    "opcion_multiple_simple": {
        "nombre": "Opción Múltiple (Respuesta Única)",
        "icono": "🔘",
        "bloom": ["Recordar", "Comprender", "Aplicar", "Analizar"],
        "descripcion": "El estudiante selecciona una única respuesta correcta entre varias opciones.",
        "recomendado": "Evaluación de conceptos, definiciones y comprensión.",
    },
    "opcion_multiple_multiple": {
        "nombre": "Opción Múltiple (Respuesta Múltiple)",
        "icono": "🔲",
        "bloom": ["Analizar", "Evaluar"],
        "descripcion": "El estudiante selecciona dos o más respuestas correctas entre varias opciones.",
        "recomendado": "Evaluación de análisis y relaciones múltiples.",
    },
    "verdadero_falso": {
        "nombre": "Verdadero / Falso",
        "icono": "⚖️",
        "bloom": ["Recordar", "Comprender"],
        "descripcion": "El estudiante determina si un enunciado es verdadero o falso.",
        "recomendado": "Verificación rápida de hechos y conceptos básicos.",
    },
    "completar_espacios": {
        "nombre": "Completar Espacios en Blanco",
        "icono": "✍️",
        "bloom": ["Recordar", "Comprender", "Aplicar"],
        "descripcion": "El estudiante completa palabras o frases faltantes en un enunciado.",
        "recomendado": "Evaluación de vocabulario técnico y retención.",
    },
    "emparejamiento": {
        "nombre": "Emparejamiento / Relación de Columnas",
        "icono": "🔗",
        "bloom": ["Recordar", "Comprender", "Analizar"],
        "descripcion": "El estudiante relaciona elementos de dos columnas.",
        "recomendado": "Asociación de conceptos, términos y definiciones.",
    },
    "ordenamiento": {
        "nombre": "Ordenamiento / Secuenciamiento",
        "icono": "🔢",
        "bloom": ["Comprender", "Aplicar"],
        "descripcion": "El estudiante ordena elementos en una secuencia lógica o cronológica.",
        "recomendado": "Procesos, procedimientos y secuencias técnicas.",
    },
    "respuesta_corta": {
        "nombre": "Respuesta Corta",
        "icono": "📝",
        "bloom": ["Comprender", "Aplicar", "Analizar"],
        "descripcion": "El estudiante responde con una palabra, frase u oración breve.",
        "recomendado": "Preguntas directas sobre hechos o procedimientos.",
    },
    "desarrollo_ensayo": {
        "nombre": "Desarrollo / Ensayo",
        "icono": "📄",
        "bloom": ["Analizar", "Evaluar", "Crear"],
        "descripcion": "El estudiante desarrolla una respuesta extensa y argumentada.",
        "recomendado": "Pensamiento crítico, análisis profundo y síntesis.",
    },
    "clasificacion": {
        "nombre": "Clasificación / Agrupación",
        "icono": "🗂️",
        "bloom": ["Comprender", "Analizar"],
        "descripcion": "El estudiante clasifica elementos en categorías predefinidas.",
        "recomendado": "Taxonomías, tipos y categorías técnicas.",
    },
    "caso_practico": {
        "nombre": "Caso Práctico / Situación Problema",
        "icono": "📋",
        "bloom": ["Aplicar", "Analizar", "Evaluar", "Crear"],
        "descripcion": "El estudiante resuelve una situación real o simulada del ámbito laboral.",
        "recomendado": "Competencias prácticas y resolución de problemas.",
    },
    "calculo_procedimiento": {
        "nombre": "Cálculo / Procedimiento Numérico",
        "icono": "🧮",
        "bloom": ["Aplicar", "Analizar"],
        "descripcion": "El estudiante resuelve un cálculo o sigue un procedimiento numérico.",
        "recomendado": "Matemáticas aplicadas, finanzas y mediciones técnicas.",
    },
    "identificacion_visual": {
        "nombre": "Identificación en Imagen / Diagrama",
        "icono": "🔍",
        "bloom": ["Recordar", "Comprender", "Aplicar"],
        "descripcion": "El estudiante identifica componentes en una imagen o diagrama.",
        "recomendado": "Anatomía, circuitos, planos y componentes técnicos.",
    },
}

NIVELES_BLOOM = ["Recordar", "Comprender", "Aplicar", "Analizar", "Evaluar", "Crear"]
NIVELES_DIFICULTAD = ["Baja", "Media", "Alta"]

# ═══════════════════════════════════════════════════════════════════════════
# ESTILOS PREMIUM
# ═══════════════════════════════════════════════════════════════════════════
st.markdown("""
<style>
@import url('https://fonts.googleapis.com/css2?family=Inter:wght@300;400;500;600;700;800;900&display=swap');

html, body, [class*="css"] {
    font-family: 'Inter', sans-serif;
    background-color: #F0F4F8;
    color: #1E293B;
}

/* Hero premium */
.banco-hero {
    background: linear-gradient(135deg, #0F172A 0%, #1E40AF 50%, #3B82F6 100%);
    color: #fff;
    padding: 2.2rem;
    border-radius: 20px;
    margin-bottom: 1.5rem;
    box-shadow: 0 25px 50px rgba(15, 23, 42, 0.25);
    position: relative;
    overflow: hidden;
}

.banco-hero::before {
    content: '';
    position: absolute;
    top: -50%;
    right: -20%;
    width: 60%;
    height: 200%;
    background: radial-gradient(circle, rgba(255,255,255,0.08) 0%, transparent 70%);
}

.banco-hero-title {
    font-size: 2.4rem;
    font-weight: 900;
    letter-spacing: -0.03em;
    margin-bottom: 0.4rem;
}

.banco-hero-sub {
    font-size: 1.05rem;
    opacity: 0.9;
    line-height: 1.5;
}

.banco-hero-badge {
    display: inline-block;
    background: rgba(255,255,255,0.15);
    border: 1px solid rgba(255,255,255,0.25);
    border-radius: 8px;
    padding: 4px 12px;
    font-size: 0.8rem;
    font-weight: 600;
    margin-top: 0.8rem;
    margin-right: 8px;
}

/* Cards de tipos de ítems */
.tipo-card {
    background: #fff;
    border: 2px solid #E2E8F0;
    border-radius: 14px;
    padding: 1.2rem;
    transition: all 0.25s ease;
    cursor: pointer;
    height: 100%;
    box-shadow: 0 2px 8px rgba(0,0,0,0.04);
}

.tipo-card:hover {
    border-color: #3B82F6;
    transform: translateY(-3px);
    box-shadow: 0 8px 20px rgba(59,130,246,0.15);
}

.tipo-card-selected {
    border-color: #2563EB;
    background: #EFF6FF;
    box-shadow: 0 4px 15px rgba(37,99,235,0.2);
}

.tipo-icono { font-size: 2rem; margin-bottom: 0.5rem; }
.tipo-nombre { font-weight: 700; font-size: 0.95rem; color: #0F172A; margin-bottom: 0.3rem; }
.tipo-desc { font-size: 0.78rem; color: #64748B; line-height: 1.4; }
.tipo-bloom { font-size: 0.7rem; color: #3B82F6; font-weight: 600; margin-top: 0.4rem; }

/* Métricas */
.stat-card {
    background: #fff;
    border-radius: 12px;
    padding: 1rem;
    text-align: center;
    border: 1px solid #E2E8F0;
    box-shadow: 0 2px 6px rgba(0,0,0,0.04);
}

.stat-value { font-size: 1.8rem; font-weight: 800; color: #1E40AF; }
.stat-label { font-size: 0.75rem; font-weight: 700; color: #94A3B8; text-transform: uppercase; letter-spacing: 0.05em; }

/* Secciones */
.section-divider {
    border: none;
    height: 3px;
    background: linear-gradient(90deg, #3B82F6, #8B5CF6, #EC4899);
    border-radius: 2px;
    margin: 1.5rem 0;
}

.form-container {
    background: #fff;
    border: 1px solid #E2E8F0;
    border-radius: 14px;
    padding: 1.5rem;
    box-shadow: 0 4px 12px rgba(0,0,0,0.05);
}
</style>
""", unsafe_allow_html=True)

# ═══════════════════════════════════════════════════════════════════════════
# ESTADO DE SESIÓN
# ═══════════════════════════════════════════════════════════════════════════
def init_estado():
    if "banco_items" not in st.session_state:
        st.session_state.banco_items = []
    if "tipo_seleccionado" not in st.session_state:
        st.session_state.tipo_seleccionado = None
    if "modo_edicion" not in st.session_state:
        st.session_state.modo_edicion = None  # None = crear, index = editar
    if "generacion_ia" not in st.session_state:
        st.session_state.generacion_ia = None

init_estado()

# ═══════════════════════════════════════════════════════════════════════════
# HERO
# ═══════════════════════════════════════════════════════════════════════════
ia.panel_sidebar_ia("Banco de Ítems PRO")

st.markdown("""
<div class="banco-hero">
    <div class="banco-hero-title">📝 Banco de Ítems ETP PRO</div>
    <div class="banco-hero-sub">
        12 tipos de ítems profesionales · Taxonomía de Bloom · Generación con IA ·
        Exportación Word institucional
    </div>
    <div>
        <span class="banco-hero-badge">🎯 12 Tipos de Ítem</span>
        <span class="banco-hero-badge">🧠 Taxonomía de Bloom</span>
        <span class="banco-hero-badge">🤖 Asistido por IA</span>
        <span class="banco-hero-badge">📄 Word Profesional</span>
    </div>
</div>
""", unsafe_allow_html=True)

# ═══════════════════════════════════════════════════════════════════════════
# MÉTRICAS DEL BANCO
# ═══════════════════════════════════════════════════════════════════════════
total_items = len(st.session_state.banco_items)
tipos_usados = len(set(item.get("tipo") for item in st.session_state.banco_items))
niveles_bloom = len(set(item.get("nivel_bloom") for item in st.session_state.banco_items))

col_s1, col_s2, col_s3, col_s4 = st.columns(4)
with col_s1:
    st.markdown(f'<div class="stat-card"><div class="stat-value">{total_items}</div><div class="stat-label">Ítems Totales</div></div>', unsafe_allow_html=True)
with col_s2:
    st.markdown(f'<div class="stat-card"><div class="stat-value">{tipos_usados}</div><div class="stat-label">Tipos de Ítem</div></div>', unsafe_allow_html=True)
with col_s3:
    st.markdown(f'<div class="stat-card"><div class="stat-value">{niveles_bloom}</div><div class="stat-label">Niveles Bloom</div></div>', unsafe_allow_html=True)
with col_s4:
    puntos_totales = sum(item.get("puntos", 0) for item in st.session_state.banco_items)
    st.markdown(f'<div class="stat-card"><div class="stat-value">{puntos_totales}</div><div class="stat-label">Puntos Totales</div></div>', unsafe_allow_html=True)

st.markdown('<hr class="section-divider">', unsafe_allow_html=True)

# ═══════════════════════════════════════════════════════════════════════════
# TAB PRINCIPAL
# ═══════════════════════════════════════════════════════════════════════════
tab_crear, tab_banco, tab_generar_ia, tab_exportar = st.tabs([
    "➕ Crear Ítem",
    "📚 Banco de Ítems",
    "🤖 Generar con IA",
    "📄 Exportar Word"
])

# ═══════════════════════════════════════════════════════════════════════════
# TAB 1: CREAR ÍTEM
# ═══════════════════════════════════════════════════════════════════════════
with tab_crear:
    st.markdown("#### 🎯 Selecciona el Tipo de Ítem")
    
    cols = st.columns(4)
    for idx, (clave, info) in enumerate(TIPOS_ITEM.items()):
        col = cols[idx % 4]
        with col:
            selected = st.session_state.tipo_seleccionado == clave
            css_class = "tipo-card tipo-card-selected" if selected else "tipo-card"
            st.markdown(f"""
            <div class="{css_class}">
                <div class="tipo-icono">{info['icono']}</div>
                <div class="tipo-nombre">{info['nombre']}</div>
                <div class="tipo-desc">{info['descripcion']}</div>
                <div class="tipo-bloom">🧠 {', '.join(info['bloom'][:3])}</div>
            </div>
            """, unsafe_allow_html=True)
            
            if st.button("Seleccionar", key=f"sel_{clave}", use_container_width=True,
                         type="primary" if selected else "secondary"):
                st.session_state.tipo_seleccionado = clave
                st.session_state.modo_edicion = None
                st.rerun()
    
    # Formulario dinámico según tipo
    if st.session_state.tipo_seleccionado:
        st.markdown('<hr class="section-divider">', unsafe_allow_html=True)
        tipo_info = TIPOS_ITEM[st.session_state.tipo_seleccionado]
        st.markdown(f"#### {tipo_info['icono']} Configurar: {tipo_info['nombre']}")
        st.info(f"💡 {tipo_info['recomendado']}")
        
        with st.form("form_item", clear_on_submit=False):
            col_f1, col_f2, col_f3 = st.columns(3)
            with col_f1:
                nivel_bloom = st.selectbox(
                    "Nivel Bloom",
                    tipo_info["bloom"],
                    help="Nivel cognitivo según Taxonomía de Bloom"
                )
            with col_f2:
                dificultad = st.selectbox("Dificultad", NIVELES_DIFICULTAD)
            with col_f3:
                puntos = st.number_input("Puntos", min_value=1, max_value=50, value=5)
            
            enunciado = st.text_area(
                "Enunciado / Pregunta",
                height=100,
                placeholder="Escribe el enunciado del ítem aquí..."
            )
            
            # Campos específicos según tipo
            st.markdown("##### 📋 Campos Específicos")
            
            tipo = st.session_state.tipo_seleccionado
            datos_extra = {}
            
            if tipo in ["opcion_multiple_simple", "opcion_multiple_multiple"]:
                opciones = []
                col_o1, col_o2 = st.columns(2)
                with col_o1:
                    opciones.append(st.text_input("Opción A", key="opt_a"))
                    opciones.append(st.text_input("Opción B", key="opt_b"))
                    opciones.append(st.text_input("Opción C", key="opt_c"))
                with col_o2:
                    opciones.append(st.text_input("Opción D", key="opt_d"))
                    opciones.append(st.text_input("Opción E (opcional)", key="opt_e"))
                
                if tipo == "opcion_multiple_simple":
                    respuesta_correcta = st.selectbox("Respuesta correcta", ["A", "B", "C", "D", "E"])
                else:
                    respuesta_correcta = st.multiselect("Respuestas correctas", ["A", "B", "C", "D", "E"])
                
                datos_extra["opciones"] = [o for o in opciones if o.strip()]
                datos_extra["respuesta_correcta"] = respuesta_correcta
                justificacion = st.text_area("Justificación de la respuesta", height=70)
                datos_extra["justificacion"] = justificacion
            
            elif tipo == "verdadero_falso":
                respuesta_correcta = st.radio("Respuesta correcta", ["Verdadero", "Falso"])
                datos_extra["respuesta_correcta"] = respuesta_correcta
                justificacion = st.text_area("Justificación", height=70)
                datos_extra["justificacion"] = justificacion
            
            elif tipo == "completar_espacios":
                respuesta_correcta = st.text_input("Palabra(s) faltante(s)", placeholder="Ej: fotosíntesis, mitocondria")
                datos_extra["respuesta_correcta"] = respuesta_correcta
                pista = st.text_input("Pista opcional")
                datos_extra["pista"] = pista
            
            elif tipo == "emparejamiento":
                st.write("Configura los pares (Columna A ↔ Columna B):")
                pares = []
                for i in range(1, 6):
                    col_e1, col_e2 = st.columns(2)
                    with col_e1:
                        elem_a = st.text_input(f"Columna A {i}", key=f"emp_a_{i}")
                    with col_e2:
                        elem_b = st.text_input(f"Columna B {i}", key=f"emp_b_{i}")
                    if elem_a.strip() and elem_b.strip():
                        pares.append({"izquierda": elem_a, "derecha": elem_b})
                datos_extra["pares"] = pares
            
            elif tipo == "ordenamiento":
                elementos = st.text_area(
                    "Elementos en orden correcto (uno por línea)",
                    height=100,
                    placeholder="Paso 1\nPaso 2\nPaso 3..."
                )
                datos_extra["elementos"] = [e.strip() for e in elementos.split("\n") if e.strip()]
            
            elif tipo == "respuesta_corta":
                respuesta_correcta = st.text_input("Respuesta esperada")
                datos_extra["respuesta_correcta"] = respuesta_correcta
                variantes = st.text_input("Variantes aceptadas (separadas por coma)")
                datos_extra["variantes"] = [v.strip() for v in variantes.split(",") if v.strip()]
            
            elif tipo == "desarrollo_ensayo":
                criterios_eval = st.text_area(
                    "Criterios de evaluación (uno por línea)",
                    height=100,
                    placeholder="Menciona al menos 3 conceptos clave\nUtiliza vocabulario técnico\nArgumenta con ejemplos"
                )
                datos_extra["criterios"] = [c.strip() for c in criterios_eval.split("\n") if c.strip()]
                extension = st.selectbox("Extensión esperada", ["1 párrafo", "2-3 párrafos", "1 página", "Libre"])
                datos_extra["extension"] = extension
            
            elif tipo == "clasificacion":
                categorias = st.text_input("Categorías (separadas por coma)", placeholder="Tipo A, Tipo B, Tipo C")
                datos_extra["categorias"] = [c.strip() for c in categorias.split(",") if c.strip()]
                elementos = st.text_area(
                    "Elementos a clasificar (formato: elemento -> categoría, uno por línea)",
                    height=100,
                    placeholder="Fotosíntesis -> Biológico\nCombustión -> Químico"
                )
                datos_extra["elementos"] = [e.strip() for e in elementos.split("\n") if e.strip()]
            
            elif tipo == "caso_practico":
                caso = st.text_area("Descripción del caso / situación", height=120)
                datos_extra["caso"] = caso
                preguntas = st.text_area("Preguntas sobre el caso (una por línea)", height=100)
                datos_extra["preguntas"] = [p.strip() for p in preguntas.split("\n") if p.strip()]
            
            elif tipo == "calculo_procedimiento":
                procedimiento = st.text_area("Procedimiento / datos del problema", height=100)
                datos_extra["procedimiento"] = procedimiento
                respuesta_correcta = st.text_input("Respuesta / resultado esperado")
                datos_extra["respuesta_correcta"] = respuesta_correcta
                unidades = st.text_input("Unidades (ej: %, metros, litros)")
                datos_extra["unidades"] = unidades
            
            elif tipo == "identificacion_visual":
                descripcion_imagen = st.text_area(
                    "Descripción de la imagen / diagrama",
                    height=80,
                    placeholder="Describe qué contiene la imagen y qué se debe identificar..."
                )
                datos_extra["descripcion_imagen"] = descripcion_imagen
                componentes = st.text_area(
                    "Componentes a identificar (uno por línea)",
                    height=80,
                    placeholder="Componente 1\nComponente 2\nComponente 3"
                )
                datos_extra["componentes"] = [c.strip() for c in componentes.split("\n") if c.strip()]
            
            st.markdown('<hr class="section-divider">', unsafe_allow_html=True)
            
            col_btn1, col_btn2 = st.columns(2)
            with col_btn1:
                submit_item = st.form_submit_button(
                    "💾 Guardar Ítem en el Banco",
                    type="primary",
                    use_container_width=True
                )
            with col_btn2:
                limpiar_form = st.form_submit_button(
                    "🗑️ Limpiar Formulario",
                    use_container_width=True
                )
            
            if submit_item:
                if not enunciado.strip():
                    st.warning("⚠️ El enunciado es obligatorio.")
                else:
                    nuevo_item = {
                        "id": str(uuid.uuid4())[:8],
                        "tipo": tipo,
                        "enunciado": enunciado,
                        "nivel_bloom": nivel_bloom,
                        "dificultad": dificultad,
                        "puntos": puntos,
                        "datos": datos_extra,
                        "creado": datetime.now().strftime("%Y-%m-%d %H:%M"),
                    }
                    
                    if st.session_state.modo_edicion is not None:
                        st.session_state.banco_items[st.session_state.modo_edicion] = nuevo_item
                        st.session_state.modo_edicion = None
                        st.success("✅ Ítem actualizado correctamente.")
                    else:
                        st.session_state.banco_items.append(nuevo_item)
                        st.success("✅ Ítem guardado en el banco.")
                    
                    st.session_state.tipo_seleccionado = None
                    st.rerun()
            
            if limpiar_form:
                st.session_state.tipo_seleccionado = None
                st.session_state.modo_edicion = None
                st.rerun()

# ═══════════════════════════════════════════════════════════════════════════
# TAB 2: BANCO DE ÍTEMS
# ═══════════════════════════════════════════════════════════════════════════
with tab_banco:
    if not st.session_state.banco_items:
        st.info("📭 El banco está vacío. Ve a la pestaña **'➕ Crear Ítem'** para agregar ítems.")
    else:
        st.markdown(f"#### 📚 Banco de Ítems ({total_items} ítems)")
        
        # Filtros
        col_f1, col_f2, col_f3 = st.columns(3)
        with col_f1:
            filtro_tipo = st.multiselect(
                "Filtrar por tipo",
                list(TIPOS_ITEM.keys()),
                format_func=lambda x: TIPOS_ITEM[x]["nombre"]
            )
        with col_f2:
            filtro_bloom = st.multiselect("Filtrar por Bloom", NIVELES_BLOOM)
        with col_f3:
            filtro_dificultad = st.multiselect("Filtrar por dificultad", NIVELES_DIFICULTAD)
        
        # Aplicar filtros
        items_filtrados = st.session_state.banco_items.copy()
        if filtro_tipo:
            items_filtrados = [i for i in items_filtrados if i["tipo"] in filtro_tipo]
        if filtro_bloom:
            items_filtrados = [i for i in items_filtrados if i["nivel_bloom"] in filtro_bloom]
        if filtro_dificultad:
            items_filtrados = [i for i in items_filtrados if i["dificultad"] in filtro_dificultad]
        
        st.write(f"Mostrando {len(items_filtrados)} de {total_items} ítems")
        
        for idx, item in enumerate(items_filtrados):
            tipo_info = TIPOS_ITEM.get(item["tipo"], {"icono": "📄", "nombre": "Desconocido"})
            with st.expander(f"{tipo_info['icono']} {item['enunciado'][:60]}... | 🧠 {item['nivel_bloom']} | ⚡ {item['dificultad']} | {item['puntos']} pts"):
                col_a1, col_a2, col_a3 = st.columns([3, 1, 1])
                with col_a1:
                    st.markdown(f"**Enunciado:** {item['enunciado']}")
                    st.markdown(f"**Tipo:** {tipo_info['nombre']}")
                    st.markdown(f"**Creado:** {item['creado']}")
                with col_a2:
                    if st.button("✏️ Editar", key=f"edit_{item['id']}"):
                        st.session_state.tipo_seleccionado = item["tipo"]
                        st.session_state.modo_edicion = idx
                        st.rerun()
                with col_a3:
                    if st.button("🗑️", key=f"del_{item['id']}"):
                        st.session_state.banco_items.pop(idx)
                        st.rerun()

# ═══════════════════════════════════════════════════════════════════════════
# TAB 3: GENERAR CON IA
# ═══════════════════════════════════════════════════════════════════════════
with tab_generar_ia:
    st.markdown("#### 🤖 Generación de Ítems con IA")
    st.info("La IA generará ítems profesionales basados en el tema, módulo y nivel cognitivo que indiques.")
    
    with st.form("form_gen_ia"):
        col_g1, col_g2 = st.columns(2)
        with col_g1:
            tema_ia = st.text_input("Tema / Contenido", placeholder="Ej: Impuestos al consumo")
            modulo_ia = st.text_input("Módulo Formativo", placeholder="Ej: MF_358_3 Impuestos al consumo")
            ra_ia = st.text_area("Resultado de Aprendizaje (RA)", height=70)
        
        with col_g2:
            tipos_ia = st.multiselect(
                "Tipos de ítems a generar",
                list(TIPOS_ITEM.keys()),
                default=["opcion_multiple_simple", "verdadero_falso", "completar_espacios"],
                format_func=lambda x: TIPOS_ITEM[x]["nombre"]
            )
            bloom_ia = st.selectbox("Nivel Bloom objetivo", NIVELES_BLOOM)
            cantidad_ia = st.slider("Cantidad de ítems por tipo", 1, 10, 3)
            dificultad_ia = st.selectbox("Dificultad", NIVELES_DIFICULTAD, index=1)
        
        st.markdown('<hr class="section-divider">', unsafe_allow_html=True)
        
        col_gen1, col_gen2 = st.columns([1, 3])
        with col_gen1:
            btn_generar = st.form_submit_button("⚡ Generar con IA", type="primary", use_container_width=True)
        with col_gen2:
            st.caption("La generación puede tardar 15-45 segundos según la cantidad de ítems.")
    
    if btn_generar:
        if not tema_ia.strip():
            st.warning("⚠️ Indica al menos el tema.")
        else:
            with st.spinner("🧠 Generando ítems profesionales con IA..."):
                try:
                    tipos_texto = "\n".join([
                        f"- {TIPOS_ITEM[t]['nombre']}: genera {cantidad_ia} ítems"
                        for t in tipos_ia
                    ])
                    
                    prompt = f"""Actúa como un Experto en Evaluación Educativa ETP del MINERD.
Genera un banco de ítems profesional con las siguientes características:

TEMA: {tema_ia}
MÓDULO: {modulo_ia if modulo_ia else 'No especificado'}
RESULTADO DE APRENDIZAJE: {ra_ia if ra_ia else 'No especificado'}
NIVEL BLOOM: {bloom_ia}
DIFICULTAD: {dificultad_ia}

TIPOS Y CANTIDADES:
{tipos_texto}

REGLAS:
- Cada ítem debe tener enunciado claro y sin ambigüedad.
- Las opciones incorrectas (distractores) deben ser plausibles.
- El nivel de dificultad debe ser {dificultad}.
- Los ítems deben ser relevantes para la Educación Técnico Profesional.
- Incluye justificación pedagógica breve para cada respuesta correcta.

Devuelve ÚNICAMENTE JSON válido con este formato exacto:
{{
  "items": [
    {{
      "tipo": "tipo_item (usa las claves: opcion_multiple_simple, opcion_multiple_multiple, verdadero_falso, completar_espacios, emparejamiento, ordenamiento, respuesta_corta, desarrollo_ensayo, clasificacion, caso_practico, calculo_procedimiento, identificacion_visual)",
      "enunciado": "...",
      "nivel_bloom": "{bloom_ia}",
      "dificultad": "{dificultad_ia}",
      "puntos": 5,
      "datos": {{
        "opciones": ["A", "B", "C", "D"],
        "respuesta_correcta": "A",
        "justificacion": "..."
      }}
    }}
  ]
}}
"""
                    resultado, flags = ia.solicitar_json(
                        prompt,
                        max_tokens=16384,
                        temperature=0.3,
                        modulo="bancoitems_ia"
                    )
                    
                    if isinstance(resultado, dict) and "items" in resultado:
                        nuevos_items = []
                        for item_ia in resultado["items"]:
                            nuevo = {
                                "id": str(uuid.uuid4())[:8],
                                "tipo": item_ia.get("tipo", "respuesta_corta"),
                                "enunciado": item_ia.get("enunciado", ""),
                                "nivel_bloom": item_ia.get("nivel_bloom", bloom_ia),
                                "dificultad": item_ia.get("dificultad", dificultad_ia),
                                "puntos": item_ia.get("puntos", 5),
                                "datos": item_ia.get("datos", {}),
                                "creado": datetime.now().strftime("%Y-%m-%d %H:%M"),
                                "origen": "IA"
                            }
                            nuevos_items.append(nuevo)
                        
                        st.session_state.banco_items.extend(nuevos_items)
                        st.session_state.generacion_ia = {
                            "cantidad": len(nuevos_items),
                            "tema": tema_ia,
                            "fecha": datetime.now().strftime("%Y-%m-%d %H:%M")
                        }
                        st.success(f"✅ Se generaron {len(nuevos_items)} ítems y se agregaron al banco.")
                    else:
                        st.warning("⚠️ La IA no devolvió el formato esperado. Intenta de nuevo.")
                
                except Exception as e:
                    ia.render_error_ia(e)

# ═══════════════════════════════════════════════════════════════════════════
# TAB 4: EXPORTAR WORD
# ═══════════════════════════════════════════════════════════════════════════
with tab_exportar:
    st.markdown("#### 📄 Exportar Banco a Word Institucional")
    
    if not st.session_state.banco_items:
        st.info("📭 No hay ítems para exportar. Primero crea o genera ítems.")
    else:
        with st.form("form_export"):
            col_e1, col_e2 = st.columns(2)
            with col_e1:
                centro = st.text_input("Centro Educativo", value="Politécnico Salesiano Arquides Calderón")
                docente = st.text_input("Docente", value="Ing. Bernardo Antonio Hernández Batista")
                modulo = st.text_input("Módulo Formativo", placeholder="Ej: MF_358_3")
            with col_e2:
                titulo_doc = st.text_input("Título del Documento", value="Banco de Ítems — Evaluación Diversificada ETP")
                incluir_respuestas = st.checkbox("Incluir hoja de respuestas (solo docente)", value=True)
                incluir_justificaciones = st.checkbox("Incluir justificaciones pedagógicas", value=True)
            
            st.markdown('<hr class="section-divider">', unsafe_allow_html=True)
            
            col_btn_e1, col_btn_e2 = st.columns(2)
            with col_btn_e1:
                btn_export = st.form_submit_button("📥 Generar y Descargar Word", type="primary", use_container_width=True)
            with col_btn_e2:
                st.caption(f"Se exportarán {total_items} ítems en formato profesional.")
        
        if btn_export:
            with st.spinner("📄 Generando documento Word profesional..."):
                doc = Document()
                
                # Configuración de estilos
                doc.styles["Normal"].font.name = "Calibri"
                doc.styles["Normal"].font.size = Pt(11)
                
                for section in doc.sections:
                    section.left_margin = Inches(0.75)
                    section.right_margin = Inches(0.75)
                
                # ─── ENCABEZADO INSTITUCIONAL ───
                p_enc = doc.add_paragraph()
                p_enc.alignment = WD_ALIGN_PARAGRAPH.CENTER
                p_enc.add_run(centro).bold = True
                
                p_titulo = doc.add_paragraph()
                p_titulo.alignment = WD_ALIGN_PARAGRAPH.CENTER
                run_t = p_titulo.add_run(titulo_doc)
                run_t.bold = True
                run_t.font.size = Pt(14)
                run_t.font.color.rgb = RGBColor(0x1E, 0x40, 0xAF)
                
                p_meta = doc.add_paragraph()
                p_meta.alignment = WD_ALIGN_PARAGRAPH.CENTER
                p_meta.add_run(f"Docente: {docente} | Módulo: {modulo} | Fecha: {datetime.now().strftime('%d/%m/%Y')}")
                
                doc.add_paragraph("_" * 70)
                
                # ─── RESUMEN DEL BANCO ───
                doc.add_heading("Resumen del Banco de Ítems", level=1)
                
                tabla_resumen = doc.add_table(rows=1, cols=3)
                tabla_resumen.style = "Table Grid"
                hdr = tabla_resumen.rows[0].cells
                for i, h in enumerate(["Tipo de Ítem", "Cantidad", "Puntos"]):
                    hdr[i].text = h
                    hdr[i].paragraphs[0].runs[0].bold = True
                
                resumen_tipos = {}
                for item in st.session_state.banco_items:
                    tipo = item["tipo"]
                    if tipo not in resumen_tipos:
                        resumen_tipos[tipo] = {"cantidad": 0, "puntos": 0}
                    resumen_tipos[tipo]["cantidad"] += 1
                    resumen_tipos[tipo]["puntos"] += item.get("puntos", 0)
                
                for tipo, datos in resumen_tipos.items():
                    row = tabla_resumen.add_row().cells
                    row[0].text = TIPOS_ITEM.get(tipo, {}).get("nombre", tipo)
                    row[1].text = str(datos["cantidad"])
                    row[2].text = str(datos["puntos"])
                
                row_total = tabla_resumen.add_row().cells
                row_total[0].text = "TOTAL"
                row_total[0].paragraphs[0].runs[0].bold = True
                row_total[1].text = str(total_items)
                row_total[2].text = str(puntos_totales)
                
                doc.add_paragraph()
                
                # ─── ÍTEMS ───
                doc.add_page_break()
                doc.add_heading("Sección de Ítems", level=1)
                
                for idx, item in enumerate(st.session_state.banco_items, 1):
                    tipo_info = TIPOS_ITEM.get(item["tipo"], {"icono": "📄", "nombre": "Desconocido"})
                    
                    # Encabezado del ítem
                    p_item = doc.add_paragraph()
                    run_num = p_item.add_run(f"Ítem {idx}. ")
                    run_num.bold = True
                    p_item.add_run(f"[{tipo_info['nombre']}] ")
                    p_item.add_run(f"({item['puntos']} pts | Bloom: {item['nivel_bloom']} | {item['dificultad']})")
                    
                    # Enunciado
                    p_enun = doc.add_paragraph()
                    p_enun.add_run(item["enunciado"])
                    
                    # Datos específicos
                    datos = item.get("datos", {})
                    
                    if item["tipo"] in ["opcion_multiple_simple", "opcion_multiple_multiple"]:
                        opciones = datos.get("opciones", [])
                        letras = ["A", "B", "C", "D", "E"]
                        for j, op in enumerate(opciones):
                            if op.strip():
                                p_op = doc.add_paragraph()
                                p_op.paragraph_format.left_indent = Inches(0.5)
                                p_op.add_run(f"{letras[j]}) {op}")
                    
                    elif item["tipo"] == "completar_espacios":
                        p_blank = doc.add_paragraph()
                        p_blank.paragraph_format.left_indent = Inches(0.5)
                        p_blank.add_run("Respuesta: _________________________________")
                    
                    elif item["tipo"] == "emparejamiento":
                        pares = datos.get("pares", [])
                        if pares:
                            t_emp = doc.add_table(rows=1, cols=2)
                            t_emp.style = "Table Grid"
                            hdr_emp = t_emp.rows[0].cells
                            hdr_emp[0].text = "Columna A"
                            hdr_emp[1].text = "Columna B"
                            for par in pares:
                                row = t_emp.add_row().cells
                                row[0].text = par.get("izquierda", "")
                                row[1].text = par.get("derecha", "")
                    
                    elif item["tipo"] == "ordenamiento":
                        elementos = datos.get("elementos", [])
                        for j, elem in enumerate(elementos, 1):
                            p_ord = doc.add_paragraph()
                            p_ord.paragraph_format.left_indent = Inches(0.5)
                            p_ord.add_run(f"( ) {elem}")
                    
                    elif item["tipo"] == "desarrollo_ensayo":
                        p_esp = doc.add_paragraph()
                        p_esp.paragraph_format.left_indent = Inches(0.5)
                        p_esp.add_run("Espacio para respuesta:")
                        for _ in range(3):
                            doc.add_paragraph("_" * 60)
                    
                    elif item["tipo"] == "caso_practico":
                        caso = datos.get("caso", "")
                        if caso:
                            p_caso = doc.add_paragraph()
                            p_caso.add_run("Caso: ").bold = True
                            p_caso.add_run(caso)
                        preguntas = datos.get("preguntas", [])
                        for j, preg in enumerate(preguntas, 1):
                            p_preg = doc.add_paragraph()
                            p_preg.paragraph_format.left_indent = Inches(0.5)
                            p_preg.add_run(f"{j}. {preg}")
                    
                    doc.add_paragraph()  # Espacio entre ítems
                
                # ─── HOJA DE RESPUESTAS ───
                if incluir_respuestas:
                    doc.add_page_break()
                    doc.add_heading("🔑 Hoja de Respuestas (Solo Docente)", level=1)
                    
                    for idx, item in enumerate(st.session_state.banco_items, 1):
                        datos = item.get("datos", {})
                        p_resp = doc.add_paragraph()
                        p_resp.add_run(f"Ítem {idx}: ").bold = True
                        
                        respuesta = datos.get("respuesta_correcta", datos.get("respuesta", "N/A"))
                        if isinstance(respuesta, list):
                            respuesta = ", ".join(respuesta)
                        p_resp.add_run(str(respuesta))
                        
                        if incluir_justificaciones and datos.get("justificacion"):
                            p_just = doc.add_paragraph()
                            p_just.paragraph_format.left_indent = Inches(0.3)
                            p_just.add_run("Justificación: ").italic = True
                            p_just.add_run(datos["justificacion"]).italic = True
                
                # Guardar
                buffer = BytesIO()
                doc.save(buffer)
                buffer.seek(0)
                
                st.download_button(
                    label="📥 Descargar Banco de Ítems (.docx)",
                    data=buffer,
                    file_name=f"Banco_Items_{ia.sanear_nombre_archivo(modulo or 'ETP')}_{datetime.now().strftime('%Y%m%d')}.docx",
                    mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                    type="primary",
                    use_container_width=True
                )
                st.success("✅ Documento Word generado. Haz clic en el botón para descargar.")