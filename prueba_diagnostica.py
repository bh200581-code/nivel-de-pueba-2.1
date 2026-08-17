"""
prueba_diagnostica.py — Centro de Pruebas Diagnósticas Fígital
═══════════════════════════════════════════════════════════════════════════
• Generación de pruebas con IA.
• Modo Digital: Link interactivo (celulares/PC).
• Modo Físico: Generación de Word imprimible.
• NUEVO: Integración directa con core/drive.py (OAuth) para subir a Drive.
• Corrector IA: Lee el examen escaneado en PDF y lo corrige automáticamente.
"""
import datetime
import json
import random
import socket
import sqlite3
import string
import re
from io import BytesIO
from urllib.parse import quote

import pandas as pd
import streamlit as st
from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import parse_xml
from docx.oxml.ns import nsdecls
from docx.shared import Inches, Pt, RGBColor

try:
    from pypdf import PdfReader
except Exception:
    try:
        from PyPDF2 import PdfReader
    except Exception:
        PdfReader = None

from core import ia
from core import drive as drive_api  # INTEGRACIÓN CON DRIVE CENTRALIZADO

# ═══════════════════════════════════════════════════════════════════════════
# BASE DE DATOS
# ═══════════════════════════════════════════════════════════════════════════
DB_NAME = "gestion_etp.db"

def _conn():
    return sqlite3.connect(DB_NAME, check_same_thread=False)

def init_db():
    conn = _conn()
    cur = conn.cursor()
    cur.execute('''CREATE TABLE IF NOT EXISTS pruebas_diagnosticas (
        id INTEGER PRIMARY KEY AUTOINCREMENT,
        codigo TEXT UNIQUE,
        titulo TEXT,
        area TEXT,
        modulo TEXT,
        docente TEXT,
        grado TEXT,
        seccion TEXT,
        tipo TEXT,
        num_preguntas INTEGER,
        preguntas_json TEXT,
        estado TEXT DEFAULT 'activa',
        fecha_creacion TEXT
    )''')
    cur.execute('''CREATE TABLE IF NOT EXISTS respuestas_diagnosticas (
        id INTEGER PRIMARY KEY AUTOINCREMENT,
        prueba_codigo TEXT,
        estudiante TEXT,
        matricula TEXT,
        respuestas_json TEXT,
        correctas INTEGER,
        total INTEGER,
        puntaje REAL,
        nivel TEXT,
        valoracion TEXT,
        fecha TEXT
    )''')
    conn.commit()
    conn.close()

init_db()

# ── CRUD pruebas ──
def guardar_prueba(datos):
    conn = _conn()
    conn.execute('''INSERT INTO pruebas_diagnosticas
        (codigo, titulo, area, modulo, docente, grado, seccion, tipo,
         num_preguntas, preguntas_json, estado, fecha_creacion)
        VALUES (?,?,?,?,?,?,?,?,?,?,?,?)''',
        (datos["codigo"], datos["titulo"], datos["area"], datos["modulo"],
         datos["docente"], datos["grado"], datos["seccion"], datos["tipo"],
         datos["num_preguntas"], json.dumps(datos["preguntas"], ensure_ascii=False),
         "activa", datetime.date.today().strftime("%Y-%m-%d")))
    conn.commit()
    conn.close()

def obtener_prueba_por_codigo(codigo):
    conn = _conn()
    cur = conn.cursor()
    cur.execute("SELECT * FROM pruebas_diagnosticas WHERE codigo=?", (codigo,))
    row = cur.fetchone()
    conn.close()
    if not row: return None
    cols = ["id", "codigo", "titulo", "area", "modulo", "docente", "grado",
            "seccion", "tipo", "num_preguntas", "preguntas_json", "estado", "fecha_creacion"]
    d = dict(zip(cols, row))
    d["preguntas"] = json.loads(d["preguntas_json"])
    return d

def obtener_mis_pruebas(docente=None):
    conn = _conn()
    cur = conn.cursor()
    if docente:
        cur.execute("""SELECT codigo,titulo,modulo,grado,seccion,num_preguntas,estado,fecha_creacion,docente
                       FROM pruebas_diagnosticas WHERE docente=? ORDER BY id DESC""", (docente,))
    else:
        cur.execute("""SELECT codigo,titulo,modulo,grado,seccion,num_preguntas,estado,fecha_creacion,docente
                       FROM pruebas_diagnosticas ORDER BY id DESC""")
    rows = cur.fetchall()
    conn.close()
    cols = ["codigo", "titulo", "modulo", "grado", "seccion",
            "num_preguntas", "estado", "fecha_creacion", "docente"]
    return [dict(zip(cols, r)) for r in rows]

def cambiar_estado_prueba(codigo, estado):
    conn = _conn()
    conn.execute("UPDATE pruebas_diagnosticas SET estado=? WHERE codigo=?", (estado, codigo))
    conn.commit()
    conn.close()

# ── CRUD respuestas ──
def guardar_respuesta(datos):
    conn = _conn()
    conn.execute('''INSERT INTO respuestas_diagnosticas
        (prueba_codigo, estudiante, matricula, respuestas_json, correctas, total,
         puntaje, nivel, valoracion, fecha)
        VALUES (?,?,?,?,?,?,?,?,?,?)''',
        (datos["codigo"], datos["estudiante"], datos["matricula"],
         json.dumps(datos["respuestas"], ensure_ascii=False), datos["correctas"],
         datos["total"], datos["puntaje"], datos["nivel"], datos["valoracion"],
         datetime.datetime.now().strftime("%Y-%m-%d %H:%M")))
    conn.commit()
    conn.close()

def obtener_respuestas(codigo):
    conn = _conn()
    cur = conn.cursor()
    cur.execute("""SELECT estudiante,matricula,correctas,total,puntaje,nivel,valoracion,fecha
                   FROM respuestas_diagnosticas WHERE prueba_codigo=? ORDER BY id DESC""", (codigo,))
    rows = cur.fetchall()
    conn.close()
    cols = ["estudiante", "matricula", "correctas", "total", "puntaje", "nivel", "valoracion", "fecha"]
    return [dict(zip(cols, r)) for r in rows]

# ═══════════════════════════════════════════════════════════════════════════
# WORD Y PDF
# ═══════════════════════════════════════════════════════════════════════════
def extraer_texto_pdf(archivo):
    if PdfReader is None: raise RuntimeError("Instala pypdf o PyPDF2.")
    archivo.seek(0)
    reader = PdfReader(archivo)
    texto = "".join((p.extract_text() or "") for p in reader.pages)
    return re.sub(r"\s+", " ", texto).strip()[:40000]

def generar_word_prueba_fisica(prueba):
    doc = Document()
    doc.styles["Normal"].font.name = "Calibri"
    doc.styles["Normal"].font.size = Pt(11)
    
    p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.add_run("POLITÉCNICO SALESIANO ARQUIDES CALDERÓN\n").bold = True
    p.add_run(f"{prueba['titulo']}\n").bold = True
    
    doc.add_paragraph(f"Módulo/Asignatura: {prueba['modulo']}   |   Docente: {prueba['docente']}")
    doc.add_paragraph(f"Grado: {prueba['grado']}   |   Sección: {prueba['seccion']}")
    doc.add_paragraph("Nombre del Estudiante: __________________________________  Fecha: _________")
    
    doc.add_heading("Instrucciones", level=3)
    doc.add_paragraph("Lee detenidamente cada enunciado y selecciona o escribe la respuesta correcta. Evita tachaduras.")
    
    for i, p in enumerate(prueba.get("preguntas", []), 1):
        doc.add_paragraph(f"{i}. {p.get('enunciado', '')}").bold = True
        if p.get("tipo") == "opcion_multiple":
            for op in p.get("opciones", []): doc.add_paragraph(f"(   ) {op}")
            doc.add_paragraph("")
        elif p.get("tipo") == "verdadero_falso":
            doc.add_paragraph("(   ) Verdadero       (   ) Falso\n")
        else:
            doc.add_paragraph("Respuesta: __________________________________________________________________\n")

    buf = BytesIO(); doc.save(buf); buf.seek(0)
    return buf

# ═══════════════════════════════════════════════════════════════════════════
# HELPERS
# ═══════════════════════════════════════════════════════════════════════════
def generar_codigo():
    return "DX-" + "".join(random.choices(string.ascii_uppercase + string.digits, k=6))

def detectar_ip_local():
    try:
        s = socket.socket(socket.AF_INET, socket.SOCK_DGRAM)
        s.connect(("8.8.8.8", 80))
        ip = s.getsockname()[0]
        s.close()
        return ip
    except Exception: return "127.0.0.1"

def construir_link(codigo):
    base = st.secrets.get("url_publica", "") if hasattr(st, "secrets") else ""
    if not base:
        import os
        base = os.environ.get("STREAMLIT_URL", "")
    if base: return f"{base.rstrip('/')}/?prueba={codigo}"
    return f"http://{detectar_ip_local()}:8501/?prueba={codigo}"

def qr_url(link):
    return f"https://api.qrserver.com/v1/create-qr-code/?size=200x200&data={quote(link)}"

def calificar(preguntas, respuestas):
    correctas = 0
    for i, p in enumerate(preguntas):
        resp = respuestas.get(str(i))
        if resp is None or resp == "": continue
        if p["tipo"] == "opcion_multiple":
            if str(resp) == str(p["correcta"]): correctas += 1
        elif p["tipo"] == "verdadero_falso":
            if (str(resp).lower() == "verdadero") == bool(p["correcta"]): correctas += 1
    total = len(preguntas)
    puntaje = round((correctas / total) * 100, 1) if total else 0.0
    return correctas, total, puntaje

def nivel_cuantitativo(puntaje):
    if puntaje >= 90: return "Excelente"
    if puntaje >= 80: return "Muy Bueno"
    if puntaje >= 70: return "Bueno"
    if puntaje >= 60: return "Suficiente"
    return "Insuficiente"

def nivel_cualitativo(puntaje):
    if puntaje >= 70: return "Logrado"
    if puntaje >= 50: return "En Proceso"
    return "Necesita Apoyo"

# ═══════════════════════════════════════════════════════════════════════════
# INTELIGENCIA ARTIFICIAL CORRECCIÓN
# ═══════════════════════════════════════════════════════════════════════════
def corregir_prueba_ia(texto_estudiante, prueba):
    import json
    resp_str = json.dumps(prueba.get("preguntas", []), ensure_ascii=False)
    prompt = f"""Actúa como un Docente ETP corrigiendo un examen.
Aquí tienes la PRUEBA ORIGINAL Y SUS RESPUESTAS CORRECTAS:
{resp_str}

Aquí tienes el TEXTO EXTRAÍDO DEL EXAMEN ESCANEADO DEL ESTUDIANTE:
{texto_estudiante}

TAREA:
1. Extrae el nombre del estudiante si está escrito en el examen.
2. Compara las marcas o respuestas del estudiante con las correctas.
3. Calcula el puntaje de 0 a 100.
4. Escribe una valoración cualitativa motivadora detallando en qué falló.

CODIFICACIÓN: salto de línea → {ia.MARKER_NL} · comilla doble → {ia.MARKER_DQ}
Devuelve ÚNICAMENTE JSON válido:
{{
 "NOMBRE_ESTUDIANTE": "...",
 "CORRECTAS": 8,
 "TOTAL_PREGUNTAS": 10,
 "PUNTAJE": 80,
 "VALORACION": "Buen dominio general, sin embargo fallaste en..."
}}"""
    datos, _ = ia.solicitar_json(prompt, max_tokens=2000, temperature=0.1, modulo="corrector_diagnostica")
    return ia.decodificar_marcadores(datos)

# ═══════════════════════════════════════════════════════════════════════════
# ESTILOS COMPARTIDOS
# ═══════════════════════════════════════════════════════════════════════════
CSS = """
<style>
@import url('https://fonts.googleapis.com/css2?family=Inter:wght@300;400;500;600;700;800;900&display=swap');
html, body, [class*="css"] { font-family: 'Inter', sans-serif; background-color: #F8FAFC; color: #1E293B; }

.pd-hero { background: linear-gradient(135deg, #0F172A 0%, #6D28D9 55%, #A78BFA 100%); color:#fff;
  padding:2.2rem; border-radius:20px; margin-bottom:1.5rem; box-shadow:0 25px 50px rgba(109,40,217,0.3);
  position:relative; overflow:hidden; }
.pd-hero::before { content:''; position:absolute; top:-50%; left:-50%; width:200%; height:200%;
  background:radial-gradient(circle, rgba(255,255,255,0.12) 0%, transparent 60%);
  animation:pdPulse 6s ease-in-out infinite; }
@keyframes pdPulse { 0%,100%{transform:scale(1);opacity:.5} 50%{transform:scale(1.1);opacity:.8} }
.pd-hero-title { font-size:2.3rem; font-weight:900; letter-spacing:-0.03em; position:relative; }
.pd-hero-sub { font-size:1.05rem; opacity:0.9; margin-top:0.4rem; position:relative; }
.pd-hero-badge { display:inline-block; background:rgba(255,255,255,0.15); border:1px solid rgba(255,255,255,0.25);
  border-radius:8px; padding:4px 12px; font-size:0.8rem; font-weight:600; margin-top:0.8rem;
  margin-right:8px; position:relative; }

.pd-section-title { color:#6D28D9; font-weight:700; font-size:1.12rem; border-bottom:2px solid #EDE9FE;
  padding-bottom:8px; margin:1.2rem 0 0.9rem 0; }
.pd-qcard { background:#fff; border:2px solid #E2E8F0; border-radius:14px; padding:1.2rem;
  margin-bottom:0.6rem; box-shadow:0 2px 8px rgba(0,0,0,0.04); }
.pd-linkbox { background:#F5F3FF; border:2px dashed #A78BFA; border-radius:12px; padding:1rem;
  margin-top:0.8rem; font-family:monospace; font-size:0.95rem; word-break:break-all; }
.pd-kpi { background:#fff; border-radius:14px; padding:1.2rem; text-align:center;
  border:1px solid #E2E8F0; box-shadow:0 4px 12px rgba(0,0,0,0.06); }
.pd-kpi-value { font-size:2rem; font-weight:800; color:#6D28D9; }
.pd-kpi-label { font-size:0.75rem; font-weight:700; color:#94A3B8; text-transform:uppercase;
  letter-spacing:0.05em; margin-top:0.3rem; }
.score-card { background: #1E293B; color: white; border-radius: 16px; padding: 25px; text-align: center; border-bottom: 5px solid #10B981; margin-top: 20px;}
.score-number { font-size: 4rem; font-weight: 900; line-height: 1; color: #10B981; }
</style>
"""
st.markdown(CSS, unsafe_allow_html=True)


# ═══════════════════════════════════════════════════════════════════════════
# VISTA ESTUDIANTE (activada con ?prueba=CODIGO — sin login)
# ═══════════════════════════════════════════════════════════════════════════
def render_vista_estudiante(codigo):
    st.markdown("<style>[data-testid='stSidebar'] {display:none !important;}</style>", unsafe_allow_html=True)

    prueba = obtener_prueba_por_codigo(codigo)
    if not prueba:
        st.error("❌ Prueba no encontrada. Verifica el enlace con tu docente.")
        st.stop()
    if prueba["estado"] != "activa":
        st.warning("⏸️ Esta prueba ya fue cerrada por el docente.")
        st.stop()

    clave_resultado = f"resultado_prueba_{codigo}"

    # ── Si ya respondió, mostrar resultado ──
    if st.session_state.get(clave_resultado):
        res = st.session_state[clave_resultado]
        st.markdown(f"""
        <div class="pd-hero">
            <div class="pd-hero-title">🎉 ¡Prueba Enviada!</div>
            <div class="pd-hero-sub">Gracias por completar la prueba diagnóstica, {res.get('nombre','')}.</div>
        </div>
        """, unsafe_allow_html=True)
        c1, c2, c3 = st.columns(3)
        with c1:
            st.markdown(f'<div class="pd-kpi"><div class="pd-kpi-value">{res["puntaje"]}</div><div class="pd-kpi-label">Puntaje / 100</div></div>', unsafe_allow_html=True)
        with c2:
            st.markdown(f'<div class="pd-kpi"><div class="pd-kpi-value" style="color:#059669;">{res["correctas"]}/{res["total"]}</div><div class="pd-kpi-label">Correctas</div></div>', unsafe_allow_html=True)
        with c3:
            st.markdown(f'<div class="pd-kpi"><div class="pd-kpi-value" style="color:#D97706;font-size:1.4rem;">{res["nivel"]}</div><div class="pd-kpi-label">Nivel de logro</div></div>', unsafe_allow_html=True)
        st.markdown('<div class="pd-section-title">💬 Valoración cualitativa</div>', unsafe_allow_html=True)
        st.info(res["valoracion"])
        st.markdown('<div class="pd-section-title">📤 Comprobante</div>', unsafe_allow_html=True)
        st.caption("Guarda esta pantalla como comprobante. Tu docente y el coordinador ya tienen tu resultado.")
        if st.button("🔄 Enviar otra respuesta (nuevo estudiante)", use_container_width=True):
            del st.session_state[clave_resultado]
            st.rerun()
        st.stop()

    # ── Formulario de la prueba ──
    st.markdown(f"""
    <div class="pd-hero">
        <div class="pd-hero-title">🩺 {prueba['titulo']}</div>
        <div class="pd-hero-sub">{prueba['modulo']} · {prueba['area']} · {prueba['grado']} "{prueba['seccion']}"</div>
        <div class="pd-hero-sub">Docente: {prueba['docente']}</div>
    </div>
    """, unsafe_allow_html=True)

    with st.form(f"form_estudiante_{codigo}"):
        st.markdown('<div class="pd-section-title">👤 Identifícate</div>', unsafe_allow_html=True)
        col1, col2 = st.columns(2)
        with col1:
            nombre_est = st.text_input("Nombre completo *", placeholder="Ej: Juan Pérez")
        with col2:
            matricula = st.text_input("Matrícula / Cédula", placeholder="Opcional")

        preguntas = prueba["preguntas"]
        st.markdown(f'<div class="pd-section-title">📝 Responde ({len(preguntas)} preguntas)</div>', unsafe_allow_html=True)

        respuestas_ui = {}
        for i, p in enumerate(preguntas, 1):
            st.markdown(f'<div class="pd-qcard"><b>Pregunta {i}.</b> {p["enunciado"]}</div>', unsafe_allow_html=True)
            if p["tipo"] == "opcion_multiple":
                opciones = p.get("opciones", [])
                resp = st.radio(f"Elige una opción (Pregunta {i})", opciones,
                                key=f"q_{codigo}_{i}", index=None)
                respuestas_ui[str(i - 1)] = ("opcion_multiple", resp, opciones)
            elif p["tipo"] == "verdadero_falso":
                resp = st.radio(f"Verdadero o Falso (Pregunta {i})", ["Verdadero", "Falso"],
                                key=f"q_{codigo}_{i}", index=None)
                respuestas_ui[str(i - 1)] = ("verdadero_falso", resp, None)

        submit = st.form_submit_button("✅ Enviar mis respuestas", type="primary", use_container_width=True)

    if submit:
        if not nombre_est.strip():
            st.error("⚠️ Debes escribir tu nombre completo.")
            st.stop()

        respuestas = {}
        sin_responder = 0
        for idx, (tipo, resp, opciones) in respuestas_ui.items():
            if resp is None:
                sin_responder += 1
                respuestas[idx] = None
            elif tipo == "opcion_multiple":
                respuestas[idx] = opciones.index(resp)
            else:
                respuestas[idx] = resp

        if sin_responder > 0:
            st.warning(f"📌 Dejaste {sin_responder} pregunta(s) sin responder. Se contarán como incorrectas.")

        correctas, total, puntaje = calificar(preguntas, respuestas)
        nivel = nivel_cualitativo(puntaje)

        valoracion = ""
        cfg = ia.config_ia()
        if cfg.get("api_key"):
            try:
                prompt_val = f"""Eres un docente experto en evaluación diagnóstica ETP del MINERD.
El estudiante {nombre_est} obtuvo {puntaje}/100 ({correctas} de {total} correctas) en una prueba diagnóstica de {prueba['modulo']}.
Nivel de logro: {nivel}.
Escribe UNA valoración cualitativa breve (3-4 líneas), constructiva y motivadora, dirigida al estudiante y al docente."""
                valoracion, _ = ia.solicitar_texto(prompt_val, max_tokens=1024, temperature=0.4, modulo="prueba_diagnostica")
            except Exception: pass
        if not valoracion:
            valoracion = f"El estudiante {nombre_est} alcanzó {puntaje}/100. Nivel: '{nivel}'."

        guardar_respuesta({
            "codigo": codigo, "estudiante": nombre_est.strip(), "matricula": matricula.strip(),
            "respuestas": respuestas, "correctas": correctas, "total": total,
            "puntaje": puntaje, "nivel": nivel, "valoracion": valoracion,
        })

        st.session_state[clave_resultado] = {
            "nombre": nombre_est.strip(), "puntaje": puntaje, "correctas": correctas,
            "total": total, "nivel": nivel, "valoracion": valoracion,
        }
        st.rerun()

# ═══════════════════════════════════════════════════════════════════════════
# DETECCIÓN DE MODO (estudiante vs gestión)
# ═══════════════════════════════════════════════════════════════════════════
_codigo_qs = st.query_params.get("prueba")
if _codigo_qs:
    render_vista_estudiante(_codigo_qs)
    st.stop()

# ═══════════════════════════════════════════════════════════════════════════
# VISTA DE GESTIÓN (docente / coordinador)
# ═══════════════════════════════════════════════════════════════════════════
if not (st.session_state.get("coordinador_autenticado") or st.session_state.get("docente_autenticado")):
    st.error("🔒 Inicia sesión para gestionar pruebas diagnósticas.")
    st.stop()

es_coordinador = st.session_state.get("coordinador_autenticado", False)
nombre_usuario = st.session_state.get("coordinador_nombre") or st.session_state.get("nombre_docente") or "Docente"

ia.panel_sidebar_ia("Pruebas Diagnósticas")

st.markdown(f"""
<div class="pd-hero">
    <div class="pd-hero-title">🩺 Centro de Pruebas Diagnósticas Fígital</div>
    <div class="pd-hero-sub">Genera pruebas con IA. Aplícalas enviando un enlace digital, o imprime el Word, aplícala en papel y deja que la IA la escanee y la corrija por ti.</div>
    <div>
        <span class="pd-hero-badge">🤖 Pruebas con IA</span>
        <span class="pd-hero-badge">🔗 Enlace Móvil Digital</span>
        <span class="pd-hero-badge">🖨️ Word Imprimible</span>
        <span class="pd-hero-badge">📄 Escáner Inteligente</span>
    </div>
</div>
""", unsafe_allow_html=True)

tab1, tab2, tab3, tab4 = st.tabs(["➕ 1. Crear Prueba", "🔗 2. Compartir Digital / Imprimir", "📸 3. Corregir Escaneo IA", "📊 4. Resultados"])

# ─────────────────────────────────────────────────────────────────────────────
# TAB 1: CREAR PRUEBA
# ─────────────────────────────────────────────────────────────────────────────
with tab1:
    st.markdown('<div class="pd-section-title">📋 1. Datos de la Prueba</div>', unsafe_allow_html=True)
    col_a, col_b = st.columns(2)
    with col_a:
        tipo_prueba = st.radio("Tipo de evaluación", ["Módulo Formativo (ETP)", "Malla Curricular Académica"], horizontal=True)
        titulo_prueba = st.text_input("Título de la prueba *", placeholder="Ej: Diagnóstico de Impuestos al Consumo")
        docente_prueba = st.text_input("Docente", value=nombre_usuario)
    with col_b:
        area_prueba = st.text_input("Área / Asignatura *", placeholder="Ej: MF 358-3 / Matemáticas")
        modulo_prueba = st.text_input("Módulo / Tema *", placeholder="Ej: Impuestos al Consumo")
        col_g, col_s = st.columns(2)
        with col_g: grado_prueba = st.text_input("Grado", placeholder="Ej: 5to")
        with col_s: seccion_prueba = st.text_input("Sección", placeholder="Ej: B")

    st.markdown('<div class="pd-section-title">🧠 2. Generar Preguntas con IA</div>', unsafe_allow_html=True)
    contenido_base = st.text_area("Contenido / RA a evaluar", height=120, placeholder="Ej: Identificar los medios de presentación de documentación...")
    
    col_n, col_d = st.columns(2)
    with col_n: num_preguntas = st.slider("Cantidad de preguntas", 3, 20, 8)
    with col_d: dificultad = st.selectbox("Dificultad", ["Básica", "Intermedia", "Avanzada"], index=1)

    if st.button("✨ Generar Preguntas con IA", type="primary", use_container_width=True):
        if not contenido_base.strip() or not titulo_prueba.strip():
            st.warning("⚠️ Completa el título y el contenido.")
        else:
            with st.spinner("🧠 Generando preguntas..."):
                try:
                    prompt_gen = f"""Eres experto en evaluación diagnóstica ETP (MINERD).
Genera {num_preguntas} preguntas ({dificultad}) sobre {modulo_prueba} ({area_prueba}).
Contenido base: {contenido_base}
Usa opción múltiple (4 opciones, 1 correcta) y verdadero/falso.
Devuelve SOLO JSON válido:
{{
  "preguntas": [
    {{"tipo": "opcion_multiple", "enunciado": "...", "opciones": ["a","b","c","d"], "correcta": 0}}
  ]
}}"""
                    texto, flags = ia.solicitar_ia(prompt_gen, modo="json", max_tokens=8192, temperature=0.4, modulo="prueba_diagnostica")
                    datos = ia.decodificar_marcadores(ia.parsear_json_robusto(texto))
                    st.session_state.preguntas_borrador = datos.get("preguntas", [])
                    st.success(f"✅ {len(st.session_state.preguntas_borrador)} preguntas creadas. Revísalas abajo y publica.")
                except Exception as e: ia.render_error_ia(e, None)

    if st.session_state.get("preguntas_borrador"):
        st.markdown('<div class="pd-section-title">✏️ 3. Revisar y Editar Preguntas</div>', unsafe_allow_html=True)
        borrador = st.session_state.preguntas_borrador
        for i, p in enumerate(borrador):
            with st.expander(f"Pregunta {i+1}: {p.get('enunciado','')[:80]}...", expanded=(i == 0)):
                nuevo_enunciado = st.text_area(f"Enunciado (P{i+1})", value=p.get("enunciado", ""), key=f"edit_enun_{i}")
                if p["tipo"] == "opcion_multiple":
                    opciones = p.get("opciones", ["", "", "", ""])
                    col_o1, col_o2 = st.columns(2)
                    with col_o1:
                        op0 = st.text_input(f"Op A", value=opciones[0], key=f"o0_{i}")
                        op1 = st.text_input(f"Op B", value=opciones[1] if len(opciones)>1 else "", key=f"o1_{i}")
                    with col_o2:
                        op2 = st.text_input(f"Op C", value=opciones[2] if len(opciones)>2 else "", key=f"o2_{i}")
                        op3 = st.text_input(f"Op D", value=opciones[3] if len(opciones)>3 else "", key=f"o3_{i}")
                    correcta = st.selectbox("Correcta", [0, 1, 2, 3], format_func=lambda x: ["A", "B", "C", "D"][x], index=p.get("correcta", 0), key=f"cor_{i}")
                    borrador[i].update({"enunciado": nuevo_enunciado, "opciones": [op0, op1, op2, op3], "correcta": correcta})
                else:
                    correcta_vf = st.selectbox("Correcta", ["Verdadero", "Falso"], index=0 if p.get("correcta") else 1, key=f"cor_{i}")
                    borrador[i].update({"enunciado": nuevo_enunciado, "correcta": correcta_vf == "Verdadero"})

        st.markdown('<div class="pd-section-title">🚀 4. Guardar y Publicar Prueba</div>', unsafe_allow_html=True)
        if st.button("🚀 PUBLICAR PRUEBA Y GENERAR ENLACES", type="primary", use_container_width=True):
            codigo = generar_codigo()
            guardar_prueba({
                "codigo": codigo, "titulo": titulo_prueba.strip(), "area": area_prueba.strip(),
                "modulo": modulo_prueba.strip(), "docente": docente_prueba.strip(),
                "grado": grado_prueba.strip(), "seccion": seccion_prueba.strip(),
                "tipo": tipo_prueba, "num_preguntas": len(borrador), "preguntas": borrador,
            })
            st.session_state.preguntas_borrador = None
            st.session_state.prueba_recien_creada = codigo
            st.toast("✅ Prueba publicada", icon="🎉")
            st.rerun()

# ─────────────────────────────────────────────────────────────────────────────
# TAB 2: MIS PRUEBAS Y ENLACES (FÍGITAL)
# ─────────────────────────────────────────────────────────────────────────────
with tab2:
    st.markdown('<div class="pd-section-title">🔗 Compartir Digital o Imprimir Físico</div>', unsafe_allow_html=True)
    pruebas = obtener_mis_pruebas(None if es_coordinador else nombre_usuario)

    if not pruebas:
        st.info("📭 Aún no has creado pruebas.")
    else:
        for p in pruebas:
            estado_emoji = "🟢 Activa" if p["estado"] == "activa" else "🔴 Cerrada"
            with st.expander(f"{estado_emoji} · {p['titulo']} · {p['modulo']} ({p['grado']} '{p['seccion']}')"):
                link = construir_link(p["codigo"])
                
                st.markdown("#### 💻 Modo Digital (Estudiantes en vivo)")
                col_qr, col_info = st.columns([1, 2])
                with col_qr:
                    st.image(qr_url(link), width=150)
                with col_info:
                    st.markdown(f"**Código:** `{p['codigo']}`")
                    st.markdown(f'<div class="pd-linkbox">{link}</div>', unsafe_allow_html=True)

                st.markdown("#### 🖨️ Modo Físico (Impreso)")
                c_d1, c_d2 = st.columns(2)
                
                prueba_completa = obtener_prueba_por_codigo(p["codigo"])
                buf_word = generar_word_prueba_fisica(prueba_completa)
                
                with c_d1:
                    st.download_button("📥 Descargar Word para Imprimir", data=buf_word, 
                                       file_name=f"Prueba_{p['codigo']}.docx",
                                       mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                                       use_container_width=True, key=f"dl_{p['codigo']}")
                with c_d2:
                    if st.button("☁️ Subir a Google Drive", use_container_width=True, key=f"dr_{p['codigo']}"):
                        with st.spinner("Conectando con Google Drive (OAuth)..."):
                            try:
                                servicio = drive_api._servicio()
                                from googleapiclient.http import MediaIoBaseUpload
                                import io
                                buf_word.seek(0)
                                media = MediaIoBaseUpload(io.BytesIO(buf_word.getvalue()), mimetype=drive_api.MIME_WORD, resumable=True)
                                meta = {"name": f"Prueba_{p['codigo']}.docx", "parents": [drive_api.DRIVE_CARPETA_ID]}
                                archivo = servicio.files().create(body=meta, media_body=media, fields="id, webViewLink").execute()
                                
                                # Hacemos el archivo público para que los estudiantes lo puedan descargar
                                servicio.permissions().create(fileId=archivo.get('id'), body={'type': 'anyone', 'role': 'reader'}).execute()
                                
                                drive_link = archivo.get('webViewLink')
                                st.success(f"✅ ¡Subido! [Ver en Drive]({drive_link})")
                            except Exception as e:
                                st.error(f"⚠️ Error conectando a Drive: {e}")
                                st.info("Asegúrate de tener el archivo 'credentials.json' configurado como indica core/drive.py")

# ─────────────────────────────────────────────────────────────────────────────
# TAB 3: CORREGIR ESCANEO IA
# ─────────────────────────────────────────────────────────────────────────────
with tab3:
    st.markdown('<div class="pd-section-title">🤖 Corrector de Exámenes Escaneados con IA</div>', unsafe_allow_html=True)
    st.info("¿Diste la prueba en físico? Sube el PDF escaneado del estudiante. La IA la leerá, la comparará con tu plantilla de preguntas y la calificará automáticamente.")
    
    if not pruebas:
        st.warning("Crea una prueba primero en la pestaña 1.")
    else:
        opciones_pruebas = {f"{p['titulo']} ({p['codigo']})": p["codigo"] for p in pruebas}
        sel_corr = st.selectbox("Selecciona la plantilla de la prueba", list(opciones_pruebas.keys()))
        archivo_estudiante = st.file_uploader("Sube el examen escaneado (PDF)", type=["pdf"], key="pdf_est_corr")
        
        if st.button("🤖 Leer y Calificar Examen", type="primary", use_container_width=True, disabled=not archivo_estudiante):
            if not st.session_state.get("api_key_global", ""):
                st.error("🔒 Falta tu API Key.")
            else:
                prueba_data = obtener_prueba_por_codigo(opciones_pruebas[sel_corr])
                with st.spinner("Leyendo caligrafía y comparando respuestas con la plantilla..."):
                    try:
                        texto_escaneado = extraer_texto_pdf(archivo_estudiante)
                        resultado_ia = corregir_prueba_ia(texto_escaneado, prueba_data)
                        
                        nota = int(resultado_ia.get("PUNTAJE", 0))
                        est = resultado_ia.get("NOMBRE_ESTUDIANTE", "Desconocido")
                        nivel = nivel_cualitativo(nota)
                        retro = resultado_ia.get("VALORACION", "")
                        
                        guardar_respuesta({
                            "codigo": prueba_data["codigo"], "estudiante": est, "matricula": "",
                            "respuestas": {}, "correctas": int(resultado_ia.get("CORRECTAS", 0)), 
                            "total": int(resultado_ia.get("TOTAL_PREGUNTAS", 10)),
                            "puntaje": nota, "nivel": nivel, "valoracion": retro
                        })
                        
                        st.session_state.ultima_correccion_pd = resultado_ia
                        st.toast("✅ Examen calificado correctamente", icon="🎯")
                    except Exception as e:
                        ia.render_error_ia(e)

        if st.session_state.get("ultima_correccion_pd"):
            corr = st.session_state.ultima_correccion_pd
            c_nota, c_retro = st.columns([1, 2])
            with c_nota:
                color_nota = "#10B981" if int(corr.get('PUNTAJE',0)) >= 70 else "#EF4444"
                st.markdown(f"""
                <div class="score-card" style="border-bottom-color: {color_nota};">
                    <div style="font-size: 1.1rem; margin-bottom: 5px;">{corr.get('NOMBRE_ESTUDIANTE', '')}</div>
                    <div class="score-number" style="color: {color_nota};">{corr.get('PUNTAJE', 0)}</div>
                    <div>PUNTOS OBTENIDOS</div>
                </div>
                """, unsafe_allow_html=True)
            with c_retro:
                st.info(f"**Retroalimentación Pedagógica IA:**\n{corr.get('VALORACION', '')}")
                st.success("✅ Este resultado ha sido enviado automáticamente al Historial General.")

# ─────────────────────────────────────────────────────────────────────────────
# TAB 4: RESULTADOS
# ─────────────────────────────────────────────────────────────────────────────
with tab4:
    st.markdown('<div class="pd-section-title">📊 Resultados Consolidados (Físicos y Digitales)</div>', unsafe_allow_html=True)
    if not pruebas:
        st.info("📭 No hay pruebas.")
    else:
        sel_res = st.selectbox("Ver resultados de la prueba:", list(opciones_pruebas.keys()), key="sel_resultados")
        respuestas = obtener_respuestas(opciones_pruebas[sel_res])

        if not respuestas:
            st.info("📨 Aún no hay calificaciones registradas para esta prueba.")
        else:
            df = pd.DataFrame(respuestas)
            c1, c2, c3, c4 = st.columns(4)
            with c1: st.markdown(f'<div class="pd-kpi"><div class="pd-kpi-value">{len(df)}</div><div class="pd-kpi-label">Estudiantes Evaluados</div></div>', unsafe_allow_html=True)
            with c2: st.markdown(f'<div class="pd-kpi"><div class="pd-kpi-value" style="color:#059669;">{round(df["puntaje"].mean(),1)}</div><div class="pd-kpi-label">Promedio de la Clase</div></div>', unsafe_allow_html=True)
            
            st.markdown("##### 📋 Listado de Calificaciones")
            st.dataframe(df[["estudiante", "puntaje", "nivel", "fecha"]], use_container_width=True, hide_index=True)

            if st.button("📥 Generar Informe Consolidado (.docx)", type="primary", use_container_width=True):
                doc = Document()
                doc.styles["Normal"].font.name = "Calibri"
                doc.styles["Normal"].font.size = Pt(10)

                doc.add_heading("INFORME DE RESULTADOS DIAGNÓSTICOS", level=1)
                doc.add_paragraph(f"Prueba: {sel_res}\nFecha: {datetime.date.today().strftime('%d/%m/%Y')}")
                
                tabla = doc.add_table(rows=1, cols=4)
                tabla.style = "Table Grid"
                hdr = tabla.rows[0].cells
                for i, h in enumerate(["Estudiante", "Puntaje", "Nivel", "Fecha"]):
                    hdr[i].text = h; hdr[i].paragraphs[0].runs[0].bold = True

                for _, row in df.iterrows():
                    rc = tabla.add_row().cells
                    rc[0].text = str(row["estudiante"])
                    rc[1].text = str(row["puntaje"])
                    rc[2].text = str(row["nivel"])
                    rc[3].text = str(row["fecha"])

                buf = BytesIO(); doc.save(buf); buf.seek(0)
                st.download_button("⬇️ Descargar Informe (.docx)", data=buf, file_name=f"Resultados_{opciones_pruebas[sel_res]}.docx",
                                   mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document", type="primary", use_container_width=True)