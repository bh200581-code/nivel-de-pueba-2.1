"""
evaluador_planes.py — Auditoría de Planificaciones Diarias ETP (MEJORADO · Paquete 1-8)
• 14 criterios oficiales en 4 categorías + esquema de 20 secciones.
• NUEVO: historial SQLite, auditoría por lotes, gráfico por categoría,
  integración con Acompañamiento, matching robusto, comparativa por docente,
  paginación Word y export CSV/Word.
"""
import datetime
import difflib
import re
import sqlite3
import unicodedata
from io import BytesIO

import pandas as pd
import streamlit as st
from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement, parse_xml
from docx.oxml.ns import nsdecls, qn
from docx.shared import Inches, Pt, RGBColor

try:
    from pypdf import PdfReader
except Exception:
    try:
        from PyPDF2 import PdfReader
    except Exception:
        PdfReader = None

from core import ia

# ═══════════════════════════════════════════════════════════════════════════
# FUENTE ÚNICA DE VERDAD (sin cambios)
# ═══════════════════════════════════════════════════════════════════════════
CRITERIOS_OFICIALES = [
    ("I. Datos de Identificación y Contextualización", "Identificación Completa del Docente y del Centro",
     "Nombre completo, cédula, regional, distrito, centro educativo con código, grado y sección, modalidad y tanda."),
    ("I. Datos de Identificación y Contextualización", "Contextualización del Módulo",
     "Área/Asignatura/Módulo Formativo con código (MF), bachillerato técnico y fecha de aplicación."),
    ("I. Datos de Identificación y Contextualización", "Temporalidad Viable",
     "Duración de 50 minutos según instructivo MINERD y momentos dosificados (10 inicio / 30 desarrollo / 10 cierre)."),
    ("II. Alineación Curricular y Trazabilidad", "Transcripción del Resultado de Aprendizaje (RA)",
     "Código y texto completo del RA tal como aparece en el diseño curricular."),
    ("II. Alineación Curricular y Trazabilidad", "Criterios de Evaluación (CE) Explícitos",
     "Códigos y descripciones de los CE asociados al RA."),
    ("II. Alineación Curricular y Trazabilidad", "Elemento de Capacidad (EC) Alineado",
     "EC coherente con el RA y los CE, con redacción observable."),
    ("II. Alineación Curricular y Trazabilidad", "Intención Educativa del Día",
     "Propósito redactado con claridad, orientado al perfil profesional y al contexto laboral del módulo."),
    ("III. Secuencia Didáctica y Momentos Pedagógicos", "Inicio de Clase (3 fases)",
     "Motivación/activación, recuperación de saberes previos y presentación de la intención educativa."),
    ("III. Secuencia Didáctica y Momentos Pedagógicos", "Desarrollo con Metodologías Activas ETP (3 fases)",
     "Estudio de casos, trabajo colaborativo, modelado técnico y práctica del estudiante con roles diferenciados."),
    ("III. Secuencia Didáctica y Momentos Pedagógicos", "Cierre con Metacognición (3 fases)",
     "Actividad de cierre interactivo, preguntas de metacognición y reflexión sobre la aplicación real/laboral."),
    ("III. Secuencia Didáctica y Momentos Pedagógicos", "Tridimensionalidad de Contenidos",
     "Componentes conceptuales, procedimentales y actitudinales articulados y coherentes con el módulo."),
    ("IV. Recursos, Evaluación y Atención a la Diversidad", "Pertinencia de Recursos Técnicos",
     "Recursos verídicos del entorno técnico (PDI, fichas técnicas, documentos simulados, plataformas interactivas)."),
    ("IV. Recursos, Evaluación y Atención a la Diversidad", "Instrumento de Evaluación y Escala L/EP/NA",
     "Lista de cotejo u otro instrumento con criterios observables y escala de valoración L / EP / NA."),
    ("IV. Recursos, Evaluación y Atención a la Diversidad", "Adaptaciones NEAE y Plan Alternativo",
     "Adaptaciones para estudiantes con necesidades específicas y observaciones/plan alternativo ante contingencias."),
]
CATEGORIAS_ORDEN = [
    "I. Datos de Identificación y Contextualización",
    "II. Alineación Curricular y Trazabilidad",
    "III. Secuencia Didáctica y Momentos Pedagógicos",
    "IV. Recursos, Evaluación y Atención a la Diversidad",
]
CATEGORIA_ICONOS = {c: i for c, i in zip(CATEGORIAS_ORDEN, ["🗂️", "🧭", "‍", "📊"])}
TOTAL_CRITERIOS = len(CRITERIOS_OFICIALES)
UMBRAL_EXCELENTE = 13
UMBRAL_ACEPTABLE = 10
ESQUEMA_OFICIAL = [
    "Datos Generales", "Características del grupo", "Módulo Formativo (MF)", "RA con código y texto",
    "Criterios de Evaluación (CE)", "Elemento de Capacidad (EC)", "Tipo / Tiempo / Estrategias / Valor",
    "Componentes Curriculares", "Enunciado de la Actividad", "Intención Educativa",
    "Momento INICIO (3 fases)", "Momento DESARROLLO (3 fases)", "Momento CIERRE (3 fases)",
    "Recursos", "Instrumento e indicadores", "Adaptaciones NEAE", "Observaciones / plan alternativo",
    "Lista de Cotejo", "Escala de Valoración (L/EP/NA)", "Firmas",
]

# ═══════════════════════════════════════════════════════════════════════════
# PERSISTENCIA SQLite (mejora 1 y 8)
# ═══════════════════════════════════════════════════════════════════════════
DB_NAME = "gestion_etp.db"


def _conn():
    return sqlite3.connect(DB_NAME, check_same_thread=False)


def asegurar_tabla_auditorias():
    conn = _conn()
    cur = conn.cursor()
    cur.execute('''CREATE TABLE IF NOT EXISTS auditorias_planes (
        id INTEGER PRIMARY KEY AUTOINCREMENT,
        fecha TEXT, docente TEXT, modulo TEXT, evaluador TEXT,
        porcentaje INTEGER, cumplidos INTEGER, dictamen TEXT, completitud INTEGER)''')
    conn.commit()
    conn.close()


def insertar_auditoria(fecha, docente, modulo, evaluador, porcentaje, cumplidos, dictamen, completitud):
    conn = _conn()
    conn.execute('''INSERT INTO auditorias_planes
        (fecha, docente, modulo, evaluador, porcentaje, cumplidos, dictamen, completitud)
        VALUES (?,?,?,?,?,?,?,?)''',
        (fecha, docente, modulo, evaluador, porcentaje, cumplidos, dictamen, completitud))
    conn.commit()
    conn.close()


def listar_auditorias():
    conn = _conn()
    cur = conn.cursor()
    cur.execute("SELECT id, fecha, docente, modulo, evaluador, porcentaje, cumplidos, dictamen, completitud "
                "FROM auditorias_planes ORDER BY id DESC")
    cols = ["id", "fecha", "docente", "modulo", "evaluador", "porcentaje", "cumplidos", "dictamen", "completitud"]
    rows = [dict(zip(cols, r)) for r in cur.fetchall()]
    conn.close()
    return rows


def registrar_acompanamiento(docente, modulo, dictamen):
    """Mejora 2: registra acompañamiento si el plan es No Aceptable."""
    try:
        conn = _conn()
        conn.execute("INSERT INTO acompanamientos (fecha, docente, modulo, estado) VALUES (?,?,?,?)",
                     (datetime.date.today().isoformat(), docente, modulo, f"Auditoría: {dictamen}"))
        conn.commit()
        conn.close()
        return True
    except Exception:
        return False


asegurar_tabla_auditorias()

# ═══════════════════════════════════════════════════════════════════════════
# HELPERS
# ═══════════════════════════════════════════════════════════════════════════
def construir_texto_criterios():
    lineas, cat_actual = [], None
    for idx, (cat, nombre, desc) in enumerate(CRITERIOS_OFICIALES, 1):
        if cat != cat_actual:
            lineas.append(f"\n{cat}")
            cat_actual = cat
        lineas.append(f"{idx}. {nombre}: {desc}")
    return "\n".join(lineas)


def validar_criterios(criterios_ia):
    """Mejora 7: matching por NO, luego por similitud de texto, luego enumeración."""
    por_numero = {}
    for item in criterios_ia or []:
        try:
            n = int(item.get("NO"))
        except (TypeError, ValueError):
            continue
        if 1 <= n <= TOTAL_CRITERIOS and n not in por_numero:
            por_numero[n] = item
    if not por_numero and criterios_ia:
        for item in criterios_ia:
            texto_item = str(item.get("CRITERIO", "")).lower()
            mejor_idx, mejor_ratio = None, 0.0
            for idx, (_, nombre, _) in enumerate(CRITERIOS_OFICIALES, 1):
                ratio = difflib.SequenceMatcher(None, texto_item, nombre.lower()).ratio()
                if ratio > mejor_ratio:
                    mejor_idx, mejor_ratio = idx, ratio
            if mejor_idx and mejor_ratio >= 0.6 and mejor_idx not in por_numero:
                por_numero[mejor_idx] = item
    if not por_numero and criterios_ia:
        por_numero = {i + 1: item for i, item in enumerate(criterios_ia[:TOTAL_CRITERIOS])}
    validados, faltantes = [], []
    for idx, (cat, nombre, _) in enumerate(CRITERIOS_OFICIALES, 1):
        item = por_numero.get(idx)
        if item:
            validados.append({"NO": idx, "CATEGORIA": cat,
                              "CRITERIO": item.get("CRITERIO") or nombre,
                              "CUMPLE": bool(item.get("CUMPLE", False)),
                              "OBSERVACION": item.get("OBSERVACION") or "Sin observación registrada."})
        else:
            faltantes.append(nombre)
            validados.append({"NO": idx, "CATEGORIA": cat, "CRITERIO": nombre, "CUMPLE": False,
                              "OBSERVACION": "⚠️ La IA no evaluó este criterio — pendiente de revisión manual."})
    return validados, faltantes


def calcular_dictamen(cumplidos):
    if cumplidos >= UMBRAL_EXCELENTE:
        return "Excelente", "#10B981"
    if cumplidos >= UMBRAL_ACEPTABLE:
        return "Aceptable con Mejoras", "#F59E0B"
    return "No Aceptable", "#EF4444"


def extraer_texto_plan(archivo):
    nombre = archivo.name.lower()
    if nombre.endswith(".pdf"):
        if PdfReader is None:
            raise RuntimeError("No hay librería PDF. Instala pypdf o PyPDF2.")
        reader = PdfReader(archivo)
        return " ".join(p.extract_text() or "" for p in reader.pages)
    doc = Document(archivo)
    texto = "\n".join(p.text for p in doc.paragraphs)
    for tabla in doc.tables:
        for fila in tabla.rows:
            texto += "\n" + " | ".join(c.text for c in fila.cells)
    return texto


# ═══════════════════════════════════════════════════════════════════════════
# UTILIDADES WORD (con paginación — mejora 5)
# ═══════════════════════════════════════════════════════════════════════════
def shade_cell(cell, color):
    cell._tc.get_or_add_tcPr().append(parse_xml(r'<w:shd {} w:fill="{}"/>'.format(nsdecls('w'), color)))


def fijar_anchos_columna(tabla, anchos):
    tabla.autofit = False
    for row in tabla.rows:
        for idx, ancho in enumerate(anchos):
            if idx < len(row.cells):
                row.cells[idx].width = Inches(ancho)


def agregar_linea_divisoria(doc, color="2563EB"):
    p = doc.add_paragraph()
    pPr = p._p.get_or_add_pPr()
    pBdr = OxmlElement("w:pBdr")
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), "8")
    bottom.set(qn("w:space"), "1")
    bottom.set(qn("w:color"), color)
    pBdr.append(bottom)
    pPr.append(pBdr)


def agregar_numeracion_pagina(doc):
    pie = doc.sections[0].footer.paragraphs[0]
    pie.alignment = WD_ALIGN_PARAGRAPH.CENTER
    pie.text = ""
    run = pie.add_run()

    def _campo(instr):
        i = OxmlElement("w:fldChar"); i.set(qn("w:fldCharType"), "begin")
        t = OxmlElement("w:instrText"); t.set(qn("xml:space"), "preserve"); t.text = instr
        f = OxmlElement("w:fldChar"); f.set(qn("w:fldCharType"), "end")
        run._r.append(i); run._r.append(t); run._r.append(f)

    run.add_text("Página ")
    _campo("PAGE")
    run.add_text(" de ")
    _campo("NUMPAGES")
    run.font.size = Pt(8)
    run.font.color.rgb = RGBColor(0x64, 0x74, 0x8B)


def generar_reporte_word(criterios_validados, puntos_fuertes, puntos_mejorar,
                         esquema_detectado, porcentaje, cumplidos, dictamen, color_hex,
                         completitud, meta):
    doc = Document()
    doc.styles["Normal"].font.name = "Arial"
    doc.styles["Normal"].font.size = Pt(10)
    for s in doc.sections:
        s.left_margin = s.right_margin = s.top_margin = s.bottom_margin = Inches(0.5)
    p_enc = doc.add_paragraph()
    p_enc.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r_c = p_enc.add_run(f"{meta['centro']}\n"); r_c.bold = True; r_c.font.size = Pt(12)
    r_t = p_enc.add_run("INSTRUMENTO DE VALIDACIÓN TÉCNICO-PEDAGÓGICA\n"); r_t.bold = True; r_t.font.size = Pt(13)
    p_enc.add_run("Auditoría de Planificación de Clase Diaria (ETP) · Marco MINERD\n").italic = True
    agregar_linea_divisoria(doc)
    color_fondo = {"#10B981": "D1FAE5", "#F59E0B": "FEF3C7", "#EF4444": "FEE2E2"}.get(color_hex, "F1F5F9")
    t_res = doc.add_table(rows=1, cols=1); t_res.style = "Table Grid"
    celda = t_res.cell(0, 0); shade_cell(celda, color_fondo)
    p = celda.paragraphs[0]; p.add_run("RESUMEN EJECUTIVO\n").bold = True
    r_d = p.add_run(f"Puntaje: {porcentaje}% · Criterios: {cumplidos}/{TOTAL_CRITERIOS} · "
                    f"Completitud: {completitud}% · Dictamen: {dictamen}")
    r_d.bold = True; r_d.font.color.rgb = RGBColor.from_string(color_hex.lstrip("#"))
    fijar_anchos_columna(t_res, [7.5])
    doc.add_paragraph()
    t_datos = doc.add_table(rows=2, cols=4); t_datos.style = "Table Grid"
    t_datos.cell(0, 0).text = "Módulo:"; t_datos.cell(0, 1).text = meta["modulo"]
    t_datos.cell(0, 2).text = "Docente:"; t_datos.cell(0, 3).text = meta["docente"]
    t_datos.cell(1, 0).text = "Validador:"; t_datos.cell(1, 1).text = meta["evaluador"]
    t_datos.cell(1, 2).text = "Dictamen:"; t_datos.cell(1, 3).text = dictamen
    fijar_anchos_columna(t_datos, [1.6, 2.15, 1.6, 2.15])
    doc.add_paragraph()
    doc.add_heading("I. CORRESPONDENCIA CON EL ESQUEMA OFICIAL", level=2)
    t_esq = doc.add_table(rows=1, cols=3); t_esq.style = "Table Grid"
    for i, h in enumerate(["Sección del Esquema", "Estado", "Nota de Auditoría"]):
        t_esq.rows[0].cells[i].text = h
        t_esq.rows[0].cells[i].paragraphs[0].runs[0].bold = True
        shade_cell(t_esq.rows[0].cells[i], "DBEAFE")
    for item in esquema_detectado:
        row = t_esq.add_row().cells
        row[0].text = str(item.get("SECCION", ""))
        presente = bool(item.get("PRESENTE", False))
        row[1].text = "✅ Presente" if presente else "❌ Ausente"
        shade_cell(row[1], "D1FAE5" if presente else "FEE2E2")
        row[2].text = str(item.get("NOTA", ""))
    fijar_anchos_columna(t_esq, [3.2, 1.0, 3.3])
    doc.add_paragraph()
    doc.add_heading(f"II. MATRIZ DE CRITERIOS ({TOTAL_CRITERIOS} Puntos)", level=2)
    for cat in CATEGORIAS_ORDEN:
        items = [c for c in criterios_validados if c["CATEGORIA"] == cat]
        doc.add_paragraph(f"{cat} ({sum(1 for c in items if c['CUMPLE'])}/{len(items)})").runs[0].bold = True
        t_crit = doc.add_table(rows=1, cols=4); t_crit.style = "Table Grid"
        for i, h in enumerate(["No.", "CRITERIO", "CUMPLE", "OBSERVACIONES"]):
            t_crit.rows[0].cells[i].text = h
            t_crit.rows[0].cells[i].paragraphs[0].runs[0].bold = True
            shade_cell(t_crit.rows[0].cells[i], "DBEAFE")
        for item in items:
            row = t_crit.add_row().cells
            row[0].text = str(item["NO"]); row[1].text = item["CRITERIO"]
            row[2].text = "✅ SÍ" if item["CUMPLE"] else "❌ NO"
            shade_cell(row[2], "D1FAE5" if item["CUMPLE"] else "FEE2E2")
            row[3].text = item["OBSERVACION"]
        fijar_anchos_columna(t_crit, [0.4, 3.5, 0.8, 2.8])
        doc.add_paragraph()
    doc.add_heading("III. RETROALIMENTACIÓN", level=2)
    p_ret = doc.add_paragraph(); p_ret.add_run("Puntos Fuertes:\n").bold = True
    for pf in puntos_fuertes: p_ret.add_run(f"• {pf}\n")
    p_ret.add_run("\nÁreas de Mejora:\n").bold = True
    for pm in puntos_mejorar: p_ret.add_run(f"• {pm}\n")
    doc.add_paragraph("\n\n")
    t_f = doc.add_table(rows=2, cols=3)
    t_f.cell(0, 0).text = t_f.cell(0, 1).text = t_f.cell(0, 2).text = "_________________________"
    t_f.cell(1, 0).text = "Docente (Planificador)"; t_f.cell(1, 1).text = "Coordinación (Validador)"
    t_f.cell(1, 2).text = "Dirección Académica"
    for row in t_f.rows:
        for c in row.cells:
            c.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
    agregar_numeracion_pagina(doc)   # ← mejora 5
    buf = BytesIO(); doc.save(buf); buf.seek(0)
    return buf


# ═══════════════════════════════════════════════════════════════════════════
# INTERFAZ
# ═══════════════════════════════════════════════════════════════════════════
st.markdown("""
<style>
@import url('https://fonts.googleapis.com/css2?family=Inter:wght@300;400;500;600;700;800&display=swap');
html, body, [class*="css"] { font-family: 'Inter', sans-serif; background-color: #F8FAFC; color: #1E293B; }
.eval-hero { background: linear-gradient(135deg, #0F172A 0%, #1E3A8A 55%, #7C3AED 100%); color: #fff;
padding: 1.8rem; border-radius: 18px; margin-bottom: 1.25rem; box-shadow: 0 18px 35px rgba(15,23,42,.18); }
.eval-title { font-size: 2rem; font-weight: 800; }
.eval-sub { opacity: .88; font-size: 1rem; margin-top: .3rem; }
.section-title { color: #1D4ED8; font-weight: 700; font-size: 1.12rem; border-bottom: 2px solid #DBEAFE;
padding-bottom: 8px; margin: 1.2rem 0 .9rem 0; }
.metric-card { background:#fff; border:1px solid #E2E8F0; border-top:4px solid #2563EB; border-radius:12px;
padding:16px; box-shadow:0 4px 12px rgba(15,23,42,.06); text-align:center; }
.metric-value { font-size:2.2rem; font-weight:800; color:#0F172A; }
.metric-label { font-size:.8rem; font-weight:800; color:#64748B; text-transform:uppercase; }
.info-box { background:#EFF6FF; border-left:4px solid #2563EB; padding:12px 16px; border-radius:6px; margin:12px 0; }
.cat-card { background:#fff; border:1px solid #E2E8F0; border-left:4px solid #2563EB; border-radius:6px; padding:12px 16px; margin-bottom:10px; }
.cat-title { font-weight:700; margin-bottom:6px; }
.crit-ok { color:#059669; font-size:.85rem; margin-top:3px; }
.crit-no { color:#DC2626; font-size:.85rem; margin-top:3px; }
</style>
""", unsafe_allow_html=True)

ia.panel_sidebar_ia("Auditor de Planificaciones")
st.markdown("""
<div class="eval-hero">
    <div class="eval-title">✅ Auditoría de Planificaciones Diarias ETP</div>
    <div class="eval-sub">Valida el esquema oficial MINERD · 14 criterios · dictamen · historial persistente</div>
</div>
""", unsafe_allow_html=True)

if "audit_results" not in st.session_state:
    st.session_state.audit_results = []

tab_audit, tab_hist = st.tabs(["✅ Auditar Planificación", "📋 Historial"])

with tab_audit:
    with st.form("form_evaluacion", clear_on_submit=False):
        st.markdown('<div class="section-title">👤 1. Datos del Docente y Módulo</div>', unsafe_allow_html=True)
        col1, col2 = st.columns(2)
        with col1:
            # Mejora 5: auto-relleno desde sesión
            docente = st.text_input("Nombre del Docente Evaluado",
                                    value=st.session_state.get("nombre_docente", "") if st.session_state.get("docente_autenticado") else "")
            centro = st.text_input("Centro Educativo", value="Politécnico Salesiano Arquides Calderón")
        with col2:
            modulo = st.text_input("Módulo Formativo", placeholder="Ej: MF 358-3 ...")
            evaluador = st.text_input("Validador Técnico / Coordinador",
                                      value=st.session_state.get("coordinador_nombre", "Ing. Bernardo Hernández"))
        st.markdown('<div class="section-title">📄 2. Cargar Planificación (una o varias)</div>', unsafe_allow_html=True)
        archivos_plan = st.file_uploader("Sube el documento (PDF o Word) — puedes seleccionar varios",
                                         type=["pdf", "docx"], accept_multiple_files=True)
        planificacion_texto = st.text_area("Texto de la planificación (opcional si no subes archivo):", height=100)
        max_tokens, temperature = ia.control_avanzado(default_tokens=16384, tope=32000, default_temp=0.1)
        modo_debug = st.checkbox("🐛 Modo depuración", value=False)
        submit_button = st.form_submit_button("⚙️ Auditar Planificación(es)", type="primary", use_container_width=True)

    if submit_button:
        cfg = ia.config_ia()
        fuentes = []
        if archivos_plan:
            for a in archivos_plan:
                try:
                    fuentes.append((a.name, extraer_texto_plan(a)))
                except Exception as e:
                    st.warning(f"⚠️ No se pudo leer {a.name}: {e}")
        elif planificacion_texto.strip():
            fuentes.append(("Texto manual", planificacion_texto))
        if not cfg["api_key"]:
            st.error("🔒 Debes ingresar tu API Key en la barra lateral.")
        elif not docente or not modulo:
            st.warning("📝 Completa el nombre del docente y el módulo.")
        elif not fuentes:
            st.warning("⚠️ Sube al menos un archivo o pega el texto de la planificación.")
        else:
            resultados = []
            prog = st.progress(0.0)
            for i, (nombre_fuente, texto_evaluar) in enumerate(fuentes):
                if len(texto_evaluar) > 60000:
                    texto_evaluar = texto_evaluar[:60000]
                with st.spinner(f"🧠 Auditando {nombre_fuente} ({i+1}/{len(fuentes)})..."):
                    texto_crudo = None
                    try:
                        lista_esquema = "\n".join(f"- {s}" for s in ESQUEMA_OFICIAL)
                        prompt_maestro = f"""Actúa como un Validador Técnico-Pedagógico Nivel Máster de la ETP (MINERD).
Audita UNA PLANIFICACIÓN DE CLASE DIARIA (ETP) contra su ESQUEMA OFICIAL y los {TOTAL_CRITERIOS} criterios oficiales.
CONTENIDO EXTRAÍDO DE LA PLANIFICACIÓN:
{texto_evaluar}
ESQUEMA OFICIAL (verifica PRESENTE/AUSENTE):
{lista_esquema}
{TOTAL_CRITERIOS} CRITERIOS A EVALUAR:
{construir_texto_criterios()}
INSTRUCCIONES:
- Para cada sección del esquema indica PRESENTE true/false y una NOTA breve.
- Evalúa CUMPLE true/false cada criterio con observación constructiva.
- Devuelve los {TOTAL_CRITERIOS} criterios en orden con campo "NO" del 1 al {TOTAL_CRITERIOS}.
- Extrae 3 puntos fuertes y los puntos a mejorar.
CODIFICACIÓN: salto de línea → {ia.MARKER_NL} · comilla doble → {ia.MARKER_DQ} · tabulación → {ia.MARKER_TAB}.
FORMATO (JSON NATIVO):
{{
 "ESQUEMA_DETECTADO": [{{"SECCION": "...", "PRESENTE": true, "NOTA": "..."}}],
 "PUNTOS_FUERTES": ["..."],
 "PUNTOS_A_MEJORAR": ["..."],
 "EVALUACION_CRITERIOS": [{{"NO": 1, "CRITERIO": "...", "CUMPLE": true, "OBSERVACION": "..."}}]
}}
"""
                        texto_crudo, flags = ia.solicitar_ia(prompt_maestro, modo="json",
                                                             max_tokens=max_tokens, temperature=temperature,
                                                             modulo="evaluador_planes")
                        datos = ia.decodificar_marcadores(ia.parsear_json_robusto(texto_crudo))
                        criterios_validados, faltantes = validar_criterios(datos.get("EVALUACION_CRITERIOS", []))
                        puntos_fuertes = datos.get("PUNTOS_FUERTES", []) or []
                        puntos_mejorar = datos.get("PUNTOS_A_MEJORAR", []) or []
                        esquema_detectado = datos.get("ESQUEMA_DETECTADO", []) or []
                        cumplidos = sum(1 for c in criterios_validados if c["CUMPLE"])
                        porcentaje = round((cumplidos / TOTAL_CRITERIOS) * 100)
                        dictamen, color_dictamen = calcular_dictamen(cumplidos)
                        presentes = sum(1 for s in esquema_detectado if bool(s.get("PRESENTE", False)))
                        completitud = round((presentes / len(ESQUEMA_OFICIAL)) * 100) if esquema_detectado else 0
                        insertar_auditoria(datetime.date.today().isoformat(), docente, modulo, evaluador,
                                           porcentaje, cumplidos, dictamen, completitud)
                        buffer = generar_reporte_word(criterios_validados, puntos_fuertes, puntos_mejorar,
                                                      esquema_detectado, porcentaje, cumplidos, dictamen,
                                                      color_dictamen, completitud,
                                                      {"centro": centro, "modulo": modulo,
                                                       "docente": docente, "evaluador": evaluador})
                        resultados.append({
                            "fuente": nombre_fuente, "docente": docente, "modulo": modulo,
                            "porcentaje": porcentaje, "cumplidos": cumplidos, "dictamen": dictamen,
                            "color": color_dictamen, "completitud": completitud,
                            "criterios": criterios_validados, "buffer": buffer.getvalue(),
                        })
                    except ValueError as ve:
                        ia.render_error_ia(ve, texto_crudo)
                    except Exception as e:
                        ia.render_error_ia(e, texto_crudo)
                prog.progress((i + 1) / len(fuentes))
            if resultados:
                st.session_state.audit_results = resultados
                st.toast(f"✅ {len(resultados)} auditoría(s) completada(s).", icon="✅")
                st.rerun()

    # ── Resultados en pantalla ──
    if st.session_state.audit_results:
        st.markdown('<div class="section-title">📊 Resultados de la Auditoría</div>', unsafe_allow_html=True)
        for res in st.session_state.audit_results:
            st.markdown(f"#### 📄 {res['fuente']}")
            c1, c2, c3, c4 = st.columns(4)
            with c1:
                st.markdown(f'<div class="metric-card"><div class="metric-label">Puntaje</div><div class="metric-value">{res["porcentaje"]}%</div></div>', unsafe_allow_html=True)
            with c2:
                st.markdown(f'<div class="metric-card"><div class="metric-label">Criterios</div><div class="metric-value">{res["cumplidos"]}/{TOTAL_CRITERIOS}</div></div>', unsafe_allow_html=True)
            with c3:
                st.markdown(f'<div class="metric-card" style="border-top-color:{res["color"]};"><div class="metric-label">Dictamen</div><div class="metric-value" style="color:{res["color"]};font-size:1.4rem;">{res["dictamen"]}</div></div>', unsafe_allow_html=True)
            with c4:
                st.markdown(f'<div class="metric-card"><div class="metric-label">Completitud</div><div class="metric-value">{res["completitud"]}%</div></div>', unsafe_allow_html=True)
            # Mejora 3: gráfico por categoría
            filas_cat = []
            for cat in CATEGORIAS_ORDEN:
                items = [c for c in res["criterios"] if c["CATEGORIA"] == cat]
                filas_cat.append({"Categoría": cat.split(". ")[-1][:30],
                                  "Cumplidos": sum(1 for c in items if c["CUMPLE"]),
                                  "Total": len(items)})
            st.bar_chart(pd.DataFrame(filas_cat).set_index("Categoría"))
            with st.expander("📋 Detalle por categoría"):
                for cat in CATEGORIAS_ORDEN:
                    items = [c for c in res["criterios"] if c["CATEGORIA"] == cat]
                    html_rows = "".join(
                        f"<div class='{'crit-ok' if c['CUMPLE'] else 'crit-no'}'>"
                        f"{'✅' if c['CUMPLE'] else '❌'} <b>{c['NO']}. {c['CRITERIO']}</b> — {c['OBSERVACION']}</div>"
                        for c in items)
                    st.markdown(f"<div class='cat-card'><div class='cat-title'>{CATEGORIA_ICONOS[cat]} {cat}</div>{html_rows}</div>",
                                unsafe_allow_html=True)
            # Mejora 2: acompañamiento si No Aceptable
            if res["dictamen"] == "No Aceptable":
                if st.button(f"📋 Registrar Acompañamiento para {res['docente']}", key=f"acomp_{res['fuente']}"):
                    if registrar_acompanamiento(res["docente"], res["modulo"], res["dictamen"]):
                        st.success("✅ Acompañamiento registrado en el cronograma.")
                    else:
                        st.info("ℹ️ No se pudo registrar automáticamente; usa el módulo de Acompañamiento.")
            # Mejora 6: descargas
            cd1, cd2 = st.columns(2)
            with cd1:
                st.download_button("📥 Descargar Reporte Word (.docx)", data=res["buffer"],
                                   file_name=f"Auditoria_PlanDiario_{ia.sanear_nombre_archivo(res['docente'])}.docx",
                                   mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                                   type="primary", use_container_width=True, key=f"word_{res['fuente']}")
            with cd2:
                csv = pd.DataFrame(res["criterios"]).to_csv(index=False).encode("utf-8-sig")
                st.download_button("📄 Descargar Criterios (.csv)", data=csv,
                                   file_name=f"Criterios_{ia.sanear_nombre_archivo(res['docente'])}.csv",
                                   mime="text/csv", use_container_width=True, key=f"csv_{res['fuente']}")

with tab_hist:
    st.markdown('<div class="section-title">🗄️ Historial de Auditorías</div>', unsafe_allow_html=True)
    hist = listar_auditorias()
    if not hist:
        st.info("Aún no hay auditorías registradas.")
    else:
        df_hist = pd.DataFrame(hist)
        col_f1, col_f2 = st.columns(2)
        with col_f1:
            f_doc = st.selectbox("Filtrar por docente", ["Todos"] + sorted(df_hist["docente"].unique().tolist()))
        with col_f2:
            f_dic = st.selectbox("Filtrar por dictamen", ["Todos", "Excelente", "Aceptable con Mejoras", "No Aceptable"])
        df_f = df_hist.copy()
        if f_doc != "Todos":
            df_f = df_f[df_f["docente"] == f_doc]
        if f_dic != "Todos":
            df_f = df_f[df_f["dictamen"] == f_dic]
        st.dataframe(df_f.drop(columns=["id"]), use_container_width=True, hide_index=True)
        # Mejora 8: comparativa por docente
        st.markdown("#### 🏆 Ranking histórico por docente")
        ranking = df_hist.groupby("docente").agg(
            Auditorías=("id", "count"),
            Puntaje_Promedio=("porcentaje", "mean")).sort_values("Puntaje_Promedio", ascending=False).round(1)
        st.dataframe(ranking, use_container_width=True)
        csv_hist = df_f.to_csv(index=False).encode("utf-8-sig")
        st.download_button("📄 Exportar historial (.csv)", data=csv_hist,
                           file_name="Historial_Auditorias.csv", mime="text/csv", use_container_width=True)